"""
Project service — pure Python, no PyQt6 imports.

Provides high-level operations for scanning folders, parsing notes,
reading/writing project metadata, and querying file metadata.

v2: SQLite as primary store. JSON on network drive as portable export.
The DB lives at %APPDATA%/PhysioMetrics/PhysioMetrics.db (local only).
JSON files are auto-merged on open and exported on save.
"""

import re
import json
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple, Set
from datetime import datetime

from core.project_manager import ProjectManager


# Metadata column keys for the master file list
METADATA_COLUMNS = [
    'file_name', 'file_path', 'experiment', 'strain', 'stim_type',
    'power', 'sex', 'animal_id', 'protocol', 'channel', 'stim_channel',
    'events_channel', 'channel_count', 'sweep_count', 'keywords',
    'status', 'exports', 'linked_notes',
]


class ProjectService:
    """
    Pure-Python service for project metadata operations.

    Uses SQLite as primary store (ProjectStoreSQLite in %APPDATA%).
    JSON .physiometrics files on network drives are portable exports,
    auto-merged on open and exported on save.
    """

    def __init__(self, project_manager: Optional[ProjectManager] = None):
        self._pm = project_manager or ProjectManager()
        self._store = None  # Lazy-loaded ProjectStoreSQLite
        self._project_id: Optional[int] = None
        self._project_path: Optional[Path] = None  # .physiometrics JSON path
        self._project_name: str = ""
        self._data_directory: Optional[Path] = None
        self._experiments: List[Dict] = []
        self._notes_files_data: List[Dict] = []
        self._dirty = False
        self._last_merge_report: Optional[Dict] = None

        # Legacy compatibility: in-memory file list (populated from DB)
        self._files: List[Dict[str, Any]] = []

    # ------------------------------------------------------------------
    # Store (lazy-loaded)
    # ------------------------------------------------------------------

    @property
    def store(self):
        """Lazy-load the SQLite store."""
        if self._store is None:
            from core.adapters.project_store_sqlite import ProjectStoreSQLite
            self._store = ProjectStoreSQLite()
        return self._store

    # ------------------------------------------------------------------
    # Project open / save / load
    # ------------------------------------------------------------------

    def open_project(self, folder: Path) -> Dict[str, Any]:
        """
        Open or create a project for a folder.

        1. Upsert project in SQLite DB by data_directory
        2. If .physiometrics JSON exists, merge it into DB
        3. Populate in-memory _files from DB

        Returns:
            Summary dict with project_name, file_count, data_directory, merge_report.
        """
        folder = Path(folder)
        if not folder.is_dir():
            raise ValueError(f"Not a directory: {folder}")

        # Look for existing .physiometrics file
        pm_files = list(folder.glob("*.physiometrics"))
        json_path = None
        if pm_files:
            pm_files.sort(key=lambda p: p.stat().st_mtime, reverse=True)
            json_path = pm_files[0]

        project_name = folder.name
        if json_path:
            try:
                with open(json_path, 'r') as f:
                    json_data = json.load(f)
                project_name = json_data.get("project_name", folder.name)
            except Exception:
                json_data = None
        else:
            json_data = None

        # Upsert project in DB
        self._project_id = self.store.upsert_project(
            name=project_name,
            data_directory=str(folder),
            json_path=str(json_path) if json_path else None,
        )
        self._project_name = project_name
        self._data_directory = folder
        self._project_path = json_path

        # Merge JSON into DB if it exists and is newer
        merge_report = None
        if json_path and json_data:
            merge_report = self._sync_json(json_path, json_data)

        # Populate in-memory file list from DB
        self._refresh_files_from_db()
        self._dirty = False
        self._last_merge_report = merge_report

        result = {
            "project_name": self._project_name,
            "data_directory": str(folder),
            "file_count": len(self._files),
            "status": "loaded" if json_path else "created_new",
        }
        if merge_report:
            result["merge_report"] = merge_report

        return result

    def load_project(self, project_path: Path) -> Dict[str, Any]:
        """Load a project from a .physiometrics file (legacy compat)."""
        project_path = Path(project_path)
        folder = project_path.parent
        return self.open_project(folder)

    def save_project(self, name: Optional[str] = None) -> str:
        """
        Save the current project.

        1. Flush any in-memory changes to DB
        2. Export DB -> JSON on network drive
        3. Update JSON sync metadata

        Returns path to saved .physiometrics file.
        """
        if self._data_directory is None:
            raise RuntimeError("No project open — call open_project() first")

        project_name = name or self._project_name or self._data_directory.name
        self._project_name = project_name

        # Update project name in DB
        if self._project_id:
            self.store.upsert_project(
                name=project_name,
                data_directory=str(self._data_directory),
                json_path=str(self._project_path) if self._project_path else None,
            )

        # Export DB -> JSON
        json_data = self.store.export_to_json(self._project_id)
        json_data["experiments"] = self._experiments
        json_data["notes_files"] = self._notes_files_data

        # Write JSON to network drive via ProjectManager
        self._project_path = self._pm.save_project(
            project_name=project_name,
            data_directory=self._data_directory,
            files_data=json_data["files"],
            experiments=self._experiments,
            notes_files=self._notes_files_data,
        )

        # Update sync metadata
        if self._project_path and self._project_path.exists():
            mtime = self._project_path.stat().st_mtime
            file_hash = self.file_hash(self._project_path)
            self.store.update_project_json_sync(
                self._project_id, mtime, file_hash,
            )

        self._dirty = False
        return str(self._project_path)

    def _sync_json(
        self, json_path: Path, json_data: Optional[Dict] = None,
    ) -> Optional[Dict[str, Any]]:
        """
        Sync a .physiometrics JSON file into the DB.

        Compares mtime/hash to skip unnecessary re-imports.
        Returns merge report or None if no sync needed.
        """
        if not json_path.exists():
            return None

        project = self.store.get_project(self._project_id)
        if not project:
            return None

        try:
            current_mtime = json_path.stat().st_mtime
            current_hash = self.file_hash(json_path)
        except OSError:
            return None

        # Skip if already synced
        if (project["json_hash"] == current_hash and
                project["json_mtime"] == current_mtime):
            return None

        # Backup before merge
        try:
            self.store.backup(trigger_event="merge")
        except Exception as e:
            print(f"[project-service] Backup before merge failed: {e}")

        # Load JSON if not provided
        if json_data is None:
            try:
                with open(json_path, 'r') as f:
                    json_data = json.load(f)
            except Exception as e:
                print(f"[project-service] Error reading JSON: {e}")
                return None

        # Import/merge
        report = self.store.import_from_json(
            self._project_id, json_data, self._data_directory,
        )

        # Store experiments and notes from JSON
        self._experiments = json_data.get("experiments", [])
        self._notes_files_data = json_data.get("notes_files", [])

        # Update sync metadata
        self.store.update_project_json_sync(
            self._project_id, current_mtime, current_hash,
        )

        return report

    def sync_json(self) -> Optional[Dict[str, Any]]:
        """Public method to trigger JSON sync (called by ViewModel on file change)."""
        if self._project_path and self._project_path.exists():
            report = self._sync_json(self._project_path)
            if report:
                self._refresh_files_from_db()
            return report
        return None

    def _refresh_files_from_db(self):
        """Reload in-memory _files list from DB."""
        if self._project_id is None:
            self._files = []
            return

        files, _ = self.store.get_project_files(self._project_id)

        # Convert to format expected by legacy code (absolute paths)
        for f in files:
            rel_path = f.get('file_path', '')
            if self._data_directory and rel_path:
                abs_path = self._data_directory / rel_path
                f['file_path'] = str(abs_path)

        self._files = files

    @property
    def is_dirty(self) -> bool:
        return self._dirty

    @property
    def project_path(self) -> Optional[Path]:
        return self._project_path

    @property
    def data_directory(self) -> Optional[Path]:
        return self._data_directory

    @property
    def project_id(self) -> Optional[int]:
        return self._project_id

    @property
    def last_merge_report(self) -> Optional[Dict]:
        return self._last_merge_report

    # ------------------------------------------------------------------
    # Relative / absolute path helpers
    # ------------------------------------------------------------------

    def _to_relative(self, abs_path: str) -> str:
        """Convert an absolute path to relative from data_directory."""
        if not self._data_directory or not abs_path:
            return abs_path
        try:
            rel = Path(abs_path).relative_to(self._data_directory)
            return str(rel).replace("\\", "/")
        except (ValueError, TypeError):
            return abs_path

    def _to_absolute(self, rel_or_abs: str) -> str:
        """Resolve a relative-or-absolute path against data_directory."""
        if not rel_or_abs:
            return rel_or_abs
        p = Path(rel_or_abs)
        if p.is_absolute():
            return str(p)
        if self._data_directory:
            return str(self._data_directory / p)
        return rel_or_abs

    def _resolve_file_id(self, file_path: str) -> Optional[int]:
        """Resolve a file path (relative or absolute) to a file_id in the DB."""
        if self._project_id is None:
            return None

        # Try relative path first
        rel_path = self._to_relative(file_path)
        fid = self.store.get_file_id_by_path(self._project_id, rel_path)
        if fid:
            return fid

        # Try with forward slashes
        rel_path_fwd = rel_path.replace("\\", "/")
        if rel_path_fwd != rel_path:
            fid = self.store.get_file_id_by_path(self._project_id, rel_path_fwd)
            if fid:
                return fid

        # Try resolving absolute path then making relative
        abs_path = self._to_absolute(file_path)
        rel_path2 = self._to_relative(abs_path)
        if rel_path2 != rel_path:
            fid = self.store.get_file_id_by_path(self._project_id, rel_path2)
            if fid:
                return fid

        return None

    # ------------------------------------------------------------------
    # Scanning
    # ------------------------------------------------------------------

    def scan_folder(
        self,
        root: Optional[Path] = None,
        recursive: bool = True,
        file_types: Optional[List[str]] = None,
        progress_callback=None,
    ) -> List[Dict[str, Any]]:
        """
        Discover recording files in a folder tree.

        Returns list of newly discovered file dicts (not yet added to project).
        """
        from core import project_builder
        from core.fast_abf_reader import extract_path_keywords

        root = Path(root) if root else self._data_directory
        if root is None:
            raise RuntimeError("No folder to scan — pass root or open a project first")

        files = project_builder.discover_files(
            str(root), recursive=recursive, file_types=file_types
        )

        abf_files = files.get("abf_files", [])
        smrx_files = files.get("smrx_files", [])
        edf_files = files.get("edf_files", [])
        photometry_files = files.get("photometry_files", [])

        all_data_files = abf_files + smrx_files + edf_files

        # Build set of existing paths (from DB)
        existing_paths = set()
        if self._project_id:
            db_files, _ = self.store.get_project_files(self._project_id)
            for f in db_files:
                fp = f.get("file_path", "")
                if fp and self._data_directory:
                    abs_p = str((self._data_directory / fp).resolve())
                    existing_paths.add(abs_p)

        new_entries = []
        total = len(all_data_files) + len(photometry_files)
        current = 0

        for data_path in all_data_files:
            normalized = str(data_path.resolve())
            if normalized in existing_paths:
                current += 1
                continue

            file_type = data_path.suffix.lower().lstrip(".")
            path_info = extract_path_keywords(data_path, root)

            keywords_display = []
            if path_info.get("relative_path"):
                keywords_display.append(path_info["relative_path"])
            elif path_info["subdirs"]:
                keywords_display.append("/".join(path_info["subdirs"]))
            if path_info["power_levels"]:
                keywords_display.extend(path_info["power_levels"])
            if path_info["animal_ids"]:
                keywords_display.extend([f"ID:{aid}" for aid in path_info["animal_ids"]])

            entry = {
                "file_path": str(data_path),
                "file_name": data_path.name,
                "file_type": file_type,
                "protocol": "",
                "channel_count": 0,
                "sweep_count": 0,
                "stim_channels": [],
                "stim_frequency": "",
                "path_keywords": path_info,
                "keywords_display": ", ".join(keywords_display) if keywords_display else "",
                "channel": "",
                "stim_channel": "",
                "events_channel": "",
                "experiment": "",
                "strain": "",
                "stim_type": "",
                "power": path_info["power_levels"][0] if path_info["power_levels"] else "",
                "sex": "",
                "animal_id": path_info["animal_ids"][0] if path_info["animal_ids"] else "",
                "status": "pending",
                "exports": {},
                "linked_notes": "",
            }
            new_entries.append(entry)
            existing_paths.add(normalized)
            current += 1

            if progress_callback and current % 20 == 0:
                progress_callback(current, total, f"Discovering files... {current}/{total}")

        # Photometry files
        for fp_path in photometry_files:
            normalized = str(fp_path.resolve())
            if normalized in existing_paths:
                current += 1
                continue

            fp_info = project_builder.extract_photometry_info(fp_path)
            if fp_info:
                path_info = extract_path_keywords(fp_path, root)
                keywords_display = []
                if path_info.get("relative_path"):
                    keywords_display.append(path_info["relative_path"])

                regions_str = ", ".join(fp_info.get("signal_columns", [])) or "Unknown"
                channel_info = f"{fp_info.get('region_count', 0)} regions ({regions_str})"

                entry = {
                    "file_path": str(fp_path),
                    "file_name": fp_info.get("file_name", fp_path.name),
                    "file_type": "photometry",
                    "protocol": fp_info.get("protocol", "Neurophotometrics"),
                    "channel_count": fp_info.get("region_count", 0),
                    "sweep_count": 1,
                    "channel": channel_info,
                    "stim_channel": "",
                    "events_channel": "",
                    "experiment": "",
                    "strain": "",
                    "stim_type": "",
                    "power": path_info["power_levels"][0] if path_info["power_levels"] else "",
                    "sex": "",
                    "animal_id": path_info["animal_ids"][0] if path_info["animal_ids"] else "",
                    "status": "pending",
                    "exports": {},
                    "linked_notes": "",
                    "keywords_display": ", ".join(keywords_display) if keywords_display else "",
                }
                new_entries.append(entry)
                existing_paths.add(normalized)

            current += 1
            if progress_callback:
                progress_callback(current, total, f"Discovering files... {current}/{total}")

        if progress_callback:
            progress_callback(total, total, f"Found {len(new_entries)} new files")

        return new_entries

    def load_file_metadata(
        self,
        file_entries: Optional[List[Dict]] = None,
        progress_callback=None,
    ) -> int:
        """
        Load detailed metadata (protocol, channels, sweeps) for file entries.

        Operates on self._files if file_entries is None.
        """
        from core.fast_abf_reader import read_file_metadata_fast

        targets = file_entries if file_entries is not None else self._files
        updated = 0
        total = len(targets)

        for i, entry in enumerate(targets):
            file_path = Path(entry.get("file_path", ""))
            if not file_path.exists():
                continue
            if entry.get("file_type") == "photometry":
                continue
            if entry.get("protocol") and entry["protocol"] not in ("", "Loading..."):
                continue

            metadata = read_file_metadata_fast(file_path)
            if metadata:
                entry["protocol"] = metadata.get("protocol", "Unknown")
                entry["channel_count"] = metadata.get("channel_count", 0)
                entry["sweep_count"] = metadata.get("sweep_count", 0)
                entry["stim_channels"] = metadata.get("stim_channels", [])
                entry["stim_frequency"] = metadata.get("stim_frequency", "")
                if metadata.get("stim_frequency") and not entry.get("stim_type"):
                    entry["stim_type"] = metadata["stim_frequency"]
                if metadata.get("stim_channels") and not entry.get("events_channel"):
                    entry["events_channel"] = ", ".join(metadata["stim_channels"])
                updated += 1
                self._dirty = True

                # Also update DB
                rel_path = self._to_relative(str(file_path))
                fid = self._resolve_file_id(str(file_path))
                if fid:
                    self.store.update_file(fid, {
                        "protocol": entry["protocol"],
                        "channel_count": entry["channel_count"],
                        "sweep_count": entry["sweep_count"],
                        "stim_type": entry.get("stim_type", ""),
                        "events_channel": entry.get("events_channel", ""),
                    }, record_field_timestamps=False)

            if progress_callback and (i % 10 == 0 or i == total - 1):
                progress_callback(i + 1, total, f"Loading metadata... {i + 1}/{total}")

        return updated

    # ------------------------------------------------------------------
    # Notes file reading and matching
    # ------------------------------------------------------------------

    def read_notes_file(self, path: Path, max_rows: int = 5) -> List[Dict[str, Any]]:
        """
        Read and parse a notes file (Excel/CSV/TXT/DOCX).

        Checks DB notes cache first. Otherwise parses and caches.
        """
        path = Path(path)

        # Check cache
        if max_rows == 5 and self._project_id:
            cached = self._get_cached_notes(path)
            if cached is not None:
                return cached

        suffix = path.suffix.lower()
        entries = []

        try:
            if suffix in (".xlsx", ".xls"):
                entries = self._parse_excel_notes(path, max_rows=max_rows)
            elif suffix == ".csv":
                entries = self._parse_csv_notes(path, max_rows=max_rows)
            elif suffix == ".txt":
                entries = self._parse_text_notes(path)
            elif suffix == ".docx":
                entries = self._parse_docx_notes(path)
        except Exception as e:
            print(f"[project-service] Error reading notes file {path}: {e}")
            return [{"error": str(e), "path": str(path)}]

        # Cache the result
        if self._project_id:
            all_refs = []
            for entry in entries:
                all_refs.extend(entry.get("abf_references", []))
            current_hash = self.file_hash(path)
            self.store.save_notes_cache(
                self._project_id, str(path), current_hash,
                entries, sorted(set(all_refs)),
            )

        return entries

    def _get_cached_notes(self, path: Path) -> Optional[List[Dict]]:
        """Check DB cache for parsed notes."""
        try:
            current_hash = self.file_hash(path)
            cached = self.store.get_notes_cache(self._project_id, current_hash)
            if cached:
                return cached["content"]
        except (OSError, Exception):
            pass
        return None

    def _parse_excel_notes(self, path: Path, max_rows: int = 5) -> List[Dict]:
        """Parse Excel file, returning rows as dicts with ABF refs detected."""
        import pandas as pd

        entries = []
        sheets = pd.read_excel(path, sheet_name=None, nrows=1000)

        for sheet_name, df in sheets.items():
            abf_refs = set()
            for col in df.columns:
                for val in df[col].astype(str):
                    refs = self._extract_abf_references(val)
                    abf_refs.update(refs)

            if max_rows <= 0:
                preview = df.to_dict(orient="records")
            else:
                preview = df.head(max_rows).to_dict(orient="records")

            entries.append({
                "source_file": str(path),
                "source_name": path.name,
                "sheet_name": sheet_name,
                "row_count": len(df),
                "columns": list(df.columns),
                "abf_references": sorted(abf_refs),
                "content_preview": preview,
            })

        return entries

    def _parse_csv_notes(self, path: Path, max_rows: int = 5) -> List[Dict]:
        """Parse CSV file."""
        import pandas as pd

        df = pd.read_csv(path, nrows=1000, encoding="utf-8", on_bad_lines="skip")

        abf_refs = set()
        for col in df.columns:
            for val in df[col].astype(str):
                refs = self._extract_abf_references(val)
                abf_refs.update(refs)

        if max_rows <= 0:
            preview = df.to_dict(orient="records")
        else:
            preview = df.head(max_rows).to_dict(orient="records")

        return [{
            "source_file": str(path),
            "source_name": path.name,
            "row_count": len(df),
            "columns": list(df.columns),
            "abf_references": sorted(abf_refs),
            "content_preview": preview,
        }]

    def _parse_text_notes(self, path: Path) -> List[Dict]:
        """Parse plain text file."""
        content = path.read_text(encoding="utf-8", errors="replace")
        abf_refs = self._extract_abf_references(content)

        return [{
            "source_file": str(path),
            "source_name": path.name,
            "content": content[:5000],
            "line_count": content.count("\n") + 1,
            "abf_references": sorted(abf_refs),
        }]

    def _parse_docx_notes(self, path: Path) -> List[Dict]:
        """Parse Word document."""
        try:
            from docx import Document
        except ImportError:
            return [{"error": "python-docx not installed", "path": str(path)}]

        doc = Document(str(path))

        content_parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                content_parts.append(p.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                content_parts.append(row_text)

        content = "\n".join(content_parts)
        abf_refs = self._extract_abf_references(content)

        return [{
            "source_file": str(path),
            "source_name": path.name,
            "content": content[:5000],
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "abf_references": sorted(abf_refs),
        }]

    @staticmethod
    def _extract_abf_references(text: str) -> Set[str]:
        """
        Extract ABF file stem references from text.

        Supports multiple reference styles:
        - Full filename: "2024_03_21_0017.abf"
        - 6-10 digit numbers: "25708007"
        - Date_sequence patterns: "2024_03_21_0017"
        - Short numeric references: 3-5 digit numbers near ABF context
        """
        refs = set()

        abf_matches = re.findall(r"(\w+)\.abf", text, re.IGNORECASE)
        refs.update(abf_matches)

        date_seq = re.findall(r"\b(\d{4}_\d{2}_\d{2}_\d{4})\b", text)
        refs.update(date_seq)

        num_matches = re.findall(r"\b(\d{6,10})\b", text)
        refs.update(num_matches)

        short_matches = re.findall(r"(?:file|abf|recording|#)\s*(\d{2,5})\b", text, re.IGNORECASE)
        refs.update(short_matches)

        range_matches = re.findall(r"(?:file|recording|abf)s?\s+(\d+)\s*[-\u2013]\s*(\d+)", text, re.IGNORECASE)
        for start, end in range_matches:
            try:
                s, e = int(start), int(end)
                if e - s < 50:
                    for i in range(s, e + 1):
                        refs.add(str(i).zfill(len(start)))
            except ValueError:
                pass

        return refs

    def match_notes_to_files(
        self,
        notes_entries: List[Dict],
        fuzzy_range: int = 5,
    ) -> List[Dict[str, Any]]:
        """Match notes entries to recording files."""
        matches = []

        stem_to_files: Dict[str, List[Dict]] = {}
        suffix_to_files: Dict[str, List[Dict]] = {}

        for f in self._files:
            stem = Path(f.get("file_name", "")).stem
            if stem:
                stem_to_files.setdefault(stem, []).append(f)
                trailing = re.search(r"(\d{3,5})$", stem)
                if trailing:
                    suffix = trailing.group(1)
                    suffix_to_files.setdefault(suffix, []).append(f)
                    if len(suffix) >= 4:
                        suffix_to_files.setdefault(suffix[-3:], []).append(f)
                    if len(suffix) >= 5:
                        suffix_to_files.setdefault(suffix[-4:], []).append(f)

        seen_matches = set()

        for note in notes_entries:
            abf_refs = note.get("abf_references", [])
            source = note.get("source_name", "unknown")

            for ref in abf_refs:
                if ref in stem_to_files:
                    for f in stem_to_files[ref]:
                        key = (f.get("file_path", ""), ref)
                        if key not in seen_matches:
                            seen_matches.add(key)
                            matches.append({
                                "file_path": f.get("file_path", ""),
                                "file_name": f.get("file_name", ""),
                                "match_type": "exact",
                                "matched_ref": ref,
                                "notes_source": source,
                            })
                    continue

                if ref.isdigit() and len(ref) <= 5 and ref in suffix_to_files:
                    for f in suffix_to_files[ref]:
                        key = (f.get("file_path", ""), ref)
                        if key not in seen_matches:
                            seen_matches.add(key)
                            matches.append({
                                "file_path": f.get("file_path", ""),
                                "file_name": f.get("file_name", ""),
                                "match_type": "suffix",
                                "matched_ref": ref,
                                "notes_source": source,
                            })
                    continue

                numeric_parts = re.findall(r"\d+", ref)
                if not numeric_parts:
                    continue

                main_numeric = max(numeric_parts, key=len)
                numeric_idx = ref.find(main_numeric)

                try:
                    base_num = int(main_numeric)
                except ValueError:
                    continue

                for delta in range(-fuzzy_range, fuzzy_range + 1):
                    if delta == 0:
                        continue
                    candidate_num = base_num + delta
                    if candidate_num < 0:
                        continue
                    candidate_str = str(candidate_num).zfill(len(main_numeric))
                    candidate_stem = ref[:numeric_idx] + candidate_str + ref[numeric_idx + len(main_numeric):]

                    if candidate_stem in stem_to_files:
                        for f in stem_to_files[candidate_stem]:
                            key = (f.get("file_path", ""), candidate_stem)
                            if key not in seen_matches:
                                seen_matches.add(key)
                                matches.append({
                                    "file_path": f.get("file_path", ""),
                                    "file_name": f.get("file_name", ""),
                                    "match_type": "fuzzy",
                                    "matched_ref": candidate_stem,
                                    "original_ref": ref,
                                    "notes_source": source,
                                })

        return matches

    # ------------------------------------------------------------------
    # CRUD operations
    # ------------------------------------------------------------------

    def get_project_files(self) -> List[Dict[str, Any]]:
        """Return the current list of files with all metadata."""
        return [dict(f) for f in self._files]

    def get_file_by_path(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Get a single file entry by its path."""
        file_path_resolved = str(Path(file_path).resolve())
        for f in self._files:
            if str(Path(f.get("file_path", "")).resolve()) == file_path_resolved:
                return dict(f)
        return None

    def update_file_metadata(
        self, file_path: str, updates: Dict[str, Any],
        provenance: Optional[Dict[str, str]] = None,
    ) -> Optional[Dict[str, Any]]:
        """
        Update metadata for a single file.

        Accepts ANY field — standard fields update columns directly,
        unknown fields go to custom_values. No EDITABLE_COLUMNS restriction.
        """
        fid = self._resolve_file_id(file_path)
        if fid is None:
            return None

        # Update DB
        self.store.update_file(fid, updates)

        # Record provenance if available
        if provenance:
            for key, value in updates.items():
                self.store.add_provenance(
                    file_id=fid,
                    field=key,
                    value=str(value),
                    source_type=provenance.get("source_type", "user"),
                    source_detail=provenance.get("source_detail", ""),
                    source_preview=provenance.get("source_preview", ""),
                    confidence=float(provenance.get("confidence", 0.8)),
                    reason=provenance.get("reason", ""),
                )

        # Update in-memory list
        file_path_resolved = str(Path(self._to_absolute(file_path)).resolve())
        for f in self._files:
            if str(Path(f.get("file_path", "")).resolve()) == file_path_resolved:
                for key, value in updates.items():
                    f[key] = value
                self._dirty = True
                return dict(f)

        # File in DB but not in memory — refresh
        self._refresh_files_from_db()
        self._dirty = True
        return self.get_file_by_path(file_path)

    def batch_update(self, file_paths: List[str], updates: Dict[str, Any],
                     provenance: Optional[Dict[str, str]] = None) -> int:
        """Apply the same metadata updates to multiple files."""
        count = 0
        for fp in file_paths:
            result = self.update_file_metadata(fp, updates, provenance=provenance)
            if result is not None:
                count += 1
        return count

    def add_files(self, file_entries: List[Dict[str, Any]]) -> int:
        """Add new file entries to the project (both DB and in-memory)."""
        if self._project_id is None:
            return 0

        added = 0
        for entry in file_entries:
            fp = entry.get("file_path", "")
            if not fp:
                continue

            rel_path = self._to_relative(fp)
            # Check if already in DB
            existing = self.store.get_file_id_by_path(self._project_id, rel_path)
            if existing:
                continue

            # Convert to relative for DB storage
            db_entry = dict(entry)
            db_entry['file_path'] = rel_path
            self.store.upsert_file(self._project_id, rel_path, db_entry)
            added += 1
            self._dirty = True

        if added > 0:
            self._refresh_files_from_db()

        return added

    def remove_file(self, file_path: str) -> bool:
        """Remove a file from the project."""
        fid = self._resolve_file_id(file_path)
        if fid is None:
            return False

        result = self.store.delete_file(fid)
        if result:
            self._refresh_files_from_db()
            self._dirty = True
        return result

    # ------------------------------------------------------------------
    # Subrow support (multi-animal recordings)
    # ------------------------------------------------------------------

    def add_subrow(
        self, file_path: str, channel: str, animal_id: str = "",
        sex: str = "", group: str = "", **extra
    ) -> Optional[Dict[str, Any]]:
        """Create a child entry for multi-animal recordings."""
        fid = self._resolve_file_id(file_path)
        if fid is None:
            return None

        # Get parent data for inheritance
        parent = self.store.get_file(fid)
        if not parent:
            return None

        subrow_data = {
            "channel": channel,
            "animal_id": animal_id,
            "sex": sex,
            "group": group,
            "protocol": parent.get("protocol", ""),
            "stim_type": parent.get("stim_type", ""),
            "power": parent.get("power", ""),
            "experiment": parent.get("experiment", ""),
            "strain": parent.get("strain", ""),
        }
        subrow_data.update(extra)

        subrow_id = self.store.add_subrow(fid, subrow_data)
        self._dirty = True

        # Return the subrow data with ID
        subrow_data["subrow_id"] = subrow_id
        return subrow_data

    def get_subrows(self, file_path: str) -> List[Dict[str, Any]]:
        """Get subrows for a file."""
        fid = self._resolve_file_id(file_path)
        if fid is None:
            return []
        return self.store.get_subrows(fid)

    def update_subrow(
        self, file_path: str, subrow_id: int, updates: Dict[str, Any]
    ) -> Optional[Dict[str, Any]]:
        """Update a specific subrow."""
        result = self.store.update_subrow(subrow_id, updates)
        if result:
            self._dirty = True
            return updates
        return None

    # ------------------------------------------------------------------
    # Provenance tracking
    # ------------------------------------------------------------------

    def get_provenance(
        self, file_path: str, field: Optional[str] = None
    ) -> List[Dict[str, Any]]:
        """Get provenance records for a file."""
        fid = self._resolve_file_id(file_path)
        if fid is None:
            return []
        return self.store.get_provenance(fid, field)

    # ------------------------------------------------------------------
    # Custom columns
    # ------------------------------------------------------------------

    def add_custom_column(
        self, column_key: str, display_name: str,
        column_type: str = 'text', sort_order: int = 0,
    ) -> Optional[int]:
        """Define a custom metadata column for the current project."""
        if self._project_id is None:
            return None
        return self.store.add_custom_column(
            self._project_id, column_key, display_name,
            column_type, sort_order,
        )

    def get_custom_columns(self) -> List[Dict[str, Any]]:
        """Get custom column definitions for the current project."""
        if self._project_id is None:
            return []
        return self.store.get_custom_columns(self._project_id)

    # ------------------------------------------------------------------
    # Analyses
    # ------------------------------------------------------------------

    def record_analysis(
        self, file_path: str, analysis_type: str,
        output_path: str = '', parameters: Optional[Dict] = None,
    ) -> Optional[int]:
        """Record an analysis output for a file."""
        fid = self._resolve_file_id(file_path)
        if fid is None:
            return None
        return self.store.record_analysis(fid, analysis_type, output_path, parameters)

    # ------------------------------------------------------------------
    # Backup
    # ------------------------------------------------------------------

    def backup(self, trigger_event: str = "manual") -> str:
        """Create a backup of the database."""
        return self.store.backup(trigger_event)

    # ------------------------------------------------------------------
    # Analysis helpers
    # ------------------------------------------------------------------

    def get_metadata_completeness(self) -> Dict[str, Any]:
        """Get percentage filled for each metadata column."""
        if self._project_id is None:
            return {"per_column": {}, "overall_pct": 0.0, "total_files": 0}
        return self.store.get_metadata_completeness(self._project_id)

    def get_unique_values(self, field: str) -> List[str]:
        """Get sorted unique non-empty values for a metadata field."""
        if self._project_id is None:
            return []
        return self.store.get_unique_values(self._project_id, field)

    def get_file_count(self) -> int:
        """Get total number of files in the project."""
        if self._project_id is None:
            return 0
        return self.store.get_file_count(self._project_id)

    # ------------------------------------------------------------------
    # Grouped file listing
    # ------------------------------------------------------------------

    def get_files_grouped(self) -> Dict[str, Any]:
        """Return files grouped by folder with summary stats."""
        if not self._files:
            return {
                "project_root": str(self._data_directory) if self._data_directory else "",
                "total_files": 0,
                "folders": [],
            }

        folder_map: Dict[str, List[Dict]] = {}
        for f in self._files:
            fp = f.get("file_path", "")
            rel = self._to_relative(fp)
            parent = str(Path(rel).parent).replace("\\", "/")
            if parent == ".":
                parent = ""
            folder_map.setdefault(parent, []).append(f)

        fields_to_check = ["experiment", "strain", "stim_type", "power", "sex",
                           "animal_id", "channel", "protocol"]

        folders = []
        for folder_path in sorted(folder_map.keys()):
            files_in_folder = folder_map[folder_path]
            count = len(files_in_folder)

            total_cells = count * len(fields_to_check)
            filled_cells = 0
            for f in files_in_folder:
                for field in fields_to_check:
                    val = f.get(field, "")
                    if val and str(val).strip() and str(val) not in ("Loading...", "Unknown", "pending"):
                        filled_cells += 1
            filled_pct = round((filled_cells / total_cells * 100) if total_cells else 0, 1)

            common = {}
            for field in ["strain", "animal_id", "stim_type", "power", "experiment"]:
                vals = set()
                for f in files_in_folder:
                    v = f.get(field, "")
                    if v and str(v).strip():
                        vals.add(str(v).strip())
                if len(vals) == 1:
                    common[field] = vals.pop()

            entry = {
                "path": folder_path,
                "count": count,
                "filled_pct": filled_pct,
            }
            entry.update(common)
            entry["files"] = [f.get("file_name", "") for f in files_in_folder]
            folders.append(entry)

        return {
            "project_root": str(self._data_directory) if self._data_directory else "",
            "total_files": len(self._files),
            "folders": folders,
        }

    # ------------------------------------------------------------------
    # File hash
    # ------------------------------------------------------------------

    @staticmethod
    def file_hash(path: Path) -> str:
        """Compute SHA-256 hash of a file for change detection."""
        h = hashlib.sha256()
        with open(path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                h.update(chunk)
        return h.hexdigest()
