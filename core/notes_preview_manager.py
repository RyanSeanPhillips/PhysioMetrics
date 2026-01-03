"""
Notes Preview Manager for PhysioMetrics.

Handles notes file preview, highlighting, and linking to ABF files.
Extracted from main.py to improve code maintainability.
"""

from pathlib import Path
import re


class NotesPreviewManager:
    """Manages notes file preview, highlighting, and linking functionality.

    Provides:
    - Preview dialogs for notes files (Excel, CSV, TXT, DOCX)
    - Dual highlighting (ABF stem in green, search term in orange)
    - Fuzzy matching for ABF file references
    - Linked notes column updates
    """

    def __init__(self, main_window):
        """Initialize the NotesPreviewManager.

        Args:
            main_window: Reference to the MainWindow instance
        """
        self.mw = main_window

    def show_linked_notes_for_file(self, file_name: str):
        """Show a dialog with notes files that reference the given data file.

        Uses fuzzy matching (+-5 numeric range) if no exact matches are found.
        """
        # Get the file stem for matching
        file_stem = Path(file_name).stem

        # Find notes files that reference this file (exact match)
        linked_notes = self.get_notes_for_abf(file_stem)
        is_fuzzy = False
        fuzzy_stems = []

        # If no exact matches, try fuzzy matching
        if not linked_notes:
            linked_notes, fuzzy_stems = self.get_fuzzy_notes_for_abf(file_stem)
            if linked_notes:
                is_fuzzy = True

        if not linked_notes:
            self.mw._log_status_message(f"No notes found referencing '{file_name}'. Run 'Scan References' first.", 3000)
            return

        # Build info text with fuzzy match warning if applicable
        if is_fuzzy:
            # Find the closest matching stem numerically
            closest_stem = fuzzy_stems[0] if fuzzy_stems else file_stem

            # Extract numeric part from original stem
            orig_nums = re.findall(r'\d+', file_stem)
            if orig_nums and fuzzy_stems:
                orig_num = int(max(orig_nums, key=len))
                min_diff = float('inf')
                for stem in fuzzy_stems:
                    stem_nums = re.findall(r'\d+', stem)
                    if stem_nums:
                        stem_num = int(max(stem_nums, key=len))
                        diff = abs(stem_num - orig_num)
                        if diff < min_diff:
                            min_diff = diff
                            closest_stem = stem

            stems_str = ', '.join(fuzzy_stems[:3])  # Show first 3 matches
            if len(fuzzy_stems) > 3:
                stems_str += f" +{len(fuzzy_stems) - 3} more"
            info_text = (
                f'<span style="color: #FFA500;">⚠ FUZZY MATCH:</span> '
                f'No exact match for <b>{file_name}</b><br>'
                f'Found {len(linked_notes)} file(s) referencing nearby: <b>{stems_str}</b>'
            )
            title = f"Notes for: {file_name} (Fuzzy Match)"
            # Highlight the closest matching stem
            highlight_stem = closest_stem
        else:
            info_text = f"Found {len(linked_notes)} notes file(s) referencing <b>{file_name}</b>"
            title = f"Notes for: {file_name}"
            highlight_stem = file_stem

        # Use shared preview dialog
        self.show_notes_preview_dialog(
            files=[{'name': n['name'], 'path': n['path']} for n in linked_notes],
            title=title,
            info_text=info_text,
            highlight_stem=highlight_stem
        )

    def show_notes_preview_dialog(self, files: list, title: str, info_text: str = None, highlight_stem: str = None):
        """Show a unified preview dialog for notes files.

        Args:
            files: List of dicts with 'name' and 'path' keys
            title: Dialog window title
            info_text: Optional HTML text shown above the preview (e.g., file count)
            highlight_stem: Optional ABF filename stem to highlight in previews
        """
        from PyQt6.QtWidgets import (QDialog, QVBoxLayout, QHBoxLayout, QLabel,
                                     QPushButton, QTabWidget)

        if not files:
            return

        # Dark theme dialog stylesheet
        dark_dialog_style = """
            QDialog {
                background-color: #1e1e1e;
                color: #d4d4d4;
            }
            QLabel {
                color: #d4d4d4;
            }
            QPushButton {
                background-color: #3c3c3c;
                color: #d4d4d4;
                border: 1px solid #555555;
                padding: 6px 16px;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #4a4a4a;
            }
            QPushButton:pressed {
                background-color: #2a2a2a;
            }
            QTabWidget::pane {
                border: 1px solid #3e3e42;
                background-color: #1e1e1e;
            }
            QTabBar::tab {
                background-color: #2d2d30;
                color: #d4d4d4;
                padding: 8px 16px;
                border: 1px solid #3e3e42;
                border-bottom: none;
                margin-right: 2px;
            }
            QTabBar::tab:selected {
                background-color: #1e1e1e;
                border-bottom-color: #1e1e1e;
            }
            QTabBar::tab:hover:!selected {
                background-color: #3e3e42;
            }
        """

        # Create dialog
        dialog = QDialog(self.mw)
        dialog.setWindowTitle(title)
        dialog.resize(1100, 750)
        dialog.setStyleSheet(dark_dialog_style)

        # Enable dark title bar on Windows
        self.mw._enable_dark_title_bar(dialog)

        layout = QVBoxLayout(dialog)
        layout.setContentsMargins(8, 8, 8, 8)

        # Info label (optional)
        if info_text:
            info_label = QLabel(info_text)
            info_label.setStyleSheet("font-size: 11pt; padding: 4px;")
            layout.addWidget(info_label)

        # Store file info for "Open in External App" button
        self.mw._preview_dialog_files = {}

        # Single file: no tabs needed, just show the preview directly
        if len(files) == 1:
            file_info = files[0]
            preview_widget = self.create_note_preview_widget(
                {'name': file_info['name'], 'path': file_info['path']},
                highlight_stem or ''
            )
            layout.addWidget(preview_widget, stretch=1)
            self.mw._preview_dialog_files[0] = file_info
            current_file_getter = lambda: self.mw._preview_dialog_files.get(0)
        else:
            # Multiple files: use tabs
            tab_widget = QTabWidget()

            for idx, file_info in enumerate(files):
                preview_widget = self.create_note_preview_widget(
                    {'name': file_info['name'], 'path': file_info['path']},
                    highlight_stem or ''
                )
                tab_name = file_info['name']
                if len(tab_name) > 30:
                    tab_name = tab_name[:27] + "..."
                tab_widget.addTab(preview_widget, tab_name)
                self.mw._preview_dialog_files[idx] = file_info

            layout.addWidget(tab_widget, stretch=1)
            current_file_getter = lambda: self.mw._preview_dialog_files.get(tab_widget.currentIndex())

        # Bottom button row
        button_layout = QHBoxLayout()

        open_btn = QPushButton("Open in External App")
        open_btn.setToolTip("Open the file in its default application")

        def open_current():
            import os
            file_info = current_file_getter()
            if file_info and file_info.get('path'):
                os.startfile(file_info['path'])

        open_btn.clicked.connect(open_current)
        button_layout.addWidget(open_btn)

        button_layout.addStretch()

        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dialog.reject)
        button_layout.addWidget(close_btn)

        layout.addLayout(button_layout)

        dialog.exec()

    def create_note_preview_widget(self, note_info: dict, abf_stem: str, search_term: str = None):
        """Create a preview widget for a single note file with ABF highlighting.

        Args:
            note_info: Dict with 'path' key pointing to the notes file
            abf_stem: Primary highlight term (ABF filename) - shown in green
            search_term: Optional secondary search term - shown in orange/yellow
        """
        from PyQt6.QtWidgets import QWidget, QVBoxLayout, QTabWidget, QTextEdit, QLabel

        container = QWidget()
        layout = QVBoxLayout(container)
        layout.setContentsMargins(0, 0, 0, 0)

        file_path = Path(note_info['path'])

        # Add filename header
        filename_label = QLabel(f"<b>File:</b> {file_path.name}")
        filename_label.setStyleSheet("QLabel { color: #9cdcfe; padding: 4px 8px; background-color: #2d2d30; border-radius: 3px; }")
        filename_label.setToolTip(str(file_path))
        layout.addWidget(filename_label)

        if not file_path.exists():
            text_edit = QTextEdit()
            text_edit.setPlainText(f"File not found: {file_path}")
            text_edit.setReadOnly(True)
            text_edit.setStyleSheet("QTextEdit { background-color: #1e1e1e; color: #d4d4d4; border: none; }")
            layout.addWidget(text_edit)
            return container

        suffix = file_path.suffix.lower()

        try:
            if suffix in ['.xlsx', '.xls']:
                import pandas as pd
                sheets = pd.read_excel(file_path, sheet_name=None, nrows=500)

                if sheets:
                    # Multiple sheets -> nested tabs
                    sheet_tabs = QTabWidget()
                    sheet_tabs.setStyleSheet("""
                        QTabWidget::pane { border: 1px solid #3e3e42; background-color: #1e1e1e; }
                        QTabBar::tab { background-color: #2d2d30; color: #d4d4d4; padding: 6px 12px; border: 1px solid #3e3e42; }
                        QTabBar::tab:selected { background-color: #1e1e1e; }
                    """)

                    # Sort sheets: matching sheets first, then others
                    sheet_items = list(sheets.items())
                    matching_sheets = []
                    other_sheets = []
                    for sheet_name, df in sheet_items:
                        if self.df_contains_stem(df, abf_stem):
                            matching_sheets.append((sheet_name, df, True))
                        else:
                            other_sheets.append((sheet_name, df, False))

                    # Add matching sheets first, then others
                    for sheet_name, df, has_match in matching_sheets + other_sheets:
                        table = self.create_highlighted_table(df, abf_stem, search_term)
                        tab_name = f"★ {sheet_name}" if has_match else sheet_name
                        sheet_tabs.addTab(table, tab_name)

                    layout.addWidget(sheet_tabs)
                else:
                    text_edit = QTextEdit()
                    text_edit.setPlainText("[Workbook is empty]")
                    text_edit.setReadOnly(True)
                    layout.addWidget(text_edit)

            elif suffix == '.csv':
                import pandas as pd
                df = pd.read_csv(file_path, nrows=500, encoding='utf-8', on_bad_lines='skip')
                table = self.create_highlighted_table(df, abf_stem, search_term)
                layout.addWidget(table)

            elif suffix == '.txt':
                content = file_path.read_text(encoding='utf-8', errors='replace')
                text_edit = self.create_highlighted_text(content, abf_stem, search_term)
                layout.addWidget(text_edit)

            elif suffix == '.docx':
                try:
                    from docx import Document
                    doc = Document(str(file_path))

                    # Collect paragraphs
                    content_parts = []
                    for p in doc.paragraphs:
                        if p.text.strip():
                            content_parts.append(p.text)

                    # Also collect table contents (where ABF references often appear)
                    for table in doc.tables:
                        table_rows = []
                        for row in table.rows:
                            row_cells = [cell.text.strip() for cell in row.cells]
                            table_rows.append(' | '.join(row_cells))
                        if table_rows:
                            content_parts.append('\n--- Table ---')
                            content_parts.extend(table_rows)
                            content_parts.append('--- End Table ---\n')

                    content = '\n\n'.join(content_parts)
                    text_edit = self.create_highlighted_text(content, abf_stem, search_term)
                    layout.addWidget(text_edit)
                except ImportError:
                    text_edit = QTextEdit()
                    text_edit.setPlainText("python-docx not installed. Cannot preview .docx files.\n\nUse 'Open in External App' to view.")
                    text_edit.setReadOnly(True)
                    text_edit.setStyleSheet("QTextEdit { background-color: #1e1e1e; color: #d4d4d4; border: none; }")
                    layout.addWidget(text_edit)

            else:
                text_edit = QTextEdit()
                text_edit.setPlainText(f"Preview not supported for {suffix} files.\n\nUse 'Open in External App' to view.")
                text_edit.setReadOnly(True)
                text_edit.setStyleSheet("QTextEdit { background-color: #1e1e1e; color: #d4d4d4; border: none; }")
                layout.addWidget(text_edit)

        except Exception as e:
            text_edit = QTextEdit()
            text_edit.setPlainText(f"Error loading preview: {e}")
            text_edit.setReadOnly(True)
            text_edit.setStyleSheet("QTextEdit { background-color: #1e1e1e; color: #d4d4d4; border: none; }")
            layout.addWidget(text_edit)

        return container

    def create_highlighted_table(self, df, abf_stem: str, search_term: str = None):
        """Create a QTableWidget from DataFrame with dual highlighting.

        Args:
            df: DataFrame to display
            abf_stem: Primary highlight term (ABF filename) - shown in green
            search_term: Optional secondary search term - shown in orange/yellow
        """
        from PyQt6.QtWidgets import QTableWidget, QTableWidgetItem, QHeaderView
        from PyQt6.QtCore import Qt
        from PyQt6.QtGui import QColor, QBrush

        table = QTableWidget()
        table.setRowCount(len(df))
        table.setColumnCount(len(df.columns))
        table.setHorizontalHeaderLabels([str(col) for col in df.columns])

        # Style the table
        table.setStyleSheet("""
            QTableWidget {
                background-color: #1e1e1e;
                color: #d4d4d4;
                gridline-color: #3e3e42;
                border: none;
            }
            QTableWidget::item {
                padding: 4px;
            }
            QHeaderView::section {
                background-color: #2d2d30;
                color: #d4d4d4;
                padding: 4px;
                border: 1px solid #3e3e42;
            }
        """)

        # Colors
        PRIMARY_BG = QColor("#3d5a3d")  # Green background for ABF match
        PRIMARY_TEXT = QColor("#90ff90")  # Bright green text
        SECONDARY_BG = QColor("#5a4a2a")  # Orange background for search term
        SECONDARY_TEXT = QColor("#ffcc66")  # Orange text
        BOTH_BG = QColor("#4a5a4a")  # Mixed color when both match

        scroll_rows = []  # Track rows with matches for scrolling

        for row_idx in range(len(df)):
            for col_idx in range(len(df.columns)):
                value = str(df.iloc[row_idx, col_idx])
                item = QTableWidgetItem(value)

                has_primary = abf_stem and abf_stem.lower() in value.lower()
                has_secondary = search_term and search_term.lower() in value.lower()

                if has_primary or has_secondary:
                    if row_idx not in scroll_rows:
                        scroll_rows.append(row_idx)

                if has_primary and has_secondary:
                    item.setBackground(QBrush(BOTH_BG))
                    item.setForeground(QBrush(PRIMARY_TEXT))
                elif has_primary:
                    item.setBackground(QBrush(PRIMARY_BG))
                    item.setForeground(QBrush(PRIMARY_TEXT))
                elif has_secondary:
                    item.setBackground(QBrush(SECONDARY_BG))
                    item.setForeground(QBrush(SECONDARY_TEXT))

                table.setItem(row_idx, col_idx, item)

        # Auto-resize columns
        header = table.horizontalHeader()
        header.setSectionResizeMode(QHeaderView.ResizeMode.ResizeToContents)
        header.setStretchLastSection(True)

        # Scroll to first match
        if scroll_rows:
            first_row = min(scroll_rows)
            table.scrollToItem(table.item(first_row, 0))

        return table

    def create_highlighted_text(self, content: str, abf_stem: str, search_term: str = None):
        """Create a QTextEdit with dual highlighting and scrolled to first match.

        Args:
            content: Text content to display
            abf_stem: Primary highlight term (ABF filename) - shown in green
            search_term: Optional secondary search term - shown in orange/yellow
        """
        from PyQt6.QtWidgets import QTextEdit
        from PyQt6.QtCore import QTimer

        text_edit = QTextEdit()
        text_edit.setReadOnly(True)

        # Create patterns for both highlight terms
        abf_pattern = re.compile(re.escape(abf_stem), re.IGNORECASE) if abf_stem else None
        search_pattern = re.compile(re.escape(search_term), re.IGNORECASE) if search_term else None

        # Process line by line
        lines = content.split('\n')
        highlighted_lines = []

        for line in lines:
            has_primary = abf_stem and abf_stem.lower() in line.lower()
            has_secondary = search_term and search_term.lower() in line.lower()

            if has_primary or has_secondary:
                line_html = line

                # Apply primary highlighting (green) - ABF stem
                if has_primary and abf_pattern:
                    line_html = abf_pattern.sub(
                        r'<span style="color: #90ff90; font-weight: bold;">\g<0></span>',
                        line_html
                    )

                # Apply secondary highlighting (orange) - search term
                # Need to be careful not to double-highlight if terms overlap
                if has_secondary and search_pattern:
                    # Only highlight search term if it doesn't overlap with ABF stem
                    if not has_primary or (abf_stem.lower() not in search_term.lower() and search_term.lower() not in abf_stem.lower()):
                        line_html = search_pattern.sub(
                            r'<span style="color: #ffcc66; font-weight: bold; background-color: #5a4a2a;">\g<0></span>',
                            line_html
                        )

                # Determine background color based on what matched
                if has_primary and has_secondary:
                    # Both match - use a mixed color
                    bg_color = "#4a5a4a"
                elif has_primary:
                    bg_color = "#3d5a3d"
                else:
                    bg_color = "#5a4a2a"

                highlighted_lines.append(
                    f'<div style="background-color: {bg_color}; padding: 2px 4px; margin: 1px 0;">{line_html}</div>'
                )
            else:
                highlighted_lines.append(f'<div>{line}</div>')

        highlighted = ''.join(highlighted_lines)
        text_edit.setHtml(f"<div style='color: #d4d4d4; font-family: Consolas, monospace; font-size: 10pt;'>{highlighted}</div>")

        text_edit.setStyleSheet("""
            QTextEdit {
                background-color: #1e1e1e;
                border: none;
                font-size: 10pt;
            }
        """)

        # Scroll to first match after widget is shown (prefer primary, fallback to secondary)
        def scroll_to_match():
            # Safety check: widget may have been deleted if user is rapidly typing in search
            try:
                # Check if widget is still valid (accessing any property will throw if deleted)
                if not text_edit.isVisible():
                    return
            except RuntimeError:
                # Widget has been deleted (C/C++ object no longer exists)
                return

            # Find position in original content (case-insensitive)
            match_pos = -1
            if abf_stem:
                match_pos = content.lower().find(abf_stem.lower())
            if match_pos < 0 and search_term:
                match_pos = content.lower().find(search_term.lower())

            if match_pos >= 0:
                # Calculate approximate scroll position as percentage of document
                total_len = len(content)
                if total_len > 0:
                    scroll_percent = match_pos / total_len
                    try:
                        scrollbar = text_edit.verticalScrollBar()
                        max_scroll = scrollbar.maximum()
                        # Scroll to position (with some offset to show context above)
                        target_pos = int(max_scroll * scroll_percent)
                        # Subtract some pixels to show context above the match
                        target_pos = max(0, target_pos - 50)
                        scrollbar.setValue(target_pos)
                    except RuntimeError:
                        # Widget deleted during scroll operation
                        pass

        # Use longer delay to ensure dialog is fully rendered
        QTimer.singleShot(300, scroll_to_match)

        return text_edit

    def df_contains_stem(self, df, abf_stem: str) -> bool:
        """Check if a DataFrame contains the ABF stem anywhere."""
        for col in df.columns:
            for val in df[col].astype(str):
                if abf_stem.lower() in val.lower():
                    return True
        return False

    def get_notes_for_abf(self, abf_stem: str) -> list:
        """Find notes files that reference a specific ABF file.

        Args:
            abf_stem: The ABF filename stem (without extension)

        Returns:
            List of note_info dicts that reference this ABF file
        """
        linked = []

        if not hasattr(self.mw, '_notes_files_data') or not self.mw._notes_files_data:
            return linked

        for note_info in self.mw._notes_files_data:
            matches = note_info.get('matches', [])
            if abf_stem in matches:
                linked.append(note_info)

        return linked

    def get_fuzzy_notes_for_abf(self, abf_stem: str, search_range: int = 5) -> tuple:
        """Find notes files that reference numeric neighbors of an ABF file.

        Used when exact match fails. Searches for files with similar numeric patterns
        within +-search_range of the original filename's numeric portion.

        Args:
            abf_stem: The ABF filename stem (without extension), e.g., "25708007"
            search_range: How far to search (default +-5)

        Returns:
            Tuple of (list of note_info dicts, list of matched_stems that were fuzzy matched)
        """
        if not hasattr(self.mw, '_notes_files_data') or not self.mw._notes_files_data:
            return [], []

        # Extract numeric portion from filename
        # Handle patterns like: "25708007", "mouse1_25708007", "25708007_trial1"
        # We want the longest numeric sequence (likely the file ID)
        numeric_matches = re.findall(r'\d+', abf_stem)
        if not numeric_matches:
            return [], []

        # Use the longest numeric sequence as the ID to match
        main_numeric = max(numeric_matches, key=len)
        numeric_idx = abf_stem.find(main_numeric)

        # Generate candidate stems within range
        candidates = set()
        try:
            base_num = int(main_numeric)
            for delta in range(-search_range, search_range + 1):
                if delta == 0:
                    continue  # Skip exact match (already checked)
                candidate_num = base_num + delta
                if candidate_num >= 0:
                    # Preserve leading zeros and surrounding text
                    candidate_str = str(candidate_num).zfill(len(main_numeric))
                    candidate_stem = abf_stem[:numeric_idx] + candidate_str + abf_stem[numeric_idx + len(main_numeric):]
                    candidates.add(candidate_stem)
        except ValueError:
            return [], []

        # Search notes files for these candidates
        linked = []
        matched_stems = []

        for note_info in self.mw._notes_files_data:
            matches = note_info.get('matches', [])
            for match in matches:
                if match in candidates:
                    if note_info not in linked:
                        linked.append(note_info)
                    if match not in matched_stems:
                        matched_stems.append(match)

        return linked, matched_stems

    def update_linked_notes_column(self):
        """Update the linked_notes column in file table based on scanned notes.

        Uses fuzzy matching (+-5 numeric range) when exact matches aren't found.
        Fuzzy matches are displayed with a '~' prefix to indicate caution.
        """
        if not hasattr(self.mw, '_file_table_model') or not self.mw._file_table_model:
            return

        if not hasattr(self.mw, '_notes_files_data') or not self.mw._notes_files_data:
            return

        # Build a reverse index: ABF stem -> count of notes referencing it
        abf_notes_count = {}
        for note_info in self.mw._notes_files_data:
            matches = note_info.get('matches', [])
            for abf_stem in matches:
                abf_notes_count[abf_stem] = abf_notes_count.get(abf_stem, 0) + 1

        fuzzy_match_count = 0

        # Update each row in the file table
        for row in range(self.mw._file_table_model.rowCount()):
            row_data = self.mw._file_table_model.get_row_data(row)
            if row_data:
                file_name = row_data.get('file_name', '')
                if file_name:
                    file_stem = Path(file_name).stem
                    exact_count = abf_notes_count.get(file_stem, 0)

                    # Check for fuzzy match if no exact matches
                    is_fuzzy = False
                    fuzzy_stems = []
                    if exact_count == 0:
                        fuzzy_notes, fuzzy_stems = self.get_fuzzy_notes_for_abf(file_stem)
                        if fuzzy_notes:
                            exact_count = len(fuzzy_notes)
                            is_fuzzy = True
                            fuzzy_match_count += 1

                    # Format display: "~N" for fuzzy, "N" for exact, "0" for no matches
                    if exact_count > 0:
                        display_value = f"~{exact_count}" if is_fuzzy else str(exact_count)
                    else:
                        display_value = '0'

                    # Update model data (including fuzzy info for tooltips)
                    self.mw._file_table_model.set_cell_value(row, 'linked_notes', display_value)
                    self.mw._file_table_model.set_cell_value(row, 'linked_notes_fuzzy', is_fuzzy)
                    self.mw._file_table_model.set_cell_value(row, 'linked_notes_fuzzy_stems', fuzzy_stems)

                    # Also update master file list with detailed info
                    if row < len(self.mw._master_file_list):
                        self.mw._master_file_list[row]['linked_notes'] = exact_count
                        self.mw._master_file_list[row]['linked_notes_fuzzy'] = is_fuzzy
                        self.mw._master_file_list[row]['linked_notes_fuzzy_stems'] = fuzzy_stems

        fuzzy_msg = f" ({fuzzy_match_count} fuzzy)" if fuzzy_match_count > 0 else ""
        print(f"[notes] Updated linked_notes for {self.mw._file_table_model.rowCount()} files{fuzzy_msg}")

    def populate_notes_table_from_data(self):
        """Populate notes table from stored notes files data."""
        if not hasattr(self.mw, '_notes_model') or not self.mw._notes_files_data:
            return

        self.mw._notes_model.removeRows(0, self.mw._notes_model.rowCount())

        from PyQt6.QtGui import QStandardItem
        from PyQt6.QtCore import Qt

        for file_info in self.mw._notes_files_data:
            # Create checkbox item for "Use" column
            use_item = QStandardItem()
            use_item.setCheckable(True)
            use_as_notes = file_info.get('use_as_notes', True)
            use_item.setCheckState(Qt.CheckState.Checked if use_as_notes else Qt.CheckState.Unchecked)
            use_item.setEditable(False)

            # File name with path stored
            name_item = QStandardItem(file_info.get('name', ''))
            name_item.setData(file_info.get('path', ''), Qt.ItemDataRole.UserRole)

            # Matches column
            match_count = file_info.get('match_count', 0)
            matches_item = QStandardItem(str(match_count) if match_count > 0 else "—")
            matches_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)

            type_item = QStandardItem(file_info.get('type', ''))
            location_item = QStandardItem(file_info.get('location', ''))
            size_item = QStandardItem(file_info.get('size', ''))
            modified_item = QStandardItem(file_info.get('modified', ''))

            self.mw._notes_model.appendRow([use_item, name_item, matches_item, type_item, location_item, size_item, modified_item])
