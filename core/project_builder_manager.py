"""
Project Builder Manager - Handles project file management, notes files, and scanning.

Extracted from main.py to improve maintainability.
Contains all Project Builder tab functionality including:
- File scanning and discovery
- Notes file management
- Project save/load operations
- Conflict resolution
- Table management
"""

from pathlib import Path
from typing import TYPE_CHECKING, List, Dict, Any, Optional

from PyQt6.QtCore import Qt, QThread, pyqtSignal, QTimer
from PyQt6.QtWidgets import QApplication, QProgressDialog

if TYPE_CHECKING:
    from PyQt6.QtWidgets import QMainWindow


class ProjectBuilderManager:
    """Manages Project Builder tab functionality.

    Handles file scanning, notes management, project save/load,
    and table operations.
    """

    def __init__(self, main_window: 'QMainWindow'):
        """Initialize ProjectBuilderManager.

        Args:
            main_window: Reference to MainWindow for widget access
        """
        self.mw = main_window

        # Notes tab state
        self._notes_hide_unmatched = False
        self._scan_worker = None

        # Project state (references to main window state)
        # These are accessed via self.mw.* when needed

    # =========================================================================
    # Notes Files Tab - Initialization
    # =========================================================================

    def init_notes_files_model(self):
        """Initialize the notes files table model."""
        from PyQt6.QtGui import QStandardItemModel

        if not hasattr(self.mw, 'notesFilesTable'):
            return

        self.mw._notes_model = QStandardItemModel()
        self.mw._notes_model.setHorizontalHeaderLabels(
            ['Use', 'File Name', '', 'Type', 'Matches', 'Location', 'Size', 'Modified']
        )
        self.mw.notesFilesTable.setModel(self.mw._notes_model)

        # Set up Actions column delegate
        from core.file_table_delegates import NotesActionsDelegate
        self.mw._notes_actions_delegate = NotesActionsDelegate(self.mw.notesFilesTable)
        self.mw._notes_actions_delegate.folder_clicked.connect(self._on_notes_action_folder)
        self.mw._notes_actions_delegate.open_clicked.connect(self._on_notes_action_open)
        self.mw._notes_actions_delegate.preview_clicked.connect(self._on_notes_action_preview)
        self.mw.notesFilesTable.setItemDelegateForColumn(2, self.mw._notes_actions_delegate)

        # Enable mouse tracking for hover effects
        self.mw.notesFilesTable.setMouseTracking(True)

        # Set column widths
        from PyQt6.QtWidgets import QHeaderView
        header = self.mw.notesFilesTable.horizontalHeader()
        header.setStretchLastSection(True)
        header.setSectionResizeMode(0, QHeaderView.ResizeMode.Fixed)
        header.setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)
        header.setSectionResizeMode(2, QHeaderView.ResizeMode.Fixed)
        header.setSectionResizeMode(3, QHeaderView.ResizeMode.ResizeToContents)
        header.setSectionResizeMode(4, QHeaderView.ResizeMode.ResizeToContents)
        header.setSectionResizeMode(5, QHeaderView.ResizeMode.ResizeToContents)
        header.setSectionResizeMode(6, QHeaderView.ResizeMode.ResizeToContents)
        header.setSectionResizeMode(7, QHeaderView.ResizeMode.ResizeToContents)
        self.mw.notesFilesTable.setColumnWidth(0, 50)
        self.mw.notesFilesTable.setColumnWidth(2, 82)

        # Double-click to open
        self.mw.notesFilesTable.doubleClicked.connect(self.on_notes_open)

        # Connect model changes
        self.mw._notes_model.itemChanged.connect(self._on_notes_item_changed)

        # Hide bottom buttons since actions are in table
        if hasattr(self.mw, 'openNoteButton'):
            self.mw.openNoteButton.setVisible(False)
        if hasattr(self.mw, 'previewNoteButton'):
            self.mw.previewNoteButton.setVisible(False)

    def add_notes_action_buttons(self):
        """Add additional action buttons to the Notes tab programmatically."""
        from PyQt6.QtWidgets import QPushButton

        if not hasattr(self.mw, 'notesFilesTable'):
            return

        if not hasattr(self.mw, 'linkNoteButton'):
            return

        # Find the layout containing linkNoteButton
        parent_layout = None
        if hasattr(self.mw, 'notesActionsLayout'):
            parent_layout = self.mw.notesActionsLayout
        else:
            parent_widget = self.mw.linkNoteButton.parentWidget()
            if parent_widget:
                layout = parent_widget.layout()
                if layout:
                    for i in range(layout.count()):
                        item = layout.itemAt(i)
                        if item and item.layout():
                            inner_layout = item.layout()
                            for j in range(inner_layout.count()):
                                widget = inner_layout.itemAt(j).widget()
                                if widget == self.mw.linkNoteButton:
                                    parent_layout = inner_layout
                                    break
                            if parent_layout:
                                break

        if not parent_layout:
            return

        # Button styles
        button_style = """
            QPushButton {
                background-color: #6c757d;
                color: white;
                border-radius: 4px;
                padding: 4px 8px;
            }
            QPushButton:hover {
                background-color: #5a6268;
            }
            QPushButton:pressed {
                background-color: #545b62;
            }
        """

        ai_button_style = """
            QPushButton {
                background-color: #0d6efd;
                color: white;
                border: 1px solid #0b5ed7;
                border-radius: 4px;
                padding: 4px 8px;
            }
            QPushButton:hover {
                background-color: #0b5ed7;
            }
            QPushButton:pressed {
                background-color: #0a58ca;
            }
        """

        # Scan References button
        self.mw.scanReferencesButton = QPushButton("Scan for Refs")
        self.mw.scanReferencesButton.setToolTip("Scan notes files for references to data file names")
        self.mw.scanReferencesButton.setMinimumSize(70, 26)
        self.mw.scanReferencesButton.setMaximumWidth(100)
        self.mw.scanReferencesButton.setStyleSheet(button_style)
        self.mw.scanReferencesButton.clicked.connect(self.on_notes_scan_references)
        parent_layout.insertWidget(4, self.mw.scanReferencesButton)

        # Toggle Unmatched button
        self.mw.toggleUnmatchedButton = QPushButton("Hide Unmatched")
        self.mw.toggleUnmatchedButton.setToolTip("Hide/show files with no data file matches")
        self.mw.toggleUnmatchedButton.setMinimumSize(100, 26)
        self.mw.toggleUnmatchedButton.setMaximumWidth(120)
        self.mw.toggleUnmatchedButton.setStyleSheet(button_style)
        self.mw.toggleUnmatchedButton.clicked.connect(self.on_notes_toggle_unmatched)
        parent_layout.insertWidget(5, self.mw.toggleUnmatchedButton)

        # AI Index button
        self.mw.buildAIIndexButton = QPushButton("AI Index")
        self.mw.buildAIIndexButton.setToolTip("Build AI-powered metadata lookup table")
        self.mw.buildAIIndexButton.setMinimumSize(80, 26)
        self.mw.buildAIIndexButton.setMaximumWidth(100)
        self.mw.buildAIIndexButton.setStyleSheet(ai_button_style)
        self.mw.buildAIIndexButton.clicked.connect(self.on_notes_build_ai_index)
        self.mw.buildAIIndexButton.setVisible(False)  # Not currently used
        parent_layout.insertWidget(6, self.mw.buildAIIndexButton)

        # Move filter widgets if they exist
        if hasattr(self.mw, 'notesFilterEdit') and hasattr(self.mw, 'notesFilterCountLabel'):
            filter_edit = self.mw.notesFilterEdit
            count_label = self.mw.notesFilterCountLabel

            old_layout = filter_edit.parentWidget().layout() if filter_edit.parentWidget() else None
            if old_layout:
                for i in range(old_layout.count()):
                    item = old_layout.itemAt(i)
                    if item and item.widget() == filter_edit:
                        old_layout.takeAt(i)
                        break
                for i in range(old_layout.count()):
                    item = old_layout.itemAt(i)
                    if item and item.widget() == count_label:
                        old_layout.takeAt(i)
                        break

                insert_pos = parent_layout.count() - 1
                if insert_pos < 0:
                    insert_pos = 0
                parent_layout.insertWidget(insert_pos, filter_edit)
                parent_layout.insertWidget(insert_pos + 1, count_label)

        print("[notes] Added Scan Refs and Unmatched buttons to action row")

    # =========================================================================
    # Notes Files Tab - Browse and Search
    # =========================================================================

    def on_notes_browse(self):
        """Browse for a folder to search for notes files."""
        from PyQt6.QtWidgets import QFileDialog

        start_dir = ""
        if self.mw._notes_directory:
            start_dir = str(self.mw._notes_directory)
        elif self.mw._project_directory:
            start_dir = str(Path(self.mw._project_directory).parent)

        folder = QFileDialog.getExistingDirectory(
            self.mw,
            "Select Notes Folder",
            start_dir
        )

        if folder:
            self.mw._notes_directory = Path(folder)
            if hasattr(self.mw, 'notesFolderEdit'):
                self.mw.notesFolderEdit.setText(str(folder))
            self.mw._log_status_message(f"Notes folder set to: {folder}", 2000)

    def set_default_notes_directory(self):
        """Set default notes directory to parent of project directory."""
        if self.mw._project_directory:
            parent = Path(self.mw._project_directory).parent
            self.mw._notes_directory = parent
            if hasattr(self.mw, 'notesFolderEdit'):
                self.mw.notesFolderEdit.setText(str(parent))

    def on_notes_search(self):
        """Search for notes files in project folder and parent directories."""
        from PyQt6.QtGui import QStandardItem, QCursor
        from PyQt6.QtWidgets import QApplication
        import os
        from datetime import datetime

        if not self.mw._project_directory:
            self.mw._log_status_message("Please select a project directory first", 2000)
            return

        project_path = Path(self.mw._project_directory)

        # Get selected file types
        extensions = []
        if hasattr(self.mw, 'notesWordCheckbox') and self.mw.notesWordCheckbox.isChecked():
            extensions.extend(['.docx', '.doc'])
        if hasattr(self.mw, 'notesExcelCheckbox') and self.mw.notesExcelCheckbox.isChecked():
            extensions.extend(['.xlsx', '.xls'])
        if hasattr(self.mw, 'notesCsvCheckbox') and self.mw.notesCsvCheckbox.isChecked():
            extensions.append('.csv')
        if hasattr(self.mw, 'notesTxtCheckbox') and self.mw.notesTxtCheckbox.isChecked():
            extensions.append('.txt')
        if hasattr(self.mw, 'notesPdfCheckbox') and self.mw.notesPdfCheckbox.isChecked():
            extensions.append('.pdf')

        if not extensions:
            extensions = ['.docx', '.doc', '.xlsx', '.xls', '.csv', '.txt']

        # Show wait cursor
        QApplication.setOverrideCursor(QCursor(Qt.CursorShape.WaitCursor))

        try:
            found_files = []

            # Search project folder recursively
            for ext in extensions:
                for f in project_path.rglob(f'*{ext}'):
                    if f.is_file():
                        found_files.append(f)

            # Search parent and grandparent (non-recursive)
            for parent_level in [project_path.parent, project_path.parent.parent]:
                if parent_level.exists():
                    for ext in extensions:
                        for f in parent_level.glob(f'*{ext}'):
                            if f.is_file() and f not in found_files:
                                found_files.append(f)

            # Clear existing
            self.mw._notes_model.removeRows(0, self.mw._notes_model.rowCount())
            self.mw._notes_files_data = []

            # Add found files to model
            for file_path in sorted(found_files, key=lambda x: x.name.lower()):
                file_info = {
                    'name': file_path.name,
                    'path': str(file_path),
                    'type': file_path.suffix.upper()[1:],
                    'matches': [],
                    'match_count': 0,
                    'use_as_notes': True
                }

                try:
                    stat = file_path.stat()
                    file_info['size'] = stat.st_size
                    file_info['modified'] = datetime.fromtimestamp(stat.st_mtime)
                except:
                    file_info['size'] = 0
                    file_info['modified'] = None

                self.mw._notes_files_data.append(file_info)

                # Create row items
                use_item = QStandardItem()
                use_item.setCheckable(True)
                use_item.setCheckState(Qt.CheckState.Checked)

                name_item = QStandardItem(file_path.name)
                name_item.setData(str(file_path), Qt.ItemDataRole.UserRole)

                actions_item = QStandardItem()
                type_item = QStandardItem(file_info['type'])
                matches_item = QStandardItem("0")

                # Location relative to project
                try:
                    rel_path = file_path.parent.relative_to(project_path)
                    location = str(rel_path) if str(rel_path) != '.' else '.'
                except ValueError:
                    location = str(file_path.parent)
                location_item = QStandardItem(location)

                # Size formatting
                size_bytes = file_info['size']
                if size_bytes < 1024:
                    size_str = f"{size_bytes} B"
                elif size_bytes < 1024 * 1024:
                    size_str = f"{size_bytes / 1024:.1f} KB"
                else:
                    size_str = f"{size_bytes / (1024 * 1024):.1f} MB"
                size_item = QStandardItem(size_str)

                # Modified date
                if file_info['modified']:
                    mod_str = file_info['modified'].strftime('%Y-%m-%d %H:%M')
                else:
                    mod_str = ""
                modified_item = QStandardItem(mod_str)

                self.mw._notes_model.appendRow([
                    use_item, name_item, actions_item, type_item,
                    matches_item, location_item, size_item, modified_item
                ])

            total_count = len(found_files)
            self.mw._log_status_message(f"Found {total_count} notes files", 2000)

            if hasattr(self.mw, 'notesFilterCountLabel'):
                self.mw.notesFilterCountLabel.setText(f"{total_count} files")

        finally:
            QApplication.restoreOverrideCursor()

    # =========================================================================
    # Notes Files Tab - Item Events
    # =========================================================================

    def _on_notes_item_changed(self, item):
        """Handle changes to notes table items (checkbox state)."""
        if item.column() != 0:
            return

        row = item.row()
        if row < len(self.mw._notes_files_data):
            is_checked = item.checkState() == Qt.CheckState.Checked
            self.mw._notes_files_data[row]['use_as_notes'] = is_checked

    def _on_notes_action_folder(self, row: int):
        """Open containing folder for notes file at given row."""
        import os
        import subprocess
        import sys

        if not hasattr(self.mw, '_notes_model'):
            return

        path_item = self.mw._notes_model.item(row, 1)
        if path_item:
            file_path = path_item.data(Qt.ItemDataRole.UserRole)
            if file_path and Path(file_path).exists():
                folder_path = Path(file_path).parent
                try:
                    if sys.platform == 'win32':
                        subprocess.run(['explorer', '/select,', str(file_path)])
                    elif sys.platform == 'darwin':
                        subprocess.run(['open', '-R', str(file_path)])
                    else:
                        subprocess.run(['xdg-open', str(folder_path)])
                    self.mw._log_status_message(f"Opened folder: {folder_path.name}", 1500)
                except Exception as e:
                    self.mw._log_status_message(f"Error opening folder: {e}", 3000)

    def _on_notes_action_open(self, row: int):
        """Open notes file at given row."""
        import os
        import subprocess
        import sys

        if not hasattr(self.mw, '_notes_model'):
            return

        path_item = self.mw._notes_model.item(row, 1)
        if path_item:
            file_path = path_item.data(Qt.ItemDataRole.UserRole)
            if file_path and Path(file_path).exists():
                try:
                    if sys.platform == 'win32':
                        os.startfile(file_path)
                    elif sys.platform == 'darwin':
                        subprocess.run(['open', file_path])
                    else:
                        subprocess.run(['xdg-open', file_path])
                    self.mw._log_status_message(f"Opened: {Path(file_path).name}", 1500)
                except Exception as e:
                    self.mw._log_status_message(f"Error opening file: {e}", 3000)

    def _on_notes_action_preview(self, row: int):
        """Preview notes file at given row."""
        if not hasattr(self.mw, '_notes_model'):
            return

        path_item = self.mw._notes_model.item(row, 1)
        if not path_item:
            return

        file_path = Path(path_item.data(Qt.ItemDataRole.UserRole))
        if not file_path.exists():
            self.mw._log_status_message("File not found", 1500)
            return

        self.mw._show_notes_preview_dialog(
            files=[{'name': file_path.name, 'path': str(file_path)}],
            title=f"Preview: {file_path.name}"
        )

    # =========================================================================
    # Notes Files Tab - Open, Preview, Link
    # =========================================================================

    def on_notes_open(self):
        """Open selected notes file in default application."""
        import os
        import subprocess
        import sys

        if not hasattr(self.mw, 'notesFilesTable'):
            return

        indexes = self.mw.notesFilesTable.selectionModel().selectedRows()
        if not indexes:
            self.mw._log_status_message("Please select a file to open", 1500)
            return

        row = indexes[0].row()
        path_item = self.mw._notes_model.item(row, 1)
        if path_item:
            file_path = path_item.data(Qt.ItemDataRole.UserRole)
            if file_path and Path(file_path).exists():
                try:
                    if sys.platform == 'win32':
                        os.startfile(file_path)
                    elif sys.platform == 'darwin':
                        subprocess.run(['open', file_path])
                    else:
                        subprocess.run(['xdg-open', file_path])
                    self.mw._log_status_message(f"Opened: {Path(file_path).name}", 1500)
                except Exception as e:
                    self.mw._log_status_message(f"Error opening file: {e}", 3000)

    def on_notes_preview(self):
        """Preview contents of selected notes file."""
        if not hasattr(self.mw, 'notesFilesTable'):
            return

        indexes = self.mw.notesFilesTable.selectionModel().selectedRows()
        if not indexes:
            self.mw._log_status_message("Please select a file to preview", 1500)
            return

        row = indexes[0].row()
        path_item = self.mw._notes_model.item(row, 1)
        if not path_item:
            return

        file_path = Path(path_item.data(Qt.ItemDataRole.UserRole))
        if not file_path.exists():
            self.mw._log_status_message("File not found", 1500)
            return

        self.mw._show_notes_preview_dialog(
            files=[{'name': file_path.name, 'path': str(file_path)}],
            title=f"Preview: {file_path.name}"
        )

    def on_notes_link(self):
        """Link selected notes file to a data file in the project."""
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QLabel, QListWidget, QDialogButtonBox

        if not hasattr(self.mw, 'notesFilesTable'):
            return

        indexes = self.mw.notesFilesTable.selectionModel().selectedRows()
        if not indexes:
            self.mw._log_status_message("Please select a notes file to link", 1500)
            return

        row = indexes[0].row()
        path_item = self.mw._notes_model.item(row, 1)
        if not path_item:
            return

        notes_path = Path(path_item.data(Qt.ItemDataRole.UserRole))

        if not self.mw._master_file_list:
            self.mw._log_status_message("No data files in project. Scan for files first.", 2000)
            return

        # Create link dialog
        dialog = QDialog(self.mw)
        dialog.setWindowTitle(f"Link: {notes_path.name}")
        dialog.resize(500, 400)
        layout = QVBoxLayout(dialog)

        layout.addWidget(QLabel(f"Link '{notes_path.name}' to which data file?"))

        file_list = QListWidget()
        for task in self.mw._master_file_list:
            if not task.get('is_sub_row'):
                file_list.addItem(task.get('file_name', str(task.get('file_path', 'Unknown'))))
        layout.addWidget(file_list)

        buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        buttons.accepted.connect(dialog.accept)
        buttons.rejected.connect(dialog.reject)
        layout.addWidget(buttons)

        if dialog.exec() == QDialog.DialogCode.Accepted:
            selected = file_list.currentItem()
            if selected:
                self.mw._log_status_message(f"Linked '{notes_path.name}' to '{selected.text()}'", 2000)
            else:
                self.mw._log_status_message("No file selected", 1500)

    # =========================================================================
    # Notes Files Tab - Scan References
    # =========================================================================

    def on_notes_scan_references(self):
        """Scan notes files for references to data files."""
        import tempfile
        import shutil

        if not self.mw._notes_files_data:
            self.mw._log_status_message("No notes files to scan. Search for files first.", 2000)
            return

        # Filter to files with 'Use' checkbox checked
        files_to_scan = [
            (i, f) for i, f in enumerate(self.mw._notes_files_data)
            if f.get('use_as_notes', False)
        ]

        if not files_to_scan:
            self.mw._log_status_message("No files selected. Check 'Use' boxes.", 2000)
            return

        # Get list of data filenames
        abf_filenames = self._get_project_abf_filenames()
        if not abf_filenames:
            self.mw._log_status_message("No data files in project to match against.", 2000)
            return

        print(f"[notes-scan] Scanning {len(files_to_scan)} selected files for {len(abf_filenames)} references...")

        # Create worker thread
        class ScanWorker(QThread):
            progress_update = pyqtSignal(int, int, int, str, list)
            status_update = pyqtSignal(str)
            finished = pyqtSignal(int, int)

            def __init__(self, files_to_scan, abf_names, scan_func):
                super().__init__()
                self.files_to_scan = files_to_scan
                self.abf_names = abf_names
                self.scan_func = scan_func
                self._cancelled = False
                self.cache_dir = None

            def cancel(self):
                self._cancelled = True

            def run(self):
                import hashlib
                import time

                total_matches = 0
                files_with_matches = 0

                # Use persistent cache directory
                cache_base = Path(tempfile.gettempdir()) / "physiometrics_notes_cache"
                cache_base.mkdir(exist_ok=True)
                self.cache_dir = cache_base

                # Stage 1: Copy files locally
                self.status_update.emit("Checking cache / copying files...")
                local_files = []

                for progress_idx, (original_idx, file_info) in enumerate(self.files_to_scan):
                    if self._cancelled:
                        break

                    src_path = Path(file_info['path'])
                    path_hash = hashlib.md5(str(src_path).encode()).hexdigest()[:8]
                    cache_name = f"{path_hash}_{src_path.name}"
                    cache_path = cache_base / cache_name

                    try:
                        if cache_path.exists():
                            src_mtime = src_path.stat().st_mtime
                            cache_mtime = cache_path.stat().st_mtime
                            if cache_mtime >= src_mtime:
                                local_files.append((original_idx, str(cache_path), file_info))
                                self.progress_update.emit(0, progress_idx, original_idx, f"(cached) {file_info['name']}", [])
                                continue

                        shutil.copy2(src_path, cache_path)
                        local_files.append((original_idx, str(cache_path), file_info))
                        self.progress_update.emit(0, progress_idx, original_idx, file_info['name'], [])
                    except Exception as e:
                        print(f"[notes-scan] Failed to copy {src_path.name}: {e}")

                if self._cancelled:
                    return

                # Stage 2: Scan local copies
                self.status_update.emit("Scanning files...")

                def scan_single_file(args):
                    original_idx, local_path, file_info = args
                    try:
                        matches = self.scan_func(local_path, self.abf_names)
                    except Exception as e:
                        print(f"[notes-scan] Error scanning {file_info['name']}: {e}")
                        matches = []
                    return (original_idx, file_info, matches, local_path)

                try:
                    from concurrent.futures import ThreadPoolExecutor, as_completed
                    import os

                    max_workers = min(os.cpu_count() or 4, 8)

                    with ThreadPoolExecutor(max_workers=max_workers) as executor:
                        futures = {
                            executor.submit(scan_single_file, (orig_idx, path, info)): orig_idx
                            for orig_idx, path, info in local_files
                        }

                        completed = 0
                        for future in as_completed(futures):
                            if self._cancelled:
                                executor.shutdown(wait=False, cancel_futures=True)
                                break

                            original_idx, file_info, matches, local_path = future.result()

                            if matches:
                                files_with_matches += 1
                                total_matches += len(matches)

                            self.progress_update.emit(1, completed, original_idx, file_info['name'], matches)
                            completed += 1

                except Exception as parallel_error:
                    # Fallback to sequential
                    print(f"[notes-scan] Parallel failed, using sequential: {parallel_error}")
                    self.status_update.emit("Scanning files (sequential)...")

                    for progress_idx, (original_idx, local_path, file_info) in enumerate(local_files):
                        if self._cancelled:
                            break

                        try:
                            matches = self.scan_func(local_path, self.abf_names)
                        except Exception as e:
                            matches = []

                        if matches:
                            files_with_matches += 1
                            total_matches += len(matches)

                        self.progress_update.emit(1, progress_idx, original_idx, file_info['name'], matches)

                self.finished.emit(files_with_matches, total_matches)

        # Create progress dialog
        progress = QProgressDialog("Copying files locally...", "Cancel", 0, len(files_to_scan) * 2, self.mw)
        progress.setWindowModality(Qt.WindowModality.WindowModal)
        progress.setMinimumDuration(0)
        progress.setValue(0)

        # Create and store worker
        self._scan_worker = ScanWorker(files_to_scan, abf_filenames, self._scan_file_for_abf_references)

        def on_status(msg):
            progress.setLabelText(msg)

        def on_progress(stage, progress_idx, original_idx, filename, matches):
            if progress.wasCanceled():
                self._scan_worker.cancel()
                return

            total_progress = progress_idx + (len(files_to_scan) if stage == 1 else 0)
            progress.setValue(total_progress + 1)

            if stage == 0:
                progress.setLabelText(f"Copying ({progress_idx+1}/{len(files_to_scan)}): {filename}")
            else:
                progress.setLabelText(f"Scanning ({progress_idx+1}/{len(files_to_scan)}): {filename}")

                if original_idx < len(self.mw._notes_files_data):
                    self.mw._notes_files_data[original_idx]['matches'] = matches
                    self.mw._notes_files_data[original_idx]['match_count'] = len(matches)

                if original_idx < self.mw._notes_model.rowCount():
                    matches_item = self.mw._notes_model.item(original_idx, 4)
                    if matches_item:
                        count = len(matches)
                        matches_item.setText(str(count) if count > 0 else "0")

        def on_finished(files_with_matches, total_matches):
            progress.close()
            msg = f"Scan complete: {files_with_matches} files with {total_matches} total references"
            print(f"[notes-scan] {msg}")
            self.mw._log_status_message(msg, 3000)

            # Update linked_notes column
            self.mw._update_linked_notes_column()

            # Auto-hide unmatched
            self._notes_hide_unmatched = True
            self._apply_hide_unmatched(show_status=False)

        # Connect signals
        self._scan_worker.status_update.connect(on_status)
        self._scan_worker.progress_update.connect(on_progress)
        self._scan_worker.finished.connect(on_finished)
        progress.canceled.connect(self._scan_worker.cancel)

        # Start scanning
        self._scan_worker.start()

    def _get_project_abf_filenames(self) -> list:
        """Get list of data file identifiers from the project file table."""
        abf_names = []

        if hasattr(self.mw, '_master_file_list') and self.mw._master_file_list:
            for file_info in self.mw._master_file_list:
                if isinstance(file_info, dict):
                    file_name = file_info.get('file_name', '')
                    file_type = file_info.get('file_type', '')
                    file_path = file_info.get('file_path', '')

                    if file_name:
                        stem = Path(file_name).stem
                        if stem not in abf_names:
                            abf_names.append(stem)

                        # For photometry files, add parent folder identifiers
                        if file_type == 'photometry' and file_path:
                            fp_path = Path(file_path)
                            for parent in fp_path.parents:
                                parent_name = parent.name
                                if (parent_name and
                                    not parent_name.lower().startswith('fp_data') and
                                    parent_name not in abf_names and
                                    len(parent_name) > 2):
                                    abf_names.append(parent_name)
                                    break

                elif isinstance(file_info, (str, Path)):
                    stem = Path(file_info).stem
                    if stem not in abf_names:
                        abf_names.append(stem)

        return abf_names

    def _scan_file_for_abf_references(self, file_path: str, abf_filenames: list) -> list:
        """Scan a single file for references to ABF filenames."""
        import re

        file_path = Path(file_path)
        if not file_path.exists():
            return []

        matches = []
        suffix = file_path.suffix.lower()

        try:
            content = ""

            if suffix in ['.txt', '.csv']:
                with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                    content = f.read()

            elif suffix == '.xlsx':
                try:
                    from openpyxl import load_workbook

                    escaped_names = [re.escape(name) for name in abf_filenames]
                    pattern = re.compile('|'.join(escaped_names), re.IGNORECASE)
                    name_lookup = {name.lower(): name for name in abf_filenames}
                    found_set = set()

                    wb = load_workbook(file_path, read_only=True, data_only=True)

                    for sheet in wb.worksheets:
                        sheet_text_parts = []
                        for row in sheet.iter_rows(values_only=True):
                            for cell in row:
                                if cell is not None:
                                    sheet_text_parts.append(str(cell))

                        if sheet_text_parts:
                            sheet_text = ' '.join(sheet_text_parts)
                            for match in pattern.findall(sheet_text):
                                found_set.add(name_lookup.get(match.lower(), match))

                        if len(found_set) >= len(abf_filenames):
                            break

                    wb.close()
                    return list(found_set)

                except Exception as e:
                    print(f"[notes-scan] Could not read Excel file {file_path.name}: {e}")
                    return []

            elif suffix == '.xls':
                try:
                    import pandas as pd
                    xl = pd.ExcelFile(file_path)
                    parts = []
                    for sheet_name in xl.sheet_names:
                        df = pd.read_excel(xl, sheet_name=sheet_name, dtype=str)
                        parts.append(df.fillna('').astype(str).values.flatten())
                    content = ' '.join(' '.join(p) for p in parts)
                except Exception as e:
                    print(f"[notes-scan] Could not read .xls file {file_path.name}: {e}")
                    return []

            elif suffix in ['.docx']:
                try:
                    from docx import Document
                    doc = Document(file_path)
                    paragraphs = [p.text for p in doc.paragraphs]
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                paragraphs.append(cell.text)
                    content = ' '.join(paragraphs)
                except ImportError:
                    print(f"[notes-scan] python-docx not installed")
                    return []
                except Exception as e:
                    print(f"[notes-scan] Could not read Word file {file_path.name}: {e}")
                    return []

            elif suffix == '.doc':
                return []

            elif suffix == '.pdf':
                return []

            # Search for references
            content_lower = content.lower()
            for abf_name in abf_filenames:
                if abf_name.lower() in content_lower:
                    if abf_name not in matches:
                        matches.append(abf_name)

        except Exception as e:
            print(f"[notes-scan] Error scanning {file_path.name}: {e}")

        return matches

    # =========================================================================
    # Notes Files Tab - Toggle/Hide Unmatched
    # =========================================================================

    def on_notes_toggle_unmatched(self):
        """Toggle visibility of notes files with no ABF matches."""
        self._notes_hide_unmatched = not self._notes_hide_unmatched
        self._apply_hide_unmatched()

    def _apply_hide_unmatched(self, show_status=True):
        """Apply the current hide/show unmatched setting."""
        table = self.mw.notesFilesTable
        hidden_count = 0

        for row in range(self.mw._notes_model.rowCount()):
            if row < len(self.mw._notes_files_data):
                match_count = self.mw._notes_files_data[row].get('match_count', 0)
                if match_count == 0 and self._notes_hide_unmatched:
                    table.setRowHidden(row, True)
                    hidden_count += 1
                else:
                    table.setRowHidden(row, False)

        # Update button text
        if hasattr(self.mw, 'toggleUnmatchedButton'):
            if self._notes_hide_unmatched:
                self.mw.toggleUnmatchedButton.setText("Show Unmatched")
            else:
                self.mw.toggleUnmatchedButton.setText("Hide Unmatched")

        if show_status:
            if self._notes_hide_unmatched:
                self.mw._log_status_message(f"Hidden {hidden_count} files with no matches", 2000)
            else:
                self.mw._log_status_message("Showing all files", 2000)

    def on_notes_build_ai_index(self):
        """Build AI-powered metadata index from selected notes files."""
        from PyQt6.QtWidgets import QMessageBox

        use_files = [f for f in self.mw._notes_files_data if f.get('use_as_notes', False)]

        if not use_files:
            QMessageBox.warning(
                self.mw,
                "No Notes Selected",
                "Please check the 'Use' column for notes files to include."
            )
            return

        abf_filenames = self._get_project_abf_filenames()

        QMessageBox.information(
            self.mw,
            "AI Index - Coming Soon",
            f"Ready to index {len(use_files)} notes files.\n\n"
            f"This feature will:\n"
            f"- Read selected notes files\n"
            f"- Send to AI for metadata extraction\n"
            f"- Build lookup table for {len(abf_filenames)} data files\n\n"
            f"Implementation in progress..."
        )

    # =========================================================================
    # Project Directory & File Scanning
    # =========================================================================

    def on_project_browse_directory(self):
        """Browse for directory containing recordings."""
        from PyQt6.QtWidgets import QFileDialog

        directory = QFileDialog.getExistingDirectory(
            self.mw,
            "Select Directory with Recordings",
            "",
            QFileDialog.Option.ShowDirsOnly
        )

        if directory:
            self.mw._project_directory = directory
            self.mw.directoryPathEdit.setText(directory)
            self.mw._log_status_message(f"Selected directory: {directory}", 2000)
            print(f"[project-builder] Selected directory: {directory}")

    def on_project_clear_files(self):
        """Clear the discovered files table and master file list with confirmation."""
        from PyQt6.QtWidgets import QMessageBox

        file_count = len(self.mw._master_file_list) if self.mw._master_file_list else 0

        if file_count == 0:
            self.mw._log_status_message("No files to clear", 1500)
            return

        msg = QMessageBox(self.mw)
        msg.setIcon(QMessageBox.Icon.Warning)
        msg.setWindowTitle("Clear All Files?")
        msg.setText(f"Are you sure you want to clear {file_count} files from the list?")
        msg.setInformativeText(
            "This will remove all files from the table.\n"
            "Your original data files will NOT be deleted.\n\n"
            "Any unsaved metadata changes will be lost."
        )
        msg.setStandardButtons(QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        msg.setDefaultButton(QMessageBox.StandardButton.No)

        if msg.exec() != QMessageBox.StandardButton.Yes:
            return

        # Clear everything
        self.mw._file_table_model.clear()
        self.mw._discovered_files_data = []
        self.mw._master_file_list = []
        self.mw.summaryLabel.setText("Summary: No files scanned")
        self.mw.projectProgressBar.setVisible(False)
        self.mw.projectProgressBar.setValue(0)

        if hasattr(self.mw, 'tableFilterEdit'):
            self.mw.tableFilterEdit.clear()
        if hasattr(self.mw, 'filterCountLabel'):
            self.mw.filterCountLabel.setText("")

        self.mw._log_status_message("Cleared all files", 1500)

    def get_selected_file_types(self) -> list:
        """Get list of file types to scan based on checkbox states."""
        file_types = []
        if hasattr(self.mw, 'scanAbfCheckbox') and self.mw.scanAbfCheckbox.isChecked():
            file_types.append('abf')
        if hasattr(self.mw, 'scanSmrxCheckbox') and self.mw.scanSmrxCheckbox.isChecked():
            file_types.append('smrx')
        if hasattr(self.mw, 'scanEdfCheckbox') and self.mw.scanEdfCheckbox.isChecked():
            file_types.append('edf')
        if hasattr(self.mw, 'scanPhotometryCheckbox') and self.mw.scanPhotometryCheckbox.isChecked():
            file_types.append('photometry')
        if hasattr(self.mw, 'scanNotesCheckbox') and self.mw.scanNotesCheckbox.isChecked():
            file_types.append('notes')
        return file_types if file_types else None

    # =========================================================================
    # Project Save/Load
    # =========================================================================

    def on_project_new(self):
        """Create a new project - prompt for name and directory."""
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QHBoxLayout, QLabel, QLineEdit, QPushButton, QFileDialog

        dialog = QDialog(self.mw)
        dialog.setWindowTitle("New Project")
        dialog.setMinimumWidth(500)

        layout = QVBoxLayout(dialog)

        layout.addWidget(QLabel("Project Name:"))
        name_edit = QLineEdit()
        name_edit.setPlaceholderText("Enter project name...")
        layout.addWidget(name_edit)

        layout.addSpacing(10)

        layout.addWidget(QLabel("Data Directory:"))
        dir_layout = QHBoxLayout()
        dir_edit = QLineEdit()
        dir_edit.setPlaceholderText("Select directory containing data files...")
        dir_edit.setReadOnly(True)
        dir_layout.addWidget(dir_edit)

        browse_btn = QPushButton("Browse...")
        browse_btn.setMaximumWidth(100)
        def browse_dir():
            directory = QFileDialog.getExistingDirectory(dialog, "Select Data Directory")
            if directory:
                dir_edit.setText(directory)
                if not name_edit.text().strip():
                    name_edit.setText(Path(directory).name)
        browse_btn.clicked.connect(browse_dir)
        dir_layout.addWidget(browse_btn)
        layout.addLayout(dir_layout)

        layout.addSpacing(20)

        button_layout = QHBoxLayout()
        button_layout.addStretch()

        cancel_btn = QPushButton("Cancel")
        cancel_btn.clicked.connect(dialog.reject)
        button_layout.addWidget(cancel_btn)

        create_btn = QPushButton("Create Project")
        create_btn.setDefault(True)
        create_btn.clicked.connect(dialog.accept)
        button_layout.addWidget(create_btn)

        layout.addLayout(button_layout)

        if dialog.exec() == QDialog.DialogCode.Accepted:
            project_name = name_edit.text().strip()
            directory = dir_edit.text().strip()

            if not project_name:
                self.mw._show_warning("Missing Information", "Please enter a project name.")
                return

            if not directory:
                self.mw._show_warning("Missing Information", "Please select a data directory.")
                return

            # Clear everything
            self.mw._file_table_model.clear()
            self.mw._discovered_files_data = []
            self.mw._master_file_list = []

            # Set new project info
            self.mw._project_directory = directory
            self.mw.directoryPathEdit.setText(directory)
            self.mw._current_project_name = project_name
            if hasattr(self.mw, 'projectNameCombo'):
                self.mw.projectNameCombo.setCurrentText(project_name)
            self.mw.summaryLabel.setText("Summary: No files scanned")

            self.mw._log_status_message(f"New project created: {project_name}", 2000)

    def on_project_save(self):
        """Save current project to data directory."""
        from PyQt6.QtWidgets import QInputDialog

        if not self.mw._project_directory:
            self.mw._show_warning("No Directory", "Please select a directory first.")
            return

        if not self.mw._discovered_files_data:
            self.mw._show_warning("No Files", "Please scan for files first.")
            return

        project_name = getattr(self.mw, '_current_project_name', '') or ''
        if not project_name.strip():
            project_name, ok = QInputDialog.getText(
                self.mw, "Save Project", "Enter project name:",
                text=Path(self.mw._project_directory).name
            )
            if not ok or not project_name.strip():
                return
            project_name = project_name.strip()
            self.mw._current_project_name = project_name
            if hasattr(self.mw, 'projectNameCombo'):
                self.mw.projectNameCombo.setCurrentText(project_name)

        try:
            files_to_save = self.mw._master_file_list if self.mw._master_file_list else self.mw._discovered_files_data
            print(f"[project-save] Saving {len(files_to_save)} files from master list")

            project_path = self.mw.project_manager.save_project(
                project_name,
                Path(self.mw._project_directory),
                files_to_save,
                [],
                notes_directory=self.mw._notes_directory,
                notes_files=self.mw._notes_files_data
            )

            self.mw._show_info("Project Saved", f"Project saved to:\n{project_path}")
            self.mw._log_status_message(f"Project saved: {project_name}", 3000)

            self.mw._populate_load_project_combo()

        except Exception as e:
            self.mw._show_error("Save Failed", f"Failed to save project:\n{e}")
            print(f"[project] Error saving: {e}")
            import traceback
            traceback.print_exc()

    def cancel_pending_autosave(self):
        """Cancel any pending autosave operation."""
        if hasattr(self.mw, '_autosave_timer') and self.mw._autosave_timer:
            self.mw._autosave_timer.stop()
            print("[autosave] Cancelled pending autosave")

    def project_autosave(self):
        """Silently autosave the project if it has a name and directory."""
        if not self.mw._project_directory or not self.mw._master_file_list:
            return

        project_name = getattr(self.mw, '_current_project_name', '') or ''
        if not project_name.strip():
            return

        try:
            if not hasattr(self.mw, '_autosave_timer'):
                self.mw._autosave_timer = QTimer()
                self.mw._autosave_timer.setSingleShot(True)
                self.mw._autosave_timer.timeout.connect(self._do_autosave)

            self.mw._autosave_timer.start(1000)
        except Exception as e:
            print(f"[autosave] Error setting up autosave: {e}")

    def _do_autosave(self):
        """Actually perform the autosave."""
        if not self.mw._project_directory or not self.mw._master_file_list:
            return

        project_name = getattr(self.mw, '_current_project_name', '') or ''
        if not project_name.strip():
            return

        try:
            self.mw.project_manager.save_project(
                project_name,
                Path(self.mw._project_directory),
                self.mw._master_file_list,
                [],
                notes_directory=self.mw._notes_directory,
                notes_files=self.mw._notes_files_data
            )
            print(f"[autosave] Project autosaved: {project_name}")
        except Exception as e:
            print(f"[autosave] Error: {e}")

    # =========================================================================
    # Table Management Methods
    # =========================================================================

    def mark_active_analysis_complete(self, channel_used: str = None, stim_channel_used: str = None,
                                       events_channel_used: str = None, export_info: dict = None):
        """
        Mark the currently active master list row as completed.
        Called after successful export/save of analysis data.

        Args:
            channel_used: The channel that was analyzed (e.g., "AD0")
            stim_channel_used: The stim channel used (e.g., "AD1")
            events_channel_used: The events channel used (e.g., "AD2")
            export_info: Dict with export metadata and file flags:
                - export_path: Where files were saved
                - export_date: When exported
                - export_version: App version used
                - exports: Dict of boolean flags for each export type
                - strain, stim_type, power, sex, animal_id: Metadata from dialog
        """
        # Get current file path from state
        current_file_path = str(self.mw.state.in_path) if self.mw.state.in_path else None

        # Find matching row(s) by file path if active row not set
        row = self.mw._active_master_list_row

        if row is None and current_file_path:
            # Try to find a row matching the current file
            # Priority: 1) Same file + same channel, 2) Same file + completed (different channel), 3) Same file + pending
            current_normalized = str(Path(current_file_path).resolve())

            matching_rows = []  # [(row_idx, task), ...]
            for i, task in enumerate(self.mw._master_file_list):
                task_path = task.get('file_path', '')
                if task_path:
                    task_normalized = str(Path(task_path).resolve())
                    if task_normalized == current_normalized:
                        matching_rows.append((i, task))

            if matching_rows:
                # Priority 1: Find row with same channel (to update existing)
                for i, task in matching_rows:
                    if task.get('channel', '') == channel_used:
                        row = i
                        print(f"[master-list] Found matching row {i} with same channel {channel_used}")
                        break

                # Priority 2: Find completed row with different channel (to create sub-row)
                if row is None:
                    for i, task in matching_rows:
                        if task.get('status', 'pending') == 'completed':
                            row = i
                            print(f"[master-list] Found completed row {i} (channel: {task.get('channel', '')})")
                            break

                # Priority 3: Find any pending row (first time analysis)
                if row is None:
                    for i, task in matching_rows:
                        if task.get('status', 'pending') == 'pending':
                            row = i
                            print(f"[master-list] Found pending row {i}")
                            break

                # Fallback: Use first matching row
                if row is None and matching_rows:
                    row = matching_rows[0][0]
                    print(f"[master-list] Using first matching row {row}")

        if row is None:
            print(f"[master-list] No matching row found for file: {current_file_path}")
            return

        if row >= len(self.mw._master_file_list):
            return

        task = self.mw._master_file_list[row]

        # Check if this is a parent row or sub-row
        is_sub_row = task.get('is_sub_row', False)
        existing_channel = task.get('channel', '')
        existing_status = task.get('status', 'pending')

        # Debug logging to understand the flow
        print(f"[master-list] mark_active_analysis_complete called:")
        print(f"  - Row: {row}")
        print(f"  - File: {task.get('file_name', 'unknown')}")
        print(f"  - Is sub-row: {is_sub_row}")
        print(f"  - Existing channel in row: '{existing_channel}'")
        print(f"  - Channel just analyzed: '{channel_used}'")
        print(f"  - Existing status: '{existing_status}'")

        # For parent rows, ALWAYS create sub-rows instead of updating the parent
        # This keeps parent as a "header" row showing file-level info only
        if not is_sub_row:
            print(f"[master-list] Parent row detected - creating sub-row for {channel_used}")

            # If parent had existing completed analysis, move it to sub-row first
            if existing_channel and existing_status == 'completed':
                # Check if there's already a sub-row for the existing channel
                has_existing_subrow = False
                for t in self.mw._master_file_list:
                    if (t.get('is_sub_row') and
                        str(t.get('file_path')) == str(task.get('file_path')) and
                        t.get('channel') == existing_channel):
                        has_existing_subrow = True
                        break

                if not has_existing_subrow:
                    # Move existing channel data to a sub-row first
                    existing_export_info = {
                        'export_path': task.get('export_path', ''),
                        'export_date': task.get('export_date', ''),
                        'export_version': task.get('export_version', ''),
                        'exports': task.get('exports', {}),
                        'strain': task.get('strain', ''),
                        'stim_type': task.get('stim_type', ''),
                        'power': task.get('power', ''),
                        'sex': task.get('sex', ''),
                        'animal_id': task.get('animal_id', ''),
                    }
                    self.create_sub_row_from_analysis(
                        row, existing_channel,
                        task.get('stim_channel', ''),
                        task.get('events_channel', ''),
                        existing_export_info
                    )
                    print(f"[master-list]   - Moved existing {existing_channel} to sub-row")

            # Check if there's already a sub-row for the channel we're analyzing
            existing_subrow_idx = None
            for idx, t in enumerate(self.mw._master_file_list):
                if (t.get('is_sub_row') and
                    str(t.get('file_path')) == str(task.get('file_path')) and
                    t.get('channel') == channel_used):
                    existing_subrow_idx = idx
                    print(f"[master-list]   - Found existing sub-row at {idx} for {channel_used}")
                    break

            if existing_subrow_idx is not None:
                # Update existing sub-row instead of creating new one
                existing_task = self.mw._master_file_list[existing_subrow_idx]
                existing_task['status'] = 'completed'
                if stim_channel_used:
                    existing_task['stim_channel'] = stim_channel_used
                if events_channel_used:
                    existing_task['events_channel'] = events_channel_used
                if export_info:
                    existing_task['export_path'] = export_info.get('export_path', '')
                    existing_task['export_date'] = export_info.get('export_date', '')
                    existing_task['export_version'] = export_info.get('export_version', '')
                    existing_task['exports'] = export_info.get('exports', existing_task.get('exports', {}))
                    if export_info.get('strain'):
                        existing_task['strain'] = export_info['strain']
                    if export_info.get('stim_type'):
                        existing_task['stim_type'] = export_info['stim_type']
                    if export_info.get('power'):
                        existing_task['power'] = export_info['power']
                    if export_info.get('sex'):
                        existing_task['sex'] = export_info['sex']
                    if export_info.get('animal_id'):
                        existing_task['animal_id'] = export_info['animal_id']
                # Update the model
                if existing_subrow_idx < self.mw._file_table_model.rowCount():
                    self.mw._file_table_model.update_row(existing_subrow_idx, existing_task)
                print(f"[master-list]   - Updated existing sub-row for {channel_used}")
            else:
                # Create new sub-row for the new analysis
                self.create_sub_row_from_analysis(
                    row, channel_used, stim_channel_used, events_channel_used, export_info
                )
                print(f"[master-list]   - Created sub-row for {channel_used}")

            # Apply styling to update colors and clean parent
            self.apply_all_row_styling()

            self.mw._active_master_list_row = None
            return

        # For sub-rows, update the existing sub-row
        task['status'] = 'completed'
        if channel_used:
            task['channel'] = channel_used
        if stim_channel_used:
            task['stim_channel'] = stim_channel_used
        if events_channel_used:
            task['events_channel'] = events_channel_used

        # Update export tracking info
        if export_info:
            task['export_path'] = export_info.get('export_path', '')
            task['export_date'] = export_info.get('export_date', '')
            task['export_version'] = export_info.get('export_version', '')
            task['exports'] = export_info.get('exports', task.get('exports', {}))

            # Update metadata from save dialog
            if export_info.get('strain'):
                task['strain'] = export_info['strain']
            if export_info.get('stim_type'):
                task['stim_type'] = export_info['stim_type']
            if export_info.get('power'):
                task['power'] = export_info['power']
            if export_info.get('sex'):
                task['sex'] = export_info['sex']
            if export_info.get('animal_id'):
                task['animal_id'] = export_info['animal_id']

        # Update the model to refresh the table display
        if row < self.mw._file_table_model.rowCount():
            self.mw._file_table_model.update_row(row, task)

        print(f"[master-list] Marked row {row} as completed:")
        print(f"  - Channel: {channel_used}")
        print(f"  - Stim Ch: {stim_channel_used}")
        print(f"  - Events Ch: {events_channel_used}")
        if export_info:
            print(f"  - Export path: {export_info.get('export_path', 'N/A')}")
            exports = export_info.get('exports', {})
            saved = [k for k, v in exports.items() if v]
            print(f"  - Saved: {', '.join(saved) if saved else 'none'}")

        # Clear active row - analysis is done
        self.mw._active_master_list_row = None

    def rebuild_table_from_master_list(self):
        """Rebuild the table from the current _master_file_list order using Model/View."""
        # With Model/View, we simply update the model's data
        # The view automatically reflects the changes
        # Note: The model's _get_display_value() handles all formatting
        # (sub-row names, exports summary, status icons, etc.)

        # Sync custom columns from DB if available
        try:
            vm = getattr(self.mw, '_project_viewmodel', None)
            if vm is not None:
                db_cols = vm.get_custom_columns()
                if db_cols:
                    self.mw._file_table_model.sync_custom_columns_from_db(db_cols)
        except Exception as e:
            print(f"[project-builder] Error syncing custom columns: {e}")

        # Set data on model - this triggers view update
        self.mw._file_table_model.set_files(self.mw._master_file_list)

        # Track conflict rows for highlighting
        conflict_rows = set()
        for i, task in enumerate(self.mw._master_file_list):
            warnings = task.get('scan_warnings', {})
            if isinstance(warnings, str):
                # scan_warnings may be stored as a string repr in JSON
                try:
                    import ast
                    warnings = ast.literal_eval(warnings)
                except (ValueError, SyntaxError):
                    warnings = {}
            if isinstance(warnings, dict) and warnings.get('conflicts'):
                conflict_rows.add(i)
        self.mw._file_table_model.set_conflict_rows(conflict_rows)

        print(f"[project-builder] Table rebuilt with {len(self.mw._master_file_list)} rows")

    def get_experiment_history(self) -> list:
        """Get list of previously used experiment names from QSettings."""
        from PyQt6.QtCore import QSettings
        settings = QSettings("PhysioMetrics", "BreathAnalysis")
        history = settings.value("project_builder/experiment_history", [])
        return history if isinstance(history, list) else []

    def update_experiment_history(self, experiment_name: str):
        """Add an experiment name to history (max 50 entries, most recent first)."""
        from PyQt6.QtCore import QSettings
        settings = QSettings("PhysioMetrics", "BreathAnalysis")

        history = self.get_experiment_history()

        # Remove if already exists (to move to front)
        if experiment_name in history:
            history.remove(experiment_name)

        # Add to front
        history.insert(0, experiment_name)

        # Keep max 50
        history = history[:50]

        settings.setValue("project_builder/experiment_history", history)

    def auto_fit_table_columns(self):
        """
        Auto-fit column widths for QTableView.
        Uses column definitions from the model for sizing hints.
        """
        table = self.mw.discoveredFilesTable
        header = table.horizontalHeader()
        model = self.mw._file_table_model

        if model.columnCount() == 0:
            return

        # Check if full content mode is enabled
        full_content_mode = getattr(self.mw, 'tableFullContentCheckbox', None)
        full_content_mode = full_content_mode.isChecked() if full_content_mode else False

        # Get column definitions from model
        visible_columns = model.get_visible_columns()

        # First, resize to contents to get content-based widths
        table.resizeColumnsToContents()

        if full_content_mode:
            # Full content mode: show all content, enable horizontal scrolling
            table.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
            # Just use content widths, but ensure minimum widths from column defs
            for i, col_def in enumerate(visible_columns):
                current_width = table.columnWidth(i)
                min_width = col_def.min_width
                if current_width < min_width:
                    table.setColumnWidth(i, min_width)
        else:
            # Fit-to-view mode: constrain to visible width
            table.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)

            # Get available width
            available_width = table.viewport().width()
            if available_width < 100:
                available_width = table.width() - 20

            # Calculate column widths based on column definitions
            widths = []
            expandable_indices = []
            total_fixed = 0

            for i, col_def in enumerate(visible_columns):
                content_width = table.columnWidth(i)
                base_width = max(col_def.min_width, min(content_width, col_def.width + 30))

                if col_def.fixed:
                    # Fixed columns use their defined width
                    widths.append(col_def.width)
                    total_fixed += col_def.width
                elif col_def.expandable:
                    widths.append(base_width)
                    expandable_indices.append(i)
                else:
                    widths.append(base_width)

            total_width = sum(widths)

            # Distribute extra space or shrink as needed
            if total_width < available_width and expandable_indices:
                # Distribute extra space to expandable columns
                extra = available_width - total_width
                per_col = extra // len(expandable_indices)
                for idx in expandable_indices:
                    widths[idx] += per_col
            elif total_width > available_width:
                # Shrink proportionally, but respect minimums
                scale = available_width / total_width
                for i, col_def in enumerate(visible_columns):
                    if not col_def.fixed:
                        widths[i] = max(col_def.min_width, int(widths[i] * scale))

            # Apply widths
            for i, width in enumerate(widths):
                table.setColumnWidth(i, width)

    def apply_row_styling(self, row: int, is_sub_row: bool = False):
        """Apply visual styling - now handled by model via data() method."""
        # Styling is handled by FileTableModel.data() for BackgroundRole, ForegroundRole, FontRole
        pass

    def apply_all_row_styling(self):
        """Apply styling to all rows - cleans parent rows that have sub-rows."""
        # Styling is now handled by the model, but we still need to clean task dicts
        self.clean_parent_rows_with_subrows()

    def clean_parent_rows_with_subrows(self):
        """Clear channel-specific fields from parent tasks that have sub-rows."""
        # Find parent rows that have sub-rows
        parent_to_subrows = {}  # parent_file_path -> [sub_row_indices]

        for row, task in enumerate(self.mw._master_file_list):
            if task.get('is_sub_row'):
                parent_path = str(task.get('file_path', ''))
                if parent_path not in parent_to_subrows:
                    parent_to_subrows[parent_path] = []
                parent_to_subrows[parent_path].append(row)

        # Clear fields in parent tasks that have sub-rows
        for row, task in enumerate(self.mw._master_file_list):
            if task.get('is_sub_row'):
                continue

            file_path = str(task.get('file_path', ''))
            if file_path in parent_to_subrows:
                sub_rows = parent_to_subrows[file_path]

                # Count completed sub-rows for status display
                completed = sum(1 for sr in sub_rows
                                if self.mw._master_file_list[sr].get('status') == 'completed')

                # Clear channel-specific fields in the task dict
                task['channel'] = ''
                task['stim_channel'] = ''
                task['events_channel'] = ''
                task['animal_id'] = ''
                task['status'] = f"{completed} ✓" if completed > 0 else ''
                task['exports'] = {}

    def update_row_status_icon(self, row, task):
        """Update the status icon for a row based on its current state.

        The model's data() method handles the actual status icon display.
        This method just triggers a refresh of the row.
        """
        if row >= self.mw._file_table_model.rowCount():
            return

        # Update the model row to refresh the display
        self.mw._file_table_model.update_row(row, task)

    def create_sub_row_from_analysis(self, source_row: int, channel_used: str,
                                      stim_channel_used: str, events_channel_used: str,
                                      export_info: dict):
        """
        Create a new sub-row when user analyzes a different channel from an existing analyzed row.

        This prevents overwriting previous analysis when the user analyzes the same file
        but with a different channel.

        Args:
            source_row: Row index of the source task
            channel_used: The channel that was analyzed
            stim_channel_used: The stim channel used
            events_channel_used: The events channel used
            export_info: Export metadata dict from save dialog
        """
        if source_row >= len(self.mw._master_file_list):
            return

        source_task = self.mw._master_file_list[source_row]
        file_path = source_task.get('file_path')

        # Create new task based on source but with new channel and analysis info
        new_task = {
            'file_path': file_path,
            'file_name': source_task.get('file_name', ''),
            'protocol': source_task.get('protocol', ''),
            'channel_count': source_task.get('channel_count', 0),
            'sweep_count': source_task.get('sweep_count', 0),
            'channel_names': source_task.get('channel_names', []),
            'stim_channels': source_task.get('stim_channels', []),
            'path_keywords': source_task.get('path_keywords', {}),
            'keywords_display': source_task.get('keywords_display', ''),
            # Analysis results
            'channel': channel_used,
            'stim_channel': stim_channel_used or '',
            'events_channel': events_channel_used or '',
            # Copy or update metadata
            'strain': export_info.get('strain', '') if export_info else source_task.get('strain', ''),
            'stim_type': export_info.get('stim_type', '') if export_info else source_task.get('stim_type', ''),
            'power': export_info.get('power', '') if export_info else source_task.get('power', ''),
            'sex': export_info.get('sex', '') if export_info else source_task.get('sex', ''),
            'animal_id': export_info.get('animal_id', '') if export_info else '',
            'status': 'completed',
            'is_sub_row': True,
            'parent_file': file_path,
            # Export tracking
            'export_path': export_info.get('export_path', '') if export_info else '',
            'export_date': export_info.get('export_date', '') if export_info else '',
            'export_version': export_info.get('export_version', '') if export_info else '',
            'exports': export_info.get('exports', {}) if export_info else {
                'npz': False,
                'timeseries_csv': False,
                'breaths_csv': False,
                'events_csv': False,
                'pdf': False,
                'session_state': False,
                'ml_training': False,
            }
        }

        # Insert after the source row
        insert_row = source_row + 1
        self.mw._master_file_list.insert(insert_row, new_task)

        # Rebuild table from master list
        self.rebuild_table_from_master_list()

        print(f"[master-list] Created new sub-row at {insert_row} for channel {channel_used}")
        print(f"  - This preserves the previous analysis of {source_task.get('channel', 'unknown')} in row {source_row}")

    def create_sub_row_from_saved_data(self, source_task: dict, source_row: int,
                                        channel: str, info: dict):
        """
        Create a new sub-row from scanned saved data for an additional channel.

        Updates _master_file_list only - caller should rebuild table after all updates.

        Args:
            source_task: The source task dict to base the new row on
            source_row: Row index of the source task
            channel: The channel for this saved data
            info: Export info dict from NPZ metadata
        """
        file_path = source_task.get('file_path')

        # Track any older/duplicate NPZ files found
        older_files = info.get('older_npz_files', [])
        scan_warnings = {}
        if older_files:
            scan_warnings = {
                'conflicts': [],
                'older_npz_count': len(older_files),
                'older_npz_files': older_files,
            }

        # Create new task based on source but with new channel and analysis info
        new_task = {
            'file_path': file_path,
            'file_name': source_task.get('file_name', ''),
            'protocol': source_task.get('protocol', ''),
            'channel_count': source_task.get('channel_count', 0),
            'sweep_count': source_task.get('sweep_count', 0),
            'channel_names': source_task.get('channel_names', []),
            'stim_channels': source_task.get('stim_channels', []),
            'path_keywords': source_task.get('path_keywords', {}),
            'keywords_display': source_task.get('keywords_display', ''),
            # Analysis results
            'channel': channel,
            'stim_channel': info.get('stim_channel', ''),
            'events_channel': info.get('events_channel', ''),
            # Metadata from NPZ
            'strain': info.get('strain', '') or source_task.get('strain', ''),
            'stim_type': info.get('stim_type', '') or source_task.get('stim_type', ''),
            'power': info.get('power', '') or source_task.get('power', ''),
            'sex': info.get('sex', '') or source_task.get('sex', ''),
            'animal_id': info.get('animal_id', ''),
            'status': 'completed',
            'is_sub_row': True,
            'parent_file': file_path,
            # Export tracking
            'export_path': info.get('export_path', ''),
            'export_date': info.get('export_date', ''),
            'exports': {
                'npz': info.get('npz', False),
                'timeseries_csv': info.get('timeseries_csv', False),
                'breaths_csv': info.get('breaths_csv', False),
                'events_csv': info.get('events_csv', False),
                'pdf': info.get('pdf', False),
                'session_state': info.get('session_state', False),
                'ml_training': info.get('ml_training', False),
            }
        }
        if scan_warnings:
            new_task['scan_warnings'] = scan_warnings

        # Find where to insert - after all existing rows for this file
        insert_row = source_row + 1
        while insert_row < len(self.mw._master_file_list):
            next_task = self.mw._master_file_list[insert_row]
            if str(next_task.get('file_path')) != str(file_path):
                break
            insert_row += 1

        self.mw._master_file_list.insert(insert_row, new_task)

        print(f"[scan-saved] Created sub-row at {insert_row} for channel {channel}")

    def update_task_with_export_info(self, task: dict, info: dict, row: int):
        """
        Update a task dict with export info from saved data scan.

        Args:
            task: The task dict from _master_file_list
            info: Export info dict from NPZ metadata
            row: Table row index
        """
        # Track conflicts and multiple files
        conflicts = []
        older_files = info.get('older_npz_files', [])

        # Check for conflicts between existing task data and NPZ data
        metadata_fields = [
            ('strain', 'strain', 'Strain'),
            ('stim_type', 'stim_type', 'Stim Type'),
            ('power', 'power', 'Power'),
            ('sex', 'sex', 'Sex'),
            ('animal_id', 'animal_id', 'Animal ID'),
        ]
        for task_key, info_key, display_name in metadata_fields:
            task_val = task.get(task_key, '').strip()
            info_val = (info.get(info_key, '') or '').strip()
            if task_val and info_val and task_val != info_val:
                conflicts.append(f"{display_name}: table='{task_val}' vs NPZ='{info_val}'")

        # Store conflict info in task
        if conflicts or older_files:
            task['scan_warnings'] = {
                'conflicts': conflicts,
                'older_npz_count': len(older_files),
                'older_npz_files': older_files,
            }
            if conflicts:
                print(f"[scan-saved] Conflicts found for row {row}: {conflicts}")
            if older_files:
                print(f"[scan-saved] {len(older_files)} older NPZ file(s) found for row {row}")

        # Update task with export info
        task['export_path'] = info.get('export_path', '')
        task['export_date'] = info.get('export_date', '')
        task['status'] = 'completed'
        task['exports'] = {
            'npz': info.get('npz', False),
            'timeseries_csv': info.get('timeseries_csv', False),
            'breaths_csv': info.get('breaths_csv', False),
            'events_csv': info.get('events_csv', False),
            'pdf': info.get('pdf', False),
            'session_state': info.get('session_state', False),
            'ml_training': info.get('ml_training', False),
        }

        # Update metadata from NPZ if available and not already set (don't overwrite)
        if info.get('strain') and not task.get('strain'):
            task['strain'] = info['strain']
        if info.get('stim_type') and not task.get('stim_type'):
            task['stim_type'] = info['stim_type']
        if info.get('power') and not task.get('power'):
            task['power'] = info['power']
        if info.get('sex') and not task.get('sex'):
            task['sex'] = info['sex']
        if info.get('animal_id') and not task.get('animal_id'):
            task['animal_id'] = info['animal_id']
        if info.get('stim_channel') and not task.get('stim_channel'):
            task['stim_channel'] = info['stim_channel']
        if info.get('events_channel') and not task.get('events_channel'):
            task['events_channel'] = info['events_channel']

        # Update model row - the model will handle display
        if row < self.mw._file_table_model.rowCount():
            self.mw._file_table_model.update_row(row, task)
            # Mark conflict rows for highlighting
            sw = task.get('scan_warnings', {})
            if isinstance(sw, str):
                try:
                    import ast
                    sw = ast.literal_eval(sw)
                except (ValueError, SyntaxError):
                    sw = {}
            if isinstance(sw, dict) and sw.get('conflicts'):
                self.mw._file_table_model.add_conflict_row(row)

    def on_analyze_row(self, row):
        """Open the file from the specified row in the Analysis tab."""
        if row >= len(self.mw._master_file_list):
            return

        task = self.mw._master_file_list[row]
        file_path = task.get('file_path')

        if not file_path or not Path(file_path).exists():
            self.mw._show_warning("File Not Found", f"Cannot find file:\n{file_path}")
            return

        print(f"[master-list] Analyzing file from row {row}: {file_path}")

        # Track which row is being analyzed
        self.mw._active_master_list_row = row

        # Store pending channel selections if specified
        self.mw._pending_analysis_channel = task.get('channel', '')
        self.mw._pending_stim_channels = task.get('stim_channels', [])

        # Load the file
        self.mw.load_file(Path(file_path))

        # Switch to Analysis tab (Tab 0 = Project Builder, Tab 1 = Analysis)
        if hasattr(self.mw, 'Tabs'):
            self.mw.Tabs.setCurrentIndex(1)

        # Update status to "in progress"
        task['status'] = 'in_progress'
        # Update the model row to reflect the new status
        if row < self.mw._file_table_model.rowCount():
            self.mw._file_table_model.update_row(row, task)

    def on_add_row_for_file(self, source_row, force_override=False):
        """Add a new row for the same file (for multi-channel/multi-animal analysis)."""
        if source_row >= len(self.mw._master_file_list):
            return

        source_task = self.mw._master_file_list[source_row]
        file_path = source_task.get('file_path')

        # Find the parent task (may be source_task or a different row)
        parent_task = source_task
        if source_task.get('is_sub_row'):
            # Find the actual parent row
            for task in self.mw._master_file_list:
                if not task.get('is_sub_row') and str(task.get('file_path')) == str(file_path):
                    parent_task = task
                    break

        # Count existing SUB-ROWS for this file (parent row doesn't count as an analysis)
        existing_subrows = []
        for i, task in enumerate(self.mw._master_file_list):
            if task.get('is_sub_row') and str(task.get('file_path')) == str(file_path):
                existing_subrows.append(i)

        # Calculate available analysis channels from PARENT task (has file metadata)
        total_channels = parent_task.get('channel_count', 0)

        # Get stim channels from file metadata
        stim_channels_from_metadata = parent_task.get('stim_channels', [])
        if not isinstance(stim_channels_from_metadata, list):
            stim_channels_from_metadata = [stim_channels_from_metadata] if stim_channels_from_metadata else []

        # Also check existing sub-rows for stim/events channels they use
        stim_events_from_subrows = set()
        for i in existing_subrows:
            task = self.mw._master_file_list[i]
            stim_ch = task.get('stim_channel', '')
            events_ch = task.get('events_channel', '')
            if stim_ch:
                stim_events_from_subrows.add(stim_ch)
            if events_ch:
                stim_events_from_subrows.add(events_ch)

        # Combine all non-analysis channels
        all_stim_channels = set(stim_channels_from_metadata) | stim_events_from_subrows
        num_stim_channels = len(all_stim_channels)
        available_analysis_channels = total_channels - num_stim_channels

        # Check if we can add more rows (compare sub-row count to available channels)
        if len(existing_subrows) >= available_analysis_channels and not force_override:
            from PyQt6.QtWidgets import QMessageBox
            msg = QMessageBox(self.mw)
            msg.setWindowTitle("Channel Limit Reached")
            msg.setText(
                f"This file has {total_channels} channels.\n"
                f"Detected {num_stim_channels} stim/events channel(s): {', '.join(sorted(all_stim_channels)) if all_stim_channels else 'none'}\n"
                f"Maximum {available_analysis_channels} analysis rows allowed.\n"
                f"Already have {len(existing_subrows)} analyzed channel(s)."
            )
            msg.setInformativeText("Do you want to add a row anyway?")
            msg.setStandardButtons(QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            msg.setDefaultButton(QMessageBox.StandardButton.No)
            if msg.exec() != QMessageBox.StandardButton.Yes:
                return
            # User clicked Yes - force override
            force_override = True

        # Find the next available channel (use parent_task for channel metadata)
        channel_names = parent_task.get('channel_names', [])
        if not channel_names:
            # Generate default channel names
            channel_names = [f"AD{i}" for i in range(total_channels)]

        # Filter out stim/events channels to get analysis channels
        analysis_channels = [ch for ch in channel_names if ch not in all_stim_channels]

        # Find which channels are already used (from sub-rows only)
        used_channels = set()
        for i in existing_subrows:
            ch = self.mw._master_file_list[i].get('channel', '')
            if ch:
                used_channels.add(ch)

        # Get next available channel
        next_channel = ''
        for ch in analysis_channels:
            if ch not in used_channels:
                next_channel = ch
                break

        # If no analysis channels available but user forced override, pick any unused channel
        if not next_channel and force_override:
            for ch in channel_names:
                if ch not in used_channels:
                    next_channel = ch
                    break

        # Create a new task based on parent task (for file metadata) and source task (for user metadata)
        new_task = {
            'file_path': file_path,
            'file_name': parent_task.get('file_name', ''),
            'protocol': parent_task.get('protocol', ''),
            'channel_count': total_channels,
            'sweep_count': parent_task.get('sweep_count', 0),
            'channel_names': channel_names,
            'stim_channels': list(all_stim_channels),
            'path_keywords': parent_task.get('path_keywords', {}),
            'keywords_display': parent_task.get('keywords_display', ''),
            # Auto-populated channel
            'channel': next_channel,
            'stim_channel': '',  # User can set this
            'events_channel': '',  # Filled on save
            # Copy metadata from source (may have user edits)
            'strain': source_task.get('strain', ''),
            'stim_type': source_task.get('stim_type', ''),
            'power': source_task.get('power', ''),
            'sex': source_task.get('sex', ''),
            'animal_id': '',  # Different animal - leave blank
            'status': 'pending',
            'is_sub_row': True,
            'parent_file': file_path,
            # Export tracking (filled on save)
            'export_path': '',
            'export_date': '',
            'export_version': '',
            'exports': {
                'npz': False,
                'timeseries_csv': False,
                'breaths_csv': False,
                'events_csv': False,
                'pdf': False,
                'session_state': False,
                'ml_training': False,
            }
        }

        # Insert after the source row
        insert_row = source_row + 1
        self.mw._master_file_list.insert(insert_row, new_task)

        # Rebuild table from master list (model handles all formatting)
        self.rebuild_table_from_master_list()

        remaining = available_analysis_channels - len(existing_subrows) - 1
        print(f"[master-list] Added sub-row at {insert_row} with channel {next_channel} ({remaining} more available)")

        # Autosave after adding row
        self.project_autosave()

    def on_remove_sub_row(self, row):
        """Remove a sub-row from the master list."""
        if row >= len(self.mw._master_file_list):
            return

        task = self.mw._master_file_list[row]

        # Only allow removing sub-rows, not primary rows
        if not task.get('is_sub_row', False):
            self.mw._show_warning("Cannot Remove",
                "Cannot remove the primary row for a file.\n"
                "Use 'Omit' to hide files you don't want to analyze.")
            return

        # Confirm deletion if this row has been analyzed
        if task.get('status') == 'completed':
            from PyQt6.QtWidgets import QMessageBox

            # Get export info for confirmation message
            export_summary = self.mw._format_exports_summary(task)
            export_path = task.get('export_path', '')
            channel = task.get('channel', 'unknown channel')

            msg = QMessageBox(self.mw)
            msg.setIcon(QMessageBox.Icon.Warning)
            msg.setWindowTitle("Delete Analyzed Row?")
            msg.setText(f"This row ({channel}) has been analyzed.")

            details = []
            if export_summary:
                details.append(f"Exports: {export_summary}")
            if export_path:
                details.append(f"Saved to: {export_path}")
            if details:
                msg.setInformativeText("\n".join(details) + "\n\nNote: The exported files will NOT be deleted.")

            msg.setDetailedText(
                "This will remove the row from the Project Builder list only.\n\n"
                "The exported data files on disk remain unchanged. "
                "You can add a new row and re-analyze if needed."
            )
            msg.setStandardButtons(QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.Cancel)
            msg.setDefaultButton(QMessageBox.StandardButton.Cancel)

            if msg.exec() != QMessageBox.StandardButton.Yes:
                return  # User cancelled

        # Remove from data and table
        table = self.mw.discoveredFilesTable
        table.blockSignals(True)

        del self.mw._master_file_list[row]
        table.removeRow(row)

        table.blockSignals(False)

        print(f"[master-list] Removed sub-row at {row}")

    def on_master_list_context_menu(self, position):
        """Show context menu for bulk editing selected rows."""
        from PyQt6.QtWidgets import QMenu, QInputDialog
        from PyQt6.QtGui import QAction

        table = self.mw.discoveredFilesTable
        selected_rows = set(index.row() for index in table.selectedIndexes())

        if not selected_rows:
            return

        menu = QMenu(self.mw)

        # Bulk edit options
        edit_menu = menu.addMenu(f"Set for {len(selected_rows)} selected rows")

        # Experiment (column 5) - with history suggestions
        experiment_menu = edit_menu.addMenu("Set Experiment")
        exp_history = self.get_experiment_history()
        if exp_history:
            for exp_name in exp_history[:10]:  # Show top 10
                exp_action = QAction(exp_name, self.mw)
                exp_action.triggered.connect(lambda checked, e=exp_name: self.bulk_set_column(selected_rows, 5, 'experiment', "Set Experiment", e))
                experiment_menu.addAction(exp_action)
            experiment_menu.addSeparator()
        custom_exp_action = QAction("Custom...", self.mw)
        custom_exp_action.triggered.connect(lambda: self.bulk_set_column(selected_rows, 5, 'experiment', "Set Experiment"))
        experiment_menu.addAction(custom_exp_action)

        # Strain (column 9)
        strain_action = QAction("Set Strain...", self.mw)
        strain_action.triggered.connect(lambda: self.bulk_set_column(selected_rows, 9, 'strain', "Set Strain"))
        edit_menu.addAction(strain_action)

        # Stim Type (column 10)
        stim_action = QAction("Set Stim Type...", self.mw)
        stim_action.triggered.connect(lambda: self.bulk_set_column(selected_rows, 10, 'stim_type', "Set Stim Type"))
        edit_menu.addAction(stim_action)

        # Power (column 11)
        power_action = QAction("Set Power...", self.mw)
        power_action.triggered.connect(lambda: self.bulk_set_column(selected_rows, 11, 'power', "Set Power"))
        edit_menu.addAction(power_action)

        # Sex (column 12)
        sex_action = QAction("Set Sex...", self.mw)
        sex_action.triggered.connect(lambda: self.bulk_set_column(selected_rows, 12, 'sex', "Set Sex"))
        edit_menu.addAction(sex_action)

        # Animal ID (column 13)
        animal_action = QAction("Set Animal ID...", self.mw)
        animal_action.triggered.connect(lambda: self.bulk_set_column(selected_rows, 13, 'animal_id', "Set Animal ID"))
        edit_menu.addAction(animal_action)

        # Channel (column 6)
        channel_action = QAction("Set Channel...", self.mw)
        channel_action.triggered.connect(lambda: self.bulk_set_column(selected_rows, 6, 'channel', "Set Channel"))
        edit_menu.addAction(channel_action)

        menu.addSeparator()

        # Combine files option (if multiple files selected)
        if len(selected_rows) > 1:
            combine_action = QAction(f"Combine {len(selected_rows)} files for analysis...", self.mw)
            combine_action.triggered.connect(lambda: self.combine_selected_files(selected_rows))
            menu.addAction(combine_action)

        # Check if any selected rows have scan warnings (conflicts or multiple NPZ files)
        rows_with_warnings = []
        rows_with_conflicts = []
        for row in selected_rows:
            if row < len(self.mw._master_file_list):
                task = self.mw._master_file_list[row]
                warnings = task.get('scan_warnings', {})
                if warnings:
                    rows_with_warnings.append(row)
                    if warnings.get('conflicts'):
                        rows_with_conflicts.append(row)

        if rows_with_warnings:
            menu.addSeparator()
            conflict_menu = menu.addMenu(f"⚠ Resolve Warnings ({len(rows_with_warnings)} rows)")

            if rows_with_conflicts:
                # Use NPZ values - overwrite table with NPZ data
                use_npz_action = QAction("Use NPZ values (overwrite table)", self.mw)
                use_npz_action.triggered.connect(lambda: self.resolve_conflicts_use_npz(rows_with_conflicts))
                conflict_menu.addAction(use_npz_action)

                # Keep table values - dismiss the warning
                keep_table_action = QAction("Keep table values (dismiss warning)", self.mw)
                keep_table_action.triggered.connect(lambda: self.resolve_conflicts_keep_table(rows_with_warnings))
                conflict_menu.addAction(keep_table_action)

                conflict_menu.addSeparator()

            # View details
            view_details_action = QAction("View warning details...", self.mw)
            view_details_action.triggered.connect(lambda: self.show_conflict_details(rows_with_warnings))
            conflict_menu.addAction(view_details_action)

            # Clear all warnings
            clear_action = QAction("Clear all warnings", self.mw)
            clear_action.triggered.connect(lambda: self.clear_scan_warnings(rows_with_warnings))
            conflict_menu.addAction(clear_action)

        # Export options
        menu.addSeparator()
        export_menu = menu.addMenu("Export Table")

        export_selected_action = QAction(f"Export {len(selected_rows)} selected rows to CSV...", self.mw)
        export_selected_action.triggered.connect(lambda: self.export_table_to_csv(selected_rows))
        export_menu.addAction(export_selected_action)

        export_all_action = QAction("Export all rows to CSV...", self.mw)
        export_all_action.triggered.connect(lambda: self.export_table_to_csv(None))
        export_menu.addAction(export_all_action)

        menu.exec(table.viewport().mapToGlobal(position))

    def export_table_to_csv(self, selected_rows=None):
        """Export table data to CSV file.

        Args:
            selected_rows: Set of row indices to export, or None for all rows
        """
        import csv
        from PyQt6.QtWidgets import QFileDialog

        # Determine which rows to export
        if selected_rows is None:
            rows_to_export = list(range(len(self.mw._master_file_list)))
        else:
            rows_to_export = sorted(selected_rows)

        if not rows_to_export:
            self.mw._show_warning("No Data", "No rows to export.")
            return

        # Get visible columns
        visible_cols = self.mw._file_table_model.get_visible_columns()

        # Build header row (skip actions column)
        headers = []
        col_keys = []
        for col_def in visible_cols:
            if col_def.key != 'actions':
                headers.append(col_def.header or col_def.key)
                col_keys.append(col_def.key)

        # Get save file path
        default_name = f"{getattr(self.mw, '_current_project_name', None) or 'project'}_table.csv"
        file_path, _ = QFileDialog.getSaveFileName(
            self.mw, "Export Table to CSV",
            str(Path(self.mw._project_directory) / default_name) if self.mw._project_directory else default_name,
            "CSV files (*.csv);;All files (*.*)"
        )

        if not file_path:
            return

        try:
            with open(file_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)

                # Write header
                writer.writerow(headers)

                # Write data rows
                for row_idx in rows_to_export:
                    if row_idx >= len(self.mw._master_file_list):
                        continue
                    task = self.mw._master_file_list[row_idx]

                    row_data = []
                    for key in col_keys:
                        value = task.get(key, '')

                        # Format special values
                        if key == 'file_name' and task.get('is_sub_row'):
                            channel = task.get('channel', '')
                            value = f"↳ {channel}"
                        elif key == 'exports' and isinstance(value, dict):
                            # Format exports dict
                            parts = []
                            csv_count = sum([
                                1 if value.get('timeseries_csv') else 0,
                                1 if value.get('breaths_csv') else 0,
                                1 if value.get('events_csv') else 0,
                            ])
                            if value.get('pdf'):
                                parts.append('PDF')
                            if csv_count > 0:
                                parts.append(f'{csv_count} CSV')
                            if value.get('npz'):
                                parts.append('NPZ')
                            if value.get('session_state'):
                                parts.append('Session')
                            value = ', '.join(parts) if parts else ''
                        elif key == 'file_path':
                            value = str(value) if value else ''

                        row_data.append(str(value) if value is not None else '')

                    writer.writerow(row_data)

            self.mw._log_status_message(f"✓ Exported {len(rows_to_export)} rows to {Path(file_path).name}", 3000)
            self.mw._show_info("Export Complete", f"Exported {len(rows_to_export)} rows to:\n{file_path}")

        except Exception as e:
            self.mw._show_error("Export Failed", f"Failed to export table:\n{e}")
            print(f"[export] Error: {e}")

    def bulk_set_column(self, rows, column, field_name, dialog_title, preset_value=None):
        """Set a column value for multiple rows at once.

        Args:
            rows: Set of row indices to update
            column: Column index in the table
            field_name: Field name in the task dict
            dialog_title: Title for the input dialog
            preset_value: If provided, use this value without showing dialog
        """
        from PyQt6.QtWidgets import QInputDialog

        if preset_value is not None:
            value = preset_value
            ok = True
        else:
            # Get existing values to suggest
            existing_values = set()
            for row in rows:
                if row < len(self.mw._master_file_list):
                    val = self.mw._master_file_list[row].get(field_name, '')
                    if val:
                        existing_values.add(val)

            # Default to first existing value or empty
            default_value = list(existing_values)[0] if existing_values else ''

            value, ok = QInputDialog.getText(
                self.mw, dialog_title,
                f"Enter value for {len(rows)} rows:",
                text=default_value
            )

        if ok:
            for row in rows:
                if row < len(self.mw._master_file_list):
                    task = self.mw._master_file_list[row]
                    task[field_name] = value
                    # Update the model row to reflect the change
                    if row < self.mw._file_table_model.rowCount():
                        self.mw._file_table_model.update_row(row, task)

            # Update experiment history if this is the experiment field
            if field_name == 'experiment' and value:
                self.update_experiment_history(value)

            print(f"[master-list] Bulk set {field_name} = '{value}' for {len(rows)} rows")

    def resolve_conflicts_use_npz(self, rows):
        """Resolve conflicts by overwriting table values with NPZ values."""
        resolved_count = 0
        for row in rows:
            if row >= len(self.mw._master_file_list):
                continue

            task = self.mw._master_file_list[row]
            warnings = task.get('scan_warnings', {})
            conflicts = warnings.get('conflicts', [])

            if not conflicts:
                continue

            # Parse conflicts and apply NPZ values
            # Conflicts are in format: "Field: table='X' vs NPZ='Y'"
            for conflict in conflicts:
                try:
                    # Extract field name and NPZ value
                    field_part = conflict.split(':')[0].strip()
                    npz_part = conflict.split("NPZ='")[1].split("'")[0]

                    # Map display name to field
                    field_map = {
                        'Strain': 'strain',
                        'Stim Type': 'stim_type',
                        'Power': 'power',
                        'Sex': 'sex',
                        'Animal ID': 'animal_id',
                    }

                    if field_part in field_map:
                        field_name = field_map[field_part]
                        task[field_name] = npz_part
                        print(f"[conflict-resolve] Row {row}: Set {field_name} = '{npz_part}' from NPZ")
                except Exception as e:
                    print(f"[conflict-resolve] Error parsing conflict '{conflict}': {e}")

            # Clear the warnings after resolving
            task.pop('scan_warnings', None)
            resolved_count += 1

            # Update the model row to reflect changes
            self.update_row_status_icon(row, task)

        self.mw._log_status_message(f"✓ Resolved conflicts for {resolved_count} rows (used NPZ values)", 3000)
        self.mw._update_resolve_conflicts_button()

    def resolve_conflicts_keep_table(self, rows):
        """Resolve conflicts by keeping table values and dismissing warnings."""
        resolved_count = 0
        for row in rows:
            if row >= len(self.mw._master_file_list):
                continue

            task = self.mw._master_file_list[row]
            if task.get('scan_warnings'):
                task.pop('scan_warnings', None)
                resolved_count += 1
                self.update_row_status_icon(row, task)

        self.mw._log_status_message(f"✓ Dismissed warnings for {resolved_count} rows (kept table values)", 3000)
        self.mw._update_resolve_conflicts_button()

    def clear_scan_warnings(self, rows):
        """Clear all scan warnings for selected rows."""
        cleared_count = 0
        for row in rows:
            if row >= len(self.mw._master_file_list):
                continue

            task = self.mw._master_file_list[row]
            if task.get('scan_warnings'):
                task.pop('scan_warnings', None)
                cleared_count += 1
                self.update_row_status_icon(row, task)

        self.mw._log_status_message(f"✓ Cleared warnings for {cleared_count} rows", 3000)
        self.mw._update_resolve_conflicts_button()

    def show_conflict_details(self, rows):
        """Show a dialog with detailed conflict information."""
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QTextEdit, QPushButton, QHBoxLayout

        dialog = QDialog(self.mw)
        dialog.setWindowTitle("Scan Warning Details")
        dialog.setMinimumSize(600, 400)

        layout = QVBoxLayout(dialog)

        text_edit = QTextEdit()
        text_edit.setReadOnly(True)

        # Build details text
        details = []
        for row in rows:
            if row >= len(self.mw._master_file_list):
                continue

            task = self.mw._master_file_list[row]
            warnings = task.get('scan_warnings', {})
            if not warnings:
                continue

            file_name = task.get('file_name', 'Unknown')
            channel = task.get('channel', '')
            details.append(f"═══ Row {row}: {file_name} ({channel}) ═══")

            conflicts = warnings.get('conflicts', [])
            if conflicts:
                details.append("⚠ Data Conflicts:")
                for c in conflicts:
                    details.append(f"   • {c}")

            older_files = warnings.get('older_npz_files', [])
            if older_files:
                details.append(f"📁 {len(older_files)} Older NPZ File(s):")
                for older in older_files:
                    details.append(f"   • {Path(older['file']).name}")
                    details.append(f"     Date: {older['date']}")
                    details.append(f"     Path: {older['file']}")

            details.append("")

        text_edit.setText('\n'.join(details) if details else "No warning details found.")
        layout.addWidget(text_edit)

        # Buttons
        btn_layout = QHBoxLayout()

        btn_use_npz = QPushButton("Use NPZ Values")
        btn_use_npz.clicked.connect(lambda: (self.resolve_conflicts_use_npz(rows), dialog.accept()))
        btn_layout.addWidget(btn_use_npz)

        btn_keep_table = QPushButton("Keep Table Values")
        btn_keep_table.clicked.connect(lambda: (self.resolve_conflicts_keep_table(rows), dialog.accept()))
        btn_layout.addWidget(btn_keep_table)

        btn_close = QPushButton("Close")
        btn_close.clicked.connect(dialog.reject)
        btn_layout.addWidget(btn_close)

        layout.addLayout(btn_layout)
        dialog.exec()

    def combine_selected_files(self, rows):
        """Combine selected files for multi-file analysis."""
        from PyQt6.QtWidgets import QMessageBox

        # Get file paths for selected rows
        file_paths = []
        for row in sorted(rows):
            if row < len(self.mw._master_file_list):
                task = self.mw._master_file_list[row]
                file_path = task.get('file_path')
                if file_path and Path(file_path).exists():
                    file_paths.append(Path(file_path))

        if len(file_paths) < 2:
            self.mw._show_warning("Cannot Combine", "Need at least 2 valid files to combine.")
            return

        # Confirm with user
        file_names = [p.name for p in file_paths]
        msg = QMessageBox(self.mw)
        msg.setIcon(QMessageBox.Icon.Question)
        msg.setWindowTitle("Combine Files")
        msg.setText(f"Combine {len(file_paths)} files for analysis?")
        msg.setInformativeText("Files will be concatenated in order:\n• " + "\n• ".join(file_names[:5]))
        if len(file_names) > 5:
            msg.setInformativeText(msg.informativeText() + f"\n... and {len(file_names) - 5} more")
        msg.setStandardButtons(QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.Cancel)

        if msg.exec() == QMessageBox.StandardButton.Yes:
            # Load files using multi-file loader
            print(f"[master-list] Combining {len(file_paths)} files for analysis")
            self.mw.load_multiple_files(file_paths)
            # Switch to Analysis tab (Tab 0 = Project Builder, Tab 1 = Analysis)
            if hasattr(self.mw, 'Tabs'):
                self.mw.Tabs.setCurrentIndex(1)
