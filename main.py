from PyQt6 import uic
from PyQt6.QtWidgets import QApplication, QMainWindow, QFileDialog, QMessageBox, QListWidgetItem, QAbstractItemView, QTreeWidgetItem, QTableView
from PyQt6.QtCore import QSettings, QTimer, Qt
from PyQt6.QtGui import QIcon
from consolidation import ConsolidationManager
from core.file_table_model import FileTableModel, ColumnDef, DEFAULT_COLUMNS
from core.file_table_delegates import ButtonDelegate, StatusDelegate, AutoCompleteDelegate

import re
from PyQt6.QtWidgets import (
    QDialog, QFormLayout, QLineEdit, QComboBox, QLabel,
    QDialogButtonBox, QPushButton, QHBoxLayout, QCheckBox
)

import csv, json



from pathlib import Path
from typing import List
import sys
import os

# Fix KMeans memory leak warning on Windows
os.environ['OMP_NUM_THREADS'] = '1'

import numpy as np
import pandas as pd

from core.state import AppState
from core import abf_io, filters
from core.plotting import PlotHost
from core import stim as stimdet   # stim detection

# Peak detection: Switch between standard and downsampled versions
from core import peaks as peakdet              # Standard (original)
# from core import peaks_downsampled as peakdet    # Downsampled (BROKEN - messes up peaks)

from core import metrics  # calculation of breath metrics
from core.navigation_manager import NavigationManager
from core import telemetry  # Anonymous usage tracking
from plotting import PlotManager
from export import ExportManager


# Import editing modes
from editing import EditingModes
# Import dialogs
from dialogs import SpectralAnalysisDialog, SaveMetaDialog, PeakNavigatorDialog
from dialogs.advanced_peak_editor_dialog import AdvancedPeakEditorDialog
from dialogs.photometry_import_dialog import PhotometryImportDialog
from core import photometry
from core.channel_manager import ChannelManagerWidget

# Import version
from version_info import VERSION_STRING


ORG = "PhysioMetrics"
APP = "PhysioMetrics"

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()

        # Initialize managers
        self.consolidation_manager = ConsolidationManager(self)
        from core.project_manager import ProjectManager
        self.project_manager = ProjectManager()

        ui_file = Path(__file__).parent / "ui" / "pleth_app_layout_05_horizontal.ui"
        uic.loadUi(ui_file, self)

        # icon_path = Path(__file__).parent / "assets" / "plethapp_thumbnail_light_02.ico"
        icon_path = Path(__file__).parent / "assets" / "plethapp_thumbnail_dark_round.ico"
        self.setWindowIcon(QIcon(str(icon_path)))
        # after uic.loadUi(ui_file, self)
        from PyQt6.QtWidgets import QWidget, QPushButton
        for w in self.findChildren(QWidget):
            if w.property("startHidden") is True:
                w.hide()

        self.setWindowTitle(f"PhysioMetrics v{VERSION_STRING}")

        # Enable dark title bar on Windows 11
        self._enable_dark_title_bar()

        # Tab order is now set in the UI file:
        # 0: Notes Files, 1: Project Files, 2: Consolidation, 3: Code Notebook

        # Style status bar to match dark theme
        self.statusBar().setStyleSheet("""
            QStatusBar {
                background-color: #1e1e1e;
                color: #d4d4d4;
                border-top: 1px solid #3e3e42;
            }
            QStatusBar::item {
                border: none;
            }
        """)
        # Disable the resize grip (removes the dots on the right side)
        self.statusBar().setSizeGripEnabled(False)

        # Add message history tracking and dropdown
        self._status_message_history = []
        self._setup_status_history_dropdown()

        # Add filename display to status bar (right side, before history button)
        self._setup_filename_display()

        # Add editing instructions label to status bar (left side, supports rich text)
        self._setup_editing_instructions_label()

        self.settings = QSettings(ORG, APP)
        self.state = AppState()
        self.single_panel_mode = False  # flips True after stim channel selection

        # Notch filter parameters
        self.notch_filter_lower = None
        self.notch_filter_upper = None

        # Filter order
        self.filter_order = 4  # Default Butterworth filter order

        # Auto-GMM refresh (default OFF for better performance)
        self.auto_gmm_enabled = False
        # Track if eupnea/sniffing detection is out of date
        self.eupnea_sniffing_out_of_date = False

        # Z-score normalization
        self.use_zscore_normalization = True  # Default: enabled
        self.zscore_global_mean = None  # Global mean across all sweeps (cached)
        self.zscore_global_std = None   # Global std across all sweeps (cached)

        # GMM clustering cache (for fast dialog loading)
        self._cached_gmm_results = None

        # Outlier detection metrics (default set)
        self.outlier_metrics = ["if", "amp_insp", "amp_exp", "ti", "te", "area_insp", "area_exp"]

        # Cross-sweep outlier detection
        self.global_outlier_stats = None  # Dict[metric_key, (mean, std)] - computed across all sweeps
        self.metrics_by_sweep = {}  # Dict[sweep_idx, Dict[metric_key, metric_array]]
        self.onsets_by_sweep = {}  # Dict[sweep_idx, onsets_array]

        # Eupnea detection parameters
        self.eupnea_freq_threshold = 5.0  # Hz - frequency threshold for eupnea (used in frequency mode)
        self.eupnea_min_duration = 2.0  # seconds - minimum sustained duration for eupnea region
        self.eupnea_detection_mode = "gmm"  # "gmm" or "frequency" - default to GMM-based detection

        # --- Embed Matplotlib into MainPlot (QFrame in Designer) ---
        self.plot_host = PlotHost(self.MainPlot)
        layout = self.MainPlot.layout()
        if layout is None:
            from PyQt6.QtWidgets import QVBoxLayout
            layout = QVBoxLayout(self.MainPlot)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.addWidget(self.plot_host)

        # --- Setup Channel Manager Widget ---
        self._setup_channel_manager()

        saved_geom = self.settings.value("geometry")
        if saved_geom:
            self.restoreGeometry(saved_geom)

        # --- Wire browse ---
        self.BrowseButton.clicked.connect(self.on_browse_clicked)

        # Add Ctrl+O shortcut - triggers different buttons based on active tab
        from PyQt6.QtGui import QShortcut, QKeySequence
        ctrl_o_shortcut = QShortcut(QKeySequence("Ctrl+O"), self)
        ctrl_o_shortcut.activated.connect(self.on_ctrl_o_pressed)

        # Add F1 shortcut for Help
        f1_shortcut = QShortcut(QKeySequence("F1"), self)
        f1_shortcut.activated.connect(self.on_help_clicked)

        # Add Z/X shortcuts for navigation with ApplicationShortcut context
        from PyQt6.QtCore import Qt
        z_shortcut = QShortcut(QKeySequence("Z"), self)
        z_shortcut.setContext(Qt.ShortcutContext.ApplicationShortcut)
        z_shortcut.activated.connect(self._on_z_pressed)

        x_shortcut = QShortcut(QKeySequence("X"), self)
        x_shortcut.setContext(Qt.ShortcutContext.ApplicationShortcut)
        x_shortcut.activated.connect(self._on_x_pressed)

        # --- Wire channel selection (immediate application) ---
        self.AnalyzeChanSelect.currentIndexChanged.connect(self.on_analyze_channel_changed)
        self.StimChanSelect.currentIndexChanged.connect(self.on_stim_channel_changed)
        self.EventsChanSelect.currentIndexChanged.connect(self.on_events_channel_changed)

        # --- Wire classifier selection ---
        # Replace "ML Model" with specific algorithms for all three dropdowns
        self.peak_detec_combo.clear()
        self.peak_detec_combo.addItems(["Threshold", "XGBoost", "Random Forest", "MLP"])

        self.eup_sniff_combo.clear()
        self.eup_sniff_combo.addItems(["None (Clear)", "All Eupnea", "GMM", "XGBoost", "Random Forest", "MLP"])

        self.digh_combo.clear()
        self.digh_combo.addItems(["None (Clear)", "Manual", "XGBoost", "Random Forest", "MLP"])

        # Auto-load ML models from last used directory (if available)
        # This must happen BEFORE setting defaults and connecting signals
        self._auto_load_ml_models_on_startup()

        # Update dropdown states based on loaded models (will disable ML options if models not loaded)
        self._update_classifier_dropdowns()

        # NOW connect signals AFTER models are loaded and dropdowns are configured
        self.peak_detec_combo.currentTextChanged.connect(self.on_classifier_changed)
        self.eup_sniff_combo.currentTextChanged.connect(self.on_eupnea_sniff_classifier_changed)
        self.digh_combo.currentTextChanged.connect(self.on_sigh_classifier_changed)

        # Set defaults - XGBoost for peak detection and sigh, GMM for eupnea/sniff
        # If models not loaded, _update_classifier_dropdowns() already fell back to Threshold/GMM/Manual
        self.peak_detec_combo.setCurrentText("XGBoost")
        self.eup_sniff_combo.setCurrentText("GMM")  # GMM is the default for eupnea/sniff classification
        self.digh_combo.setCurrentText("XGBoost")

        # --- Wire filter controls ---
        self._redraw_timer = QTimer(self)
        self._redraw_timer.setSingleShot(True)
        self._redraw_timer.setInterval(150)       # ms
        self._redraw_timer.timeout.connect(self.redraw_main_plot)

        # filters: commit-on-finish, not per key
        self.LowPassVal.editingFinished.connect(self.update_and_redraw)
        self.HighPassVal.editingFinished.connect(self.update_and_redraw)
        self.FilterOrderSpin.valueChanged.connect(self.update_and_redraw)

        # checkboxes toggled immediately, but we debounce the draw
        self.LowPass_checkBox.toggled.connect(self.update_and_redraw)
        self.HighPass_checkBox.toggled.connect(self.update_and_redraw)
        self.InvertSignal_checkBox.toggled.connect(self.update_and_redraw)

        # Re-enable Apply button when filters change (peaks need to be recalculated)
        self.LowPassVal.editingFinished.connect(self._on_filter_changed)
        self.HighPassVal.editingFinished.connect(self._on_filter_changed)
        self.FilterOrderSpin.valueChanged.connect(self._on_filter_changed)
        self.LowPass_checkBox.toggled.connect(self._on_filter_changed)
        self.HighPass_checkBox.toggled.connect(self._on_filter_changed)
        self.InvertSignal_checkBox.toggled.connect(self._on_filter_changed)

        # Spectral Analysis button
        self.SpectralAnalysisButton.clicked.connect(self.on_spectral_analysis_clicked)

        # Outlier Threshold button - REMOVED (moved to Analysis Options dialog)
        # self.OutlierThreshButton.clicked.connect(self.on_outlier_thresh_clicked)

        # Eupnea Threshold button

        # GMM Clustering button - REMOVED (moved to Analysis Options dialog)
        # self.GMMClusteringButton.clicked.connect(self.on_gmm_clustering_clicked)

        # Peak Navigator/Editor button
        self.editor_pushButton.clicked.connect(self.on_peak_navigator_clicked)

        # Auto-Update GMM checkbox
        # Auto-Update GMM checkbox moved to GMM dialog
        # Connect the manual update button
        self.UpdateEupneaSniffingButton.clicked.connect(self.on_update_eupnea_sniffing_clicked)

        # --- Initialize Navigation Manager ---
        self.navigation_manager = NavigationManager(self)
        self.WindowRangeValue.setText("10")  # Default window length

        # --- Initialize Plot Manager (BEFORE signal connections that may trigger plotting) ---
        self.plot_manager = PlotManager(self)

        # --- Initialize Export Manager ---
        self.export_manager = ExportManager(self)

        # --- Initialize Telemetry Heartbeat Timer ---
        # Send periodic user_engagement events to help GA4 recognize active users
        self.telemetry_heartbeat_timer = QTimer(self)
        self.telemetry_heartbeat_timer.timeout.connect(telemetry.log_user_engagement)
        self.telemetry_heartbeat_timer.start(45000)  # Send engagement event every 45 seconds

        # --- Peak-detect UI wiring ---
        # Prominence field and Apply button already exist in UI file
        # Connect the "More Options" button to open multi-tabbed analysis options dialog
        self.ThreshOptions.clicked.connect(self._open_analysis_options)

        # Detect button applies peaks with current parameters
        self.ApplyPeakFindPushButton.setEnabled(False)  # stays disabled until prominence detected
        self.ApplyPeakFindPushButton.clicked.connect(self._apply_peak_detection)

        # PeakPromValueSpinBox removed - prominence/threshold now set in Analysis Options dialog
        # Re-enable Apply button when user manually edits prominence (use spinbox)
        # self.PeakPromValueSpinBox.valueChanged.connect(lambda: self.ApplyPeakFindPushButton.setEnabled(True) if self.state.analyze_chan else None)

        # Connect spinbox to update threshold line on main plot
        # self.PeakPromValueSpinBox.valueChanged.connect(self._on_prominence_spinbox_changed)

        # Configure spinbox for fine-grained control (0.01 increments for second decimal place)
        # self.PeakPromValueSpinBox.setSingleStep(0.01)
        # self.PeakPromValueSpinBox.setDecimals(4)  # Show 4 decimal places
        # self.PeakPromValueSpinBox.setMinimum(0.0001)  # Minimum prominence value
        # self.PeakPromValueSpinBox.setMaximum(1000.0)  # Maximum prominence value

        # Store peak detection parameters (auto-populated when channel selected)
        self.peak_prominence = None
        self.peak_height_threshold = None  # Same as prominence by default
        self.peak_min_dist = 0.05  # Default minimum peak distance in seconds

        # Default values for eupnea and apnea thresholds
        self.ApneaThresh.setText("0.5")   # seconds - gaps longer than this are apnea

        # OutlierSD moved to Analysis Options dialog (Outlier Detection tab)
        # Create a simple object to store the value (mimics QLineEdit.text() interface)
        class OutlierSDHolder:
            def __init__(self, value="3.0"):
                self._value = value
            def text(self):
                return self._value
            def setText(self, value):
                self._value = str(value)

        self.OutlierSD = OutlierSDHolder("3.0")  # Default: 3.0 SD for outlier detection

        # Connect signals for apnea threshold changes to trigger redraw
        self.ApneaThresh.textChanged.connect(self._on_region_threshold_changed)
        # OutlierSD changes are handled in the Analysis Options dialog

        # --- y2 metric dropdown (choices only; plotting later) ---
        self.y2plot_dropdown.clear()
        self.state.y2_values_by_sweep.clear()
        self.plot_host.clear_y2()

        self.y2plot_dropdown.addItem("None", userData=None)
        for label, key in metrics.METRIC_SPECS:
            self.y2plot_dropdown.addItem(label, userData=key)

        # Make dropdown searchable by typing
        self.y2plot_dropdown.setEditable(True)
        self.y2plot_dropdown.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)

        # Add completer for case-insensitive searching
        from PyQt6.QtCore import Qt
        from PyQt6.QtWidgets import QCompleter
        completer = QCompleter([self.y2plot_dropdown.itemText(i) for i in range(self.y2plot_dropdown.count())])
        completer.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
        completer.setFilterMode(Qt.MatchFlag.MatchContains)  # Match anywhere in the string
        self.y2plot_dropdown.setCompleter(completer)

        # ADD/DELETE Peak Mode: track selection in state
        self.state.y2_metric_key = None
        self.y2plot_dropdown.currentIndexChanged.connect(self.on_y2_metric_changed)

        # --- Display Control Widgets ---
        # Y-axis autoscale checkbox (percentile vs min/max)
        if hasattr(self, 'yautoscale_checkBox'):
            self.yautoscale_checkBox.setChecked(self.state.use_percentile_autoscale)
            self.yautoscale_checkBox.toggled.connect(self.on_yautoscale_toggled)

        # Y-axis padding spinbox (for percentile mode)
        if hasattr(self, 'ypadding_SpinBox'):
            self.ypadding_SpinBox.setMinimum(0.0)
            self.ypadding_SpinBox.setMaximum(1.0)
            self.ypadding_SpinBox.setSingleStep(0.05)
            self.ypadding_SpinBox.setDecimals(2)
            self.ypadding_SpinBox.setValue(self.state.autoscale_padding)
            self.ypadding_SpinBox.valueChanged.connect(self.on_ypadding_changed)

        # Region display toggle checkboxes (line vs shade)
        if hasattr(self, 'eupnea_checkBox'):
            self.eupnea_checkBox.setChecked(self.state.eupnea_use_shade)
            self.eupnea_checkBox.toggled.connect(self.on_eupnea_display_toggled)

        if hasattr(self, 'sniffing_checkBox'):
            self.sniffing_checkBox.setChecked(self.state.sniffing_use_shade)
            self.sniffing_checkBox.toggled.connect(self.on_sniffing_display_toggled)

        if hasattr(self, 'Apnea_checkBox'):
            self.Apnea_checkBox.setChecked(self.state.apnea_use_shade)
            self.Apnea_checkBox.toggled.connect(self.on_apnea_display_toggled)

        if hasattr(self, 'Outliers_checkBox'):
            self.Outliers_checkBox.setChecked(self.state.outliers_use_shade)
            self.Outliers_checkBox.toggled.connect(self.on_outliers_display_toggled)

        # Dark mode toggle (for plot background)
        if hasattr(self, 'checkBox'):  # Generic name from UI - should be renamed to DarkModeCheckBox
            # Load saved dark mode preference (default to True = dark mode)
            dark_mode_enabled = self.settings.value("plot_dark_mode", True, type=bool)

            # Apply the theme to plot_host
            theme = "dark" if dark_mode_enabled else "light"
            self.plot_host.set_plot_theme(theme)

            # Set checkbox state to match loaded preference
            self.checkBox.setChecked(dark_mode_enabled)
            self.checkBox.toggled.connect(self.on_dark_mode_toggled)

        # Initialize editing modes manager
        self.editing_modes = EditingModes(self)

        # --- Mark Events button (event detection settings) ---
        self.MarkEventsButton.clicked.connect(self.on_mark_events_clicked)

        # --- Sigh overlay artists ---
        self._sigh_artists = []         # matplotlib artists for sigh overlay

        # --- Wire omit button ---
        self.OmitSweepButton.setCheckable(True)
        self.OmitSweepButton.clicked.connect(self.on_omit_sweep_clicked)

        # --- Move Point mode ---
        self._is_dragging = False  # Track if currently dragging a point

        self.movePointButton.setCheckable(True)

        # Mark Sniff button

        #wire save analyzed data button
        self.SaveAnalyzedDataButton.clicked.connect(self.on_save_analyzed_clicked)

        # Wire save options button to configure export metrics
        self.saveoptions_pushButton.clicked.connect(self.on_save_options_clicked)

        # Wire view summary button to show PDF preview
        self.ViewSummary_pushButton.clicked.connect(self.on_view_summary_clicked)

        # Wire Help button (from UI file)
        self.helpbutton.clicked.connect(self.on_help_clicked)

        # Set pointer cursor for update notification label (defined in UI file)
        from PyQt6.QtGui import QCursor
        from PyQt6.QtCore import Qt
        self.update_notification_label.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))

        # Store update info for later use
        self.update_info = None

        # Start background update check
        self._check_for_updates_on_startup()



        # Defaults: 0.5–20 Hz band, all off initially
        self.HighPassVal.setText("0.5")
        self.LowPassVal.setText("20")
        self.HighPass_checkBox.setChecked(True)  # Default ON to remove baseline drift
        self.LowPass_checkBox.setChecked(True)
        self.InvertSignal_checkBox.setChecked(False)

        # Push defaults into state (no-op if no data yet)
        self.update_and_redraw()
        self._refresh_omit_button_label()

        # Connect matplotlib toolbar to turn off edit modes
        self.plot_host.set_toolbar_callback(self.editing_modes.turn_off_all_edit_modes)


        # --- Curation tab wiring ---
        self.FilePathButton.clicked.connect(self.on_curation_choose_dir_clicked)
        # Enable multiple selection for both list widgets
        self.FileList.setSelectionMode(QAbstractItemView.SelectionMode.ExtendedSelection)
        self.FilestoConsolidateList.setSelectionMode(QAbstractItemView.SelectionMode.ExtendedSelection)
        # Note: Move buttons and search filter are connected in NavigationManager
        self.FileListSearchBox.setPlaceholderText("Filter by keywords (e.g., 'gfp 2.5mW' or 'gfp, chr2')...")
        # Wire consolidate button
        self.ConsolidateSaveDataButton.clicked.connect(self.on_consolidate_save_data_clicked)

        # Wire new consolidation tab in Project Builder
        self.consolidationSourceList.setSelectionMode(QAbstractItemView.SelectionMode.ExtendedSelection)
        self.consolidationFilesList.setSelectionMode(QAbstractItemView.SelectionMode.ExtendedSelection)
        self.consolidationSaveButton.clicked.connect(self._on_consolidation_save_clicked)
        # Connect move buttons
        self.consolidationMoveAllRight.clicked.connect(self._consolidation_move_all_right)
        self.consolidationMoveSingleRight.clicked.connect(self._consolidation_move_selected_right)
        self.consolidationMoveSingleLeft.clicked.connect(self._consolidation_move_selected_left)
        self.consolidationMoveAllLeft.clicked.connect(self._consolidation_move_all_left)
        # Connect search filter
        self.consolidationSearchBox.textChanged.connect(self._filter_consolidation_source_list)
        # Connect browse button
        self.consolidationBrowseButton.clicked.connect(self._on_consolidation_browse_clicked)
        # Connect reset button (returns to project files)
        self.consolidationResetButton.clicked.connect(self._on_consolidation_reset_clicked)
        self.consolidationResetButton.hide()  # Hidden initially, shown when browsing custom folder
        # Track consolidation source mode (None = project files, path = custom folder)
        self._consolidation_custom_folder = None
        # Auto-refresh consolidation source when switching to the tab
        self.leftColumnTabs.currentChanged.connect(self._on_left_column_tab_changed)

        # ========================================
        # TESTING MODE: Auto-load file and set channels
        # ========================================
        # To enable: set environment variable PLETHAPP_TESTING=1
        # Windows CMD (two commands):
        #   set PLETHAPP_TESTING=1
        #   python run_debug.py
        # Windows PowerShell:
        #   $env:PLETHAPP_TESTING="1"; python run_debug.py
        # Linux/Mac:
        #   PLETHAPP_TESTING=1 python run_debug.py
        # Silently check environment variables (only print if set)
        # print(f"[DEBUG] PLETHAPP_TESTING environment variable = '{os.environ.get('PLETHAPP_TESTING')}'")
        # print(f"[DEBUG] PLETHAPP_PULSE_TEST environment variable = '{os.environ.get('PLETHAPP_PULSE_TEST')}'")
        if os.environ.get('PLETHAPP_TESTING') == '1':
            print("[TESTING MODE] Auto-loading test file...")
            # Check if pulse test mode
            if os.environ.get('PLETHAPP_PULSE_TEST') == '1':
                test_file = Path(r"C:\Users\rphil2\Dropbox\python scripts\breath_analysis\pyqt6\examples\R2 R5 R1\25122001.abf")
                print("[TESTING MODE] Pulse test mode - loading 25122001.abf (25ms pulse experiment)")
            else:
                test_file = Path(r"C:\Users\rphil2\Dropbox\python scripts\breath_analysis\pyqt6\examples\25121004.abf")
                print("[TESTING MODE] Standard test mode - loading 25121004.abf (30Hz stim)")
            print(f"[TESTING MODE] Checking if file exists: {test_file}")
            print(f"[TESTING MODE] File exists: {test_file.exists()}")
            if test_file.exists():
                print(f"[TESTING MODE] Scheduling auto-load in 100ms...")
                QTimer.singleShot(100, lambda: self._auto_load_test_file(test_file))
            else:
                print(f"[TESTING MODE] Warning: Test file not found: {test_file}")

        # === Project Builder Connections ===
        self.browseDirectoryButton.clicked.connect(self.on_project_browse_directory)
        self.scanFilesButton.clicked.connect(self.on_project_scan_files)
        # NOTE: experiment-based workflow buttons were removed from .ui file
        self.clearFilesButton.clicked.connect(self.on_project_clear_files)
        self.newProjectButton.clicked.connect(self.on_project_new)
        self.saveProjectButton.clicked.connect(self.on_project_save)

        # Set up the new unified project name combo (replaces loadProjectCombo + projectNameEdit)
        self._setup_project_name_combo()

        # Notes Files tab connections
        if hasattr(self, 'searchNotesButton'):
            self.searchNotesButton.clicked.connect(self.on_notes_search)
        if hasattr(self, 'browseNotesButton'):
            self.browseNotesButton.clicked.connect(self.on_notes_browse)
        if hasattr(self, 'openNoteButton'):
            self.openNoteButton.clicked.connect(self.on_notes_open)
        if hasattr(self, 'previewNoteButton'):
            self.previewNoteButton.clicked.connect(self.on_notes_preview)
        if hasattr(self, 'linkNoteButton'):
            self.linkNoteButton.clicked.connect(self.on_notes_link)
        if hasattr(self, 'notesFilterEdit'):
            self.notesFilterEdit.textChanged.connect(self._on_notes_filter_changed)

        # Initialize notes files model
        self._init_notes_files_model()

        # Add extra notes action buttons (programmatically added)
        self._add_notes_action_buttons()

        # Connect Project Builder buttons (defined in .ui file)
        self._connect_project_builder_buttons()

        # Project Builder state
        self._project_directory = None
        self._notes_directory = None  # Separate folder for notes files
        self._discovered_files_data = []  # Store file metadata from scan
        self._notes_files_data = []  # Store notes file metadata

        # === Master File List Setup ===
        # Hide the Project Organization section (right column) - we're using a flat master list instead
        self._setup_master_file_list()

        # Store file list data (each row = one analysis task: file + channel + animal)
        self._master_file_list = []  # List of task dicts
        self._active_master_list_row = None  # Track which row is being analyzed

        # optional: keep a handle to the chosen dir
        self._curation_dir = None

    def _auto_load_test_file(self, file_path: Path):
        """Helper function for testing mode - auto-loads file and sets channels."""
        print(f"[TESTING MODE] _auto_load_test_file() called with: {file_path}")
        print(f"[TESTING MODE] Calling load_file()...")
        self.load_file(file_path)
        print(f"[TESTING MODE] load_file() returned, starting polling timer...")

        # Poll until file is loaded, then set channels
        self._check_file_loaded_timer = QTimer()
        self._check_file_loaded_timer.timeout.connect(self._check_if_file_loaded)
        self._check_file_loaded_timer.start(100)  # Check every 100ms
        print(f"[TESTING MODE] Polling timer started")

    def _check_if_file_loaded(self):
        """Poll to see if file has finished loading."""
        # Check if state has been populated with data AND combos have been populated
        print(f"[TESTING MODE] Polling: t={self.state.t is not None}, sweeps={self.state.sweeps is not None}, "
              f"AnalyzeChanSelect.count={self.AnalyzeChanSelect.count()}, StimChanSelect.count={self.StimChanSelect.count()}")
        if (self.state.t is not None and
            self.state.sweeps is not None and
            len(self.state.sweeps) > 0 and
            self.AnalyzeChanSelect.count() > 0 and
            self.StimChanSelect.count() > 0):
            # File is loaded and combos are populated!
            print(f"[TESTING MODE] File loaded! Stopping timer and setting test channels...")
            self._check_file_loaded_timer.stop()
            self._set_test_channels()

    def _set_test_channels(self):
        """Helper function for testing mode - sets analysis and stim channels."""
        is_pulse_test = os.environ.get('PLETHAPP_PULSE_TEST') == '1'

        if is_pulse_test:
            print("[TESTING MODE] Setting analyze channel to 0 (first), stim channel to last...")
        else:
            print("[TESTING MODE] Setting analyze channel to 0, stim channel to 7...")

        print(f"[TESTING MODE] AnalyzeChanSelect has {self.AnalyzeChanSelect.count()} items")
        print(f"[TESTING MODE] StimChanSelect has {self.StimChanSelect.count()} items")

        # Set analyze channel to index 1 (channel 0, since index 0 is "All Channels")
        if self.AnalyzeChanSelect.count() > 1:
            self.AnalyzeChanSelect.setCurrentIndex(1)  # Channel 0 is at index 1
            print(f"[TESTING MODE] Set analyze channel to index 1: {self.AnalyzeChanSelect.currentText()}")
            # Manually trigger the channel change handler
            self.on_analyze_channel_changed(1)

        # Set stim channel (last channel for pulse test, channel 7 for standard test)
        stim_channel_set = False
        if is_pulse_test:
            # Select the last channel in the dropdown
            last_idx = self.StimChanSelect.count() - 1
            if last_idx >= 0:
                self.StimChanSelect.setCurrentIndex(last_idx)
                print(f"[TESTING MODE] Set stim channel to last (index {last_idx}): {self.StimChanSelect.currentText()}")
                stim_channel_set = True
                # Manually trigger the channel change handler
                self.on_stim_channel_changed(last_idx)
        else:
            # Set stim channel to 7 (need to find the index)
            for i in range(self.StimChanSelect.count()):
                item_text = self.StimChanSelect.itemText(i)
                # Check if this item contains "7" (handles "7", "IN 7", "Channel 7", etc.)
                # Extract the number from the item text
                if "7" in item_text.split():  # Split by whitespace and check if "7" is one of the words
                    print(f"[TESTING MODE] Found stim channel at index {i}: '{item_text}'")
                    self.StimChanSelect.setCurrentIndex(i)
                    stim_channel_set = True
                    # Manually trigger the channel change handler
                    self.on_stim_channel_changed(i)
                    break

        if not stim_channel_set:
            if is_pulse_test:
                print("[TESTING MODE] Warning: Could not find last stim channel")
            else:
                print("[TESTING MODE] Warning: Could not find stim channel '7'")

        # Wait a bit for channels to be processed, then click Apply Peak Detection
        QTimer.singleShot(200, self._click_apply_peak_detection)

    def _click_apply_peak_detection(self):
        """Helper for testing mode - clicks the Apply button for peak detection."""
        if self.ApplyPeakFindPushButton.isEnabled():
            print("[TESTING MODE] Clicking Apply Peak Detection button...")
            self.ApplyPeakFindPushButton.click()
            print("[TESTING MODE] Auto-load complete!")
        else:
            print("[TESTING MODE] Warning: Apply Peak Detection button is not enabled")
            print("[TESTING MODE] Auto-load complete (but peaks not detected)")

        # Log main screen view for telemetry
        telemetry.log_screen_view('Main Analysis Screen', screen_class='main')

    # ---------- File browse ----------
    def closeEvent(self, event):
        """Save window geometry on close."""
        self.settings.setValue("geometry", self.saveGeometry())

        # Log telemetry session end
        telemetry.log_session_end()

        super().closeEvent(event)

    def _setup_status_history_dropdown(self):
        """Add a subtle dropdown button to the status bar for viewing message history."""
        from PyQt6.QtWidgets import QPushButton, QMenu
        from PyQt6.QtGui import QIcon
        from PyQt6.QtCore import QSize

        # Create a small button with just a "▼" symbol
        self.history_button = QPushButton("📋", self)
        self.history_button.setFixedSize(24, 20)
        self.history_button.setToolTip("View message history")
        self.history_button.setStyleSheet("""
            QPushButton {
                background-color: transparent;
                border: none;
                color: #d4d4d4;
                font-size: 14px;
                padding: 0px;
            }
            QPushButton:hover {
                background-color: #3e3e42;
                border-radius: 3px;
            }
        """)
        self.history_button.clicked.connect(self._show_message_history)

        # Add to status bar (right side)
        self.statusBar().addPermanentWidget(self.history_button)

    def _setup_filename_display(self):
        """Add a label to the status bar to display the current filename."""
        from PyQt6.QtWidgets import QLabel
        from PyQt6.QtCore import Qt

        # Create label for filename display
        self.filename_label = QLabel("No file loaded", self)
        self.filename_label.setStyleSheet("""
            QLabel {
                color: #d4d4d4;
                padding: 2px 8px;
                margin-right: 10px;
            }
        """)
        self.filename_label.setAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)

        # Add to status bar (right side, before history button)
        self.statusBar().addPermanentWidget(self.filename_label)

    def _setup_editing_instructions_label(self):
        """Add a label to the status bar for rich text editing instructions."""
        from PyQt6.QtWidgets import QLabel
        from PyQt6.QtCore import Qt

        # Create label for editing mode instructions (supports HTML)
        self.editing_instructions_label = QLabel("", self)
        self.editing_instructions_label.setStyleSheet("""
            QLabel {
                color: #d4d4d4;
                padding: 2px 8px;
                font-size: 11px;
            }
        """)
        self.editing_instructions_label.setTextFormat(Qt.TextFormat.RichText)
        self.editing_instructions_label.setAlignment(Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter)

        # Add to status bar (left side - use addWidget, not addPermanentWidget)
        self.statusBar().addWidget(self.editing_instructions_label, 1)  # stretch=1 to take available space

    def _setup_channel_manager(self):
        """Setup the Channel Manager using UI elements from the .ui file."""
        # Create the ChannelManagerWidget, passing it the UI elements to use
        self.channel_manager = ChannelManagerWidget(
            parent=self,
            summary_label=self.channelManagerSummaryLabel,
            expand_btn=self.channelManagerExpandBtn,
            preview_label=self.channelManagerPreviewLabel
        )

        # Connect signals
        self.channel_manager.apply_requested.connect(self._on_channel_manager_apply)
        self.channel_manager.settings_requested.connect(self._on_channel_settings_requested)

    def _on_channel_manager_apply(self):
        """Called when channel manager requests to apply changes."""
        st = self.state

        # Get the Pleth channel from channel manager
        pleth_channel = self.channel_manager.get_pleth_channel()

        # Check if Pleth channel changed - if so, clear analysis data
        channel_changed = (pleth_channel != st.analyze_chan)

        if pleth_channel:
            if channel_changed:

                # Clear peaks and breath data (same as on_analyze_channel_changed)
                st.peaks_by_sweep.clear()
                st.sigh_by_sweep.clear()
                if hasattr(st, 'breath_by_sweep'):
                    st.breath_by_sweep.clear()

                # Clear omitted sweeps and regions
                st.omitted_sweeps.clear()
                st.omitted_ranges.clear()
                self._refresh_omit_button_label()

                # Clear sniffing regions
                if hasattr(st, 'sniff_regions_by_sweep'):
                    st.sniff_regions_by_sweep.clear()

                # Clear event markers (bout annotations)
                if hasattr(st, 'bout_annotations'):
                    st.bout_annotations.clear()

                # Clear z-score global statistics cache
                self.zscore_global_mean = None
                self.zscore_global_std = None

            st.analyze_chan = pleth_channel

            # Clear processing cache so filters are re-applied
            st.proc_cache.clear()

            # Also update the old dropdown to stay in sync (until we remove it)
            idx = self.AnalyzeChanSelect.findText(pleth_channel)
            if idx >= 0:
                self.AnalyzeChanSelect.blockSignals(True)
                self.AnalyzeChanSelect.setCurrentIndex(idx)
                self.AnalyzeChanSelect.blockSignals(False)

            # Switch to single panel mode (required for proper navigation)
            if not self.single_panel_mode:
                self.single_panel_mode = True

            # Clear saved view to force fresh autoscale
            self.plot_host.clear_saved_view("single")
            self.plot_host.clear_saved_view("grid")

            # Auto-detect prominence threshold (required for peak detection)
            self._auto_detect_prominence_silent()

            # Enable peak detection button (always enable when channel changes or is set)
            self.ApplyPeakFindPushButton.setEnabled(True)
        else:
            # No Pleth channel - still clear saved view to show full data range
            self.plot_host.clear_saved_view("single")
            self.plot_host.clear_saved_view("grid")

        # Redraw the plot with new channel configuration
        self.redraw_main_plot()

    def _on_channel_settings_requested(self, channel_name: str):
        """Called when user clicks settings (gear) icon for a channel."""
        # TODO: Open settings dialog for computed channels (e.g., ΔF/F parameters)
        pass

    def _populate_channel_manager(self, ch_names: list):
        """Populate the channel manager widget with available channels."""
        from core.channel_manager import ChannelConfig

        # Get the currently selected analyze channel (from old dropdown)
        analyze_chan = self.state.analyze_chan

        configs = []
        for i, name in enumerate(ch_names):
            # Auto-detect channel type from name
            name_lower = name.lower()
            if any(kw in name_lower for kw in ['pleth', 'resp', 'breath', 'flow', 'airway']):
                channel_type = "Pleth"
            elif any(kw in name_lower for kw in ['opto', 'stim', 'laser', 'led', 'light']):
                channel_type = "Opto Stim"
            else:
                channel_type = "Raw Signal"

            configs.append(ChannelConfig(
                name=name,
                visible=True,  # All channels visible by default as preview
                channel_type=channel_type,
                source="file",
                order=i
            ))

        self.channel_manager.set_channels(configs)

    def _show_editing_instructions(self, html: str):
        """Show rich text editing instructions in the status bar."""
        self.editing_instructions_label.setText(html)
        self._editing_instructions_html = html  # Track for restoration (separate from plain text messages)

    def _clear_editing_instructions(self):
        """Clear the editing instructions from status bar."""
        self.editing_instructions_label.setText("")
        self._editing_instructions_html = None

    def _restore_editing_instructions(self):
        """Restore the editing instructions if any are active."""
        if getattr(self, '_editing_instructions_html', None):
            self.editing_instructions_label.setText(self._editing_instructions_html)

    def _update_filename_display(self):
        """Update the filename display in status bar with file metadata."""
        st = self.state
        if not hasattr(st, 'file_info') or not st.file_info:
            self.filename_label.setText("No file loaded")
            return

        file_info = st.file_info[0]
        filename = file_info['path'].name

        # Build metadata string for ABF files
        metadata_parts = []
        if file_info.get('protocol'):
            metadata_parts.append(file_info['protocol'])
        if file_info.get('n_channels'):
            metadata_parts.append(f"{file_info['n_channels']} ch")

        if len(st.file_info) == 1:
            if metadata_parts:
                self.filename_label.setText(f"File: {filename}  [{', '.join(metadata_parts)}]")
            else:
                self.filename_label.setText(f"File: {filename}")
        else:
            if metadata_parts:
                self.filename_label.setText(f"Files: {len(st.file_info)} files ({filename}, ...)  [{', '.join(metadata_parts)}]")
            else:
                self.filename_label.setText(f"Files: {len(st.file_info)} files ({filename}, ...)")

    def _enable_dark_title_bar(self, widget=None):
        """Enable dark title bar on Windows 10/11.

        Args:
            widget: QWidget to apply dark title bar to. If None, applies to self.
        """
        if sys.platform == "win32":
            try:
                from ctypes import windll, byref, sizeof, c_int

                # DWMWA_USE_IMMERSIVE_DARK_MODE
                DWMWA_USE_IMMERSIVE_DARK_MODE = 20

                # Get window handle
                target = widget if widget else self
                hwnd = int(target.winId())

                # Set dark mode (1 = dark, 0 = light)
                value = c_int(1)
                windll.dwmapi.DwmSetWindowAttribute(
                    hwnd,
                    DWMWA_USE_IMMERSIVE_DARK_MODE,
                    byref(value),
                    sizeof(value)
                )
            except Exception as e:
                # Silently fail if not supported (Windows 10 older builds)
                pass

    def _show_message_history(self):
        """Show a menu with recent status bar messages."""
        from PyQt6.QtWidgets import QMenu
        from PyQt6.QtGui import QAction

        menu = QMenu(self)
        menu.setStyleSheet("""
            QMenu {
                background-color: #2d2d2d;
                color: #d4d4d4;
                border: 1px solid #3e3e42;
            }
            QMenu::item:selected {
                background-color: #3e3e42;
            }
        """)

        if not self._status_message_history:
            action = QAction("No messages yet", self)
            action.setEnabled(False)
            menu.addAction(action)
        else:
            # Show last 20 messages, most recent first
            for i, (timestamp, message) in enumerate(reversed(self._status_message_history[-20:])):
                action = QAction(f"{timestamp} - {message}", self)
                action.setEnabled(False)  # Not clickable, just for display
                menu.addAction(action)

        # Show menu below the button
        menu.exec(self.history_button.mapToGlobal(self.history_button.rect().bottomLeft()))

    def _show_error(self, title: str, message: str):
        """Show an error dialog with selectable text for easy copying."""
        msg = QMessageBox(QMessageBox.Icon.Critical, title, message,
                         QMessageBox.StandardButton.Ok, self)
        msg.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse |
                                   Qt.TextInteractionFlag.TextSelectableByKeyboard)
        msg.exec()

    def _show_warning(self, title: str, message: str):
        """Show a warning dialog with selectable text for easy copying."""
        msg = QMessageBox(QMessageBox.Icon.Warning, title, message,
                         QMessageBox.StandardButton.Ok, self)
        msg.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse |
                                   Qt.TextInteractionFlag.TextSelectableByKeyboard)
        msg.exec()

    def _show_info(self, title: str, message: str):
        """Show an info dialog with selectable text for easy copying."""
        msg = QMessageBox(QMessageBox.Icon.Information, title, message,
                         QMessageBox.StandardButton.Ok, self)
        msg.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse |
                                   Qt.TextInteractionFlag.TextSelectableByKeyboard)
        msg.exec()

    def _log_status_message(self, message: str, timeout: int = 0):
        """Log a status message and show it on the status bar."""
        import datetime
        timestamp = datetime.datetime.now().strftime("%H:%M:%S")
        self._status_message_history.append((timestamp, message))

        # Keep only last 100 messages to avoid memory growth
        if len(self._status_message_history) > 100:
            self._status_message_history = self._status_message_history[-100:]

        # Show on status bar
        self.statusBar().showMessage(message, timeout)

    def keyPressEvent(self, event):
        """Handle keyboard events - delegate to editing modes first."""
        from PyQt6.QtCore import Qt

        # Try editing modes handler first
        if self.editing_modes.handle_key_press_event(event):
            event.accept()
            return

        # Fall back to default handling
        super().keyPressEvent(event)

    def on_ctrl_o_pressed(self):
        """Handle Ctrl+O shortcut - triggers different buttons based on active tab."""
        current_tab = self.Tabs.currentIndex()
        if current_tab == 0:  # Analysis tab
            self.on_browse_clicked()
        elif current_tab == 1:  # Curation tab
            self.on_curation_choose_dir_clicked()

    def _is_focus_in_text_widget(self) -> bool:
        """Check if focus is in a text input widget (to avoid shortcuts while typing)."""
        from PyQt6.QtWidgets import QLineEdit, QTextEdit, QPlainTextEdit, QSpinBox, QDoubleSpinBox, QComboBox
        focus_widget = QApplication.focusWidget()
        if focus_widget is None:
            return False
        # Check if it's a text input type
        if isinstance(focus_widget, (QLineEdit, QTextEdit, QPlainTextEdit, QSpinBox, QDoubleSpinBox)):
            return True
        # Check if it's an editable combobox
        if isinstance(focus_widget, QComboBox) and focus_widget.isEditable():
            return True
        return False

    def _on_z_pressed(self):
        """Handle Z key - navigate to previous sweep/window."""
        if self._is_focus_in_text_widget():
            return  # Don't navigate while typing
        # Analysis tab is at index 1 (Project Builder=0, Analysis=1, Curation=2)
        if self.Tabs.currentIndex() == 1:
            self.navigation_manager.on_unified_prev()

    def _on_x_pressed(self):
        """Handle X key - navigate to next sweep/window."""
        if self._is_focus_in_text_widget():
            return  # Don't navigate while typing
        # Analysis tab is at index 1 (Project Builder=0, Analysis=1, Curation=2)
        if self.Tabs.currentIndex() == 1:
            self.navigation_manager.on_unified_next()

    def on_browse_clicked(self):
        last_dir = self.settings.value("last_dir", str(Path.home()))
        # Skip exists() check - it's slow on network drives and the dialog handles it gracefully
        if last_dir and not (last_dir.startswith('\\\\') or (len(last_dir) > 1 and last_dir[1] == ':')):
            # Only check exists() for non-network, non-drive paths
            if not Path(str(last_dir)).exists():
                last_dir = str(Path.home())

        # Show wait cursor while file dialog loads (can be slow on network drives)
        self.setCursor(Qt.CursorShape.WaitCursor)
        self.statusBar().showMessage("Opening file browser...")
        QApplication.processEvents()

        try:
            paths, _ = QFileDialog.getOpenFileNames(
                self, "Select File(s)", last_dir,
                "All Supported (*.abf *.smrx *.edf *.pleth.npz *.csv);;Data Files (*.abf *.smrx *.edf);;PhysioMetrics Sessions (*.pleth.npz);;Photometry/CSV (*.csv);;ABF Files (*.abf);;SMRX Files (*.smrx);;EDF Files (*.edf);;All Files (*.*)"
            )
        finally:
            self.setCursor(Qt.CursorShape.ArrowCursor)
            self.statusBar().clearMessage()

        if not paths:
            return

        # Convert to Path objects
        file_paths = [Path(p) for p in paths]

        # Store the directory of the first file
        self.settings.setValue("last_dir", str(file_paths[0].parent))

        # Update filename display in status bar
        if len(file_paths) == 1:
            self.filename_label.setText(f"File: {file_paths[0].name}")
        else:
            self.filename_label.setText(f"Files: {len(file_paths)} files ({file_paths[0].name}, ...)")

        # Check if any files are .pleth.npz (session files)
        npz_files = [f for f in file_paths if f.suffix == '.npz' or f.name.endswith('.pleth.npz')]

        if npz_files:
            # Can only load one NPZ session at a time
            if len(file_paths) > 1:
                self._show_warning("Cannot Mix File Types",
                    "Cannot load session files (.pleth.npz) together with data files.\n\n"
                    "Please select either:\n"
                    "• One or more data files (.abf, .smrx, .edf) for concatenation, OR\n"
                    "• One session file (.pleth.npz) to restore analysis"
                )
                return

            # Load the NPZ session
            self.load_npz_state(file_paths[0])

        # Check if this is a photometry file (CSV with FP_data pattern)
        elif len(file_paths) == 1 and photometry.detect_photometry_file(file_paths[0]):
            # Show photometry import dialog
            self._show_photometry_import_dialog(file_paths[0])

        else:
            # Load data files (ABF, SMRX, EDF)
            if len(file_paths) == 1:
                self.load_file(file_paths[0])
            else:
                self.load_multiple_files(file_paths)

    def _show_photometry_import_dialog(self, initial_path: Path):
        """
        Show the photometry import wizard dialog.

        Args:
            initial_path: Path to the detected photometry file (FP_data*.csv)
        """
        dialog = PhotometryImportDialog(self, initial_path)
        result = dialog.exec()

        if result == QDialog.DialogCode.Accepted:
            # Get the selected files
            files = dialog.get_selected_files()
            print(f"[Photometry] Files selected: {files}")

            # TODO: In later phases, this will:
            # 1. Process the photometry data
            # 2. Load processed signals into the app
            # For now, just log the selection

            self._show_info(
                "Photometry Import",
                f"Selected files:\n\n"
                f"• FP Data: {files.get('fp_data', 'None')}\n"
                f"• AI Data: {files.get('ai_data', 'None')}\n\n"
                "Full processing will be available in the next update."
            )

    def load_file(self, path: Path):
        import time
        from PyQt6.QtWidgets import QProgressDialog, QApplication
        from PyQt6.QtCore import Qt

        t_start = time.time()

        # Determine file type for progress dialog title
        file_type = path.suffix.upper()[1:]  # .abf -> ABF, .smrx -> SMRX

        # Create progress dialog
        progress = QProgressDialog(f"Loading file...\n{path.name}", None, 0, 100, self)
        progress.setWindowTitle(f"Opening {file_type} File")
        progress.setWindowModality(Qt.WindowModality.WindowModal)
        progress.setMinimumDuration(0)  # Show immediately
        progress.setCancelButton(None)  # No cancel button
        progress.setValue(0)
        progress.show()
        QApplication.processEvents()

        def update_progress(current, total, message):
            """Callback to update progress dialog."""
            progress.setValue(current)
            progress.setLabelText(f"{message}\n{path.name}")
            QApplication.processEvents()

        try:
            # Load data file (supports .abf, .smrx, .edf)
            sr, sweeps_by_ch, ch_names, t, file_metadata = abf_io.load_data_file(path, progress_callback=update_progress)
        except Exception as e:
            progress.close()
            self._show_error("Load error", str(e))
            return
        finally:
            progress.close()


        st = self.state
        st.in_path = path
        # Set file_info for single file (for consistency with multi-file loading)
        n_sweeps = next(iter(sweeps_by_ch.values())).shape[1]
        st.file_info = [{
            'path': path,
            'sweep_start': 0,
            'sweep_end': n_sweeps - 1,
            **file_metadata  # Include protocol, n_channels, file_type, etc.
        }]
        st.sr_hz = sr
        st.sweeps = sweeps_by_ch
        st.channel_names = ch_names
        st.t = t
        st.sweep_idx = 0
        self.navigation_manager.reset_window_state()

        # Update filename display with metadata
        self._update_filename_display()

        # Log telemetry: file loaded with enhanced metrics
        file_ext = path.suffix.lower()[1:]  # .abf -> abf
        load_duration = time.time() - t_start
        file_size_mb = path.stat().st_size / (1024 * 1024)
        duration_minutes = (t[-1] - t[0]) / 60 if len(t) > 1 else 0

        telemetry.log_file_loaded(
            file_type=file_ext,
            num_sweeps=n_sweeps,
            num_breaths=None,  # Not detected yet
            file_size_mb=round(file_size_mb, 2),
            sampling_rate_hz=int(sr),
            duration_minutes=round(duration_minutes, 1),
            num_channels=len(ch_names)
        )

        # Log file loading timing
        telemetry.log_timing('file_load', load_duration,
                            file_size_mb=round(file_size_mb, 2),
                            num_sweeps=n_sweeps)

        # Reset peak results and trace cache
        if not hasattr(st, "peaks_by_sweep"):
            st.peaks_by_sweep = {}
            st.breath_by_sweep = {}
        else:
            st.peaks_by_sweep.clear()
            self.state.sigh_by_sweep.clear()
            st.breath_by_sweep.clear()
            self.state.omitted_sweeps.clear()
            self.state.omitted_ranges.clear()
            self._refresh_omit_button_label()

        # Clear cross-sweep outlier detection data
        self.metrics_by_sweep.clear()
        self.onsets_by_sweep.clear()
        self.global_outlier_stats = None

        # Clear export metric cache when loading new file
        # This cache stores computed metric traces during export for reuse in PDF generation
        self._export_metric_cache = {}

        # Clear z-score global statistics cache
        self.zscore_global_mean = None
        self.zscore_global_std = None

        # Clear event markers (bout annotations)
        if hasattr(st, 'bout_annotations'):
            st.bout_annotations.clear()




        # Reset Apply button
        self.ApplyPeakFindPushButton.setEnabled(False)


        


        # Fill combos safely (no signal during population)
        self.AnalyzeChanSelect.blockSignals(True)
        self.AnalyzeChanSelect.clear()
        self.AnalyzeChanSelect.addItem("All Channels")  # First option for grid view
        self.AnalyzeChanSelect.addItems(ch_names)
        self.AnalyzeChanSelect.setCurrentIndex(0)  # default = "All Channels" (grid mode)
        self.AnalyzeChanSelect.blockSignals(False)

        self.StimChanSelect.blockSignals(True)
        self.StimChanSelect.clear()
        self.StimChanSelect.addItem("None")        # default
        self.StimChanSelect.addItems(ch_names)
        self.StimChanSelect.setCurrentIndex(0)     # select "None"
        self.StimChanSelect.blockSignals(False)

        # Populate Events Channel dropdown
        self.EventsChanSelect.blockSignals(True)
        self.EventsChanSelect.clear()
        self.EventsChanSelect.addItem("None")      # default - no event channel
        self.EventsChanSelect.addItems(ch_names)
        # Restore previous selection if it exists
        if st.event_channel and st.event_channel in ch_names:
            idx = ch_names.index(st.event_channel) + 1  # +1 because "None" is at index 0
            self.EventsChanSelect.setCurrentIndex(idx)
        else:
            self.EventsChanSelect.setCurrentIndex(0)  # select "None"
        self.EventsChanSelect.blockSignals(False)

        # Populate Channel Manager widget
        self._populate_channel_manager(ch_names)

        #Clear peaks
        self.state.peaks_by_sweep.clear()
        self.state.sigh_by_sweep.clear()
        self.state.breath_by_sweep.clear()

        #Clear omitted sweeps and regions
        self.state.omitted_sweeps.clear()
        self.state.omitted_ranges.clear()
        self._refresh_omit_button_label()

        # Auto-detect stimulus and analysis channels
        # Note: auto_analysis is only returned if there's exactly one non-stim channel
        auto_stim, auto_analysis = abf_io.auto_select_channels(st.sweeps, ch_names)

        if auto_stim:
            # Found a stimulus channel - select it
            stim_idx = ch_names.index(auto_stim) + 1  # +1 because "None" is at index 0
            self.StimChanSelect.setCurrentIndex(stim_idx)
            st.stim_chan = auto_stim

            # Update channel manager to show this as Opto Stim
            self.channel_manager.set_channel_type(auto_stim, "Opto Stim")

            # Trigger stim detection
            self.on_stim_channel_changed(stim_idx)

            if auto_analysis:
                # Only one analysis channel available - auto-select it
                analysis_idx = ch_names.index(auto_analysis) + 1  # +1 because "All Channels" is at index 0
                self.AnalyzeChanSelect.setCurrentIndex(analysis_idx)
                st.analyze_chan = auto_analysis
                self.single_panel_mode = True

                # Trigger analysis channel change
                self.on_analyze_channel_changed(analysis_idx)

                # Show completion message with auto-detection info
                t_elapsed = time.time() - t_start
                self._log_status_message(
                    f"File loaded ({t_elapsed:.1f}s) - Auto-detected: stim={auto_stim}, analysis={auto_analysis}",
                    4000
                )
            else:
                # Multiple analysis channels available - start in grid mode, let user choose
                st.analyze_chan = None
                self.single_panel_mode = False
                self.plot_host.clear_saved_view("grid")
                self.plot_all_channels()

                t_elapsed = time.time() - t_start
                self._log_status_message(
                    f"File loaded ({t_elapsed:.1f}s) - Auto-detected stim: {auto_stim}",
                    4000
                )
        else:
            # No stimulus channel detected - start in grid mode
            st.analyze_chan = None
            self.single_panel_mode = False

            st.stim_chan = None
            st.stim_onsets_by_sweep.clear()
            st.stim_offsets_by_sweep.clear()
            st.stim_spans_by_sweep.clear()
            st.stim_metrics_by_sweep.clear()

            self.plot_host.clear_saved_view("grid")
            self.plot_all_channels()

            # Show completion message
            t_elapsed = time.time() - t_start
            self._log_status_message(f"File loaded ({t_elapsed:.1f}s)", 3000)

        # Apply pending channel selections from Project Builder (if any)
        self._apply_pending_channel_selections(ch_names)

        # Refresh Analysis Options dialog tabs if open (file data now available)
        if hasattr(self, '_analysis_options_dialog') and self._analysis_options_dialog is not None and self._analysis_options_dialog.isVisible():
            # Note: Don't refresh Peak Detection tab here since no channel selected yet (in grid mode)
            # Peak Detection tab will update when user selects a channel
            pass

    def _apply_pending_channel_selections(self, ch_names):
        """
        Apply pending channel selections from Project Builder after file loads.

        When a file is opened from the Project Builder, we may have pre-selected
        analysis and stim channels that should be applied after loading.
        """
        st = self.state

        # Check for pending analysis channel
        pending_analysis = getattr(self, '_pending_analysis_channel', '')
        pending_stim = getattr(self, '_pending_stim_channels', [])

        if not pending_analysis and not pending_stim:
            return  # Nothing to apply

        print(f"[channel-selection] Applying pending selections: analysis={pending_analysis}, stim={pending_stim}")

        # Apply pending stim channel (first one if multiple)
        if pending_stim:
            # Take first stim channel from list
            stim_channel = pending_stim[0] if isinstance(pending_stim, list) else pending_stim

            # Find the channel in the dropdown
            # StimChanSelect has "None" at index 0, then channel names
            if stim_channel in ch_names:
                stim_idx = ch_names.index(stim_channel) + 1  # +1 for "None"
                self.StimChanSelect.setCurrentIndex(stim_idx)
                st.stim_chan = stim_channel
                self.on_stim_channel_changed(stim_idx)
                print(f"[channel-selection] Applied stim channel: {stim_channel}")

        # Apply pending analysis channel
        if pending_analysis:
            # Find the channel in the dropdown
            # AnalyzeChanSelect has "All Channels" at index 0, then channel names
            if pending_analysis in ch_names:
                analysis_idx = ch_names.index(pending_analysis) + 1  # +1 for "All Channels"
                self.AnalyzeChanSelect.setCurrentIndex(analysis_idx)
                st.analyze_chan = pending_analysis
                self.single_panel_mode = True
                self.on_analyze_channel_changed(analysis_idx)
                print(f"[channel-selection] Applied analysis channel: {pending_analysis}")

                self._log_status_message(
                    f"Loaded with pre-selected channels: analysis={pending_analysis}"
                    + (f", stim={pending_stim[0]}" if pending_stim else ""),
                    4000
                )

        # Clear pending selections
        self._pending_analysis_channel = ''
        self._pending_stim_channels = []

    def load_multiple_files(self, file_paths: List[Path]):
        """Load and concatenate multiple ABF files."""
        from PyQt6.QtWidgets import QProgressDialog, QApplication, QMessageBox
        from PyQt6.QtCore import Qt

        # Validate files first
        valid, messages = abf_io.validate_files_for_concatenation(file_paths)

        if not valid:
            # Show error dialog
            self._show_error("File Validation Failed", "\n".join(messages))
            return
        elif messages:  # Warnings
            # Show warning dialog with option to proceed
            reply = QMessageBox.question(
                self,
                "File Validation Warnings",
                "The following warnings were detected:\n\n" + "\n".join(messages) + "\n\nDo you want to proceed anyway?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.No:
                return

        # Create progress dialog
        progress = QProgressDialog(f"Loading {len(file_paths)} files...", None, 0, 100, self)
        progress.setWindowTitle("Loading Multiple Files")
        progress.setWindowModality(Qt.WindowModality.WindowModal)
        progress.setMinimumDuration(0)  # Show immediately
        progress.setCancelButton(None)  # No cancel button
        progress.setValue(0)
        progress.show()
        QApplication.processEvents()

        def update_progress(current, total, message):
            """Callback to update progress dialog."""
            progress.setValue(current)
            progress.setLabelText(message)
            QApplication.processEvents()

        try:
            # Load and concatenate files
            sr, sweeps_by_ch, ch_names, t, file_info = abf_io.load_and_concatenate_abf_files(
                file_paths, progress_callback=update_progress
            )
        except Exception as e:
            progress.close()
            self._show_error("Load error", str(e))
            return
        finally:
            progress.close()

        # Update state (similar to load_file, but with file_info)
        st = self.state
        st.in_path = file_paths[0]  # Store first file path for display
        st.file_info = file_info  # Store multi-file metadata
        st.sr_hz = sr
        st.sweeps = sweeps_by_ch
        st.channel_names = ch_names
        st.t = t
        st.sweep_idx = 0
        self.navigation_manager.reset_window_state()

        # Reset peak results and trace cache
        if not hasattr(st, "peaks_by_sweep"):
            st.peaks_by_sweep = {}
            st.breath_by_sweep = {}
        else:
            st.peaks_by_sweep.clear()
            self.state.sigh_by_sweep.clear()
            st.breath_by_sweep.clear()
            self.state.omitted_sweeps.clear()
            self.state.omitted_ranges.clear()
            self._refresh_omit_button_label()

        # Clear cross-sweep outlier detection data
        self.metrics_by_sweep.clear()
        self.onsets_by_sweep.clear()
        self.global_outlier_stats = None

        # Clear export metric cache when loading new file
        # This cache stores computed metric traces during export for reuse in PDF generation
        self._export_metric_cache = {}

        # Reset Apply button
        self.ApplyPeakFindPushButton.setEnabled(False)

        # Fill combos safely (no signal during population)
        self.AnalyzeChanSelect.blockSignals(True)
        self.AnalyzeChanSelect.clear()
        self.AnalyzeChanSelect.addItem("All Channels")  # First option for grid view
        self.AnalyzeChanSelect.addItems(ch_names)
        self.AnalyzeChanSelect.setCurrentIndex(0)  # default = "All Channels" (grid mode)
        self.AnalyzeChanSelect.blockSignals(False)

        self.StimChanSelect.blockSignals(True)
        self.StimChanSelect.clear()
        self.StimChanSelect.addItem("None")        # default
        self.StimChanSelect.addItems(ch_names)
        self.StimChanSelect.setCurrentIndex(0)     # select "None"
        self.StimChanSelect.blockSignals(False)

        # Populate Events Channel dropdown
        self.EventsChanSelect.blockSignals(True)
        self.EventsChanSelect.clear()
        self.EventsChanSelect.addItem("None")      # default - no event channel
        self.EventsChanSelect.addItems(ch_names)
        # Restore previous selection if it exists
        if st.event_channel and st.event_channel in ch_names:
            idx = ch_names.index(st.event_channel) + 1  # +1 because "None" is at index 0
            self.EventsChanSelect.setCurrentIndex(idx)
        else:
            self.EventsChanSelect.setCurrentIndex(0)  # select "None"
        self.EventsChanSelect.blockSignals(False)

        # Populate Channel Manager widget
        self._populate_channel_manager(ch_names)

        #Clear peaks
        self.state.peaks_by_sweep.clear()
        self.state.sigh_by_sweep.clear()
        self.state.breath_by_sweep.clear()

        #Clear omitted sweeps and regions
        self.state.omitted_sweeps.clear()
        self.state.omitted_ranges.clear()
        self._refresh_omit_button_label()

        # Clear z-score global statistics cache
        self.zscore_global_mean = None
        self.zscore_global_std = None

        # Auto-detect stimulus and analysis channels
        # Note: auto_analysis is only returned if there's exactly one non-stim channel
        auto_stim, auto_analysis = abf_io.auto_select_channels(st.sweeps, ch_names)

        if auto_stim:
            # Found a stimulus channel - select it
            stim_idx = ch_names.index(auto_stim) + 1  # +1 because "None" is at index 0
            self.StimChanSelect.setCurrentIndex(stim_idx)
            st.stim_chan = auto_stim

            # Update channel manager to show this as Opto Stim
            self.channel_manager.set_channel_type(auto_stim, "Opto Stim")

            # Trigger stim detection
            self.on_stim_channel_changed(stim_idx)

            if auto_analysis:
                # Only one analysis channel available - auto-select it
                analysis_idx = ch_names.index(auto_analysis) + 1  # +1 because "All Channels" is at index 0
                self.AnalyzeChanSelect.setCurrentIndex(analysis_idx)
                st.analyze_chan = auto_analysis
                self.single_panel_mode = True

                # Trigger analysis channel change
                self.on_analyze_channel_changed(analysis_idx)
            else:
                # Multiple analysis channels available - start in grid mode, let user choose
                st.analyze_chan = None
                self.single_panel_mode = False
                self.plot_host.clear_saved_view("grid")
                self.plot_all_channels()
        else:
            # No stimulus channel detected - start in grid mode
            st.analyze_chan = None
            self.single_panel_mode = False

            st.stim_chan = None
            st.stim_onsets_by_sweep.clear()
            st.stim_offsets_by_sweep.clear()
            st.stim_spans_by_sweep.clear()
            st.stim_metrics_by_sweep.clear()

            self.plot_host.clear_saved_view("grid")
            self.plot_all_channels()

        # Show success message with file info
        total_sweeps = next(iter(sweeps_by_ch.values())).shape[1]

        # Build file summary with padding information
        file_lines = []
        for i, info in enumerate(file_info):
            line = f"  {i+1}. {info['path'].name}: sweeps {info['sweep_start']}-{info['sweep_end']}"
            if info.get('padded', False):
                orig_dur = info['original_samples'] / sr
                padded_dur = info['padded_samples'] / sr
                line += f" (padded: {orig_dur:.2f}s → {padded_dur:.2f}s)"
            file_lines.append(line)

        file_summary = "\n".join(file_lines)

        # Check if any files were padded
        padded_count = sum(1 for info in file_info if info.get('padded', False))

        message = f"Loaded {len(file_paths)} files with {total_sweeps} total sweeps:\n\n{file_summary}"
        if padded_count > 0:
            message += f"\n\nNote: {padded_count} file(s) had different sweep lengths and were padded with NaN values."

        # Add auto-detection info to message
        if auto_stim:
            message += f"\n\nAuto-detected stimulus channel: {auto_stim}"
            if auto_analysis:
                message += f"\nAuto-selected analysis channel: {auto_analysis}"

        self._show_info("Files Loaded Successfully", message)

    # ---------- Session Save/Load ----------
    def keyPressEvent(self, event):
        """Handle keyboard shortcuts globally."""
        from PyQt6.QtCore import Qt
        from PyQt6.QtGui import QKeyEvent

        # Ctrl+S - Save Data (same as Save Data button)
        if event.modifiers() == Qt.KeyboardModifier.ControlModifier and event.key() == Qt.Key.Key_S:
            self.on_save_analyzed_clicked()
            event.accept()
        else:
            # Pass event to parent for default handling
            super().keyPressEvent(event)

    def save_session_state(self):
        """Save current analysis state to .pleth.npz file (Ctrl+S)."""
        from core.npz_io import save_state_to_npz

        # Check if we have data loaded
        if not self.state.in_path:
            self._show_warning("No Data Loaded",
                "Please load a data file before saving session state."
            )
            return

        # Check if channel is selected
        if not self.state.analyze_chan:
            self._show_warning("No Channel Selected",
                "Please select a channel to analyze before saving.\n\n"
                "(Session state is saved per-channel, allowing independent analysis of multi-channel files)"
            )
            return

        # Default to analysis folder with simple naming (no metadata required for quick save)
        analysis_folder = self.state.in_path.parent / "Pleth_App_Analysis"
        analysis_folder.mkdir(exist_ok=True)

        # Use simple default name: {datafile}_{channel}_session.npz
        safe_channel = self.state.analyze_chan.replace(' ', '_').replace('/', '_').replace('\\', '_')
        default_filename = f"{self.state.in_path.stem}_{safe_channel}_session.npz"
        default_path = analysis_folder / default_filename

        # Let user modify path if desired
        save_path, _ = QFileDialog.getSaveFileName(
            self, "Save Session State",
            str(default_path),
            "PhysioMetrics Session (*.pleth.npz);;All Files (*)"
        )

        if not save_path:
            return  # User cancelled

        save_path = Path(save_path)

        # Ask about including raw data (for portability vs file size)
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QLabel, QCheckBox, QDialogButtonBox

        dialog = QDialog(self)
        dialog.setWindowTitle("Save Options")
        layout = QVBoxLayout()

        label = QLabel(
            "Choose what to include in the session file:\n\n"
            "• Analysis results (peaks, edits, filters) - Always included\n"
            "• Raw signal data - Optional (makes file portable but larger)"
        )
        layout.addWidget(label)

        include_raw_checkbox = QCheckBox("Include raw signal data (for portability)")
        include_raw_checkbox.setToolTip(
            "If checked: File can be loaded without original .abf file (larger file ~65MB)\n"
            "If unchecked: File will reload from original .abf file (smaller file ~5-10MB)"
        )
        include_raw_checkbox.setChecked(False)  # Default: don't include (smaller files)
        layout.addWidget(include_raw_checkbox)

        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        buttons.accepted.connect(dialog.accept)
        buttons.rejected.connect(dialog.reject)
        layout.addWidget(buttons)

        dialog.setLayout(layout)

        if dialog.exec() != QDialog.DialogCode.Accepted:
            return  # User cancelled

        include_raw = include_raw_checkbox.isChecked()

        # Save to NPZ
        try:
            import time
            t_start = time.time()

            # Pass GMM cache to preserve user's cluster assignments
            gmm_cache = getattr(self, '_cached_gmm_results', None)

            # Collect app-level settings to preserve
            app_settings = {
                'filter_order': self.filter_order,
                'use_zscore_normalization': self.use_zscore_normalization,
                'notch_filter_lower': self.notch_filter_lower,
                'notch_filter_upper': self.notch_filter_upper,
                'apnea_threshold': self._parse_float(self.ApneaThresh) or 0.5,
                'active_eupnea_sniff_classifier': self.state.active_eupnea_sniff_classifier
            }

            save_state_to_npz(self.state, save_path, include_raw_data=include_raw, gmm_cache=gmm_cache, app_settings=app_settings)

            t_elapsed = time.time() - t_start
            file_size_mb = save_path.stat().st_size / (1024 * 1024)

            self._log_status_message(
                f"Session saved: {save_path.name} ({file_size_mb:.1f} MB, {t_elapsed:.1f}s)",
                timeout=5000
            )

            # Count eupnea and sniffing breaths for telemetry
            eupnea_count = 0
            sniff_count = 0
            for s in self.state.sweeps.keys():
                breath_data = self.state.breath_by_sweep.get(s, {})
                onsets = breath_data.get('onsets', [])
                sniff_regions = self.state.sniff_regions_by_sweep.get(s, [])

                for i in range(len(onsets) - 1):
                    # Check if breath midpoint falls in any sniffing region
                    t_start = self.state.t[onsets[i]]
                    t_end = self.state.t[onsets[i + 1]]
                    t_mid = (t_start + t_end) / 2.0

                    is_sniff = False
                    for (region_start, region_end) in sniff_regions:
                        if region_start <= t_mid <= region_end:
                            is_sniff = True
                            break

                    if is_sniff:
                        sniff_count += 1
                    else:
                        eupnea_count += 1

            # Log telemetry: file saved with per-file edit metrics (for ML evaluation)
            telemetry.log_file_saved(
                save_type='npz',
                eupnea_count=eupnea_count,
                sniff_count=sniff_count,
                file_size_mb=round(file_size_mb, 2),
                include_raw_data=include_raw,
                num_sweeps=len(self.state.sweeps)
            )

            # Update settings with last save location
            self.settings.setValue("last_npz_save_dir", str(save_path.parent))

        except Exception as e:
            self._show_error("Save Error",
                f"Failed to save session state:\n\n{str(e)}"
            )

    def load_npz_state(self, npz_path: Path, alternative_data_path: Path = None):
        """Load complete analysis state from .pleth.npz file."""
        from core.npz_io import load_state_from_npz, get_npz_metadata, OriginalFileNotFoundError
        import time

        t_start = time.time()

        # Get metadata for display
        metadata = get_npz_metadata(npz_path)

        if 'error' in metadata:
            self._show_error("Load Error",
                f"Failed to read NPZ file:\n\n{metadata['error']}"
            )
            return

        # If no alternative path provided, check if stored path exists - if not, try project file list
        if alternative_data_path is None and hasattr(self, '_master_file_list') and self._master_file_list:
            # Get the original filename from metadata
            original_path_str = metadata.get('original_file', '')
            if original_path_str:
                stored_path = Path(original_path_str)

                # Only search project file list if the stored path doesn't exist
                if not stored_path.exists():
                    original_filename = stored_path.name
                    print(f"\n[Path Resolution] NPZ stored path not found: {stored_path}")
                    print(f"  Searching project file list for: {original_filename}")

                    # Search project's master file list for a matching filename
                    for task in self._master_file_list:
                        task_path = task.get('file_path', '')
                        if task_path and Path(task_path).name == original_filename:
                            candidate = Path(task_path)
                            if candidate.exists():
                                # Show detailed comparison of paths
                                print(f"  Found in project: {candidate}")

                                # Analyze the difference
                                stored_parts = stored_path.parts
                                project_parts = candidate.parts
                                if stored_parts != project_parts:
                                    # Find where they diverge
                                    for i, (s, p) in enumerate(zip(stored_parts, project_parts)):
                                        if s != p:
                                            print(f"  Path diverges at part {i}: '{s}' vs '{p}'")
                                            break
                                    if len(stored_parts) != len(project_parts):
                                        print(f"  Path depth differs: {len(stored_parts)} vs {len(project_parts)} parts")

                                alternative_data_path = candidate
                                break
                    else:
                        print(f"  Not found in project file list")

        # Show loading dialog
        from PyQt6.QtWidgets import QProgressDialog, QApplication, QFileDialog, QMessageBox
        from PyQt6.QtCore import Qt

        progress = QProgressDialog(f"Loading session...\n{npz_path.name}", None, 0, 100, self)
        progress.setWindowTitle("Loading PhysioMetrics Session")
        progress.setWindowModality(Qt.WindowModality.WindowModal)
        progress.setMinimumDuration(0)
        progress.setCancelButton(None)
        progress.setValue(10)
        progress.show()
        QApplication.processEvents()

        try:
            # Load state from NPZ
            progress.setLabelText(f"Reading session file...\n{npz_path.name}")
            progress.setValue(30)
            QApplication.processEvents()

            new_state, raw_data_loaded, gmm_cache, app_settings = load_state_from_npz(
                npz_path, reload_raw_data=True, alternative_data_path=alternative_data_path
            )
        except OriginalFileNotFoundError as e:
            progress.close()

            # Prompt user to locate the file
            reply = QMessageBox.question(
                self,
                "Original File Not Found",
                f"The original data file could not be found:\n\n"
                f"{e.original_path}\n\n"
                f"Would you like to locate the file manually?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                QMessageBox.StandardButton.Yes
            )

            if reply == QMessageBox.StandardButton.Yes:
                # Get file extension from original path
                orig_ext = e.original_path.suffix.lower()
                filter_map = {
                    '.abf': "ABF Files (*.abf)",
                    '.smrx': "Spike2 Files (*.smrx)",
                    '.edf': "EDF Files (*.edf)",
                }
                file_filter = filter_map.get(orig_ext, "All Data Files (*.abf *.smrx *.edf)")

                new_path, _ = QFileDialog.getOpenFileName(
                    self,
                    f"Locate {e.original_path.name}",
                    str(Path.home()),
                    file_filter
                )

                if new_path:
                    # Retry with the new path
                    self.load_npz_state(npz_path, alternative_data_path=Path(new_path))
                    return

            return  # User cancelled or said No

        except Exception as e:
            progress.close()
            import traceback
            self._show_error("Load Error",
                f"Failed to load session state:\n\n{str(e)}\n\n{traceback.format_exc()}"
            )
            return

        # Continue with successful load
        try:
            if not raw_data_loaded:
                progress.close()
                self._show_error("Load Error",
                    "Could not load raw data from original file or NPZ.\n\n"
                    f"Original file: {new_state.in_path}\n\n"
                    "Please ensure the original data file is accessible."
                )
                return

            progress.setLabelText("Restoring analysis state...")
            progress.setValue(50)
            QApplication.processEvents()

            # Replace current state
            self.state = new_state
            st = self.state

            # Update manager references to new state
            self.plot_manager.state = new_state
            self.navigation_manager.state = new_state
            self.editing_modes.state = new_state
            self.export_manager.state = new_state

            # ===== RESTORE UI ELEMENTS =====

            # Update filename display in status bar (with metadata if available)
            self._update_filename_display()

            progress.setValue(60)
            QApplication.processEvents()

            # Restore channel combos
            self.AnalyzeChanSelect.blockSignals(True)
            self.AnalyzeChanSelect.clear()
            self.AnalyzeChanSelect.addItem("All Channels")
            self.AnalyzeChanSelect.addItems(st.channel_names)

            if st.analyze_chan and st.analyze_chan in st.channel_names:
                idx = st.channel_names.index(st.analyze_chan) + 1  # +1 for "All Channels"
                self.AnalyzeChanSelect.setCurrentIndex(idx)
                self.single_panel_mode = True
            else:
                self.AnalyzeChanSelect.setCurrentIndex(0)  # "All Channels"
                self.single_panel_mode = False

            self.AnalyzeChanSelect.blockSignals(False)

            self.StimChanSelect.blockSignals(True)
            self.StimChanSelect.clear()
            self.StimChanSelect.addItem("None")
            self.StimChanSelect.addItems(st.channel_names)

            if st.stim_chan and st.stim_chan in st.channel_names:
                idx = st.channel_names.index(st.stim_chan) + 1
                self.StimChanSelect.setCurrentIndex(idx)
            else:
                self.StimChanSelect.setCurrentIndex(0)

            self.StimChanSelect.blockSignals(False)

            self.EventsChanSelect.blockSignals(True)
            self.EventsChanSelect.clear()
            self.EventsChanSelect.addItem("None")
            self.EventsChanSelect.addItems(st.channel_names)

            if st.event_channel and st.event_channel in st.channel_names:
                idx = st.channel_names.index(st.event_channel) + 1
                self.EventsChanSelect.setCurrentIndex(idx)
            else:
                self.EventsChanSelect.setCurrentIndex(0)

            self.EventsChanSelect.blockSignals(False)

            progress.setValue(70)
            QApplication.processEvents()

            # Restore filter settings
            self.LowPass_checkBox.setChecked(st.use_low)
            self.HighPass_checkBox.setChecked(st.use_high)
            self.InvertSignal_checkBox.setChecked(st.use_invert)
            # Note: use_mean_sub is stored but not restored (no UI checkbox)

            if st.low_hz:
                self.LowPassVal.setText(str(st.low_hz))
            if st.high_hz:
                self.HighPassVal.setText(str(st.high_hz))

            # Restore app-level settings (filter order, zscore, notch, apnea threshold)
            if app_settings is not None:
                self.filter_order = app_settings.get('filter_order', 4)
                self.use_zscore_normalization = app_settings.get('use_zscore_normalization', True)
                self.notch_filter_lower = app_settings.get('notch_filter_lower')
                self.notch_filter_upper = app_settings.get('notch_filter_upper')
                apnea_thresh = app_settings.get('apnea_threshold', 0.5)

                # Update UI elements
                if hasattr(self, 'FilterOrderSpin'):
                    self.FilterOrderSpin.setValue(self.filter_order)
                if hasattr(self, 'ApneaThresh'):
                    self.ApneaThresh.setText(str(apnea_thresh))

                # Restore eupnea/sniff classifier
                saved_classifier = app_settings.get('active_eupnea_sniff_classifier', 'gmm')
                self.state.active_eupnea_sniff_classifier = saved_classifier

                # Update dropdown to match saved classifier
                classifier_to_dropdown = {
                    'gmm': 'GMM',
                    'xgboost': 'XGBoost',
                    'rf': 'Random Forest',
                    'mlp': 'MLP',
                    'all_eupnea': 'All Eupnea',
                    'none': 'None (Clear)'
                }
                dropdown_text = classifier_to_dropdown.get(saved_classifier, 'GMM')
                self.eup_sniff_combo.blockSignals(True)
                self.eup_sniff_combo.setCurrentText(dropdown_text)
                self.eup_sniff_combo.blockSignals(False)

                print(f"[npz-load] Restored app settings: filter_order={self.filter_order}, "
                      f"zscore={self.use_zscore_normalization}, notch={self.notch_filter_lower}-{self.notch_filter_upper}, "
                      f"apnea_thresh={apnea_thresh}, eupnea_sniff_classifier={saved_classifier}")

            progress.setValue(80)
            QApplication.processEvents()

            # Clear caches (same as new file load)
            self.metrics_by_sweep.clear()
            self.onsets_by_sweep.clear()
            self.global_outlier_stats = None
            self._export_metric_cache = {}
            self.zscore_global_mean = None
            self.zscore_global_std = None

            # Update omit button label
            self._refresh_omit_button_label()

            # Disable peak apply button after session load to prevent accidental re-run
            # User already has peaks loaded from session, shouldn't re-detect
            self.ApplyPeakFindPushButton.setEnabled(False)
            self.ApplyPeakFindPushButton.setToolTip("Peak detection already complete (loaded from session). Modify parameters and click to re-detect.")

            progress.setValue(90)
            QApplication.processEvents()

            # Restore navigation and plot
            self.navigation_manager.reset_window_state()

            # Redraw plot (will use single_panel_mode to determine layout)
            self.redraw_main_plot()

            # Restore navigation position (after plotting)
            # Note: Window position is restored in redraw_main_plot via state.window_start_s

            # Restore GMM cache if it was saved (preserves user's cluster assignments)
            if gmm_cache is not None:
                print("[npz-load] Restoring GMM cache from session file...")
                self._cached_gmm_results = gmm_cache
            elif st.gmm_sniff_probabilities:
                # Fallback: rebuild GMM cache if probabilities exist but no cache was saved
                # (for backwards compatibility with old session files)
                # Always rebuild regardless of auto_gmm_enabled setting
                print("[npz-load] Rebuilding GMM cache from loaded probabilities (legacy fallback)...")
                self._run_automatic_gmm_clustering()

            progress.setValue(100)
            progress.close()

            # Show success message
            t_elapsed = time.time() - t_start
            file_size_mb = npz_path.stat().st_size / (1024 * 1024)

            self._log_status_message(
                f"Session loaded: {npz_path.name} ({file_size_mb:.1f} MB, {t_elapsed:.1f}s) - "
                f"Channel: {st.analyze_chan}, {metadata['n_peaks']} peaks",
                timeout=8000
            )

            # Update last directory
            self.settings.setValue("last_dir", str(npz_path.parent))

        except Exception as e:
            progress.close()
            import traceback
            self._show_error("Load Error",
                f"Failed to load session state:\n\n{str(e)}\n\n{traceback.format_exc()}"
            )

    def _proc_key(self, chan: str, sweep: int):
        st = self.state
        return (
            chan, sweep,
            st.use_low,  st.low_hz,
            st.use_high, st.high_hz,
            st.use_mean_sub, st.mean_val,
            st.use_invert,
            self.filter_order,
            self.notch_filter_lower, self.notch_filter_upper,
            self.use_zscore_normalization
        )

    def _compute_global_zscore_stats(self):
        """
        Compute global mean and std across all sweeps for z-score normalization.
        This ensures all sweeps are normalized relative to the same baseline.
        """
        import numpy as np

        st = self.state
        if not st.analyze_chan or st.analyze_chan not in st.sweeps:
            return None, None

        Y = st.sweeps[st.analyze_chan]  # (n_samples, n_sweeps)

        # Collect all processed data across sweeps (apply filters but not z-score)
        all_data = []
        for sweep_idx in range(Y.shape[1]):
            y_raw = Y[:, sweep_idx]

            # Apply all filters EXCEPT z-score
            y = filters.apply_all_1d(
                y_raw, st.sr_hz,
                st.use_low, st.low_hz,
                st.use_high, st.high_hz,
                st.use_mean_sub, st.mean_val,
                st.use_invert,
                order=self.filter_order
            )

            # Apply notch filter if configured
            if self.notch_filter_lower is not None and self.notch_filter_upper is not None:
                y = self._apply_notch_filter(y, st.sr_hz, self.notch_filter_lower, self.notch_filter_upper)

            all_data.append(y)

        # Concatenate all sweeps
        concatenated = np.concatenate(all_data)

        # Compute global statistics (excluding NaN values)
        valid_mask = ~np.isnan(concatenated)
        if not np.any(valid_mask):
            return None, None

        global_mean = np.mean(concatenated[valid_mask])
        global_std = np.std(concatenated[valid_mask], ddof=1)

        print(f"[zscore] Computed global stats: mean={global_mean:.4f}, std={global_std:.4f}")
        return global_mean, global_std

    def plot_all_channels(self):
        """Delegate to PlotManager."""
        self.plot_manager.plot_all_channels()

    def on_analyze_channel_changed(self, idx: int):
        """Apply analyze channel selection immediately."""
        st = self.state
        if not st.channel_names:
            return

        # Check if "All Channels" was selected (idx 0)
        if idx == 0:
            # Switch to grid mode (multi-channel view)
            if self.single_panel_mode:
                self.single_panel_mode = False
                st.analyze_chan = None

                # Clear stimulus data but keep the channel selected in dropdown
                # so it will be recomputed when switching back to single channel
                st.stim_onsets_by_sweep.clear()
                st.stim_offsets_by_sweep.clear()
                st.stim_spans_by_sweep.clear()
                st.stim_metrics_by_sweep.clear()

                # Clear event markers (bout annotations)
                if hasattr(st, 'bout_annotations'):
                    st.bout_annotations.clear()

                st.proc_cache.clear()

                # Clear Y2 plot data
                st.y2_metric_key = None
                st.y2_values_by_sweep.clear()
                self.plot_host.clear_y2()
                # Reset Y2 dropdown to "None"
                self.y2plot_dropdown.blockSignals(True)
                self.y2plot_dropdown.setCurrentIndex(0)  # First item is "None"
                self.y2plot_dropdown.blockSignals(False)

                # Clear saved view to force fresh autoscale for grid mode
                self.plot_host.clear_saved_view("grid")
                self.plot_host.clear_saved_view("single")

                # Switch to grid plot
                self.plot_all_channels()
        elif 0 < idx <= len(st.channel_names):
            # Switch to single channel view
            new_chan = st.channel_names[idx - 1]  # -1 because idx 0 is "All Channels"
            if new_chan != st.analyze_chan or not self.single_panel_mode:
                # Check if session files exist for this channel in analysis folder
                if st.in_path and new_chan:
                    from core.npz_io import get_npz_metadata
                    from datetime import datetime

                    # Search analysis folder for session files matching this channel
                    analysis_folder = st.in_path.parent / "Pleth_App_Analysis"
                    session_files = []

                    if analysis_folder.exists():
                        safe_channel = new_chan.replace(' ', '_').replace('/', '_').replace('\\', '_')
                        data_stem = st.in_path.stem
                        pattern = f"*_{data_stem}_{safe_channel}_session.npz"
                        session_files = list(analysis_folder.glob(pattern))

                        # Sort by modification time (newest first)
                        session_files.sort(key=lambda p: p.stat().st_mtime, reverse=True)

                    if session_files:
                        # Found session file(s) - use most recent
                        npz_path = session_files[0]

                        # Get metadata for display
                        metadata = get_npz_metadata(npz_path)

                        # Prompt user to load existing analysis
                        reply = QMessageBox.question(
                            self,
                            "Load Existing Analysis?",
                            f"Found saved analysis for channel '{new_chan}':\n\n"
                            f"{npz_path.name}\n"
                            f"Last modified: {metadata.get('modified_time', 'unknown')}\n"
                            f"Contains: {metadata.get('n_peaks', 0)} peaks\n"
                            f"GMM clustering: {'Yes' if metadata.get('has_gmm', False) else 'No'}\n\n"
                            f"Load this analysis?",
                            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                            QMessageBox.StandardButton.Yes  # Default to Yes
                        )

                        if reply == QMessageBox.StandardButton.Yes:
                            # User wants to load existing analysis
                            # Block signals to prevent recursive calls
                            self.AnalyzeChanSelect.blockSignals(True)
                            self.load_npz_state(npz_path)
                            self.AnalyzeChanSelect.blockSignals(False)

                            # Still run auto-detect to populate prominence field for this channel
                            self._auto_detect_prominence_silent()
                            return  # Don't continue with fresh channel switch

                # User declined or no NPZ exists - continue with fresh channel
                st.analyze_chan = new_chan

                # Log telemetry: channel selection
                telemetry.log_button_click('select_analyze_channel',
                                          channel_name=new_chan,
                                          channel_index=idx)

                st.proc_cache.clear()
                # Clear z-score global statistics cache
                self.zscore_global_mean = None
                self.zscore_global_std = None
                st.peaks_by_sweep.clear()
                st.sigh_by_sweep.clear()
                if hasattr(st, 'breath_by_sweep'):
                    st.breath_by_sweep.clear()

                # Clear omitted sweeps and regions
                st.omitted_sweeps.clear()
                st.omitted_ranges.clear()
                self._refresh_omit_button_label()

                # Clear sniffing regions
                if hasattr(st, 'sniff_regions_by_sweep'):
                    st.sniff_regions_by_sweep.clear()

                # Clear event markers (bout annotations)
                if hasattr(st, 'bout_annotations'):
                    st.bout_annotations.clear()

                # Immediately clear and refresh the Analysis Options dialog if open
                if hasattr(self, '_analysis_options_dialog') and self._analysis_options_dialog is not None:
                    try:
                        if self._analysis_options_dialog.isVisible():
                            current_tab = self._analysis_options_dialog.tabWidget.currentIndex()
                            print(f"[Channel Change] Clearing Analysis Options dialog (current tab: {current_tab})")

                            # Clear the cached dialogs so they'll recreate
                            if hasattr(self._analysis_options_dialog, 'prominence_dialog'):
                                self._analysis_options_dialog.prominence_dialog = None
                            if hasattr(self._analysis_options_dialog, 'gmm_dialog'):
                                self._analysis_options_dialog.gmm_dialog = None

                            # Immediately clear the current tab's content
                            if current_tab == 0:  # Peak Detection
                                container = self._analysis_options_dialog.peak_detection_container
                            elif current_tab == 1:  # Eup/Sniff Classification (GMM)
                                container = self._analysis_options_dialog.breath_classification_container
                            else:
                                container = None

                            if container:
                                layout = container.layout()
                                if layout:
                                    # Clear all widgets from the container
                                    while layout.count():
                                        child = layout.takeAt(0)
                                        if child.widget():
                                            child.widget().deleteLater()
                                    from PyQt6.QtWidgets import QApplication
                                    QApplication.processEvents()
                                    print(f"[Channel Change] Cleared tab {current_tab} content")

                                # Show "Loading..." placeholder
                                from PyQt6.QtWidgets import QLabel
                                from PyQt6.QtCore import Qt
                                label = QLabel("Loading new channel data...")
                                label.setAlignment(Qt.AlignmentFlag.AlignCenter)
                                label.setStyleSheet("padding: 20px; font-size: 14pt; color: #888;")
                                layout.addWidget(label)

                            # Mark for refresh after peak detection completes
                            self._channel_change_needs_dialog_refresh = True
                            self._dialog_tab_to_refresh = current_tab
                    except RuntimeError:
                        # Dialog was deleted
                        self._analysis_options_dialog = None

                # Clear Y2 plot data
                st.y2_metric_key = None
                st.y2_values_by_sweep.clear()
                self.plot_host.clear_y2()
                # Reset Y2 dropdown to "None"
                self.y2plot_dropdown.blockSignals(True)
                self.y2plot_dropdown.setCurrentIndex(0)  # First item is "None"
                self.y2plot_dropdown.blockSignals(False)

                # Reset navigation to first sweep
                st.sweep_idx = 0
                st.window_start_s = 0.0

                # Switch to single panel mode
                if not self.single_panel_mode:
                    self.single_panel_mode = True

                # If a stimulus channel is selected, recompute it for the current sweep
                if st.stim_chan is not None:
                    self._compute_stim_for_current_sweep()

                # Clear saved view to force fresh autoscale for single mode
                self.plot_host.clear_saved_view("single")
                self.plot_host.clear_saved_view("grid")

                # Auto-detect optimal prominence in background when channel selected
                self._auto_detect_prominence_silent()

                # Redraw plot
                self.redraw_main_plot()

                # Redraw threshold line after plot is redrawn (it gets cleared during redraw)
                if self.peak_height_threshold is not None:
                    self.plot_host.update_threshold_line(self.peak_height_threshold)

    def on_stim_channel_changed(self, idx: int):
        """Apply stimulus channel selection immediately."""
        st = self.state
        if not st.channel_names:
            return

        # idx 0 = "None", idx 1+ = channel names
        new_stim = None if idx == 0 else st.channel_names[idx - 1]
        if new_stim != st.stim_chan:
            st.stim_chan = new_stim

            # Clear stimulus detection results
            st.stim_onsets_by_sweep.clear()
            st.stim_offsets_by_sweep.clear()
            st.stim_spans_by_sweep.clear()
            st.stim_metrics_by_sweep.clear()

            # Compute stimulus for current sweep if a channel is selected
            if new_stim is not None:
                self._compute_stim_for_current_sweep()

            # Clear saved view to force fresh autoscale when stimulus changes
            self.plot_host.clear_saved_view("single")

            st.proc_cache.clear()
            self.redraw_main_plot()

    def on_events_channel_changed(self, idx: int):
        """Apply event channel selection immediately."""
        st = self.state
        if not st.channel_names:
            return

        # idx 0 = "None", idx 1+ = channel names
        new_event = None if idx == 0 else st.channel_names[idx - 1]
        if new_event != st.event_channel:
            st.event_channel = new_event
            self.redraw_main_plot()

    def _update_classifier_dropdowns(self):
        """Update dropdown options based on loaded models."""
        models = self.state.loaded_ml_models or {}

        # Check which models are available (use prefix matching for accuracy suffix)
        has_model1_xgboost = any(k.startswith('model1_xgboost') for k in models.keys())
        has_model1_rf = any(k.startswith('model1_rf') for k in models.keys())
        has_model1_mlp = any(k.startswith('model1_mlp') for k in models.keys())

        has_model3_xgboost = any(k.startswith('model3_xgboost') for k in models.keys())
        has_model3_rf = any(k.startswith('model3_rf') for k in models.keys())
        has_model3_mlp = any(k.startswith('model3_mlp') for k in models.keys())

        has_model2_xgboost = any(k.startswith('model2_xgboost') for k in models.keys())
        has_model2_rf = any(k.startswith('model2_rf') for k in models.keys())
        has_model2_mlp = any(k.startswith('model2_mlp') for k in models.keys())

        # Update Model 1 (Breath Detection) dropdown
        # Index 0=Threshold (always enabled), 1=XGBoost, 2=RF, 3=MLP
        model = self.peak_detec_combo.model()
        model.item(0).setEnabled(True)  # Threshold always available
        model.item(1).setEnabled(has_model1_xgboost)
        model.item(2).setEnabled(has_model1_rf)
        model.item(3).setEnabled(has_model1_mlp)

        # Update Model 3 (Eupnea/Sniff) dropdown
        # Index 0=None, 1=All Eupnea, 2=GMM (always enabled), 3=XGBoost, 4=RF, 5=MLP
        model = self.eup_sniff_combo.model()
        model.item(0).setEnabled(True)   # "None (Clear)" always available
        model.item(1).setEnabled(True)   # "All Eupnea" always available
        model.item(2).setEnabled(True)   # "GMM" always available
        model.item(3).setEnabled(has_model3_xgboost)  # XGBoost
        model.item(4).setEnabled(has_model3_rf)       # Random Forest
        model.item(5).setEnabled(has_model3_mlp)      # MLP

        # Update Model 2 (Sigh) dropdown
        # Index 0=Manual (always enabled), 1=XGBoost, 2=RF, 3=MLP
        model = self.digh_combo.model()
        model.item(0).setEnabled(True)  # Manual always available
        model.item(1).setEnabled(has_model2_xgboost)
        model.item(2).setEnabled(has_model2_rf)
        model.item(3).setEnabled(has_model2_mlp)

        # If current selection is disabled, fall back to default
        self._fallback_disabled_classifiers()

    def _fallback_disabled_classifiers(self):
        """Check if current classifier selections are disabled and fall back to defaults."""
        # Model 1 (Breath Detection)
        current_text = self.peak_detec_combo.currentText()
        current_index = self.peak_detec_combo.currentIndex()
        if current_index >= 0:
            model = self.peak_detec_combo.model()
            if not model.item(current_index).isEnabled():
                print(f"[Dropdown Update] Model 1 classifier '{current_text}' not available, falling back to Threshold")
                self.peak_detec_combo.setCurrentText("Threshold")
                self.state.active_classifier = "threshold"

        # Model 3 (Eupnea/Sniff)
        current_text = self.eup_sniff_combo.currentText()
        current_index = self.eup_sniff_combo.currentIndex()
        if current_index >= 0:
            model = self.eup_sniff_combo.model()
            if not model.item(current_index).isEnabled():
                print(f"[Dropdown Update] Model 3 classifier '{current_text}' not available, falling back to GMM")
                self.eup_sniff_combo.setCurrentText("GMM")
                self.state.active_eupnea_sniff_classifier = "gmm"

        # Model 2 (Sigh)
        current_text = self.digh_combo.currentText()
        current_index = self.digh_combo.currentIndex()
        if current_index >= 0:
            model = self.digh_combo.model()
            if not model.item(current_index).isEnabled():
                print(f"[Dropdown Update] Model 2 classifier '{current_text}' not available, falling back to Manual")
                self.digh_combo.setCurrentText("Manual")
                self.state.active_sigh_classifier = "manual"

    def _auto_rerun_peak_detection_if_needed(self):
        """
        Automatically re-run peak detection if:
        1. Models were just loaded
        2. Peaks have already been detected once (state.all_peaks_by_sweep is not empty)

        This updates predictions with newly loaded models without requiring manual re-detection.
        """
        # Only re-run if peaks have been detected
        if not self.state.all_peaks_by_sweep:
            print("[Auto Re-run] Skipping - no peaks detected yet")
            return

        # Only re-run if we have models loaded
        if not self.state.loaded_ml_models:
            print("[Auto Re-run] Skipping - no models loaded")
            return

        print("[Auto Re-run] Triggering peak detection to update predictions with newly loaded models")
        self.statusBar().showMessage("Updating predictions with loaded models...", 2000)

        # Trigger peak detection (this will use newly loaded models)
        self._apply_peak_detection()

    def _auto_load_ml_models_on_startup(self):
        """Silently load ML models from last used directory on startup."""
        from pathlib import Path
        import core.ml_prediction as ml_prediction

        # Get last used models directory from settings
        last_models_dir = self.settings.value("ml_models_path", None)

        # Only auto-load if we have a saved path
        if not last_models_dir:
            print("[Auto-load] No saved models directory")
            return

        models_path = Path(last_models_dir)

        # Check if directory still exists
        if not models_path.exists() or not models_path.is_dir():
            print(f"[Auto-load] Saved models directory no longer exists: {models_path}")
            return

        try:
            # Look for model files
            model_files = list(models_path.glob("model*.pkl"))

            if not model_files:
                print(f"[Auto-load] No model files found in: {models_path}")
                return

            # Load all models
            loaded_models = {}
            for model_file in model_files:
                try:
                    model, metadata = ml_prediction.load_model(model_file)
                    model_key = model_file.stem
                    loaded_models[model_key] = {
                        'model': model,
                        'metadata': metadata,
                        'path': str(model_file)
                    }
                except Exception as e:
                    print(f"[Auto-load] Warning: Failed to load {model_file.name}: {e}")

            if loaded_models:
                # Store in state
                self.state.loaded_ml_models = loaded_models

                # Update dropdown states
                self._update_classifier_dropdowns()

                print(f"[Auto-load] Successfully loaded {len(loaded_models)} models from {models_path}")
                self.statusBar().showMessage(f"Auto-loaded {len(loaded_models)} ML models", 3000)
            else:
                print(f"[Auto-load] Failed to load any models from {models_path}")

        except Exception as e:
            print(f"[Auto-load] Error loading models: {e}")
            import traceback
            traceback.print_exc()

    def on_classifier_changed(self, text: str):
        """Handle classifier selection change from main window dropdown."""
        # Map display text to internal classifier name
        classifier_map = {
            "Threshold": "threshold",
            "XGBoost": "xgboost",
            "Random Forest": "rf",
            "MLP": "mlp"
        }

        new_classifier = classifier_map.get(text, "threshold")

        # Update state (only print if actually changing)
        if self.state.active_classifier != new_classifier:
            self.state.active_classifier = new_classifier
            print(f"[Classifier] Switched to: {new_classifier}")

            # Check if ML models are loaded for this classifier
            if new_classifier != 'threshold':
                if not self.state.loaded_ml_models:
                    print(f"[Classifier] ERROR: No ML models loaded at all!")
                    self.statusBar().showMessage(f"WARNING: No ML models loaded. Load models from ML Training tab first.", 5000)
                    # Revert to threshold
                    self.peak_detec_combo.blockSignals(True)
                    self.peak_detec_combo.setCurrentText("Threshold")
                    self.peak_detec_combo.blockSignals(False)
                    self.state.active_classifier = "threshold"
                    print(f"[Classifier] Reverted to threshold (no models)")
                    return
                else:
                    # Find the model key (it may have accuracy suffix like "model1_xgboost_100%")
                    model_key_prefix = f'model1_{new_classifier}'
                    matching_keys = [k for k in self.state.loaded_ml_models.keys() if k.startswith(model_key_prefix)]

                    if not matching_keys:
                        print(f"[Classifier] ERROR: No model matching {model_key_prefix} found in loaded models!")
                        self.statusBar().showMessage(f"WARNING: {text} models not loaded. Load models from ML Training tab first.", 5000)
                        # Revert to threshold
                        self.peak_detec_combo.blockSignals(True)
                        self.peak_detec_combo.setCurrentText("Threshold")
                        self.peak_detec_combo.blockSignals(False)
                        self.state.active_classifier = "threshold"
                        print(f"[Classifier] Reverted to threshold (model not found)")
                        return
                    else:
                        print(f"[Classifier] Model {matching_keys[0]} found! Proceeding...")

            # Re-run peak detection to apply new classifier
            # This will recompute peaks_by_sweep using the new classifier's labels
            if self.state.analyze_chan and self.state.peaks_by_sweep:
                self.statusBar().showMessage(f"Switching to {text} classifier...", 2000)
                # Don't need to re-run full detection, just update which peaks are displayed
                self._update_displayed_peaks_from_classifier()
                # Guard: Only redraw if plot_manager exists (avoid error during initialization)
                if hasattr(self, 'plot_manager'):
                    self.redraw_main_plot()

    def _update_displayed_peaks_from_classifier(self):
        """Update peaks_by_sweep and breath_by_sweep based on active classifier.

        This copies the selected classifier's read-only predictions to the user-editable
        'labels' array, then updates the display.
        """
        st = self.state

        # For each sweep, copy active classifier's predictions to 'labels' (user-editable)
        for s in st.all_peaks_by_sweep.keys():
            all_peaks_data = st.all_peaks_by_sweep[s]

            # Get read-only predictions from active classifier
            active_labels_key_ro = f'labels_{st.active_classifier}_ro'
            if active_labels_key_ro in all_peaks_data and all_peaks_data[active_labels_key_ro] is not None:
                # Copy to user-editable array (this overwrites any manual edits!)
                all_peaks_data['labels'] = all_peaks_data[active_labels_key_ro].copy()
                # Reset label_source to 'auto' since these are fresh auto-predictions
                all_peaks_data['label_source'] = np.array(['auto'] * len(all_peaks_data['labels']))
                print(f"[Classifier Update] Sweep {s}: Copied {active_labels_key_ro} to 'labels', found {np.sum(all_peaks_data['labels'] == 1)} breaths")
            else:
                # Fallback to threshold if active classifier not available
                if 'labels_threshold_ro' in all_peaks_data and all_peaks_data['labels_threshold_ro'] is not None:
                    all_peaks_data['labels'] = all_peaks_data['labels_threshold_ro'].copy()
                    all_peaks_data['label_source'] = np.array(['auto'] * len(all_peaks_data['labels']))
                    print(f"[Classifier Update] Sweep {s}: Falling back to threshold, found {np.sum(all_peaks_data['labels'] == 1)} breaths")
                else:
                    print(f"[Classifier Update] Sweep {s}: WARNING - No predictions available!")
                    continue

            # Extract labeled peaks from user-editable 'labels' array
            labeled_mask = all_peaks_data['labels'] == 1
            labeled_indices = all_peaks_data['indices'][labeled_mask]
            st.peaks_by_sweep[s] = labeled_indices
            print(f"[Classifier Update] Sweep {s}: Updated peaks_by_sweep with {len(labeled_indices)} peaks")

            # Recompute breath events for labeled peaks
            y_proc = self._get_processed_for(st.analyze_chan, s)
            import core.peaks as peakdet
            breaths = peakdet.compute_breath_events(y_proc, labeled_indices, sr_hz=st.sr_hz, exclude_sec=0.030)
            st.breath_by_sweep[s] = breaths

    def on_eupnea_sniff_classifier_changed(self, text: str):
        """Handle eupnea/sniff classifier selection change."""
        classifier_map = {
            "GMM": "gmm",
            "XGBoost": "xgboost",
            "Random Forest": "rf",
            "MLP": "mlp",
            "All Eupnea": "all_eupnea",
            "None (Clear)": "none"
        }
        # Reverse mapping for reverting
        classifier_reverse = {v: k for k, v in classifier_map.items()}

        new_classifier = classifier_map.get(text, "gmm")

        old_classifier = self.state.active_eupnea_sniff_classifier
        self.state.active_eupnea_sniff_classifier = new_classifier

        # Only print if actually switching
        if old_classifier != new_classifier:
            print(f"[Eupnea/Sniff Classifier] Switched to: {new_classifier}")

        # Handle different classifier types
        if new_classifier == 'all_eupnea':
            # Label all breaths as eupnea (class 0)
            self._set_all_breaths_eupnea_sniff_class(0)
            print(f"[Eupnea/Sniff Classifier] Set all breaths to eupnea (for anesthesia experiments)")
        elif new_classifier == 'none':
            # Clear all labels - don't save any eupnea/sniff classification
            self._clear_all_eupnea_sniff_labels()
            print(f"[Eupnea/Sniff Classifier] Cleared all eupnea/sniff labels")
        elif new_classifier == 'gmm':
            # Check if GMM has been run (gmm_class_ro should exist)
            first_sweep_peaks = self.state.all_peaks_by_sweep.get(0, {})
            if 'gmm_class_ro' not in first_sweep_peaks or first_sweep_peaks['gmm_class_ro'] is None:
                print(f"[Eupnea/Sniff Classifier] GMM clustering has not been run yet - running now with default settings...")
                self.statusBar().showMessage(f"Running GMM clustering with default settings...", 2000)

                # Run GMM clustering with default settings
                self._run_automatic_gmm_clustering()
                print(f"[Eupnea/Sniff Classifier] GMM clustering complete")

            # Update breath_type_class from GMM predictions
            self._update_eupnea_sniff_from_classifier()
        else:
            # ML model selected - check if loaded
            model_key_prefix = f'model3_{new_classifier}'
            matching_keys = [k for k in self.state.loaded_ml_models.keys() if k.startswith(model_key_prefix)]

            if not matching_keys:
                print(f"[Eupnea/Sniff Classifier] ERROR: Model {model_key_prefix} not found!")
                self.statusBar().showMessage(f"WARNING: {text} model not loaded. Load Model 3 from ML Training tab.", 5000)
                # Revert to GMM
                self.eup_sniff_combo.blockSignals(True)
                self.eup_sniff_combo.setCurrentText("GMM")
                self.eup_sniff_combo.blockSignals(False)
                self.state.active_eupnea_sniff_classifier = "gmm"
                return

            # Update breath_type_class from ML model predictions
            self._update_eupnea_sniff_from_classifier()

        # Rebuild sniff/eupnea regions for display
        try:
            import core.gmm_clustering as gmm_clustering
            gmm_clustering.build_eupnea_sniffing_regions(self.state, verbose=False)
        except Exception as e:
            print(f"[Eupnea/Sniff Classifier] Warning: Could not rebuild regions: {e}")
            import traceback
            traceback.print_exc()

        # Redraw plot
        # Guard: Only redraw if plot_manager exists (avoid error during initialization)
        if hasattr(self, 'plot_manager'):
            self.redraw_main_plot()

    def _update_eupnea_sniff_from_classifier(self):
        """Copy selected classifier's predictions to breath_type_class array."""
        st = self.state

        for s in st.all_peaks_by_sweep.keys():
            all_peaks = st.all_peaks_by_sweep[s]

            # Get read-only predictions from active classifier
            if st.active_eupnea_sniff_classifier == 'gmm':
                source_key = 'gmm_class_ro'
            else:
                source_key = f'eupnea_sniff_{st.active_eupnea_sniff_classifier}_ro'

            # Debug: Check what keys are available
            if s == 0:
                available_keys = [k for k in all_peaks.keys() if 'eupnea' in k or 'gmm' in k]
                print(f"[Eupnea/Sniff Update] Sweep {s}: Available keys: {available_keys}")
                print(f"[Eupnea/Sniff Update] Sweep {s}: Looking for: {source_key}")

            if source_key in all_peaks and all_peaks[source_key] is not None:
                # Copy to user-editable array
                old_breath_type_class = all_peaks['breath_type_class'].copy() if ('breath_type_class' in all_peaks and all_peaks['breath_type_class'] is not None) else None
                all_peaks['breath_type_class'] = all_peaks[source_key].copy()
                all_peaks['eupnea_sniff_source'] = np.array([st.active_eupnea_sniff_classifier] * len(all_peaks['indices']))

                # Debug: Show what changed
                if s == 0:
                    n_eupnea = np.sum(all_peaks['breath_type_class'] == 0)
                    n_sniff = np.sum(all_peaks['breath_type_class'] == 1)
                    n_unclass = np.sum(all_peaks['breath_type_class'] == -1)
                    print(f"[Eupnea/Sniff Update] Sweep {s}: Copied {source_key} to breath_type_class")
                    print(f"[Eupnea/Sniff Update] Sweep {s}: Eupnea: {n_eupnea}, Sniffing: {n_sniff}, Unclassified: {n_unclass}")
                    if old_breath_type_class is not None:
                        n_changed = np.sum(old_breath_type_class != all_peaks['breath_type_class'])
                        print(f"[Eupnea/Sniff Update] Sweep {s}: Changed {n_changed} classifications")
            else:
                print(f"[Eupnea/Sniff Update] Sweep {s}: WARNING - No predictions available for {st.active_eupnea_sniff_classifier}!")
                if s == 0:
                    if source_key in all_peaks:
                        print(f"[Eupnea/Sniff Update] Sweep {s}: Key exists but is None")
                    else:
                        print(f"[Eupnea/Sniff Update] Sweep {s}: Key does not exist in all_peaks")

    def _set_all_breaths_eupnea_sniff_class(self, class_value: int):
        """
        Set all breaths to a specific eupnea/sniff class.

        Args:
            class_value: 0 = eupnea, 1 = sniffing, -1 = unclassified/other
        """
        import numpy as np
        st = self.state

        for s in st.all_peaks_by_sweep.keys():
            all_peaks = st.all_peaks_by_sweep[s]

            if 'indices' in all_peaks and all_peaks['indices'] is not None:
                n_breaths = len(all_peaks['indices'])
                # Set all to the specified class
                all_peaks['breath_type_class'] = np.full(n_breaths, class_value, dtype=int)
                all_peaks['eupnea_sniff_source'] = np.array([st.active_eupnea_sniff_classifier] * n_breaths)

                if s == 0:
                    print(f"[Eupnea/Sniff] Sweep {s}: Set {n_breaths} breaths to class {class_value}")

    def _clear_all_eupnea_sniff_labels(self):
        """
        Clear all eupnea/sniff labels - used when user selects 'None (Clear)'.

        This removes the breath_type_class array entirely so labels won't be saved.
        Also clears the sniff and eupnea regions from the display.
        """
        st = self.state

        for s in st.all_peaks_by_sweep.keys():
            all_peaks = st.all_peaks_by_sweep[s]

            # Remove breath_type_class array - labels won't be saved
            if 'breath_type_class' in all_peaks:
                del all_peaks['breath_type_class']
            if 'eupnea_sniff_source' in all_peaks:
                del all_peaks['eupnea_sniff_source']

            if s == 0:
                print(f"[Eupnea/Sniff] Sweep {s}: Cleared all labels")

        # Clear the region dictionaries for display
        st.sniff_regions_by_sweep.clear()

    def _clear_all_sigh_labels(self):
        """
        Clear all sigh labels - used when user selects 'None (Clear)'.

        This removes the sigh_class array entirely so labels won't be saved.
        Also clears sigh_by_sweep for display.
        """
        st = self.state

        for s in st.all_peaks_by_sweep.keys():
            all_peaks = st.all_peaks_by_sweep[s]

            # Remove sigh_class array - labels won't be saved
            if 'sigh_class' in all_peaks:
                del all_peaks['sigh_class']
            if 'sigh_source' in all_peaks:
                del all_peaks['sigh_source']

            if s == 0:
                print(f"[Sigh] Sweep {s}: Cleared all labels")

        # Clear sigh_by_sweep for display
        st.sigh_by_sweep.clear()

    def on_sigh_classifier_changed(self, text: str):
        """Handle sigh classifier selection change."""
        classifier_map = {
            "Manual": "manual",
            "XGBoost": "xgboost",
            "Random Forest": "rf",
            "MLP": "mlp",
            "None (Clear)": "none"
        }
        # Reverse mapping for reverting
        classifier_reverse = {v: k for k, v in classifier_map.items()}

        new_classifier = classifier_map.get(text, "manual")

        old_classifier = self.state.active_sigh_classifier
        self.state.active_sigh_classifier = new_classifier

        # Only print if actually switching
        if old_classifier != new_classifier:
            print(f"[Sigh Classifier] Switched to: {new_classifier}")

        # Handle different classifier types
        if new_classifier == 'none':
            # Clear all sigh labels
            self._clear_all_sigh_labels()
            print(f"[Sigh Classifier] Cleared all sigh labels")
        elif new_classifier == 'manual':
            # Update from manual annotations
            self._update_sigh_from_classifier()
        else:
            # ML model - check if loaded
            model_key_prefix = f'model2_{new_classifier}'
            matching_keys = [k for k in self.state.loaded_ml_models.keys() if k.startswith(model_key_prefix)]

            if not matching_keys:
                print(f"[Sigh Classifier] ERROR: Model {model_key_prefix} not found!")
                self.statusBar().showMessage(f"WARNING: {text} model not loaded. Load Model 2 from ML Training tab.", 5000)
                # Revert to Manual
                self.digh_combo.blockSignals(True)
                self.digh_combo.setCurrentText("Manual")
                self.digh_combo.blockSignals(False)
                self.state.active_sigh_classifier = "manual"
                return

            # Update from ML model predictions
            self._update_sigh_from_classifier()

        # Redraw plot
        # Guard: Only redraw if plot_manager exists (avoid error during initialization)
        if hasattr(self, 'plot_manager'):
            self.redraw_main_plot()

    def _update_sigh_from_classifier(self):
        """Copy selected classifier's predictions to sigh_class array and update sigh_by_sweep."""
        import numpy as np
        st = self.state

        for s in st.all_peaks_by_sweep.keys():
            all_peaks = st.all_peaks_by_sweep[s]

            # Get read-only predictions from active classifier
            if st.active_sigh_classifier == 'manual':
                # For manual mode, restore from sigh_manual_ro (preserves only manual annotations)
                # This clears ML predictions and only shows manually-added sighs
                if 'sigh_manual_ro' in all_peaks and all_peaks['sigh_manual_ro'] is not None:
                    all_peaks['sigh_class'] = all_peaks['sigh_manual_ro'].copy()
                    all_peaks['sigh_source'] = np.array(['manual'] * len(all_peaks['indices']))
                    if s == 0:
                        n_sighs = np.sum(all_peaks['sigh_class'] == 1)
                        print(f"[Sigh Update] Sweep {s}: Restored manual annotations from sigh_manual_ro ({n_sighs} sighs)")
                else:
                    # No manual annotations saved - clear all sighs (start fresh)
                    n_breaths = len(all_peaks['indices']) if 'indices' in all_peaks else 0
                    if n_breaths > 0:
                        all_peaks['sigh_class'] = np.zeros(n_breaths, dtype=np.int8)
                        # Mark deleted peaks as unclassified
                        if 'labels' in all_peaks and all_peaks['labels'] is not None:
                            all_peaks['sigh_class'][all_peaks['labels'] == 0] = -1
                        all_peaks['sigh_source'] = np.array(['manual'] * n_breaths)
                    if s == 0:
                        print(f"[Sigh Update] Sweep {s}: Cleared all sighs (manual mode, no saved annotations)")

                # Also clear sigh_by_sweep for this sweep
                if s in st.sigh_by_sweep:
                    # Only keep sighs that are in sigh_manual_ro
                    if 'sigh_manual_ro' in all_peaks and all_peaks['sigh_manual_ro'] is not None:
                        manual_sigh_mask = all_peaks['sigh_manual_ro'] == 1
                        st.sigh_by_sweep[s] = all_peaks['indices'][manual_sigh_mask].tolist()
                    else:
                        st.sigh_by_sweep[s] = []
                continue
            else:
                source_key = f'sigh_{st.active_sigh_classifier}_ro'

            # Debug: Check what keys are available
            if s == 0 and st.active_sigh_classifier != 'manual':
                available_keys = [k for k in all_peaks.keys() if 'sigh' in k]
                print(f"[Sigh Update] Sweep {s}: Available keys: {available_keys}")
                print(f"[Sigh Update] Sweep {s}: Looking for: {source_key}")

            if st.active_sigh_classifier != 'manual' and source_key in all_peaks and all_peaks[source_key] is not None:
                # Copy to user-editable array
                old_sigh_class = all_peaks['sigh_class'].copy() if ('sigh_class' in all_peaks and all_peaks['sigh_class'] is not None) else None
                all_peaks['sigh_class'] = all_peaks[source_key].copy()
                all_peaks['sigh_source'] = np.array([st.active_sigh_classifier] * len(all_peaks['indices']))

                # Debug: Show what changed
                if s == 0:
                    n_normal = np.sum(all_peaks['sigh_class'] == 0)
                    n_sigh = np.sum(all_peaks['sigh_class'] == 1)
                    n_unclass = np.sum(all_peaks['sigh_class'] == -1)
                    print(f"[Sigh Update] Sweep {s}: Copied {source_key} to sigh_class")
                    print(f"[Sigh Update] Sweep {s}: Normal: {n_normal}, Sigh: {n_sigh}, Unclassified: {n_unclass}")
                    if old_sigh_class is not None:
                        n_changed = np.sum(old_sigh_class != all_peaks['sigh_class'])
                        print(f"[Sigh Update] Sweep {s}: Changed {n_changed} classifications")
            elif st.active_sigh_classifier != 'manual':
                print(f"[Sigh Update] Sweep {s}: WARNING - No predictions available for {st.active_sigh_classifier}!")

            # Update sigh_by_sweep from sigh_class for display (backward compatibility)
            if 'sigh_class' in all_peaks and all_peaks['sigh_class'] is not None:
                sigh_mask = all_peaks['sigh_class'] == 1
                st.sigh_by_sweep[s] = all_peaks['indices'][sigh_mask]

    def on_mark_events_clicked(self):
        """Open Event Detection Settings dialog."""
        # Check if event channel is selected
        st = self.state
        if st.event_channel is None:
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.warning(
                self,
                "No Event Channel",
                "Please select an event channel from the 'Events Chan Select' dropdown first."
            )
            return

        # Check if dialog already exists and is visible
        if hasattr(self, '_event_detection_dialog') and self._event_detection_dialog.isVisible():
            # Bring existing dialog to front
            self._event_detection_dialog.raise_()
            self._event_detection_dialog.activateWindow()
            return

        # Create and show dialog (non-modal so plot can be interacted with)
        from dialogs.event_detection_dialog import EventDetectionDialog
        self._event_detection_dialog = EventDetectionDialog(parent=self, main_window=self)
        self._event_detection_dialog.show()  # Non-modal dialog


    def _compute_stim_for_current_sweep(self, thresh: float = 1.0):
        st = self.state
        if not st.stim_chan or st.stim_chan not in st.sweeps:
            return
        Y = st.sweeps[st.stim_chan]
        s = max(0, min(st.sweep_idx, Y.shape[1] - 1))
        y = Y[:, s]
        t = st.t

        on_idx, off_idx, spans_s, metrics = stimdet.detect_threshold_crossings(y, t, thresh=thresh)
        st.stim_onsets_by_sweep[s] = on_idx
        st.stim_offsets_by_sweep[s] = off_idx
        st.stim_spans_by_sweep[s] = spans_s
        st.stim_metrics_by_sweep[s] = metrics

        # Debug print
        if metrics:
            pw = metrics.get("pulse_width_s")
            dur = metrics.get("duration_s")
            hz = metrics.get("freq_hz")
            msg = f"[stim] sweep {s}: width={pw:.6f}s, duration={dur:.6f}s"
            if hz:
                msg += f", freq={hz:.3f}Hz"
            print(msg)

    def _detect_stims_all_sweeps(self, thresh: float = 1.0):
        """Detect stimulations on all sweeps (for export/preview)."""
        st = self.state
        if not st.stim_chan or st.stim_chan not in st.sweeps:
            return

        Y = st.sweeps[st.stim_chan]
        n_sweeps = Y.shape[1]
        t = st.t

        for s in range(n_sweeps):
            # Skip if already detected
            if s in st.stim_spans_by_sweep:
                continue

            y = Y[:, s]
            on_idx, off_idx, spans_s, metrics = stimdet.detect_threshold_crossings(y, t, thresh=thresh)
            st.stim_onsets_by_sweep[s] = on_idx
            st.stim_offsets_by_sweep[s] = off_idx
            st.stim_spans_by_sweep[s] = spans_s
            st.stim_metrics_by_sweep[s] = metrics

        print(f"[stim] Detected stims for all {n_sweeps} sweeps")

    # ---------- Filters & redraw ----------
    def update_and_redraw(self, *args):
        st = self.state

        # checkboxes
        st.use_low       = self.LowPass_checkBox.isChecked()
        st.use_high      = self.HighPass_checkBox.isChecked()
        # Mean subtraction is now controlled from Spectral Analysis dialog
        # st.use_mean_sub is set directly in the dialog handlers
        st.use_invert    = self.InvertSignal_checkBox.isChecked()

        # Filter order
        self.filter_order = self.FilterOrderSpin.value()


        # Peaks/breaths no longer valid if filters change
        if hasattr(self.state, "peaks_by_sweep"):
            self.state.peaks_by_sweep.clear()
        if hasattr(self.state, "breath_by_sweep"):
            self.state.breath_by_sweep.clear()

        # Peaks/breaths/y2 no longer valid if filters change
        if hasattr(self.state, "peaks_by_sweep"):
            self.state.peaks_by_sweep.clear()
        if hasattr(self.state, "breath_by_sweep"):
            self.state.breath_by_sweep.clear()
            self.state.y2_values_by_sweep.clear()
            self.plot_host.clear_y2()

        # Clear z-score global statistics cache (filters changed)
        self.zscore_global_mean = None
        self.zscore_global_std = None



        def _val_if_enabled(line, checked: bool, cast=float, default=None):
            """Return a numeric value only if the box is checked and a value exists."""
            if not checked:
                return None
            txt = line.text().strip()
            if not txt:
                return None
            try:
                return cast(txt)
            except ValueError:
                return None

        # only take values if box is checked AND entry is valid
        st.low_hz  = _val_if_enabled(self.LowPassVal, st.use_low, float, None)
        st.high_hz = _val_if_enabled(self.HighPassVal, st.use_high, float, None)
        # Mean subtraction value is now controlled from Spectral Analysis dialog
        # st.mean_win_s is set directly in the dialog handlers

        # If the checkbox is checked but the box is empty/invalid, disable that filter automatically
        if st.use_low and st.low_hz is None:
            st.use_low = False
        if st.use_high and st.high_hz is None:
            st.use_high = False
        # Mean subtraction validation is handled in Spectral Analysis dialog

        # Invalidate processed cache
        st.proc_cache.clear()

        # Debounce redraw
        self._redraw_timer.start()

    def _current_trace(self):
        """Return (t, y_proc) for analyze channel & current sweep, using cached processing."""
        st = self.state
        if not st.analyze_chan or st.analyze_chan not in st.sweeps:
            return None, None

        Y = st.sweeps[st.analyze_chan]
        s = max(0, min(st.sweep_idx, Y.shape[1] - 1))
        key = self._proc_key(st.analyze_chan, s)

        # Fast path: reuse processed data if settings didn't change
        if key in st.proc_cache:
            return st.t, st.proc_cache[key]

        # Compute once, cache, and return
        y = Y[:, s]
        y2 = filters.apply_all_1d(
            y, st.sr_hz,
            st.use_low,  st.low_hz,
            st.use_high, st.high_hz,
            st.use_mean_sub, st.mean_val,
            st.use_invert,
            order=self.filter_order
        )

        # Apply notch filter if configured
        if self.notch_filter_lower is not None and self.notch_filter_upper is not None:
            y2 = self._apply_notch_filter(y2, st.sr_hz, self.notch_filter_lower, self.notch_filter_upper)

        # Apply z-score normalization if enabled (using global statistics)
        if self.use_zscore_normalization:
            # Compute global stats if not cached
            if self.zscore_global_mean is None or self.zscore_global_std is None:
                self.zscore_global_mean, self.zscore_global_std = self._compute_global_zscore_stats()
            y2 = filters.zscore_normalize(y2, self.zscore_global_mean, self.zscore_global_std)

        st.proc_cache[key] = y2
        return st.t, y2

    def redraw_main_plot(self):
        """Delegate to PlotManager."""
        self.plot_manager.redraw_main_plot()


    ##################################################
    ##Region threshold visualization                ##
    ##################################################
    def _on_region_threshold_changed(self, *_):
        """
        Called whenever eupnea or apnea threshold values change.
        Redraws the current sweep to update region overlays.
        """
        # Simply redraw current sweep, which will use the new threshold values
        self.redraw_main_plot()

    def _on_filter_changed(self, *_):
        """
        Called whenever filter settings change.
        Re-enables Apply button since peaks need to be recalculated with new filtering.
        """
        st = self.state

        # Log telemetry: filter settings changed
        telemetry.log_button_click('filter_changed',
                                   use_low=st.use_low,
                                   low_hz=st.low_hz if st.use_low else None,
                                   use_high=st.use_high,
                                   high_hz=st.high_hz if st.use_high else None,
                                   use_mean_sub=st.use_mean_sub,
                                   use_invert=st.use_invert)

        # Only re-enable if we have a threshold value and an analysis channel
        if self.peak_prominence is not None and self.state.analyze_chan:
            self.ApplyPeakFindPushButton.setEnabled(True)

    ##################################################
    ##Peak detection parameters                     ##
    ##################################################

    def _parse_float(self, line_edit):
        txt = line_edit.text().strip()
        if not txt:
            return None
        try:
            return float(txt)
        except ValueError:
            return None

    def _get_processed_for(self, chan: str, sweep_idx: int):
        """Return processed y for (channel, sweep_idx) using the same cache key logic."""
        st = self.state
        Y = st.sweeps[chan]
        s = max(0, min(sweep_idx, Y.shape[1]-1))
        key = (chan, s, st.use_low, st.low_hz, st.use_high, st.high_hz, st.use_mean_sub, st.mean_val, st.use_invert,
               self.notch_filter_lower, self.notch_filter_upper, self.use_zscore_normalization)
        if key in st.proc_cache:
            return st.proc_cache[key]
        y = Y[:, s]
        y2 = filters.apply_all_1d(
            y, st.sr_hz,
            st.use_low,  st.low_hz,
            st.use_high, st.high_hz,
            st.use_mean_sub, st.mean_val,
            st.use_invert
        )

        # Apply notch filter if configured
        if self.notch_filter_lower is not None and self.notch_filter_upper is not None:
            y2 = self._apply_notch_filter(y2, st.sr_hz, self.notch_filter_lower, self.notch_filter_upper)

        # Apply z-score normalization if enabled (using global statistics)
        if self.use_zscore_normalization:
            # Compute global stats if not cached
            if self.zscore_global_mean is None or self.zscore_global_std is None:
                self.zscore_global_mean, self.zscore_global_std = self._compute_global_zscore_stats()
            y2 = filters.zscore_normalize(y2, self.zscore_global_mean, self.zscore_global_std)

        st.proc_cache[key] = y2
        return y2

    def _apply_notch_filter(self, y, sr_hz, lower_freq, upper_freq):
        """Apply a notch (band-stop) filter to remove frequencies between lower_freq and upper_freq."""
        from scipy import signal
        import numpy as np

        print(f"[notch-filter] Applying notch filter: {lower_freq:.2f} - {upper_freq:.2f} Hz (sr={sr_hz} Hz)")

        # Design a butterworth band-stop filter
        nyquist = sr_hz / 2.0
        low = lower_freq / nyquist
        high = upper_freq / nyquist

        # Ensure frequencies are in valid range (0, 1)
        low = np.clip(low, 0.001, 0.999)
        high = np.clip(high, 0.001, 0.999)

        if low >= high:
            print(f"[notch-filter] Invalid frequency range: {lower_freq}-{upper_freq} Hz")
            return y

        try:
            # Design 4th order Butterworth band-stop filter
            sos = signal.butter(4, [low, high], btype='bandstop', output='sos')
            # Apply filter (sos format is more numerically stable)
            y_filtered = signal.sosfiltfilt(sos, y)
            print(f"[notch-filter] Filter applied successfully. Signal range before: [{y.min():.3f}, {y.max():.3f}], after: [{y_filtered.min():.3f}, {y_filtered.max():.3f}]")
            return y_filtered
        except Exception as e:
            print(f"[notch-filter] Error applying filter: {e}")
            return y

    def _open_analysis_options(self, tab=None):
        """
        Open the multi-tabbed Analysis Options dialog.
        Consolidates Peak Detection, GMM Clustering, Outlier Detection, and ML Settings.
        Non-blocking so user can interact with main window while dialog is open.

        Args:
            tab (str, optional): Tab to open ('peak_detection', 'gmm', 'outliers', 'ml')
        """
        from dialogs.analysis_options_dialog import AnalysisOptionsDialog
        from PyQt6.QtWidgets import QApplication
        from PyQt6.QtCore import Qt

        st = self.state

        # Check if dialog already exists (reuse to preserve settings like last training data path)
        if hasattr(self, '_analysis_options_dialog') and self._analysis_options_dialog is not None:
            # Reuse existing dialog and bring to front
            self._analysis_options_dialog.show()
            self._analysis_options_dialog.raise_()
            self._analysis_options_dialog.activateWindow()
        else:
            # Show loading cursor while creating dialog (may take a moment for tab initialization)
            QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)
            QApplication.processEvents()  # Ensure cursor change is visible
            try:
                # Create and show the dialog (non-blocking)
                # Note: Individual tabs handle their own data requirements (e.g., ML Settings works without data)
                self._analysis_options_dialog = AnalysisOptionsDialog(st, parent=self)
                self._analysis_options_dialog.finished.connect(self._on_analysis_options_closed)
                self._analysis_options_dialog.show()
            finally:
                QApplication.restoreOverrideCursor()

        # Switch to requested tab if specified
        if tab and self._analysis_options_dialog is not None:
            self._analysis_options_dialog.set_active_tab(tab)

    def _on_analysis_options_closed(self):
        """Hide dialog when closed but keep it alive for fast re-opening."""
        # Don't destroy the dialog - keep it alive for instant re-opening
        # The dialog will be reused in _open_analysis_options()
        pass

    def _prebuild_analysis_options_dialog(self):
        """Pre-build the Analysis Options dialog in background for instant opening later."""
        # Only pre-build if dialog doesn't already exist
        if hasattr(self, '_analysis_options_dialog') and self._analysis_options_dialog is not None:
            return

        # Only pre-build if we have the required data
        if not hasattr(self, 'all_peak_heights') or self.all_peak_heights is None:
            return

        try:
            from dialogs.analysis_options_dialog import AnalysisOptionsDialog
            print("[Pre-build] Creating Analysis Options dialog in background...")

            # Create dialog but don't show it
            self._analysis_options_dialog = AnalysisOptionsDialog(self.state, parent=self)
            self._analysis_options_dialog.finished.connect(self._on_analysis_options_closed)

            print("[Pre-build] Dialog ready - will open instantly when requested")
        except Exception as e:
            print(f"[Pre-build] Error creating dialog: {e}")
            self._analysis_options_dialog = None

    def _open_prominence_histogram(self):
        """
        Open the multi-tab Analysis Options dialog to the Auto-Threshold tab.
        (Replaces the old single ProminenceThresholdDialog)
        """
        self._open_analysis_options(tab='peak_detection')
        telemetry.log_screen_view('Peak Detection Options Dialog', screen_class='config_dialog')

    def _calculate_local_minimum_threshold_silent(self, peak_heights):
        """
        Calculate valley threshold using exponential + Gaussian mixture model.
        Simplified version for silent auto-detection (no UI feedback).

        Args:
            peak_heights: Array of detected peak heights

        Returns:
            tuple: (valley_threshold, model_params_dict) or (None, None) if fitting fails
            model_params_dict contains: lambda_exp, mu1, sigma1, mu2, sigma2, w_exp, w_g1, w_g2
        """
        try:
            from scipy.optimize import curve_fit

            # Use 99th percentile to exclude outliers
            percentile_95 = np.percentile(peak_heights, 99)
            peaks_for_hist = peak_heights[peak_heights <= percentile_95]

            if len(peaks_for_hist) < 10:
                return (None, None)

            # Create histogram
            hist_range = (peaks_for_hist.min(), percentile_95)
            counts, bin_edges = np.histogram(peaks_for_hist, bins=200, range=hist_range)
            bin_centers = (bin_edges[:-1] + bin_edges[1:]) / 2
            bin_width = bin_edges[1] - bin_edges[0]

            # Convert to density
            density = counts / (len(peaks_for_hist) * bin_width)

            # Try 2-Gaussian model first (for eupnea + sniffing)
            def exp_2gauss_model(x, lambda_exp, mu1, sigma1, mu2, sigma2, w_exp, w_g1):
                exp_comp = lambda_exp * np.exp(-lambda_exp * x)
                gauss1 = (1 / (np.sqrt(2 * np.pi) * sigma1)) * np.exp(-0.5 * ((x - mu1) / sigma1) ** 2)
                gauss2 = (1 / (np.sqrt(2 * np.pi) * sigma2)) * np.exp(-0.5 * ((x - mu2) / sigma2) ** 2)
                w_g2 = max(0, 1 - w_exp - w_g1)
                return w_exp * exp_comp + w_g1 * gauss1 + w_g2 * gauss2

            try:
                p0_2g = [
                    1.0 / np.mean(bin_centers),  # lambda_exp
                    np.percentile(bin_centers, 40),  # mu1 (eupnea)
                    np.std(bin_centers) * 0.3,  # sigma1
                    np.percentile(bin_centers, 70),  # mu2 (sniffing)
                    np.std(bin_centers) * 0.3,  # sigma2
                    0.3,  # w_exp
                    0.4   # w_g1
                ]
                popt, _ = curve_fit(exp_2gauss_model, bin_centers, density, p0=p0_2g, maxfev=5000)
                fitted = exp_2gauss_model(bin_centers, *popt)

                # Check fit quality
                residuals = density - fitted
                ss_res = np.sum(residuals ** 2)
                ss_tot = np.sum((density - np.mean(density)) ** 2)
                r_squared = 1 - (ss_res / ss_tot) if ss_tot > 0 else 0

                if r_squared >= 0.7 and popt[5] >= 0.05 and popt[6] >= 0.05:
                    # Find valley between 0 and first Gaussian peak
                    search_end_idx = np.argmin(np.abs(bin_centers - popt[1]))
                    valley_idx = np.argmin(fitted[:search_end_idx])
                    threshold = float(bin_centers[valley_idx])

                    # Store model parameters for probability metrics
                    model_params = {
                        'lambda_exp': float(popt[0]),
                        'mu1': float(popt[1]),
                        'sigma1': float(popt[2]),
                        'mu2': float(popt[3]),
                        'sigma2': float(popt[4]),
                        'w_exp': float(popt[5]),
                        'w_g1': float(popt[6]),
                        'w_g2': float(max(0, 1 - popt[5] - popt[6]))
                    }
                    return (threshold, model_params)
            except:
                pass

            # Fallback: 1-Gaussian model
            def exp_gauss_model(x, lambda_exp, mu_gauss, sigma_gauss, w_exp):
                exp_component = lambda_exp * np.exp(-lambda_exp * x)
                gauss_component = (1 / (np.sqrt(2 * np.pi) * sigma_gauss)) * np.exp(-0.5 * ((x - mu_gauss) / sigma_gauss) ** 2)
                return w_exp * exp_component + (1 - w_exp) * gauss_component

            p0_1g = [
                1.0 / np.mean(bin_centers),
                np.median(bin_centers),
                np.std(bin_centers),
                0.3
            ]
            popt, _ = curve_fit(exp_gauss_model, bin_centers, density, p0=p0_1g, maxfev=5000)
            fitted = exp_gauss_model(bin_centers, *popt)

            # Check fit quality
            residuals = density - fitted
            ss_res = np.sum(residuals ** 2)
            ss_tot = np.sum((density - np.mean(density)) ** 2)
            r_squared = 1 - (ss_res / ss_tot) if ss_tot > 0 else 0

            if r_squared >= 0.7:
                # Find valley between 0 and Gaussian peak
                search_end_idx = np.argmin(np.abs(bin_centers - popt[1]))
                valley_idx = np.argmin(fitted[:search_end_idx])
                threshold = float(bin_centers[valley_idx])

                # Store model parameters for probability metrics (1-Gaussian fallback)
                model_params = {
                    'lambda_exp': float(popt[0]),
                    'mu1': float(popt[1]),
                    'sigma1': float(popt[2]),
                    'mu2': float(popt[1]),  # Same as mu1 for 1-Gaussian
                    'sigma2': float(popt[2]),  # Same as sigma1
                    'w_exp': float(popt[3]),
                    'w_g1': float(1 - popt[3]),
                    'w_g2': 0.0  # No second Gaussian in fallback
                }
                return (threshold, model_params)

            return (None, None)

        except Exception as e:
            print(f"[Valley Fit] Error: {e}")
            return (None, None)

    def _auto_detect_prominence_silent(self):
        """
        Auto-detect optimal prominence using Otsu's method in background (no dialog).
        Populates prominence field and enables Apply button.
        """
        import time
        from scipy.signal import find_peaks
        import numpy as np

        st = self.state
        if not st.analyze_chan or st.analyze_chan not in st.sweeps:
            return

        try:
            # Concatenate ALL sweeps for representative auto-threshold calculation
            print("[Auto-Detect] Calculating optimal prominence...")
            t_start = time.time()

            all_sweeps_data = []
            n_sweeps = st.sweeps[st.analyze_chan].shape[1]

            for sweep_idx in range(n_sweeps):
                if sweep_idx in st.omitted_sweeps:
                    continue
                y_sweep = self._get_processed_for(st.analyze_chan, sweep_idx)
                all_sweeps_data.append(y_sweep)

            if not all_sweeps_data:
                return

            y_data = np.concatenate(all_sweeps_data)

            # Detect all peaks with minimal prominence AND above baseline (height > 0)
            # height=0 filters out rebound peaks below baseline, giving cleaner 2-population model
            min_dist_samples = int(self.peak_min_dist * st.sr_hz)
            peaks, props = find_peaks(y_data, height=0, prominence=0.001, distance=min_dist_samples)
            peak_heights = y_data[peaks]

            if len(peak_heights) < 10:
                print("[Auto-Detect] Not enough peaks found")
                return

            # Store peak heights for histogram reuse (so we don't recalculate during dragging)
            self.all_peak_heights = peak_heights

            # Otsu's method: auto-calculate optimal HEIGHT threshold
            heights_norm = ((peak_heights - peak_heights.min()) /
                        (peak_heights.max() - peak_heights.min()) * 255).astype(np.uint8)

            hist, bin_edges = np.histogram(heights_norm, bins=256, range=(0, 256))
            bin_centers = (bin_edges[:-1] + bin_edges[1:]) / 2

            weight1 = np.cumsum(hist)
            weight2 = np.cumsum(hist[::-1])[::-1]
            mean1 = np.cumsum(hist * bin_centers) / (weight1 + 1e-10)
            mean2 = (np.cumsum((hist * bin_centers)[::-1]) / (weight2 + 1e-10))[::-1]

            variance = weight1[:-1] * weight2[1:] * (mean1[:-1] - mean2[1:]) ** 2
            optimal_bin = np.argmax(variance)
            optimal_thresh_norm = bin_centers[optimal_bin]

            # Convert back to original scale
            optimal_height = float((optimal_thresh_norm / 255.0 *
                            (peak_heights.max() - peak_heights.min()) +
                            peak_heights.min()))

            # Calculate local minimum threshold (valley between noise and signal)
            # This is more robust than Otsu for breath signals
            local_min_threshold, model_params = self._calculate_local_minimum_threshold_silent(peak_heights)

            # Choose threshold: prefer local minimum if available, fallback to Otsu
            if local_min_threshold is not None:
                chosen_threshold = local_min_threshold
                print(f"[Auto-Detect] Using valley threshold: {chosen_threshold:.4f} (Otsu: {optimal_height:.4f})")

                # Store model parameters for probability metrics
                import core.metrics as metrics
                metrics.set_threshold_model_params(model_params)
                print(f"[Auto-Detect] Stored model parameters for P(noise)/P(breath) calculation")
            else:
                chosen_threshold = optimal_height
                print(f"[Auto-Detect] Using Otsu threshold: {chosen_threshold:.4f} (no valley found)")
                # Clear model parameters if valley fit failed
                import core.metrics as metrics
                metrics.set_threshold_model_params(None)

            # Store and populate spinbox with auto-detected value
            # Use same value for both height and prominence thresholds
            self.peak_prominence = chosen_threshold
            # PeakPromValueSpinBox removed - threshold now shown in Analysis Options dialog
            # self.PeakPromValueSpinBox.setValue(chosen_threshold)

            # Store the height threshold value (will be used in peak detection)
            self.peak_height_threshold = chosen_threshold

            # Draw threshold line on plot
            self.plot_host.update_threshold_line(chosen_threshold)

            # Enable Apply button
            self.ApplyPeakFindPushButton.setEnabled(True)

            t_elapsed = time.time() - t_start
            # Status message already printed above with valley/Otsu choice
            self._log_status_message(f"Auto-detected threshold: {chosen_threshold:.4f}", 3000)

            # Pre-build the Analysis Options dialog in background (deferred)
            # This makes opening the dialog instant since it's already created
            QTimer.singleShot(500, self._prebuild_analysis_options_dialog)

            # Refresh Peak Detection tab if it's the one that needs updating after channel change
            if (hasattr(self, '_channel_change_needs_dialog_refresh') and
                self._channel_change_needs_dialog_refresh and
                getattr(self, '_dialog_tab_to_refresh', -1) == 0):  # Tab 0 = Peak Detection
                if hasattr(self, '_analysis_options_dialog') and self._analysis_options_dialog is not None:
                    try:
                        if self._analysis_options_dialog.isVisible():
                            print("[Auto-Detect] Refreshing Peak Detection tab after auto-detect complete")
                            self._analysis_options_dialog._refresh_peak_detection_tab()
                            self._channel_change_needs_dialog_refresh = False  # Mark as done
                    except RuntimeError:
                        self._analysis_options_dialog = None

        except Exception as e:
            print(f"[Auto-Detect] Error: {e}")
            import traceback
            traceback.print_exc()

    # PeakPromValueSpinBox removed - prominence/threshold now set in Analysis Options dialog
    # def _on_prominence_spinbox_changed(self):
    #     """Update threshold line on plot when spinbox value changes."""
    #     new_value = self.PeakPromValueSpinBox.value()
    #     if new_value > 0:
    #         self.peak_height_threshold = new_value
    #         self.plot_host.update_threshold_line(new_value)
    #
    #         # Synchronize with auto-threshold workflow: fit model for p_noise calculation
    #         self._fit_threshold_model_for_manual_slider()

    def _fit_threshold_model_for_manual_slider(self):
        """
        Fit the exponential + 2 Gaussians model for p_noise calculation.
        This is called when the manual threshold slider changes to synchronize
        with the auto-threshold workflow.
        """
        from scipy.signal import find_peaks
        import core.metrics as metrics

        st = self.state
        if not st.analyze_chan or st.analyze_chan not in st.sweeps:
            return

        try:
            # Check if we already have peak_heights from a previous auto-detect
            if not hasattr(self, 'all_peak_heights') or self.all_peak_heights is None:
                # Need to compute peak_heights by detecting peaks from current data
                all_sweeps_data = []
                n_sweeps = st.sweeps[st.analyze_chan].shape[1]

                for sweep_idx in range(n_sweeps):
                    if sweep_idx in st.omitted_sweeps:
                        continue
                    y_sweep = self._get_processed_for(st.analyze_chan, sweep_idx)
                    if y_sweep is not None:
                        all_sweeps_data.append(y_sweep)

                if not all_sweeps_data:
                    # No data available, clear model params
                    metrics.set_threshold_model_params(None)
                    return

                y_data = np.concatenate(all_sweeps_data)

                # Detect all peaks with minimal prominence AND above baseline (height > 0)
                min_dist_samples = int(self.peak_min_dist * st.sr_hz)
                peaks, props = find_peaks(y_data, height=0, prominence=0.001, distance=min_dist_samples)
                peak_heights = y_data[peaks]

                if len(peak_heights) < 10:
                    # Not enough peaks, clear model params
                    metrics.set_threshold_model_params(None)
                    return

                # Store for reuse
                self.all_peak_heights = peak_heights

            # Fit the model using the peak heights
            local_min_threshold, model_params = self._calculate_local_minimum_threshold_silent(self.all_peak_heights)

            # Store model parameters for p_noise calculation
            if model_params is not None:
                metrics.set_threshold_model_params(model_params)
                print(f"[Manual Threshold] Fitted model for p_noise calculation")
            else:
                # Model fit failed, clear params
                metrics.set_threshold_model_params(None)
                print(f"[Manual Threshold] Model fit failed, p_noise will be unavailable")

        except Exception as e:
            print(f"[Manual Threshold] Error fitting model: {e}")
            import traceback
            traceback.print_exc()
            # Clear model params on error
            metrics.set_threshold_model_params(None)

    def _precompute_remaining_classifiers_async(self):
        """
        Pre-compute remaining ML classifiers in the background after UI is responsive.
        This enables instant classifier switching without slowing down initial display.
        """
        from PyQt6.QtCore import QTimer
        import core.ml_prediction as ml_prediction
        import time

        st = self.state

        # Determine which algorithms still need to be computed
        already_computed = set()
        need_to_compute = set()

        # Check what was already run
        if st.active_classifier in ['xgboost', 'rf', 'mlp']:
            already_computed.add(st.active_classifier)
        if st.active_eupnea_sniff_classifier in ['xgboost', 'rf', 'mlp']:
            already_computed.add(st.active_eupnea_sniff_classifier)
        if st.active_sigh_classifier in ['xgboost', 'rf', 'mlp']:
            already_computed.add(st.active_sigh_classifier)

        # Determine what still needs to be computed
        all_algorithms = {'xgboost', 'rf', 'mlp'}
        need_to_compute = all_algorithms - already_computed

        if not need_to_compute:
            print("[Background] No additional classifiers to pre-compute")
            return

        print(f"[Background] Will pre-compute {sorted(need_to_compute)} classifiers in background...")

        def compute_remaining():
            """Background worker function."""
            t_start = time.time()
            total_computed = 0

            try:
                for s in st.all_peaks_by_sweep.keys():
                    all_peaks = st.all_peaks_by_sweep[s]

                    # Get peak metrics for this sweep
                    if s not in st.peak_metrics_by_sweep:
                        continue
                    peak_metrics = st.peak_metrics_by_sweep[s]

                    # Run predictions for remaining algorithms
                    for algorithm in sorted(need_to_compute):
                        try:
                            predictions = ml_prediction.predict_with_cascade(
                                peak_metrics=peak_metrics,
                                models=st.loaded_ml_models,
                                algorithm=algorithm
                            )

                            # Store predictions as read-only (for classifier switching)
                            all_peaks[f'labels_{algorithm}_ro'] = predictions['final_labels']

                            # Store eupnea/sniff predictions (Model 3 output)
                            if 'eupnea_sniff_class' in predictions:
                                all_peaks[f'eupnea_sniff_{algorithm}_ro'] = predictions['eupnea_sniff_class']

                            # Store sigh predictions (Model 2 output)
                            if 'sigh_class' in predictions:
                                all_peaks[f'sigh_{algorithm}_ro'] = predictions['sigh_class']

                            total_computed += 1

                        except KeyError:
                            if s == 0:
                                print(f"[Background] Model {algorithm} not found, skipping")
                        except Exception as e:
                            print(f"[Background] Warning: {algorithm} prediction failed: {e}")

                t_elapsed = time.time() - t_start
                print(f"[Background] Pre-computed {len(need_to_compute)} classifiers in {t_elapsed:.1f}s")
                print(f"[Background] Classifier switching is now instant!")

            except Exception as e:
                print(f"[Background] Error during pre-computation: {e}")
                import traceback
                traceback.print_exc()

        # Schedule background computation with 100ms delay (let UI settle first)
        QTimer.singleShot(100, compute_remaining)

    def _detect_single_sweep_core(self, sweep_idx: int, y_proc: np.ndarray,
                                   thresh: float, min_dist_samples: int,
                                   direction: str = "up") -> dict:
        """
        Core peak detection logic for a single sweep (parallelizable).

        This method extracts the computationally expensive parts of peak detection
        that can be safely run in parallel. It does NOT modify shared state.

        Args:
            sweep_idx: Sweep index being processed
            y_proc: Pre-processed signal array
            thresh: Height threshold for labeling
            min_dist_samples: Minimum distance between peaks in samples
            direction: Peak direction ("up" or "down")

        Returns:
            Dict with detection results (to be merged into state later)
        """
        st = self.state

        # Step 1: Detect ALL peaks (no threshold filtering)
        all_peak_indices = peakdet.detect_peaks(
            y=y_proc, sr_hz=st.sr_hz,
            thresh=None,
            prominence=None,
            min_dist_samples=min_dist_samples,
            direction=direction,
            return_all=True
        )

        # Step 2: Compute breath features for ALL peaks
        all_breaths = peakdet.compute_breath_events(y_proc, all_peak_indices, sr_hz=st.sr_hz, exclude_sec=0.030)

        # Step 3: Label peaks using threshold
        all_peaks_data = peakdet.label_peaks_by_threshold(
            y=y_proc,
            peak_indices=all_peak_indices,
            thresh=thresh,
            direction=direction
        )
        all_peaks_data['labels_threshold_ro'] = all_peaks_data['labels'].copy()

        # Initialize ML prediction arrays as None (will be filled later if models loaded)
        all_peaks_data['labels_xgboost_ro'] = None
        all_peaks_data['labels_rf_ro'] = None
        all_peaks_data['labels_mlp_ro'] = None

        # Step 4: Compute p_noise if threshold model available
        try:
            import core.metrics as metrics_mod
            on = all_breaths.get('onsets', np.array([]))
            off = all_breaths.get('offsets', np.array([]))
            exm = all_breaths.get('expmins', np.array([]))
            exo = all_breaths.get('expoffs', np.array([]))
            p_noise_all = metrics_mod.compute_p_noise(st.t, y_proc, st.sr_hz, all_peak_indices, on, off, exm, exo)
            p_breath_all = 1.0 - p_noise_all if p_noise_all is not None else None
        except Exception:
            p_noise_all = None
            p_breath_all = None

        # Step 5: Compute peak candidate metrics
        peak_metrics = peakdet.compute_peak_candidate_metrics(
            y=y_proc,
            all_peak_indices=all_peak_indices,
            breath_events=all_breaths,
            sr_hz=st.sr_hz,
            p_noise=p_noise_all,
            p_breath=p_breath_all
        )

        # Step 6: Extract labeled peaks for display
        labeled_mask = all_peaks_data['labels'] == 1
        labeled_indices = all_peak_indices[labeled_mask]

        # Step 7: Compute breath events for labeled peaks
        labeled_breaths = peakdet.compute_breath_events(y_proc, labeled_indices, sr_hz=st.sr_hz, exclude_sec=0.030)

        # Step 8: Recalculate current metrics using labeled peaks as neighbors
        try:
            import core.metrics as metrics_mod
            p_noise_labeled = metrics_mod.compute_p_noise(
                st.t, y_proc, st.sr_hz, labeled_indices,
                labeled_breaths.get('onsets', np.array([])),
                labeled_breaths.get('offsets', np.array([])),
                labeled_breaths.get('expmins', np.array([])),
                labeled_breaths.get('expoffs', np.array([]))
            )
            p_breath_labeled = 1.0 - p_noise_labeled if p_noise_labeled is not None else None

            current_metrics = peakdet.compute_peak_candidate_metrics(
                y=y_proc,
                all_peak_indices=labeled_indices,
                breath_events=labeled_breaths,
                sr_hz=st.sr_hz,
                p_noise=p_noise_labeled,
                p_breath=p_breath_labeled
            )
        except Exception:
            current_metrics = peak_metrics

        return {
            'sweep_idx': sweep_idx,
            'all_peak_indices': all_peak_indices,
            'all_breaths': all_breaths,
            'all_peaks_data': all_peaks_data,
            'peak_metrics': peak_metrics,
            'current_metrics': current_metrics,
            'labeled_indices': labeled_indices,
            'labeled_breaths': labeled_breaths,
            'p_noise_all': p_noise_all,
            'p_breath_all': p_breath_all
        }

    def _apply_peak_detection(self):
        """
        Run peak detection on the ANALYZE channel for ALL sweeps,
        store indices per sweep, and redraw current sweep with peaks + breath markers.

        Supports parallel processing for large files (10+ sweeps) with automatic fallback
        to sequential processing if parallel execution fails.
        """
        import time
        t_start = time.time()

        st = self.state
        if not st.channel_names or not st.analyze_chan or st.analyze_chan not in st.sweeps:
            return

        self._log_status_message("Detecting peaks and breath features...")

        # Get prominence from stored value (set during auto-detect or in Analysis Options dialog)
        prom = getattr(self, 'peak_prominence', None)
        if prom is None or prom <= 0:
            self._show_warning("Invalid Prominence",
                             "No prominence threshold set. Please select an analysis channel first to auto-detect the threshold.")
            return

        # Use stored height threshold (set during auto-detect, same as prominence)
        thresh = getattr(self, 'peak_height_threshold', None)
        min_d = self.peak_min_dist
        direction = "up"  # Always detect peaks above threshold for breathing signals

        min_dist_samples = None
        if min_d is not None and min_d > 0:
            min_dist_samples = max(1, int(round(min_d * st.sr_hz)))

        # Detect on ALL sweeps for the analyze channel
        any_chan = next(iter(st.sweeps.values()))
        n_sweeps = any_chan.shape[1]
        st.peaks_by_sweep.clear()
        st.breath_by_sweep.clear()
        st.all_peaks_by_sweep.clear()  # ML training data: ALL peaks with labels
        st.all_breaths_by_sweep.clear()  # ML training data: breath events for ALL peaks
        st.peak_metrics_by_sweep.clear()  # ML training data: comprehensive peak metrics
        # st.sigh_by_sweep.clear()

        # === PARALLEL CORE DETECTION ===
        # For files with many sweeps, parallelize the core detection
        # ML predictions and state updates happen sequentially afterwards
        use_parallel = getattr(self, 'use_parallel_detection', True) and n_sweeps >= 10
        core_results = {}

        if use_parallel:
            try:
                from concurrent.futures import ThreadPoolExecutor, as_completed
                import os

                max_workers = min(os.cpu_count() or 4, 8)
                print(f"[peak-detection] Using parallel detection with {max_workers} workers for {n_sweeps} sweeps...")

                # Pre-compute all processed signals (uses cache, must be done in main thread)
                processed_signals = {}
                for s in range(n_sweeps):
                    processed_signals[s] = self._get_processed_for(st.analyze_chan, s)

                # Worker function for parallel execution
                def detect_sweep(sweep_idx):
                    y_proc = processed_signals[sweep_idx]
                    return self._detect_single_sweep_core(
                        sweep_idx, y_proc, thresh, min_dist_samples, direction
                    )

                # Execute in parallel
                with ThreadPoolExecutor(max_workers=max_workers) as executor:
                    futures = {executor.submit(detect_sweep, s): s for s in range(n_sweeps)}
                    completed = 0
                    for future in as_completed(futures):
                        sweep_idx = futures[future]
                        result = future.result()
                        core_results[sweep_idx] = result
                        completed += 1
                        if completed % 10 == 0 or completed == n_sweeps:
                            print(f"[peak-detection] Completed {completed}/{n_sweeps} sweeps...")

                print(f"[peak-detection] Parallel core detection complete!")

            except Exception as e:
                print(f"[peak-detection] Parallel detection failed: {e}")
                print(f"[peak-detection] Falling back to sequential detection...")
                import traceback
                traceback.print_exc()
                core_results = {}  # Clear partial results, fall through to sequential

        # === SEQUENTIAL PROCESSING (fallback or for small files) ===
        for s in range(n_sweeps):
            # Use parallel results if available, otherwise compute sequentially
            if s in core_results:
                result = core_results[s]
                all_peak_indices = result['all_peak_indices']
                all_breaths = result['all_breaths']
                all_peaks_data = result['all_peaks_data']
                peak_metrics = result['peak_metrics']
                labeled_indices = result['labeled_indices']
                labeled_breaths = result['labeled_breaths']
                y_proc = self._get_processed_for(st.analyze_chan, s)

                # Store results in state
                st.all_breaths_by_sweep[s] = all_breaths
                st.all_peaks_by_sweep[s] = all_peaks_data
                st.peak_metrics_by_sweep[s] = peak_metrics
                st.current_peak_metrics_by_sweep[s] = result['current_metrics']

            else:
                # Sequential detection for this sweep (fallback or parallel not used)
                y_proc = self._get_processed_for(st.analyze_chan, s)

                # Step 1: Detect ALL peaks (no threshold filtering)
                # Note: User found that thresh=0 + min_distance works best (no prominence)
                all_peak_indices = peakdet.detect_peaks(
                    y=y_proc, sr_hz=st.sr_hz,
                    thresh=None,  # Don't filter by threshold yet
                    prominence=None,  # Don't use prominence for initial detection
                    min_dist_samples=min_dist_samples,
                    direction=direction,
                    return_all=True  # Return ALL detected peaks
                )

                # Step 2: Compute breath features for ALL peaks (including noise)
                # This is needed for ML training - noise peaks need features too
                all_breaths = peakdet.compute_breath_events(y_proc, all_peak_indices, sr_hz=st.sr_hz, exclude_sec=0.030)
                st.all_breaths_by_sweep[s] = all_breaths  # Store for ML metric computation

                # Step 3: Label peaks using auto-detected threshold
                all_peaks_data = peakdet.label_peaks_by_threshold(
                    y=y_proc,
                    peak_indices=all_peak_indices,
                    thresh=thresh,
                    direction=direction
                )
                # Store threshold predictions as read-only (for classifier switching)
                all_peaks_data['labels_threshold_ro'] = all_peaks_data['labels'].copy()

                # Step 3.1: Initialize ML prediction arrays as None (will be filled after metrics)
                if st.loaded_ml_models:
                    # We need peak_metrics first, so we'll come back to this after metrics computation
                    # For now, initialize ML label arrays as None
                    all_peaks_data['labels_xgboost_ro'] = None
                    all_peaks_data['labels_rf_ro'] = None
                    all_peaks_data['labels_mlp_ro'] = None

                # NOTE: 'labels' remains as the user-editable array (backward compatible)
                # It will be initialized from the active classifier after ML predictions run

                st.all_peaks_by_sweep[s] = all_peaks_data

                # Step 3.5: Compute comprehensive metrics for ML (merge detection, noise classification)
                # Get P(noise) and P(breath) if available from metrics module
                try:
                    import core.metrics as metrics_mod
                    # Extract breath events for p_noise computation
                    on = all_breaths.get('onsets', np.array([]))
                    off = all_breaths.get('offsets', np.array([]))
                    exm = all_breaths.get('expmins', np.array([]))
                    exo = all_breaths.get('expoffs', np.array([]))

                    # Check if threshold model is available
                    if s == 0:
                        print(f"[peak-detection] Threshold model params: {metrics_mod._threshold_model_params}")

                    p_noise_all = metrics_mod.compute_p_noise(st.t, y_proc, st.sr_hz, all_peak_indices, on, off, exm, exo)
                    p_breath_all = 1.0 - p_noise_all if p_noise_all is not None else None
                    if s == 0:
                        print(f"[peak-detection] Computed p_noise for {len(all_peak_indices)} peaks, sample values: {p_noise_all[:5] if p_noise_all is not None and len(p_noise_all) > 0 else 'None'}")
                except Exception as e:
                    print(f"[peak-detection] Warning: Could not compute P(noise): {e}")
                    import traceback
                    traceback.print_exc()
                    p_noise_all = None
                    p_breath_all = None

                peak_metrics = peakdet.compute_peak_candidate_metrics(
                    y=y_proc,
                    all_peak_indices=all_peak_indices,
                    breath_events=all_breaths,
                    sr_hz=st.sr_hz,
                    p_noise=p_noise_all,
                    p_breath=p_breath_all
                )
                st.peak_metrics_by_sweep[s] = peak_metrics  # Original metrics (never modified, for ML)
                st.current_peak_metrics_by_sweep[s] = peak_metrics  # Current metrics (updated after edits, for Y2 plotting)

            # === From here on, use results from either parallel or sequential detection ===
            # Get the data for this sweep (either from parallel results or just computed)
            all_peaks_data = st.all_peaks_by_sweep[s]
            all_breaths = st.all_breaths_by_sweep[s]
            peak_metrics = st.peak_metrics_by_sweep[s]
            all_peak_indices = all_peaks_data['indices']
            y_proc = self._get_processed_for(st.analyze_chan, s)

            # Step 3.7: Run ML predictions if models are loaded (now that we have peak_metrics)
            if st.loaded_ml_models and peak_metrics:
                import core.ml_prediction as ml_prediction

                # Determine which algorithms to run based on active classifiers
                # This avoids running all 9 models (3 algorithms × 3 model types) every time
                algorithms_to_run = set()

                # Add algorithm for breath detection (Model 1)
                if st.active_classifier in ['xgboost', 'rf', 'mlp']:
                    algorithms_to_run.add(st.active_classifier)

                # Add algorithm for eupnea/sniff (Model 3)
                if st.active_eupnea_sniff_classifier in ['xgboost', 'rf', 'mlp']:
                    algorithms_to_run.add(st.active_eupnea_sniff_classifier)

                # Add algorithm for sigh (Model 2)
                if st.active_sigh_classifier in ['xgboost', 'rf', 'mlp']:
                    algorithms_to_run.add(st.active_sigh_classifier)

                # If user wants to pre-compute all for instant switching, run all
                # TODO: Add checkbox in UI to toggle this behavior
                run_all_algorithms = False  # Set to True to restore old behavior

                if run_all_algorithms:
                    algorithms_to_run = {'xgboost', 'rf', 'mlp'}

                if s == 0:
                    print(f"[ML Prediction] Running algorithms: {sorted(algorithms_to_run)}")

                # Run predictions for selected algorithms only
                for algorithm in algorithms_to_run:
                    try:
                        # Enable debug output on first sweep to diagnose issues
                        predictions = ml_prediction.predict_with_cascade(
                            peak_metrics=peak_metrics,
                            models=st.loaded_ml_models,
                            algorithm=algorithm,
                            debug=(s == 0)  # Debug on first sweep only
                        )
                        # Store predictions as read-only (for classifier switching)
                        all_peaks_data[f'labels_{algorithm}_ro'] = predictions['final_labels']

                        # Store eupnea/sniff predictions (Model 3 output)
                        if 'eupnea_sniff_class' in predictions:
                            all_peaks_data[f'eupnea_sniff_{algorithm}_ro'] = predictions['eupnea_sniff_class']

                            if s == 0:
                                n_eupnea = np.sum(predictions['eupnea_sniff_class'] == 0)
                                n_sniff = np.sum(predictions['eupnea_sniff_class'] == 1)
                                n_unclass = np.sum(predictions['eupnea_sniff_class'] == -1)
                                print(f"[ML-{algorithm}] Eupnea: {n_eupnea}, Sniffing: {n_sniff}, Unclassified: {n_unclass}")
                        else:
                            all_peaks_data[f'eupnea_sniff_{algorithm}_ro'] = None

                        # Store sigh predictions (Model 2 output)
                        if 'sigh_class' in predictions:
                            all_peaks_data[f'sigh_{algorithm}_ro'] = predictions['sigh_class']

                            if s == 0:
                                n_normal = np.sum(predictions['sigh_class'] == 0)
                                n_sigh = np.sum(predictions['sigh_class'] == 1)
                                n_unclass_sigh = np.sum(predictions['sigh_class'] == -1)
                                print(f"[ML-{algorithm}] Normal: {n_normal}, Sigh: {n_sigh}, Unclassified: {n_unclass_sigh}")
                        else:
                            all_peaks_data[f'sigh_{algorithm}_ro'] = None

                        # Debug: Print prediction summary on first sweep
                        if s == 0:
                            n_breaths = np.sum(predictions['final_labels'] == 1)
                            n_noise = np.sum(predictions['final_labels'] == 0)
                            print(f"[ML-{algorithm}] Sweep {s}: {n_breaths} breaths, {n_noise} noise (total {len(predictions['final_labels'])} peaks)")
                    except KeyError:
                        # Models for this algorithm not loaded
                        if s == 0:
                            print(f"[ML-{algorithm}] Models not found, skipping")
                        all_peaks_data[f'labels_{algorithm}_ro'] = None
                    except Exception as e:
                        print(f"[ML-{algorithm}] Warning: Prediction failed: {e}")
                        all_peaks_data[f'labels_{algorithm}_ro'] = None

            # Debug: Show sample metrics for first sweep
            if s == 0 and len(peak_metrics) > 0:
                print(f"[peak-metrics] Computed {len(peak_metrics)} peak metrics for sweep {s}")
                # Show first potential merge candidate (small normalized gap + shallow trough)
                merge_candidates = [m for m in peak_metrics
                                   if m.get('gap_to_next_norm', 1.0) is not None
                                   and m.get('gap_to_next_norm') < 0.3
                                   and m.get('trough_ratio_next', 1.0) is not None
                                   and m.get('trough_ratio_next') < 0.15]
                if merge_candidates:
                    mc = merge_candidates[0]
                    print(f"[peak-metrics] Example merge candidate at peak {mc['peak_idx']}:")
                    print(f"  gap_to_next_norm={mc['gap_to_next_norm']:.2f}, trough_ratio_next={mc['trough_ratio_next']:.2f}")
                    print(f"  onset_above_zero={mc['onset_above_zero']}, prom_asymmetry={mc.get('prom_asymmetry', 'N/A')}")

            # Step 4: Initialize user-editable 'labels' array from active classifier
            # This is the array that gets displayed AND edited (backward compatible!)
            active_labels_key_ro = f'labels_{st.active_classifier}_ro'

            # Debug: Show what's available
            if s == 0:
                available_labels = [k for k in all_peaks_data.keys() if k.startswith('labels_') and k.endswith('_ro')]
                print(f"[peak-detection] Available label arrays: {available_labels}")
                for lbl_key in available_labels:
                    lbl_val = all_peaks_data.get(lbl_key)
                    if lbl_val is not None:
                        n_breaths = np.sum(lbl_val == 1)
                        print(f"[peak-detection]   {lbl_key}: {n_breaths} breaths out of {len(lbl_val)} peaks")
                    else:
                        print(f"[peak-detection]   {lbl_key}: None")

            if active_labels_key_ro in all_peaks_data and all_peaks_data[active_labels_key_ro] is not None:
                # Copy from selected classifier's read-only predictions
                all_peaks_data['labels'] = all_peaks_data[active_labels_key_ro].copy()
                if s == 0:
                    n_breaths = np.sum(all_peaks_data['labels'] == 1)
                    print(f"[peak-detection] Initialized 'labels' from {active_labels_key_ro} ({n_breaths} breaths)")
            else:
                # Fallback to threshold if active classifier not available
                # (labels already contains threshold predictions from step 3)
                if s == 0:
                    n_breaths = np.sum(all_peaks_data['labels'] == 1)
                    print(f"[peak-detection] FALLBACK: Using threshold labels ({n_breaths} breaths) - {active_labels_key_ro} was None or missing")

            # Step 4b: Initialize user-editable 'breath_type_class' array from active eupnea/sniff classifier
            if st.active_eupnea_sniff_classifier == 'gmm':
                # GMM will be computed on-demand later, initialize as None for now
                all_peaks_data['breath_type_class'] = None
                if s == 0:
                    print(f"[peak-detection] Initialized 'breath_type_class' as None (GMM will run on-demand)")
            else:
                # Copy from selected ML classifier's read-only predictions
                active_eup_sniff_key_ro = f'eupnea_sniff_{st.active_eupnea_sniff_classifier}_ro'
                if active_eup_sniff_key_ro in all_peaks_data and all_peaks_data[active_eup_sniff_key_ro] is not None:
                    all_peaks_data['breath_type_class'] = all_peaks_data[active_eup_sniff_key_ro].copy()
                    if s == 0:
                        print(f"[peak-detection] Initialized 'breath_type_class' from {active_eup_sniff_key_ro}")
                else:
                    # Fallback to None
                    all_peaks_data['breath_type_class'] = None
                    if s == 0:
                        print(f"[peak-detection] 'breath_type_class' initialized as None (classifier {st.active_eupnea_sniff_classifier} not available)")

            # Step 4c: Initialize user-editable 'sigh_class' array from active sigh classifier
            # First, create sigh_manual_ro from existing sigh_by_sweep (preserves manual annotations)
            n_peaks = len(all_peaks_data['indices'])
            sigh_manual = np.zeros(n_peaks, dtype=np.int8)
            if 'labels' in all_peaks_data:
                sigh_manual[all_peaks_data['labels'] == 0] = -1
            if s in st.sigh_by_sweep:
                for sigh_idx in st.sigh_by_sweep[s]:
                    peak_mask = all_peaks_data['indices'] == sigh_idx
                    if np.any(peak_mask):
                        sigh_manual[peak_mask] = 1
            all_peaks_data['sigh_manual_ro'] = sigh_manual.copy()

            if st.active_sigh_classifier == 'manual':
                # Manual mode - use the manual annotations
                all_peaks_data['sigh_class'] = sigh_manual.copy()
                if s == 0:
                    n_sighs = np.sum(all_peaks_data['sigh_class'] == 1)
                    print(f"[peak-detection] Initialized 'sigh_class' from manual annotations ({n_sighs} sighs)")
            else:
                # Copy from selected ML classifier's read-only predictions
                active_sigh_key_ro = f'sigh_{st.active_sigh_classifier}_ro'
                if active_sigh_key_ro in all_peaks_data and all_peaks_data[active_sigh_key_ro] is not None:
                    all_peaks_data['sigh_class'] = all_peaks_data[active_sigh_key_ro].copy()
                    if s == 0:
                        n_sighs = np.sum(all_peaks_data['sigh_class'] == 1)
                        print(f"[peak-detection] Initialized 'sigh_class' from {active_sigh_key_ro} ({n_sighs} sighs)")
                else:
                    # Fallback to zeros
                    n_peaks = len(all_peaks_data['indices'])
                    all_peaks_data['sigh_class'] = np.zeros(n_peaks, dtype=np.int8)
                    if 'labels' in all_peaks_data:
                        all_peaks_data['sigh_class'][all_peaks_data['labels'] == 0] = -1
                    if s == 0:
                        print(f"[peak-detection] 'sigh_class' initialized as zeros (classifier {st.active_sigh_classifier} not available)")

            # Extract labeled peaks for display
            labeled_mask = all_peaks_data['labels'] == 1
            labeled_indices = all_peaks_data['indices'][labeled_mask]
            st.peaks_by_sweep[s] = labeled_indices

            # Recompute breath events for only labeled peaks (for display)
            # This is simpler than trying to filter the all_breaths dict
            breaths = peakdet.compute_breath_events(y_proc, labeled_indices, sr_hz=st.sr_hz, exclude_sec=0.030)
            st.breath_by_sweep[s] = breaths

            # Recalculate current_peak_metrics_by_sweep using ONLY labeled peaks as neighbors
            # This ensures neighbor features (next_peak_*, prev_peak_*, etc.) are calculated
            # based on the filtered peak list, not the original all_peaks list
            try:
                import core.metrics as metrics_mod

                # Recompute p_noise for labeled peaks only
                p_noise_labeled = metrics_mod.compute_p_noise(st.t, y_proc, st.sr_hz, labeled_indices,
                                                               breaths.get('onsets', np.array([])),
                                                               breaths.get('offsets', np.array([])),
                                                               breaths.get('expmins', np.array([])),
                                                               breaths.get('expoffs', np.array([])))
                p_breath_labeled = 1.0 - p_noise_labeled if p_noise_labeled is not None else None

                # Recalculate metrics with labeled peaks only
                current_metrics = peakdet.compute_peak_candidate_metrics(
                    y=y_proc,
                    all_peak_indices=labeled_indices,  # Use filtered peaks as neighbors
                    breath_events=breaths,
                    sr_hz=st.sr_hz,
                    p_noise=p_noise_labeled,
                    p_breath=p_breath_labeled
                )
                st.current_peak_metrics_by_sweep[s] = current_metrics

                if s == 0:
                    print(f"[peak-metrics] Recalculated {len(current_metrics)} current metrics using {len(labeled_indices)} labeled peaks as neighbors")

            except Exception as e:
                print(f"[peak-metrics] Warning: Could not recalculate current metrics: {e}")
                import traceback
                traceback.print_exc()
                # Keep original metrics as fallback
                st.current_peak_metrics_by_sweep[s] = st.peak_metrics_by_sweep[s]

            # Debug: Show peak detection stats
            n_all = len(all_peak_indices)
            n_labeled = len(labeled_indices)
            n_noise = n_all - n_labeled
            if s == 0:  # Only print for first sweep to avoid spam
                print(f"[peak-detection] Sweep {s}: {n_all} total peaks ({n_labeled} breaths, {n_noise} noise)")

        # Summary statistics for ML training data
        total_all_peaks = sum(len(data['indices']) for data in st.all_peaks_by_sweep.values())
        total_labeled_breaths = sum(len(pks) for pks in st.peaks_by_sweep.values())
        total_noise_peaks = total_all_peaks - total_labeled_breaths
        print(f"[peak-detection] ML training data: {total_all_peaks} total peaks ({total_labeled_breaths} breaths, {total_noise_peaks} noise)")

        # Compute normalization statistics for relative metrics (Group B)
        print("[peak-detection] Computing normalization statistics for relative metrics...")
        self._compute_and_store_normalization_stats()

        # If a Y2 metric is selected, recompute it now that peaks/breaths changed
        if getattr(self.state, "y2_metric_key", None):
            self._compute_y2_all_sweeps()
            self.plot_host.clear_y2()

        # ALWAYS run GMM clustering in background (so users can toggle between classifiers)
        print("[peak-detection] Running automatic GMM clustering (for classifier toggling)...")
        self._run_automatic_gmm_clustering()
        self.eupnea_sniffing_out_of_date = False

        # Build eupnea/sniff regions based on active classifier (BEFORE first redraw)
        if st.active_eupnea_sniff_classifier == 'gmm':
            # GMM already ran above, regions already built
            print("[peak-detection] Using GMM classifier for display")
        else:
            # Use ML classifier predictions (already initialized in breath_type_class array)
            print(f"[peak-detection] Building eupnea/sniff regions from {st.active_eupnea_sniff_classifier} predictions...")
            import core.gmm_clustering as gmm_clustering
            try:
                gmm_clustering.build_eupnea_sniffing_regions(st, verbose=False)
                print(f"[peak-detection] Eupnea/sniff regions built from ML predictions")
            except Exception as e:
                print(f"[peak-detection] Warning: Could not build eupnea/sniff regions: {e}")

        # Sync sigh_by_sweep from sigh_class for display (BEFORE first redraw)
        print("[peak-detection] Syncing sigh markers from sigh_class array...")
        total_sighs = 0
        for s in st.all_peaks_by_sweep.keys():
            all_peaks = st.all_peaks_by_sweep[s]
            if 'sigh_class' in all_peaks and all_peaks['sigh_class'] is not None:
                sigh_mask = all_peaks['sigh_class'] == 1
                st.sigh_by_sweep[s] = all_peaks['indices'][sigh_mask]  # numpy array, not set!
                n_sighs = np.sum(sigh_mask)
                total_sighs += n_sighs
                if s == 0:
                    print(f"[peak-detection] Sweep {s}: {n_sighs} sighs synced to display")

        # Verify and sync: Ensure peaks_by_sweep matches active classifier
        # This fixes an issue where XGBoost results weren't showing on first detection
        active_key = f'labels_{st.active_classifier}_ro'
        first_sweep = next(iter(st.all_peaks_by_sweep.keys()), None)
        if first_sweep is not None:
            first_all_peaks = st.all_peaks_by_sweep[first_sweep]
            if active_key in first_all_peaks and first_all_peaks[active_key] is not None:
                # Active classifier predictions are available - make sure they're being used
                current_labels = first_all_peaks.get('labels')
                active_labels = first_all_peaks[active_key]
                if current_labels is not None and not np.array_equal(current_labels, active_labels):
                    print(f"[peak-detection] SYNC FIX: 'labels' doesn't match {active_key}, updating display...")
                    self._update_displayed_peaks_from_classifier()

        # SINGLE REDRAW: Show peaks, eupnea/sniff regions, and sigh markers all at once
        print(f"[peak-detection] Redrawing plot (peaks + eupnea/sniff + {total_sighs} sighs)...")
        self.redraw_main_plot()

        # Show completion message with elapsed time
        t_elapsed = time.time() - t_start

        # Log telemetry: peak detection with results
        total_peaks = sum(len(pks) for pks in st.peaks_by_sweep.values())
        total_breaths = sum(len(b.get('onsets', [])) for b in st.breath_by_sweep.values())

        telemetry.log_peak_detection(
            method='manual_threshold' if thresh else 'auto_threshold',
            num_peaks=total_peaks,
            threshold=thresh if thresh else prom,
            prominence=prom,
            min_distance_samples=min_dist_samples if min_dist_samples else 0,
            num_sweeps=n_sweeps
        )

        telemetry.log_timing('peak_detection', t_elapsed,
                            num_peaks=total_peaks,
                            num_breaths=total_breaths,
                            num_sweeps=n_sweeps)

        # Log warning if no peaks detected
        if total_peaks == 0:
            telemetry.log_warning('No peaks detected',
                                 threshold=thresh if thresh else prom,
                                 prominence=prom)

        self._log_status_message(f"Peak detection complete ({t_elapsed:.1f}s)", 3000)

        # Refresh GMM tab if it was marked for refresh after channel change
        # (Peak Detection tab was already refreshed after auto-detect)
        if hasattr(self, '_channel_change_needs_dialog_refresh') and self._channel_change_needs_dialog_refresh:
            self._channel_change_needs_dialog_refresh = False
            if hasattr(self, '_analysis_options_dialog') and self._analysis_options_dialog is not None:
                try:
                    if self._analysis_options_dialog.isVisible():
                        tab_index = getattr(self, '_dialog_tab_to_refresh', 0)
                        if tab_index == 1:  # Eup/Sniff Classification (GMM)
                            print("[Peak Detection] Refreshing GMM tab after peak detection complete")
                            self._analysis_options_dialog._refresh_breath_classification_tab()
                        # Tab 0 (Peak Detection) was already refreshed after auto-detect, skip it here
                except RuntimeError:
                    self._analysis_options_dialog = None

        # BACKGROUND PRE-COMPUTATION: Compute remaining classifiers asynchronously
        # This happens AFTER the UI is responsive, so user doesn't notice
        self._precompute_remaining_classifiers_async()

        # Disable Apply button after successful peak detection
        # Will be re-enabled if channel, filter, or file changes
        self.ApplyPeakFindPushButton.setEnabled(False)

        # Refresh Analysis Options dialog GMM tab if open (peaks are now detected)
        if hasattr(self, '_analysis_options_dialog') and self._analysis_options_dialog is not None and self._analysis_options_dialog.isVisible():
            self._analysis_options_dialog._refresh_breath_classification_tab()

    def _compute_and_store_normalization_stats(self):
        """
        Compute global normalization statistics for relative metrics.

        This computes mean and std for key metrics across ALL detected peaks
        in all sweeps, enabling normalized (z-score) versions of metrics.
        """
        from scipy.signal import peak_prominences
        from core import metrics as core_metrics

        st = self.state
        if not st.peaks_by_sweep or not st.breath_by_sweep:
            return

        # Collect raw values across all sweeps
        all_amp_insp = []
        all_amp_exp = []
        all_peak_to_trough = []
        all_prominences = []
        all_ibi = []
        all_ti = []
        all_te = []

        for s in range(len(st.peaks_by_sweep)):
            if s not in st.peaks_by_sweep or s not in st.breath_by_sweep:
                continue

            pks = st.peaks_by_sweep[s]
            breaths = st.breath_by_sweep[s]

            if len(pks) == 0:
                continue

            y_proc = self._get_processed_for(st.analyze_chan, s)
            t = st.t  # Time vector

            onsets = breaths.get('onsets', np.array([]))
            offsets = breaths.get('offsets', np.array([]))
            expmins = breaths.get('expmins', np.array([]))

            # Compute prominences for all peaks
            if len(pks) > 0:
                proms = peak_prominences(y_proc, pks)[0]
                all_prominences.extend(proms)

            # Compute amp_insp for each breath cycle
            for i in range(len(onsets) - 1):
                onset_idx = int(onsets[i])
                next_onset_idx = int(onsets[i + 1])

                # Find peak in this cycle
                pk_mask = (pks >= onset_idx) & (pks < next_onset_idx)
                if not np.any(pk_mask):
                    continue
                pk_idx = int(pks[pk_mask][0])

                # Amp insp
                amp_insp = y_proc[pk_idx] - y_proc[onset_idx]
                all_amp_insp.append(amp_insp)

                # Amp exp (if expmin exists)
                em_mask = (expmins >= onset_idx) & (expmins < next_onset_idx)
                if np.any(em_mask):
                    em_idx = int(expmins[em_mask][0])
                    amp_exp = y_proc[next_onset_idx] - y_proc[em_idx]
                    all_amp_exp.append(amp_exp)

                    # Peak to trough
                    peak_to_trough = y_proc[pk_idx] - y_proc[em_idx]
                    all_peak_to_trough.append(peak_to_trough)

                # Ti
                if i < len(offsets):
                    offset_idx = int(offsets[i])
                    if onset_idx < offset_idx <= next_onset_idx:
                        ti = t[offset_idx] - t[onset_idx]
                        all_ti.append(ti)

                        # Te
                        te = t[next_onset_idx] - t[offset_idx]
                        all_te.append(te)

            # Compute IBI (inter-breath interval)
            for i in range(len(pks) - 1):
                ibi = t[pks[i + 1]] - t[pks[i]]
                all_ibi.append(ibi)

        # Compute global statistics
        stats = {}

        if len(all_amp_insp) > 0:
            stats['amp_insp'] = {'mean': float(np.nanmean(all_amp_insp)), 'std': float(np.nanstd(all_amp_insp))}

        if len(all_amp_exp) > 0:
            stats['amp_exp'] = {'mean': float(np.nanmean(all_amp_exp)), 'std': float(np.nanstd(all_amp_exp))}

        if len(all_peak_to_trough) > 0:
            stats['peak_to_trough'] = {'mean': float(np.nanmean(all_peak_to_trough)), 'std': float(np.nanstd(all_peak_to_trough))}

        if len(all_prominences) > 0:
            stats['prominence'] = {'mean': float(np.nanmean(all_prominences)), 'std': float(np.nanstd(all_prominences))}

        if len(all_ibi) > 0:
            stats['ibi'] = {'mean': float(np.nanmean(all_ibi)), 'std': float(np.nanstd(all_ibi))}

        if len(all_ti) > 0:
            stats['ti'] = {'mean': float(np.nanmean(all_ti)), 'std': float(np.nanstd(all_ti))}

        if len(all_te) > 0:
            stats['te'] = {'mean': float(np.nanmean(all_te)), 'std': float(np.nanstd(all_te))}

        # Store statistics in metrics module
        core_metrics.set_normalization_stats(stats)

        print(f"[normalization] Computed stats for {len(stats)} metric types")
        if 'prominence' in stats:
            print(f"[normalization]   prominence: mean={stats['prominence']['mean']:.4f}, std={stats['prominence']['std']:.4f}")

    def _compute_eupnea_from_gmm(self, sweep_idx: int, signal_length: int) -> np.ndarray:
        """
        Compute eupnea mask from GMM clustering results.

        Eupnea = breaths that are NOT sniffing (based on GMM classification).
        Groups consecutive eupnic breaths into continuous regions.

        Args:
            sweep_idx: Index of the current sweep
            signal_length: Length of the signal array

        Returns:
            Boolean array (as float 0/1) marking eupneic regions
        """
        import numpy as np

        eupnea_mask = np.zeros(signal_length, dtype=bool)

        # Check if GMM probabilities are available for this sweep
        if not hasattr(self.state, 'gmm_sniff_probabilities'):
            return eupnea_mask.astype(float)

        if sweep_idx not in self.state.gmm_sniff_probabilities:
            return eupnea_mask.astype(float)

        # Get breath data for this sweep
        breath_data = self.state.breath_by_sweep.get(sweep_idx)
        if breath_data is None:
            return eupnea_mask.astype(float)

        onsets = breath_data.get('onsets', np.array([]))
        offsets = breath_data.get('offsets', np.array([]))

        if len(onsets) == 0:
            return eupnea_mask.astype(float)

        t = self.state.t
        gmm_probs = self.state.gmm_sniff_probabilities[sweep_idx]

        # Identify eupnic breaths and group consecutive ones
        eupnic_groups = []
        current_group_start = None
        current_group_end = None
        last_eupnic_idx = None

        # Get peaks for this sweep (needed to map breath → peak)
        peaks = self.state.peaks_by_sweep.get(sweep_idx)
        if peaks is None or len(peaks) != len(onsets):
            return eupnea_mask.astype(float)  # Peaks and onsets must be aligned

        for breath_idx in range(len(onsets)):
            # Get peak sample index for this breath
            # Peaks and breaths are 1:1 aligned by array index
            peak_sample_idx = int(peaks[breath_idx])

            # Look up GMM probability using peak sample index
            if peak_sample_idx not in gmm_probs:
                # Close current group if exists
                if current_group_start is not None:
                    eupnic_groups.append((current_group_start, current_group_end))
                    current_group_start = None
                    current_group_end = None
                    last_eupnic_idx = None
                continue

            sniff_prob = gmm_probs[peak_sample_idx]

            # Eupnea if sniffing probability < 0.5 (i.e., more likely eupnea)
            if sniff_prob < 0.5:
                # Get time range for this breath
                start_idx = int(onsets[breath_idx])

                # Get offset time
                if breath_idx < len(offsets):
                    end_idx = int(offsets[breath_idx])
                else:
                    # Fallback: use next onset or end of trace
                    if breath_idx + 1 < len(onsets):
                        end_idx = int(onsets[breath_idx + 1])
                    else:
                        end_idx = signal_length

                # Check if this is consecutive with the last eupnic breath
                if last_eupnic_idx is None or breath_idx != last_eupnic_idx + 1:
                    # Not consecutive - save current group and start new one
                    if current_group_start is not None:
                        eupnic_groups.append((current_group_start, current_group_end))
                    current_group_start = start_idx
                    current_group_end = end_idx
                else:
                    # Consecutive breath - extend the current group
                    current_group_end = end_idx

                last_eupnic_idx = breath_idx
            else:
                # Non-eupnic breath - close current group if exists
                if current_group_start is not None:
                    eupnic_groups.append((current_group_start, current_group_end))
                    current_group_start = None
                    current_group_end = None
                    last_eupnic_idx = None

        # Save final group if exists
        if current_group_start is not None:
            eupnic_groups.append((current_group_start, current_group_end))

        # Mark all continuous eupnic regions
        for start_idx, end_idx in eupnic_groups:
            eupnea_mask[start_idx:end_idx] = True

        return eupnea_mask.astype(float)

    def _compute_eupnea_from_active_classifier(self, sweep_idx: int, signal_length: int) -> np.ndarray:
        """
        Compute eupnea mask using the active eupnea/sniff classifier.

        This is a generalized method that works with any classifier (GMM, XGBoost, RF, MLP)
        by reading from the breath_type_class array in all_peaks_by_sweep, which always contains
        the currently active classifier's predictions.

        Args:
            sweep_idx: Index of the current sweep
            signal_length: Length of the signal array

        Returns:
            Boolean array (as float 0/1) marking eupneic regions
        """
        import numpy as np

        eupnea_mask = np.zeros(signal_length, dtype=bool)

        # Get breath data and peak data for this sweep
        all_peaks = self.state.all_peaks_by_sweep.get(sweep_idx)
        breath_data = self.state.breath_by_sweep.get(sweep_idx)

        if all_peaks is None or breath_data is None:
            return eupnea_mask.astype(float)

        # Get breath_type_class array (contains active classifier's predictions: 0=eupnea, 1=sniffing, -1=unclassified)
        breath_type_class = all_peaks.get('breath_type_class')
        if breath_type_class is None:
            # Fall back to old GMM probability method if breath_type_class not available
            return self._compute_eupnea_from_gmm(sweep_idx, signal_length)

        onsets = breath_data.get('onsets', np.array([]))
        offsets = breath_data.get('offsets', np.array([]))

        if len(onsets) == 0 or len(breath_type_class) != len(onsets):
            return eupnea_mask.astype(float)

        # Group consecutive eupnic breaths into continuous regions
        eupnic_groups = []
        current_group_start = None
        current_group_end = None
        last_eupnic_idx = None

        for breath_idx in range(len(onsets)):
            # Eupnea = breath_type_class == 0
            is_eupnic = (breath_type_class[breath_idx] == 0)

            if is_eupnic:
                # Get time range for this breath
                start_idx = int(onsets[breath_idx])

                # Get offset time
                if breath_idx < len(offsets):
                    end_idx = int(offsets[breath_idx])
                else:
                    # Fallback: use next onset or end of trace
                    if breath_idx + 1 < len(onsets):
                        end_idx = int(onsets[breath_idx + 1])
                    else:
                        end_idx = signal_length

                # Check if this is consecutive with the last eupnic breath
                if last_eupnic_idx is None or breath_idx != last_eupnic_idx + 1:
                    # Not consecutive - save current group and start new one
                    if current_group_start is not None:
                        eupnic_groups.append((current_group_start, current_group_end))
                    current_group_start = start_idx
                    current_group_end = end_idx
                else:
                    # Consecutive breath - extend the current group
                    current_group_end = end_idx

                last_eupnic_idx = breath_idx
            else:
                # Non-eupnic breath - close current group if exists
                if current_group_start is not None:
                    eupnic_groups.append((current_group_start, current_group_end))
                    current_group_start = None
                    current_group_end = None
                    last_eupnic_idx = None

        # Save final group if exists
        if current_group_start is not None:
            eupnic_groups.append((current_group_start, current_group_end))

        # Mark all continuous eupnic regions
        for start_idx, end_idx in eupnic_groups:
            eupnea_mask[start_idx:end_idx] = True

        return eupnea_mask.astype(float)

    def _run_automatic_gmm_clustering(self):
        """
        Automatically run GMM clustering after peak detection to identify sniffing breaths.
        Uses streamlined default features (if, ti, amp_insp, max_dinsp) and 2 clusters.
        Silently marks identified sniffing breaths with purple background.
        """
        import time
        from sklearn.mixture import GaussianMixture
        from sklearn.preprocessing import StandardScaler
        from sklearn.metrics import silhouette_score
        import numpy as np

        t_start = time.time()
        st = self.state

        # Check if we have breath data
        if not st.peaks_by_sweep or len(st.peaks_by_sweep) == 0:
            print("[auto-gmm] No breath data available, skipping automatic GMM clustering")
            return

        # Streamlined default features for eupnea/sniffing separation
        feature_keys = ["if", "ti", "amp_insp", "max_dinsp"]
        n_clusters = 2

        print(f"\n[auto-gmm] Running automatic GMM clustering with {n_clusters} clusters...")
        print(f"[auto-gmm] Features: {', '.join(feature_keys)}")
        self._log_status_message("Running GMM clustering...")

        try:
            # Collect breath features from all analyzed sweeps
            feature_matrix, breath_cycles = self._collect_gmm_breath_features(feature_keys)

            if len(feature_matrix) < n_clusters:
                print(f"[auto-gmm] Not enough breaths ({len(feature_matrix)}) for {n_clusters} clusters, skipping")
                return

            # Standardize features
            scaler = StandardScaler()
            feature_matrix_scaled = scaler.fit_transform(feature_matrix)

            # Fit GMM
            gmm_model = GaussianMixture(n_components=n_clusters, random_state=42, covariance_type='full')
            cluster_labels = gmm_model.fit_predict(feature_matrix_scaled)

            # Get probability estimates for each breath
            cluster_probabilities = gmm_model.predict_proba(feature_matrix_scaled)

            # Check clustering quality
            silhouette = silhouette_score(feature_matrix_scaled, cluster_labels) if n_clusters > 1 else -1
            print(f"[auto-gmm] Silhouette score: {silhouette:.3f}")

            # Identify sniffing cluster
            sniffing_cluster_id = self._identify_gmm_sniffing_cluster(
                feature_matrix, cluster_labels, feature_keys, silhouette
            )

            if sniffing_cluster_id is None:
                print("[auto-gmm] Could not identify sniffing cluster, skipping")
                return

            # Apply GMM sniffing regions to plot (stores probabilities AND creates regions)
            self._apply_gmm_sniffing_regions(
                breath_cycles, cluster_labels, cluster_probabilities, sniffing_cluster_id
            )

            n_sniffing_breaths = np.sum(cluster_labels == sniffing_cluster_id)
            print(f"[auto-gmm] Identified {n_sniffing_breaths} sniffing breaths and applied to plot")

            # Cache results for fast dialog loading
            self._cached_gmm_results = {
                'cluster_labels': cluster_labels,
                'cluster_probabilities': cluster_probabilities,
                'feature_matrix': feature_matrix,
                'breath_cycles': breath_cycles,
                'sniffing_cluster_id': sniffing_cluster_id,
                'feature_keys': feature_keys
            }
            print("[auto-gmm] Cached GMM results for fast dialog loading")

            # Show completion message with elapsed time
            t_elapsed = time.time() - t_start

            # Log telemetry: GMM clustering success
            eupnea_count = len(cluster_labels) - n_sniffing_breaths
            telemetry.log_feature_used('gmm_clustering')
            telemetry.log_timing('gmm_clustering', t_elapsed,
                                num_breaths=len(cluster_labels),
                                num_clusters=n_clusters,
                                silhouette_score=round(silhouette, 3))

            telemetry.log_breath_statistics(
                num_breaths=len(cluster_labels),
                sniff_count=int(n_sniffing_breaths),
                eupnea_count=int(eupnea_count),
                silhouette_score=round(silhouette, 3)
            )

            self._log_status_message(f"GMM clustering complete ({t_elapsed:.1f}s)", 2000)

        except Exception as e:
            print(f"[auto-gmm] Error during automatic GMM clustering: {e}")
            t_elapsed = time.time() - t_start

            # Log telemetry: GMM clustering failure
            telemetry.log_crash(f"GMM clustering failed: {type(e).__name__}",
                               operation='gmm_clustering',
                               num_breaths=len(feature_matrix) if 'feature_matrix' in locals() else 0)

            self._log_status_message(f"GMM clustering failed ({t_elapsed:.1f}s)", 3000)
            import traceback
            traceback.print_exc()

    def _collect_gmm_breath_features(self, feature_keys):
        """Collect per-breath features for GMM clustering."""
        import numpy as np
        from core import metrics, filters

        feature_matrix = []
        breath_cycles = []
        st = self.state

        for sweep_idx in sorted(st.breath_by_sweep.keys()):
            breath_data = st.breath_by_sweep[sweep_idx]

            if sweep_idx not in st.peaks_by_sweep:
                continue

            peaks = st.peaks_by_sweep[sweep_idx]
            t = st.t
            y_raw = st.sweeps[st.analyze_chan][:, sweep_idx]

            # Apply filters
            y = filters.apply_all_1d(
                y_raw, st.sr_hz,
                st.use_low, st.low_hz,
                st.use_high, st.high_hz,
                st.use_mean_sub, st.mean_val,
                st.use_invert,
                order=self.filter_order
            )

            # Apply notch filter if configured
            if self.notch_filter_lower is not None and self.notch_filter_upper is not None:
                y = self._apply_notch_filter(y, st.sr_hz,
                                              self.notch_filter_lower,
                                              self.notch_filter_upper)

            # Apply z-score normalization if enabled (using global statistics)
            if self.use_zscore_normalization:
                # Compute global stats if not cached
                if self.zscore_global_mean is None or self.zscore_global_std is None:
                    self.zscore_global_mean, self.zscore_global_std = self._compute_global_zscore_stats()
                y = filters.zscore_normalize(y, self.zscore_global_mean, self.zscore_global_std)

            # Get breath events
            onsets = breath_data.get('onsets', np.array([]))
            offsets = breath_data.get('offsets', np.array([]))
            expmins = breath_data.get('expmins', np.array([]))
            expoffs = breath_data.get('expoffs', np.array([]))

            if len(onsets) == 0:
                continue

            # Compute metrics
            metrics_dict = {}
            for feature_key in feature_keys:
                if feature_key in metrics.METRICS:
                    metric_arr = metrics.METRICS[feature_key](
                        t, y, st.sr_hz, peaks, onsets, offsets, expmins, expoffs
                    )
                    metrics_dict[feature_key] = metric_arr

            # Extract per-breath values
            n_breaths = len(onsets)
            for breath_idx in range(n_breaths):
                start = int(onsets[breath_idx])
                breath_features = []
                valid_breath = True

                for feature_key in feature_keys:
                    if feature_key not in metrics_dict:
                        valid_breath = False
                        break

                    metric_arr = metrics_dict[feature_key]
                    if start < len(metric_arr):
                        val = metric_arr[start]
                        if np.isnan(val) or not np.isfinite(val):
                            valid_breath = False
                            break
                        breath_features.append(val)
                    else:
                        valid_breath = False
                        break

                if valid_breath and len(breath_features) == len(feature_keys):
                    feature_matrix.append(breath_features)
                    breath_cycles.append((sweep_idx, breath_idx))

        return np.array(feature_matrix), breath_cycles

    def _identify_gmm_sniffing_cluster(self, feature_matrix, cluster_labels, feature_keys, silhouette):
        """Identify which cluster represents sniffing based on IF and Ti."""
        import numpy as np

        unique_labels = np.unique(cluster_labels)
        n_clusters = len(unique_labels)

        # Get indices of IF and Ti features
        if_idx = feature_keys.index('if') if 'if' in feature_keys else None
        ti_idx = feature_keys.index('ti') if 'ti' in feature_keys else None

        if if_idx is None and ti_idx is None:
            print("[auto-gmm] Cannot identify sniffing without 'if' or 'ti' features")
            return None

        # Compute mean IF and Ti for each cluster
        cluster_stats = {}
        for cluster_id in unique_labels:
            mask = cluster_labels == cluster_id
            stats = {}
            if if_idx is not None:
                stats['mean_if'] = np.mean(feature_matrix[mask, if_idx])
            if ti_idx is not None:
                stats['mean_ti'] = np.mean(feature_matrix[mask, ti_idx])
            cluster_stats[cluster_id] = stats

        # Identify sniffing: highest IF and/or lowest Ti
        cluster_scores = {}
        for cluster_id in unique_labels:
            score = 0
            if if_idx is not None:
                if_vals = [cluster_stats[c]['mean_if'] for c in unique_labels]
                if_rank = sorted(if_vals).index(cluster_stats[cluster_id]['mean_if'])
                score += if_rank / (n_clusters - 1) if n_clusters > 1 else 0
            if ti_idx is not None:
                ti_vals = [cluster_stats[c]['mean_ti'] for c in unique_labels]
                ti_rank = sorted(ti_vals, reverse=True).index(cluster_stats[cluster_id]['mean_ti'])
                score += ti_rank / (n_clusters - 1) if n_clusters > 1 else 0
            cluster_scores[cluster_id] = score

        sniffing_cluster_id = max(cluster_scores, key=cluster_scores.get)

        # Log cluster statistics
        for cluster_id in unique_labels:
            stats_str = ", ".join([f"{k}={v:.3f}" for k, v in cluster_stats[cluster_id].items()])
            marker = " (SNIFFING)" if cluster_id == sniffing_cluster_id else ""
            print(f"[auto-gmm]   Cluster {cluster_id}: {stats_str}{marker}")

        # Validate quality (warn but don't block)
        sniff_stats = cluster_stats[sniffing_cluster_id]
        if silhouette < 0.25:
            print(f"[auto-gmm] WARNING: Low cluster separation (silhouette={silhouette:.3f})")
            print(f"[auto-gmm]   Breathing patterns may be very similar (e.g., anesthetized mouse)")
        if if_idx is not None and sniff_stats['mean_if'] < 5.0:
            print(f"[auto-gmm] WARNING: 'Sniffing' cluster has low IF ({sniff_stats['mean_if']:.2f} Hz)")
            print(f"[auto-gmm]   May be normal variation, not true sniffing (typical sniffing: 5-8 Hz)")

        return sniffing_cluster_id

    def _apply_gmm_sniffing_regions(self, breath_cycles, cluster_labels, cluster_probabilities, sniffing_cluster_id):
        """Apply GMM cluster results using shared functions from core.gmm_clustering.

        Stores classifications in all_peaks_by_sweep and builds both eupnea and sniffing regions.
        Stores probabilities for each breath (for backward compatibility).

        Args:
            breath_cycles: List of (sweep_idx, breath_idx) tuples
            cluster_labels: Hard cluster assignments
            cluster_probabilities: Probability matrix (n_breaths, n_clusters)
            sniffing_cluster_id: Which cluster is sniffing
        """
        import numpy as np
        from core import gmm_clustering

        # Store probabilities by (sweep_idx, breath_idx) for backward compatibility
        if not hasattr(self.state, 'gmm_sniff_probabilities'):
            self.state.gmm_sniff_probabilities = {}
        self.state.gmm_sniff_probabilities.clear()

        for i, (sweep_idx, breath_idx) in enumerate(breath_cycles):
            # Get probability of this breath being sniffing
            sniff_prob = cluster_probabilities[i, sniffing_cluster_id]

            # Store probability
            if sweep_idx not in self.state.gmm_sniff_probabilities:
                self.state.gmm_sniff_probabilities[sweep_idx] = {}
            self.state.gmm_sniff_probabilities[sweep_idx][breath_idx] = sniff_prob

        # Check if GMM is the active classifier
        is_gmm_active = self.state.active_eupnea_sniff_classifier == 'gmm'

        # Store classifications in all_peaks_by_sweep
        # - Always store to gmm_class_ro (for switching classifiers later)
        # - Only update breath_type_class (editable) if GMM is the active classifier
        n_classified = gmm_clustering.store_gmm_classifications_in_peaks(
            self.state, breath_cycles, cluster_labels, sniffing_cluster_id,
            cluster_probabilities, confidence_threshold=0.5,
            update_editable=is_gmm_active
        )

        # Only build regions if GMM is the active classifier
        # (otherwise, the active classifier's predictions are already in breath_type_class)
        if is_gmm_active:
            # Build BOTH eupnea AND sniffing regions from stored classifications
            results = gmm_clustering.build_eupnea_sniffing_regions(
                self.state, verbose=False, log_prefix="[auto-gmm]"
            )
        else:
            # Just return dummy results - regions will be built from ML predictions
            results = {'n_sniffing': 0, 'n_eupnea': 0, 'total_sniff_regions': 0, 'total_eupnea_regions': 0}
            print(f"[auto-gmm] GMM results cached but not applied (active classifier: {self.state.active_eupnea_sniff_classifier})")

        # Calculate probability statistics
        all_sniff_probs = []
        for sweep_idx in self.state.gmm_sniff_probabilities:
            for breath_idx in self.state.gmm_sniff_probabilities[sweep_idx]:
                prob = self.state.gmm_sniff_probabilities[sweep_idx][breath_idx]
                all_sniff_probs.append(prob)

        if all_sniff_probs:
            all_sniff_probs = np.array(all_sniff_probs)
            sniff_probs_of_sniff_breaths = all_sniff_probs[all_sniff_probs >= 0.5]  # Breaths classified as sniffing
            if len(sniff_probs_of_sniff_breaths) > 0:
                mean_conf = np.mean(sniff_probs_of_sniff_breaths)
                min_conf = np.min(sniff_probs_of_sniff_breaths)
                uncertain_count = np.sum((sniff_probs_of_sniff_breaths >= 0.5) & (sniff_probs_of_sniff_breaths < 0.7))
                print(f"[auto-gmm]   Sniffing probability: mean={mean_conf:.3f}, min={min_conf:.3f}")
                if uncertain_count > 0:
                    print(f"[auto-gmm]   WARNING: {uncertain_count} breaths have uncertain classification (50-70% sniffing probability)")

        # Report results
        print(f"[auto-gmm]   Created {results['total_sniff_regions']} sniffing region(s) across sweeps")
        print(f"[auto-gmm]   Created {results['total_eupnea_regions']} eupnea region(s) across sweeps")

        return results['n_sniffing']

    def _store_gmm_probabilities_only(self, breath_cycles, cluster_probabilities, sniffing_cluster_id):
        """Store GMM sniffing probabilities without applying regions to plot.

        This is used by automatic GMM clustering to store results without
        immediately marking sniffing regions. User must manually enable
        "Apply Sniffing Detection" in GMM dialog to see markings.

        Args:
            breath_cycles: List of (sweep_idx, breath_idx) tuples
            cluster_probabilities: Probability matrix (n_breaths, n_clusters)
            sniffing_cluster_id: Which cluster is sniffing
        """
        import numpy as np

        # Store probabilities by (sweep_idx, breath_idx)
        if not hasattr(self.state, 'gmm_sniff_probabilities'):
            self.state.gmm_sniff_probabilities = {}
        self.state.gmm_sniff_probabilities.clear()

        for i, (sweep_idx, breath_idx) in enumerate(breath_cycles):
            if sweep_idx not in self.state.gmm_sniff_probabilities:
                self.state.gmm_sniff_probabilities[sweep_idx] = {}

            # Get probability of this breath being sniffing
            sniff_prob = cluster_probabilities[i, sniffing_cluster_id]
            self.state.gmm_sniff_probabilities[sweep_idx][breath_idx] = sniff_prob

        # Calculate probability statistics
        all_sniff_probs = []
        for sweep_idx in self.state.gmm_sniff_probabilities:
            for breath_idx in self.state.gmm_sniff_probabilities[sweep_idx]:
                prob = self.state.gmm_sniff_probabilities[sweep_idx][breath_idx]
                all_sniff_probs.append(prob)

        if all_sniff_probs:
            all_sniff_probs = np.array(all_sniff_probs)
            sniff_probs_of_sniff_breaths = all_sniff_probs[all_sniff_probs >= 0.5]  # Breaths classified as sniffing
            if len(sniff_probs_of_sniff_breaths) > 0:
                mean_conf = np.mean(sniff_probs_of_sniff_breaths)
                min_conf = np.min(sniff_probs_of_sniff_breaths)
                uncertain_count = np.sum((sniff_probs_of_sniff_breaths >= 0.5) & (sniff_probs_of_sniff_breaths < 0.7))
                print(f"[auto-gmm]   Sniffing probability: mean={mean_conf:.3f}, min={min_conf:.3f}")
                if uncertain_count > 0:
                    print(f"[auto-gmm]   WARNING: {uncertain_count} breaths have uncertain classification (50-70% sniffing probability)")

    ##################################################
    ##y2 plotting                                   ##
    ##################################################
    def _compute_y2_all_sweeps(self):
        """Compute active y2 metric for ALL sweeps on the analyze channel."""
        st = self.state
        key = getattr(st, "y2_metric_key", None)
        if not key:
            st.y2_values_by_sweep.clear()
            return
        if key not in metrics.METRICS:
            st.y2_values_by_sweep.clear()
            return
        if st.t is None or st.analyze_chan not in st.sweeps:
            st.y2_values_by_sweep.clear()
            return

        fn = metrics.METRICS[key]
        any_ch = next(iter(st.sweeps.values()))
        n_sweeps = any_ch.shape[1]
        st.y2_values_by_sweep = {}

        for s in range(n_sweeps):
            y_proc = self._get_processed_for(st.analyze_chan, s)
            # pull peaks/breaths if available
            pks = getattr(st, "peaks_by_sweep", {}).get(s, None)
            # breaths = getattr(st, "breath_by_sweep", {}).get(s, {}) if hasattr(st, "breath_by_sweep") else {}
            breaths = getattr(st, "breath_by_sweep", {}).get(s, {})
            on = breaths.get("onsets", None)
            off = breaths.get("offsets", None)
            exm = breaths.get("expmins", None)
            exo = breaths.get("expoffs", None)

            # Set GMM probabilities for this sweep (if available)
            gmm_probs = None
            if hasattr(st, 'gmm_sniff_probabilities') and s in st.gmm_sniff_probabilities:
                gmm_probs = st.gmm_sniff_probabilities[s]
            metrics.set_gmm_probabilities(gmm_probs)

            # Set peak candidate metrics for this sweep (if available)
            # Prefer current_peak_metrics_by_sweep (updated after edits) for Y2 plotting,
            # fallback to peak_metrics_by_sweep (original auto-detected) for ML training
            peak_metrics = None
            if hasattr(st, 'current_peak_metrics_by_sweep') and s in st.current_peak_metrics_by_sweep:
                peak_metrics = st.current_peak_metrics_by_sweep[s]
            elif hasattr(st, 'peak_metrics_by_sweep') and s in st.peak_metrics_by_sweep:
                peak_metrics = st.peak_metrics_by_sweep[s]
            metrics.set_peak_metrics(peak_metrics)

            y2 = fn(st.t, y_proc, st.sr_hz, pks, on, off, exm, exo)
            st.y2_values_by_sweep[s] = y2

        # Clear GMM probabilities after computation
        metrics.set_gmm_probabilities(None)
        # Clear peak metrics after computation
        metrics.set_peak_metrics(None)

    def on_y2_metric_changed(self, idx: int):
        key = self.y2plot_dropdown.itemData(idx)
        self.state.y2_metric_key = key  # None or e.g. "if"

        # Recompute Y2 (needs peaks/breaths for most metrics; IF falls back to peaks)
        self._compute_y2_all_sweeps()

        # Force a redraw of current sweep
        # Also reset Y2 axis so it rescales to new data
        self.plot_host.clear_y2()
        self.redraw_main_plot()

    ##################################################
    ## Display Control Handlers ##
    ##################################################

    def on_yautoscale_toggled(self, checked: bool):
        """Toggle Y-axis autoscale between percentile and full range mode."""
        self.state.use_percentile_autoscale = checked
        mode = "percentile (99th)" if checked else "full range (min/max)"
        self._log_status_message(f"Y-axis autoscale: {mode}", 2000)
        self.redraw_main_plot()

    def on_ypadding_changed(self, value: float):
        """Adjust padding for percentile autoscale mode."""
        self.state.autoscale_padding = value
        self._log_status_message(f"Y-axis padding: {value*100:.0f}%", 2000)
        # Only redraw if in percentile mode
        if self.state.use_percentile_autoscale:
            self.redraw_main_plot()

    def on_eupnea_display_toggled(self, checked: bool):
        """Toggle eupnea region display between shade and line."""
        self.state.eupnea_use_shade = checked
        mode = "shading" if checked else "line"
        self._log_status_message(f"Eupnea display: {mode}", 2000)
        self.redraw_main_plot()

    def on_sniffing_display_toggled(self, checked: bool):
        """Toggle sniffing region display between shade and line."""
        self.state.sniffing_use_shade = checked
        mode = "shading" if checked else "line"
        self._log_status_message(f"Sniffing display: {mode}", 2000)
        self.redraw_main_plot()

    def on_apnea_display_toggled(self, checked: bool):
        """Toggle apnea region display between shade and line."""
        self.state.apnea_use_shade = checked
        mode = "shading" if checked else "line"
        self._log_status_message(f"Apnea display: {mode}", 2000)
        self.redraw_main_plot()

    def on_outliers_display_toggled(self, checked: bool):
        """Toggle outliers display between shade and line."""
        self.state.outliers_use_shade = checked
        mode = "shading" if checked else "line"
        self._log_status_message(f"Outliers display: {mode}", 2000)
        self.redraw_main_plot()

    def on_dark_mode_toggled(self, checked: bool):
        """Toggle plot dark mode (dark vs light background)."""
        theme = "dark" if checked else "light"
        self.plot_host.set_plot_theme(theme)

        # Save preference
        self.settings.setValue("plot_dark_mode", checked)

        mode_text = "Dark mode" if checked else "Light mode"
        self._log_status_message(f"Plot theme: {mode_text}", 2000)

    ##################################################
    ## Turn Off All Edit Modes ##
    ##################################################
    # Note: These lines were orphaned code that was being executed as part of on_y2_metric_changed
    # They have been removed because they were clearing the callback after we restored it

    





    def on_update_eupnea_sniffing_clicked(self):
        """Handle Update Eupnea/Sniffing Detection button - rerun current classifier and apply."""
        import time
        from PyQt6.QtCore import QTimer

        st = self.state
        if not st.peaks_by_sweep or len(st.peaks_by_sweep) == 0:
            self._log_status_message("No peaks detected yet - run peak detection first", 3000)
            return

        t_start = time.time()
        current_classifier = st.active_eupnea_sniff_classifier
        print(f"[update-eupnea] Manually updating eupnea/sniffing detection using: {current_classifier}")
        self._log_status_message(f"Updating eupnea/sniffing detection ({current_classifier})...")

        # Handle based on current classifier selection
        if current_classifier == 'all_eupnea':
            # Set all breaths to eupnea
            self._set_all_breaths_eupnea_sniff_class(0)
        elif current_classifier == 'none':
            # Clear all labels
            self._clear_all_eupnea_sniff_labels()
        elif current_classifier == 'gmm':
            # Run GMM clustering
            self._run_automatic_gmm_clustering()
        else:
            # ML model - re-run predictions and update
            self._update_eupnea_sniff_from_classifier()

        # Rebuild regions for display
        try:
            import core.gmm_clustering as gmm_clustering
            gmm_clustering.build_eupnea_sniffing_regions(self.state, verbose=False)
        except Exception as e:
            print(f"[update-eupnea] Warning: Could not rebuild regions: {e}")

        # Clear out-of-date flag
        self.eupnea_sniffing_out_of_date = False

        # LIGHTWEIGHT UPDATE: Just refresh eupnea/sniffing overlays without full redraw
        # This skips expensive outlier detection and metrics recomputation
        self._refresh_eupnea_overlays_only()

        # Show completion message with elapsed time
        t_elapsed = time.time() - t_start
        print(f"[update-eupnea] Eupnea/sniffing detection updated ({t_elapsed:.1f}s)")
        self._log_status_message(f"Eupnea/sniffing detection updated ({t_elapsed:.1f}s)", 2000)
        # Clear again after the success message disappears
        QTimer.singleShot(2100, lambda: self.statusBar().clearMessage())

    def _refresh_eupnea_overlays_only(self):
        """
        Lightweight update of eupnea/sniffing region overlays without full plot redraw.
        Used after GMM clustering to avoid expensive outlier detection recomputation.
        """
        st = self.state
        s = st.sweep_idx

        # Get current trace data
        t, y = self._current_trace()
        if t is None or y is None:
            return

        # Apply time normalization if stim channel exists
        spans = st.stim_spans_by_sweep.get(s, []) if st.stim_chan else []
        if st.stim_chan and spans:
            t0 = spans[0][0]
            t_plot = t - t0
        else:
            t0 = 0.0
            t_plot = t

        # Get breath data
        br = st.breath_by_sweep.get(s, None)
        if not br:
            return

        # Compute ONLY eupnea mask from active classifier (fast, no metrics recomputation)
        eupnea_mask = self._compute_eupnea_from_active_classifier(s, len(y))

        # Get existing apnea threshold
        apnea_thresh = self._parse_float(self.ApneaThresh) or 0.5

        # Compute apnea mask (also fast, no metrics)
        pks = st.peaks_by_sweep.get(s, [])
        on_idx = br.get("onsets", [])
        off_idx = br.get("offsets", [])
        ex_idx = br.get("expmins", [])
        exoff_idx = br.get("expoffs", [])

        from core import metrics
        apnea_mask = metrics.detect_apneas(
            t, y, st.sr_hz, pks, on_idx, off_idx, ex_idx, exoff_idx,
            min_apnea_duration_sec=apnea_thresh
        )

        # Get eupnea and sniff regions for overlay display
        # IMPORTANT: Apply same time normalization as the trace (shift by t0 if stimulus active)
        sniff_regions_raw = st.sniff_regions_by_sweep.get(s, [])
        eupnea_regions_raw = st.eupnea_regions_by_sweep.get(s, [])

        # Normalize region times to match t_plot (shift by t0)
        sniff_regions = [(start - t0, end - t0) for start, end in sniff_regions_raw]
        eupnea_regions = [(start - t0, end - t0) for start, end in eupnea_regions_raw]

        # Update ONLY the region overlays (skips outlier masks entirely - huge speedup!)
        self.plot_host.update_region_overlays(t_plot, eupnea_mask, apnea_mask,
                                              outlier_mask=None, failure_mask=None,
                                              sniff_regions=sniff_regions,
                                              eupnea_regions=eupnea_regions,
                                              state=st)

        # Update sniffing region backgrounds (purple) - Note: This may be redundant with update_region_overlays
        # self.editing_modes.update_sniff_artists(t_plot, s)

        # Refresh canvas
        self.plot_host.canvas.draw_idle()
        print("[update-eupnea] Lightweight overlay refresh complete (skipped outlier detection)")

    def on_help_clicked(self):
        """Open the help dialog (F1) - non-modal."""
        from dialogs.help_dialog import HelpDialog

        # If help already open, bring to front instead of creating new window
        if hasattr(self, 'help_dialog') and self.help_dialog is not None:
            self.help_dialog.raise_()
            self.help_dialog.activateWindow()
            return

        # Create non-modal dialog
        self.help_dialog = HelpDialog(self, update_info=self.update_info)
        self.help_dialog.finished.connect(self._on_help_dialog_closed)
        self.help_dialog.show()  # Non-modal - doesn't block
        telemetry.log_screen_view('Help Dialog', screen_class='info_dialog')

    def _on_help_dialog_closed(self):
        """Clear help dialog reference when closed so it can be reopened."""
        self.help_dialog = None

    def _check_for_updates_on_startup(self):
        """Check for updates in background and update UI if available."""
        from PyQt6.QtCore import QThread, pyqtSignal

        class UpdateChecker(QThread):
            """Background thread for checking updates."""
            update_checked = pyqtSignal(object)  # Emits update_info or None

            def run(self):
                """Run update check in background."""
                from core import update_checker
                update_info = update_checker.check_for_updates()
                self.update_checked.emit(update_info)

        def on_update_checked(update_info):
            """Handle update check result."""
            if update_info:
                # Store for help dialog
                self.update_info = update_info

                # Update main window label
                from core import update_checker
                text, url = update_checker.get_main_window_update_message(update_info)
                self.update_notification_label.setText(f'<a href="{url}" style="color: #FFD700; text-decoration: underline;">{text}</a>')
                self.update_notification_label.setVisible(True)
                print(f"[Update Check] New version available: {update_info.get('version')}")
            else:
                # No update available - keep label hidden
                print("[Update Check] You're up to date!")

        # Create and start background thread
        self.update_thread = UpdateChecker()
        self.update_thread.update_checked.connect(on_update_checked)
        self.update_thread.start()

    def on_spectral_analysis_clicked(self):
        """Open spectral analysis dialog and optionally apply notch filter."""
        st = self.state
        if st.t is None or not st.analyze_chan or st.analyze_chan not in st.sweeps:
            self._show_warning("Spectral Analysis", "Please load data and select an analyze channel first.")
            return

        # Get current sweep data
        t, y = self._current_trace()
        if t is None or y is None:
            self._show_warning("Spectral Analysis", "No data available for current sweep.")
            return

        # Get stimulation spans for current sweep if available
        s = max(0, min(st.sweep_idx, self.navigation_manager._sweep_count() - 1))
        stim_spans = st.stim_spans_by_sweep.get(s, []) if st.stim_chan else []

        # Open dialog
        dlg = SpectralAnalysisDialog(
            parent=self, t=t, y=y, sr_hz=st.sr_hz, stim_spans=stim_spans,
            parent_window=self, use_zscore=self.use_zscore_normalization
        )
        telemetry.log_screen_view('Spectral Analysis Dialog', screen_class='analysis_dialog')
        if dlg.exec() == QDialog.DialogCode.Accepted:
            # Get filter parameters
            lower, upper = dlg.get_filter_params()
            print(f"[spectral-dialog] Dialog accepted. Filter params: lower={lower}, upper={upper}")

            if lower is not None and upper is not None:
                # Apply notch filter
                self.notch_filter_lower = lower
                self.notch_filter_upper = upper
                print(f"[notch-filter] Set notch filter: {lower:.2f} - {upper:.2f} Hz")

                # Clear processing cache to force recomputation with new filter
                st.proc_cache.clear()
                print(f"[notch-filter] Cleared processing cache")

                # Redraw to show filtered signal
                self.redraw_main_plot()
                print(f"[notch-filter] Main plot redrawn")

            else:
                print("[notch-filter] No filter applied (lower or upper is None)")
        else:
            print("[spectral-dialog] Dialog was not accepted (user cancelled or closed)")

    def on_outlier_thresh_clicked(self):
        """
        Open the multi-tab Analysis Options dialog to the Outlier Detection tab.
        (Replaces the old single OutlierMetricsDialog)
        """
        self._open_analysis_options(tab='outliers')
        telemetry.log_screen_view('Outlier Detection Options Dialog', screen_class='config_dialog')

    def on_gmm_clustering_clicked(self):
        """
        Open the multi-tab Analysis Options dialog to the Eup/Sniff Classification tab.
        (Replaces the old single GMMClusteringDialog)
        """
        self._open_analysis_options(tab='gmm')
        telemetry.log_screen_view('GMM Clustering Dialog', screen_class='analysis_dialog')

    def on_peak_navigator_clicked(self):
        """Open Advanced Peak Editor dialog for edge case review and curation."""
        # Always open dialog - it will handle the "no peaks" case internally
        dlg = AdvancedPeakEditorDialog(main_window=self, parent=self)
        telemetry.log_screen_view('Advanced Peak Editor Dialog', screen_class='curation_dialog')
        dlg.exec()  # Modal dialog
        telemetry.log_feature_used('advanced_peak_editor')

        # Refresh plot after dialog closes (in case user made edits)
        if hasattr(self.state, 'all_peaks_by_sweep') and self.state.all_peaks_by_sweep:
            self.redraw_main_plot()

    def _refresh_omit_button_label(self):
        """Update Omit button text based on whether current sweep is omitted."""
        # If in omit mode, always show "Omit (ON)" regardless of sweep state
        if getattr(self.editing_modes, "_omit_region_mode", False):
            self.OmitSweepButton.setText("Omit (ON)")
            self.OmitSweepButton.setToolTip("Click to exit omit region mode")
            return

        s = max(0, min(self.state.sweep_idx, self.navigation_manager._sweep_count() - 1))
        if s in self.state.omitted_sweeps:
            self.OmitSweepButton.setText("Un-omit")
            self.OmitSweepButton.setToolTip("This sweep will be excluded from saving and stats.")
        else:
            self.OmitSweepButton.setText("Omit")
            self.OmitSweepButton.setToolTip("Mark this sweep to be excluded from saving and stats.")

    def on_omit_sweep_clicked(self):
        """Handle Omit Sweep button - simple toggle to enter/exit omit region mode."""
        st = self.state
        if self.navigation_manager._sweep_count() == 0:
            return

        # Simple toggle: button checked state determines mode
        if self.OmitSweepButton.isChecked():
            # Entering omit region mode
            self.editing_modes._enter_omit_region_mode(remove_mode=False)
        else:
            # Exiting omit region mode
            self.editing_modes._exit_omit_region_mode()

    def _dim_axes_for_omitted(self, ax, label=True):
        """Delegate to PlotManager."""
        self.plot_manager.dim_axes_for_omitted(ax, label)


    ##################################################
    ##Save Data to File                             ##
    ##################################################
    def _load_save_dialog_history(self) -> dict:
        """Load autocomplete history for the Save Data dialog from QSettings."""
        return self.export_manager._load_save_dialog_history()

    def _update_save_dialog_history(self, vals: dict):
        """Update autocomplete history with new values from the Save Data dialog."""
        return self.export_manager._update_save_dialog_history(vals)

    def _sanitize_token(self, s: str) -> str:
        """Delegate to ExportManager."""
        return self.export_manager._sanitize_token(s)

    def _suggest_stim_string(self) -> str:
        """
        Build a stim name like '20Hz10s15ms' from detected stim metrics
        or '15msPulse' / '5sPulse' for single pulses.
        Rounding:
        - freq_hz -> nearest Hz
        - duration_s -> nearest second
        - pulse_width_s -> nearest millisecond (or nearest second if >1s)
        """
        return self.export_manager._suggest_stim_string()

    def on_save_analyzed_clicked(self):
        """Save analyzed data to disk after prompting for location/name."""
        return self.export_manager.on_save_analyzed_clicked()

    def on_view_summary_clicked(self):
        """Display interactive preview of the PDF summary without saving."""
        return self.export_manager.on_view_summary_clicked()

    def on_save_options_clicked(self):
        """Open dialog to configure which metrics to export."""
        from dialogs.export_options_dialog import ExportOptionsDialog

        # Get current export options from state (or None for defaults)
        current_options = getattr(self.state, 'export_metric_options', None)

        # Create and show dialog
        dialog = ExportOptionsDialog(self, current_options)

        # Connect signal to save changes
        def on_options_changed(new_options):
            self.state.export_metric_options = new_options
            self._log_status_message(f"Export options updated ({sum(new_options.values())} metrics enabled)", 2000)

        dialog.options_changed.connect(on_options_changed)
        dialog.exec()

    def _metric_keys_in_order(self):
        """Return metric keys in the UI order (from metrics.METRIC_SPECS)."""
        return self.export_manager._metric_keys_in_order()

    def _compute_metric_trace(self, key, t, y, sr_hz, peaks, breaths):
        """
        Call the metric function, passing expoffs if it exists.
        Falls back to legacy signature when needed.
        """
        return self.export_manager._compute_metric_trace(key, t, y, sr_hz, peaks, breaths)

    def _get_stim_masks(self, s: int):
        """
        Build (baseline_mask, stim_mask, post_mask) boolean arrays over st.t for sweep s.
        Uses union of all stim spans for 'stim'.
        """
        return self.export_manager._get_stim_masks(s)

    def _nanmean_sem(self, X, axis=0):
        """
        Robust mean/SEM that avoids NumPy RuntimeWarnings when there are
        0 or 1 finite values along the chosen axis.
        """
        return self.export_manager._nanmean_sem(X, axis)

    def _export_all_analyzed_data(self, preview_only=False, progress_dialog=None):
        """
        Exports (or previews) analyzed data.

        If preview_only=True: Shows interactive PDF preview dialog without saving files.
        If preview_only=False: Prompts for location/name and exports files.

        Exports:
        1) <base>_bundle.npz
            - Downsampled processed trace (kept sweeps only)
            - Downsampled y2 metric traces (all keys)
            - Peaks/breaths/sighs per kept sweep
            - Stim spans per kept sweep
            - Meta

        2) <base>_means_by_time.csv
            - t (relative to global stim start if present)
            - For each metric: optional per-sweep traces, mean, sem
            - Then the same block normalized by per-sweep baseline window (_norm)
            - Then the same block normalized by pooled eupneic baseline (_norm_eupnea)

        3) <base>_breaths.csv
            - Wide layout:
                RAW blocks:  ALL | BASELINE | STIM | POST
                NORM blocks: ALL | BASELINE | STIM | POST (per-sweep time-based)
                NORM_EUPNEA blocks: ALL | BASELINE | STIM | POST (pooled eupneic)
            - Includes `is_sigh` column (1 if any sigh peak in that breath interval)

        4) <base>_events.csv
            - Event intervals: stimulus on/off, apnea episodes, eupnea regions
            - Columns: sweep, event_type, start_time, end_time, duration
            - Times are relative to global stim start if present

        5) <base>_summary.pdf (or preview dialog if preview_only=True)
        """
        return self.export_manager._export_all_analyzed_data(preview_only, progress_dialog)

    def _mean_sem_1d(self, arr: np.ndarray):
        """Finite-only mean and SEM (ddof=1) for a 1D array. Returns (mean, sem).
        If no finite values -> (nan, nan). If only 1 finite value -> (mean, nan)."""
        return self.export_manager._mean_sem_1d(arr)

    def _save_metrics_summary_pdf(self, pdf_path, t_ds_csv, y2_ds_by_key, keys_for_csv, label_by_key, meta, stim_zero, stim_dur):
        """Delegate to ExportManager."""
        return self.export_manager._save_metrics_summary_pdf(pdf_path, t_ds_csv, y2_ds_by_key, keys_for_csv, label_by_key, meta, stim_zero, stim_dur)

    def _show_summary_preview_dialog(self, t_ds_csv, y2_ds_by_key, keys_for_csv, label_by_key, stim_zero, stim_dur):
        """Display interactive preview dialog with the three summary figures."""
        return self.export_manager._show_summary_preview_dialog(t_ds_csv, y2_ds_by_key, keys_for_csv, label_by_key, stim_zero, stim_dur)

    def _sigh_sample_indices(self, s: int, pks: np.ndarray | None) -> set[int]:
        """
        Return a set of SAMPLE indices (into st.t / y) for sigh-marked peaks on sweep s,
        regardless of how they were originally stored.

        Accepts any of these storage patterns per sweep:
        • sample indices (ints 0..N-1)
        • indices INTO the peaks list (ints 0..len(pks)-1), which we map via pks[idx]
        • times in seconds (floats), which we map to nearest sample via searchsorted
        • numpy array / list / set in any of the above forms
        """
        return self.export_manager._sigh_sample_indices(s, pks)

    def on_curation_choose_dir_clicked(self):
        # Get last used directory from settings, default to home
        last_dir = self.settings.value("curation_last_dir", str(Path.home()))
        
        # Ensure the path exists, otherwise fall back to home
        if not Path(last_dir).exists():
            last_dir = str(Path.home())
        
        base = QFileDialog.getExistingDirectory(
            self, "Choose a folder to scan", last_dir
        )
        
        if not base:
            return
        
        # Save the selected directory for next time
        self.settings.setValue("curation_last_dir", base)
        
        groups = self._scan_csv_groups(Path(base))
        self._populate_file_list_from_groups(groups)

    def _scan_csv_groups(self, base_dir: Path):
        """
        Walk base_dir recursively and group CSVs by common root:
        root + '_breaths.csv'
        root + '_timeseries.csv'
        root + '_events.csv'
        Returns a list of dicts: {key, root, dir, breaths, means, events}
        """
        groups = {}
        for dirpath, _, filenames in os.walk(str(base_dir)):
            for fn in filenames:
                lower = fn.lower()
                if not lower.endswith(".csv"):
                    continue

                kind = None
                if lower.endswith("_breaths.csv"):
                    root = fn[:-len("_breaths.csv")]
                    kind = "breaths"
                elif lower.endswith("_timeseries.csv"):
                    root = fn[:-len("_timeseries.csv")]
                    kind = "means"
                elif lower.endswith("_means_by_time.csv"):  # Legacy support
                    root = fn[:-len("_means_by_time.csv")]
                    kind = "means"
                elif lower.endswith("_events.csv"):
                    root = fn[:-len("_events.csv")]
                    kind = "events"

                if kind is None:
                    continue

                dir_p = Path(dirpath)
                key = str((dir_p / root).resolve()).lower()  # unique per dir+root (case-insensitive on Win)
                entry = groups.get(key)
                if entry is None:
                    entry = {"key": key, "root": root, "dir": dir_p, "breaths": None, "means": None, "events": None}
                    groups[key] = entry
                entry[kind] = str(dir_p / fn)

        # Return as a stable, sorted list
        return sorted(groups.values(), key=lambda e: (str(e["dir"]).lower(), e["root"].lower()))


    def _populate_file_list_from_groups(self, groups: list[dict]):
        """
        Fill left list (FileList) with one item per root. Display only name,
        store both full paths in UserRole for later consolidation.
        """
        from PyQt6.QtCore import Qt
        from PyQt6.QtWidgets import QListWidgetItem

        self.FileList.clear()
        # Do not clear the right list automatically so users don't lose selections:
        # self.FilestoConsolidateList.clear()

        for g in groups:
            root = g["root"]
            has_b = bool(g["breaths"])
            has_m = bool(g["means"])
            has_e = bool(g["events"])

            # Build suffix showing what files are present
            parts = []
            if has_b:
                parts.append("breaths")
            if has_m:
                parts.append("timeseries")
            if has_e:
                parts.append("events")

            if parts:
                suffix = f"[{' + '.join(parts)}]"
            else:
                # Shouldn't happen; skip if nothing is present
                continue

            item = QListWidgetItem(f"{root}  {suffix}")
            tt_lines = [f"Root: {root}", f"Dir:  {g['dir']}"]
            if g["breaths"]:
                tt_lines.append(f"breaths:    {g['breaths']}")
            if g["means"]:
                tt_lines.append(f"timeseries: {g['means']}")
            if g["events"]:
                tt_lines.append(f"events:     {g['events']}")
            item.setToolTip("\n".join(tt_lines))

            # Store full metadata for later use
            item.setData(Qt.ItemDataRole.UserRole, g)  # {'key', 'root', 'dir', 'breaths', 'means', 'events'}

            self.FileList.addItem(item)

        # Optional: sort visually
        self.FileList.sortItems()


    def _curation_scan_and_fill(self, root: Path):
        """Scan for matching CSVs and fill FileList with filenames (store full paths in item data)."""
        from PyQt6.QtWidgets import QListWidgetItem
        from PyQt6.QtCore import Qt

        # Clear existing items
        self.FileList.clear()

        # Patterns to include (recursive)
        patterns = ["*_breaths.csv", "*_timeseries.csv", "*_means_by_time.csv", "*_events.csv"]

        files = []
        try:
            for pat in patterns:
                files.extend(root.rglob(pat))
        except Exception as e:
            self._show_error("Scan error", f"Failed to scan folder:\n{root}\n\n{e}")
            return

        # Deduplicate & sort (by name, then path for stability)
        uniq = {}
        for p in files:
            try:
                # Only include files (ignore dirs, weird links)
                if p.is_file():
                    # keep all—even if names clash—because display is name-only,
                    # but we keep full path in item data and tooltip
                    uniq[str(p)] = p
            except Exception:
                pass

        files_sorted = sorted(uniq.values(), key=lambda x: (x.name.lower(), str(x).lower()))

        if not files_sorted:
            try:
                self._log_status_message("No matching CSV files found in the selected folder.", 4000)
            except Exception:
                pass
            return

        for p in files_sorted:
            item = QListWidgetItem(p.name)
            item.setToolTip(str(p))  # show full path on hover
            item.setData(Qt.ItemDataRole.UserRole, str(p))  # keep full path for later use
            self.FileList.addItem(item)

        # Optional: sort in the widget (already sorted, but harmless)
        self.FileList.sortItems()

    def _list_has_path(self, lw, full_path: str) -> bool:
        """Return True if any item in lw has UserRole == full_path."""
        for i in range(lw.count()):
            it = lw.item(i)
            if it and it.data(Qt.ItemDataRole.UserRole) == full_path:
                return True
        return False


    # ========================================
    # Consolidation Tab (Project Builder) Methods
    # ========================================

    def _populate_consolidation_source_list(self):
        """
        Populate the consolidation source list by scanning for analyzed CSV files.
        Uses the same logic as the Curation tab - looks for _breaths.csv, _timeseries.csv, _events.csv
        """
        self.consolidationSourceList.clear()

        # Determine which folder to scan
        if self._consolidation_custom_folder:
            scan_folder = self._consolidation_custom_folder
            self.consolidationSourceLabel.setText("Custom Folder")
            folder_str = str(scan_folder)
            if len(folder_str) > 30:
                folder_str = "..." + folder_str[-27:]
            self.consolidationFolderLabel.setText(folder_str)
            self.consolidationFolderLabel.setToolTip(str(scan_folder))
        elif self._project_directory:
            scan_folder = Path(self._project_directory)
            self.consolidationSourceLabel.setText("Project Files")
            self.consolidationFolderLabel.setText("")
        else:
            # No folder to scan
            self.consolidationSourceLabel.setText("No Folder")
            self.consolidationFolderLabel.setText("Use Browse or load a project")
            return

        # Scan for CSV groups (same as Curation tab)
        groups = self._scan_csv_groups(scan_folder)

        if not groups:
            self._log_status_message("No analyzed files found (looking for *_breaths.csv, *_timeseries.csv, *_events.csv)", 4000)
            return

        # Populate the list
        for g in groups:
            root = g["root"]
            has_b = bool(g["breaths"])
            has_m = bool(g["means"])
            has_e = bool(g["events"])

            # Build suffix showing what files are present
            parts = []
            if has_b:
                parts.append("breaths")
            if has_m:
                parts.append("timeseries")
            if has_e:
                parts.append("events")

            if not parts:
                continue

            suffix = f"[{' + '.join(parts)}]"
            display_name = f"{root}  {suffix}"

            # Build tooltip
            tt_lines = [f"Root: {root}", f"Dir:  {g['dir']}"]
            if g["breaths"]:
                tt_lines.append(f"breaths:    {g['breaths']}")
            if g["means"]:
                tt_lines.append(f"timeseries: {g['means']}")
            if g["events"]:
                tt_lines.append(f"events:     {g['events']}")

            item = QListWidgetItem(display_name)
            item.setToolTip("\n".join(tt_lines))

            # Store the same metadata format as Curation tab for compatibility
            item.setData(Qt.ItemDataRole.UserRole, g)  # {'key', 'root', 'dir', 'breaths', 'means', 'events'}

            self.consolidationSourceList.addItem(item)

        self._log_status_message(f"Found {len(groups)} analyzed file(s) ready for consolidation", 3000)

    def _on_consolidation_browse_clicked(self):
        """Handle browse button click - select a custom folder for consolidation."""
        folder = QFileDialog.getExistingDirectory(
            self, "Select Folder with Analyzed Files",
            str(self._project_directory) if self._project_directory else ""
        )

        if folder:
            self._consolidation_custom_folder = Path(folder)
            self.consolidationResetButton.show()  # Show reset button when using custom folder
            self._populate_consolidation_source_list()

    def _on_consolidation_reset_clicked(self):
        """Handle reset button click - return to project files."""
        self._consolidation_custom_folder = None
        self.consolidationResetButton.hide()  # Hide reset button when back to project files
        self._populate_consolidation_source_list()

    def _on_left_column_tab_changed(self, index: int):
        """Handle tab changes in the left column (Project Builder sub-tabs).
        Refreshes the Consolidation source list when switching to that tab.
        """
        # Get the widget at this index to check if it's the consolidation tab
        widget = self.leftColumnTabs.widget(index)
        if widget and widget.objectName() == 'consolidationContainer':
            # Only refresh from project files if not using a custom folder
            if not self._consolidation_custom_folder:
                self._populate_consolidation_source_list()

    def _filter_consolidation_source_list(self, text: str):
        """Filter the consolidation source list based on search text.
        Searches the root name, directory path, and display text.
        """
        search_terms = text.lower().strip().split()

        for i in range(self.consolidationSourceList.count()):
            item = self.consolidationSourceList.item(i)
            if not item:
                continue

            # Get metadata for comprehensive search
            metadata = item.data(Qt.ItemDataRole.UserRole)
            if isinstance(metadata, dict):
                # Build searchable text from CSV group metadata
                searchable = " ".join([
                    str(metadata.get('root', '')),
                    str(metadata.get('dir', '')),
                    item.text(),
                ]).lower()
            else:
                # Fallback to display text
                searchable = item.text().lower()

            # All search terms must match (AND logic)
            matches = all(term in searchable for term in search_terms) if search_terms else True
            item.setHidden(not matches)

    def _consolidation_move_all_right(self):
        """Move all visible items from source to consolidation list."""
        moved = 0
        skipped = 0
        for i in range(self.consolidationSourceList.count()):
            item = self.consolidationSourceList.item(i)
            if item and not item.isHidden():
                # Check if already in destination
                if not self._item_in_list(self.consolidationFilesList, item.data(Qt.ItemDataRole.UserRole)):
                    new_item = QListWidgetItem(item.text())
                    new_item.setData(Qt.ItemDataRole.UserRole, item.data(Qt.ItemDataRole.UserRole))
                    self.consolidationFilesList.addItem(new_item)
                    moved += 1
                else:
                    skipped += 1
        if moved or skipped:
            self.statusbar.showMessage(f"Added {moved} file(s) to consolidation. Skipped {skipped} duplicate(s).", 3000)

    def _consolidation_move_selected_right(self):
        """Move selected items from source to consolidation list."""
        moved = 0
        skipped = 0
        for item in self.consolidationSourceList.selectedItems():
            if not self._item_in_list(self.consolidationFilesList, item.data(Qt.ItemDataRole.UserRole)):
                new_item = QListWidgetItem(item.text())
                new_item.setData(Qt.ItemDataRole.UserRole, item.data(Qt.ItemDataRole.UserRole))
                self.consolidationFilesList.addItem(new_item)
                moved += 1
            else:
                skipped += 1
        if moved or skipped:
            self.statusbar.showMessage(f"Added {moved} file(s) to consolidation. Skipped {skipped} duplicate(s).", 3000)

    def _consolidation_move_selected_left(self):
        """Remove selected items from consolidation list."""
        items = self.consolidationFilesList.selectedItems()
        for item in items:
            row = self.consolidationFilesList.row(item)
            self.consolidationFilesList.takeItem(row)
        if items:
            self.statusbar.showMessage(f"Removed {len(items)} file(s) from consolidation.", 3000)

    def _consolidation_move_all_left(self):
        """Remove all items from consolidation list."""
        count = self.consolidationFilesList.count()
        self.consolidationFilesList.clear()
        if count:
            self.statusbar.showMessage(f"Removed all {count} file(s) from consolidation.", 3000)

    def _item_in_list(self, list_widget, user_data) -> bool:
        """Check if an item with given user data exists in the list.
        Uses the 'key' field from CSV group metadata for comparison.
        """
        # Get the key for comparison (unique identifier for CSV group)
        if isinstance(user_data, dict):
            check_key = user_data.get('key', '')
        else:
            check_key = str(user_data) if user_data else ''

        for i in range(list_widget.count()):
            item = list_widget.item(i)
            if not item:
                continue
            item_data = item.data(Qt.ItemDataRole.UserRole)
            if isinstance(item_data, dict):
                item_key = item_data.get('key', '')
            else:
                item_key = str(item_data) if item_data else ''

            if item_key and item_key == check_key:
                return True
        return False

    def _on_consolidation_save_clicked(self):
        """Handle consolidation save from Project Builder tab.
        Uses the same data format as the Curation tab for compatibility.
        """
        from PyQt6.QtCore import Qt

        # Get all items from the consolidation list
        items = []
        for i in range(self.consolidationFilesList.count()):
            item = self.consolidationFilesList.item(i)
            if item:
                items.append(item)

        if not items:
            QMessageBox.warning(self, "No Files", "No files selected for consolidation.")
            return

        # Extract file paths in the same format as Curation tab
        means_files = []
        breaths_files = []
        events_files = []

        for item in items:
            meta = item.data(Qt.ItemDataRole.UserRole) or {}
            root = meta.get("root", item.text())
            if meta.get("means"):
                means_files.append((root, Path(meta["means"])))
            if meta.get("breaths"):
                breaths_files.append((root, Path(meta["breaths"])))
            if meta.get("events"):
                events_files.append((root, Path(meta["events"])))

        if not means_files and not breaths_files and not events_files:
            QMessageBox.warning(self, "No CSV Files", "No CSV files found in the selected items.")
            return

        # Use the consolidation manager directly with the file lists
        self.consolidation_manager.consolidate_csv_files(means_files, breaths_files, events_files)

    def _propose_consolidated_filename(self, files: list) -> tuple[str, list[str]]:
        """Delegate to ConsolidationManager."""
        return self.consolidation_manager._propose_consolidated_filename(files)

    def on_consolidate_save_data_clicked(self):
        """Delegate to ConsolidationManager."""
        return self.consolidation_manager.on_consolidate_save_data_clicked()

    def _consolidate_breaths_histograms(self, files: list[tuple[str, Path]]) -> dict:
        """Delegate to ConsolidationManager."""
        return self.consolidation_manager._consolidate_breaths_histograms(files)

    def _consolidate_events(self, files: list[tuple[str, Path]]) -> pd.DataFrame:
        """Delegate to ConsolidationManager."""
        return self.consolidation_manager._consolidate_events(files)

    def _consolidate_stimulus(self, files: list[tuple[str, Path]]) -> tuple[pd.DataFrame, list[str]]:
        """Delegate to ConsolidationManager."""
        return self.consolidation_manager._consolidate_stimulus(files)

    def _try_load_npz_v2(self, npz_path: Path) -> dict | None:
        """Delegate to ConsolidationManager."""
        return self.consolidation_manager._try_load_npz_v2(npz_path)

    def _extract_timeseries_from_npz(self, npz_data: dict, metric: str, variant: str = 'raw') -> tuple[np.ndarray, np.ndarray]:
        """Delegate to ConsolidationManager."""
        return self.consolidation_manager._extract_timeseries_from_npz(npz_data, metric, variant)

    def _consolidate_from_npz_v2(self, npz_data_by_root: dict, files: list[tuple[str, Path]], metrics: list[str]) -> dict:
        """Delegate to ConsolidationManager."""
        return self.consolidation_manager._consolidate_from_npz_v2(npz_data_by_root, files, metrics)

    def _consolidate_means_files(self, files: list[tuple[str, Path]]) -> dict:
        """Delegate to ConsolidationManager."""
        return self.consolidation_manager._consolidate_means_files(files)

    def _consolidate_breaths_sighs(self, files: list[tuple[str, Path]]) -> pd.DataFrame:
        """Delegate to ConsolidationManager."""
        return self.consolidation_manager._consolidate_breaths_sighs(files)

    def _save_consolidated_to_excel(self, consolidated: dict, save_path: Path):
        """Delegate to ConsolidationManager."""
        return self.consolidation_manager._save_consolidated_to_excel(consolidated, save_path)

    def _add_events_charts(self, ws, header_row):
        """Delegate to ConsolidationManager."""
        return self.consolidation_manager._add_events_charts(ws, header_row)

    def _add_sighs_chart(self, ws, header_row):
        """Delegate to ConsolidationManager."""
        return self.consolidation_manager._add_sighs_chart(ws, header_row)

    # ========== PROJECT BUILDER METHODS ==========

    def _setup_master_file_list(self):
        """Set up the master file list using Model/View architecture."""
        from PyQt6.QtWidgets import QHeaderView, QVBoxLayout, QWidget
        from PyQt6.QtCore import Qt

        # NOTE: Project Organization section (right column) was removed from .ui file
        # The old experiment-based workflow widgets are no longer present

        # === Set up QTableView with Model/View architecture ===
        # The QTableView is already in the .ui file - just set up the model and delegates
        table = self.discoveredFilesTable

        # Create and set the model
        self._file_table_model = FileTableModel(self)
        table.setModel(self._file_table_model)

        # Configure table view properties (some may be set in .ui, but ensure they're correct)
        table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
        table.setSelectionMode(QAbstractItemView.SelectionMode.ExtendedSelection)
        table.setAlternatingRowColors(True)
        table.setSortingEnabled(False)  # We handle sorting ourselves
        table.setWordWrap(False)
        table.verticalHeader().setVisible(False)

        # Enable drag-drop column reordering
        header = table.horizontalHeader()
        header.setSectionsMovable(True)
        header.setDragEnabled(True)
        header.setDragDropMode(QAbstractItemView.DragDropMode.InternalMove)

        # Set up delegates for special columns
        self._setup_table_delegates(table)

        # Set column widths from column definitions
        for i, col_def in enumerate(self._file_table_model.get_visible_columns()):
            table.setColumnWidth(i, col_def.width)
            if col_def.fixed:
                header.setSectionResizeMode(i, QHeaderView.ResizeMode.Fixed)
            else:
                header.setSectionResizeMode(i, QHeaderView.ResizeMode.Interactive)

        # Connect signals
        table.doubleClicked.connect(self._on_table_double_clicked)
        table.clicked.connect(self._on_table_clicked)  # Single click for Notes column
        table.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        table.customContextMenuRequested.connect(self._on_master_list_context_menu)

        # Connect model signals
        self._file_table_model.cell_edited.connect(self._on_model_cell_edited)

        # Connect header click for sorting
        header.sectionClicked.connect(self._on_header_sort_clicked)

        # Connect header context menu for column visibility
        header.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        header.customContextMenuRequested.connect(self._on_header_context_menu)

        # Install event filter for resize events
        table.viewport().installEventFilter(self)
        self._table_resize_timer = None

        print("[project-builder] Master file list configured with Model/View architecture")

    def _setup_table_delegates(self, table: QTableView):
        """Set up item delegates for special column types."""
        model = self._file_table_model

        # Button delegate for actions column
        actions_col = model.get_column_index('actions')
        if actions_col >= 0:
            self._button_delegate = ButtonDelegate(table)
            self._button_delegate.analyze_clicked.connect(self._on_analyze_row)
            self._button_delegate.delete_clicked.connect(self._on_delete_row)
            self._button_delegate.info_clicked.connect(self._on_info_row)
            table.setItemDelegateForColumn(actions_col, self._button_delegate)

        # Status delegate
        status_col = model.get_column_index('status')
        if status_col >= 0:
            self._status_delegate = StatusDelegate(table)
            table.setItemDelegateForColumn(status_col, self._status_delegate)

        # Autocomplete delegates for editable columns with history
        for col_def in model.get_visible_columns():
            if col_def.editable and col_def.history_key:
                col_idx = model.get_column_index(col_def.key)
                if col_idx >= 0:
                    delegate = AutoCompleteDelegate(table, history_key=col_def.history_key)
                    table.setItemDelegateForColumn(col_idx, delegate)

    def _on_table_double_clicked(self, index):
        """Handle double-click on table - open file for analysis."""
        if not index.isValid():
            return
        row = index.row()
        col_key = self._file_table_model.get_column_key(index.column())

        # If double-clicking on an editable column, let the delegate handle it
        col_def = self._file_table_model.get_column_def(col_key) if col_key else None
        if col_def and col_def.editable:
            return  # Let editing happen

        # Otherwise, analyze the file
        self._analyze_file_at_row(row)

    def _on_table_clicked(self, index):
        """Handle single click on table - check for Notes column click."""
        if not index.isValid():
            return

        col_key = self._file_table_model.get_column_key(index.column())
        if col_key == 'linked_notes':
            row = index.row()
            row_data = self._file_table_model.get_row_data(row)
            if row_data:
                file_name = row_data.get('file_name', '')
                if file_name:
                    self._show_linked_notes_for_file(file_name)

    def _show_linked_notes_for_file(self, file_name: str):
        """Show a dialog with notes files that reference the given data file.

        Uses fuzzy matching (±5 numeric range) if no exact matches are found.
        """
        from pathlib import Path

        # Get the file stem for matching
        file_stem = Path(file_name).stem

        # Find notes files that reference this file (exact match)
        linked_notes = self._get_notes_for_abf(file_stem)
        is_fuzzy = False
        fuzzy_stems = []

        # If no exact matches, try fuzzy matching
        if not linked_notes:
            linked_notes, fuzzy_stems = self._get_fuzzy_notes_for_abf(file_stem)
            if linked_notes:
                is_fuzzy = True

        if not linked_notes:
            self._log_status_message(f"No notes found referencing '{file_name}'. Run 'Scan References' first.", 3000)
            return

        # Build info text with fuzzy match warning if applicable
        if is_fuzzy:
            # Find the closest matching stem numerically
            import re
            closest_stem = fuzzy_stems[0] if fuzzy_stems else file_stem

            # Extract numeric part from original stem
            orig_nums = re.findall(r'\d+', file_stem)
            if orig_nums and fuzzy_stems:
                orig_num = int(max(orig_nums, key=len))
                min_diff = float('inf')
                for stem in fuzzy_stems:
                    stem_nums = re.findall(r'\d+', stem)
                    if stem_nums:
                        stem_num = int(max(stem_nums, key=len))
                        diff = abs(stem_num - orig_num)
                        if diff < min_diff:
                            min_diff = diff
                            closest_stem = stem

            stems_str = ', '.join(fuzzy_stems[:3])  # Show first 3 matches
            if len(fuzzy_stems) > 3:
                stems_str += f" +{len(fuzzy_stems) - 3} more"
            info_text = (
                f'<span style="color: #FFA500;">⚠ FUZZY MATCH:</span> '
                f'No exact match for <b>{file_name}</b><br>'
                f'Found {len(linked_notes)} file(s) referencing nearby: <b>{stems_str}</b>'
            )
            title = f"Notes for: {file_name} (Fuzzy Match)"
            # Highlight the closest matching stem
            highlight_stem = closest_stem
        else:
            info_text = f"Found {len(linked_notes)} notes file(s) referencing <b>{file_name}</b>"
            title = f"Notes for: {file_name}"
            highlight_stem = file_stem

        # Use shared preview dialog
        self._show_notes_preview_dialog(
            files=[{'name': n['name'], 'path': n['path']} for n in linked_notes],
            title=title,
            info_text=info_text,
            highlight_stem=highlight_stem
        )

    def _show_notes_preview_dialog(self, files: list, title: str, info_text: str = None, highlight_stem: str = None):
        """Show a unified preview dialog for notes files.

        Args:
            files: List of dicts with 'name' and 'path' keys
            title: Dialog window title
            info_text: Optional HTML text shown above the preview (e.g., file count)
            highlight_stem: Optional ABF filename stem to highlight in previews
        """
        from PyQt6.QtWidgets import (QDialog, QVBoxLayout, QHBoxLayout, QLabel,
                                     QPushButton, QTabWidget, QWidget)
        from pathlib import Path

        if not files:
            return

        # Dark theme dialog stylesheet
        dark_dialog_style = """
            QDialog {
                background-color: #1e1e1e;
                color: #d4d4d4;
            }
            QLabel {
                color: #d4d4d4;
            }
            QPushButton {
                background-color: #3c3c3c;
                color: #d4d4d4;
                border: 1px solid #555555;
                padding: 6px 16px;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #4a4a4a;
            }
            QPushButton:pressed {
                background-color: #2a2a2a;
            }
            QTabWidget::pane {
                border: 1px solid #3e3e42;
                background-color: #1e1e1e;
            }
            QTabBar::tab {
                background-color: #2d2d30;
                color: #d4d4d4;
                padding: 8px 16px;
                border: 1px solid #3e3e42;
                border-bottom: none;
                margin-right: 2px;
            }
            QTabBar::tab:selected {
                background-color: #1e1e1e;
                border-bottom-color: #1e1e1e;
            }
            QTabBar::tab:hover:!selected {
                background-color: #3e3e42;
            }
        """

        # Create dialog
        dialog = QDialog(self)
        dialog.setWindowTitle(title)
        dialog.resize(1100, 750)
        dialog.setStyleSheet(dark_dialog_style)

        # Enable dark title bar on Windows
        self._enable_dark_title_bar(dialog)

        layout = QVBoxLayout(dialog)
        layout.setContentsMargins(8, 8, 8, 8)

        # Info label (optional)
        if info_text:
            info_label = QLabel(info_text)
            info_label.setStyleSheet("font-size: 11pt; padding: 4px;")
            layout.addWidget(info_label)

        # Store file info for "Open in External App" button
        self._preview_dialog_files = {}

        # Single file: no tabs needed, just show the preview directly
        if len(files) == 1:
            file_info = files[0]
            preview_widget = self._create_note_preview_widget(
                {'name': file_info['name'], 'path': file_info['path']},
                highlight_stem or ''
            )
            layout.addWidget(preview_widget, stretch=1)
            self._preview_dialog_files[0] = file_info
            current_file_getter = lambda: self._preview_dialog_files.get(0)
        else:
            # Multiple files: use tabs
            tab_widget = QTabWidget()

            for idx, file_info in enumerate(files):
                preview_widget = self._create_note_preview_widget(
                    {'name': file_info['name'], 'path': file_info['path']},
                    highlight_stem or ''
                )
                tab_name = file_info['name']
                if len(tab_name) > 30:
                    tab_name = tab_name[:27] + "..."
                tab_widget.addTab(preview_widget, tab_name)
                self._preview_dialog_files[idx] = file_info

            layout.addWidget(tab_widget, stretch=1)
            current_file_getter = lambda: self._preview_dialog_files.get(tab_widget.currentIndex())

        # Bottom button row
        button_layout = QHBoxLayout()

        open_btn = QPushButton("Open in External App")
        open_btn.setToolTip("Open the file in its default application")

        def open_current():
            import os
            file_info = current_file_getter()
            if file_info and file_info.get('path'):
                os.startfile(file_info['path'])

        open_btn.clicked.connect(open_current)
        button_layout.addWidget(open_btn)

        button_layout.addStretch()

        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dialog.reject)
        button_layout.addWidget(close_btn)

        layout.addLayout(button_layout)

        dialog.exec()

    def _create_note_preview_widget(self, note_info: dict, abf_stem: str, search_term: str = None):
        """Create a preview widget for a single note file with ABF highlighting.

        Args:
            note_info: Dict with 'path' key pointing to the notes file
            abf_stem: Primary highlight term (ABF filename) - shown in green
            search_term: Optional secondary search term - shown in orange/yellow
        """
        from PyQt6.QtWidgets import QWidget, QVBoxLayout, QTabWidget, QTableWidget, QTableWidgetItem, QHeaderView, QTextEdit, QLabel
        from PyQt6.QtCore import Qt
        from PyQt6.QtGui import QColor, QBrush
        from pathlib import Path

        container = QWidget()
        layout = QVBoxLayout(container)
        layout.setContentsMargins(0, 0, 0, 0)

        file_path = Path(note_info['path'])

        # Add filename header
        filename_label = QLabel(f"<b>File:</b> {file_path.name}")
        filename_label.setStyleSheet("QLabel { color: #9cdcfe; padding: 4px 8px; background-color: #2d2d30; border-radius: 3px; }")
        filename_label.setToolTip(str(file_path))
        layout.addWidget(filename_label)

        if not file_path.exists():
            text_edit = QTextEdit()
            text_edit.setPlainText(f"File not found: {file_path}")
            text_edit.setReadOnly(True)
            text_edit.setStyleSheet("QTextEdit { background-color: #1e1e1e; color: #d4d4d4; border: none; }")
            layout.addWidget(text_edit)
            return container

        suffix = file_path.suffix.lower()

        try:
            if suffix in ['.xlsx', '.xls']:
                import pandas as pd
                sheets = pd.read_excel(file_path, sheet_name=None, nrows=500)

                if sheets:
                    # Multiple sheets -> nested tabs
                    sheet_tabs = QTabWidget()
                    sheet_tabs.setStyleSheet("""
                        QTabWidget::pane { border: 1px solid #3e3e42; background-color: #1e1e1e; }
                        QTabBar::tab { background-color: #2d2d30; color: #d4d4d4; padding: 6px 12px; border: 1px solid #3e3e42; }
                        QTabBar::tab:selected { background-color: #1e1e1e; }
                    """)

                    # Sort sheets: matching sheets first, then others
                    sheet_items = list(sheets.items())
                    matching_sheets = []
                    other_sheets = []
                    for sheet_name, df in sheet_items:
                        if self._df_contains_stem(df, abf_stem):
                            matching_sheets.append((sheet_name, df, True))
                        else:
                            other_sheets.append((sheet_name, df, False))

                    # Add matching sheets first, then others
                    for sheet_name, df, has_match in matching_sheets + other_sheets:
                        table = self._create_highlighted_table(df, abf_stem, search_term)
                        tab_name = f"★ {sheet_name}" if has_match else sheet_name
                        sheet_tabs.addTab(table, tab_name)

                    layout.addWidget(sheet_tabs)
                else:
                    text_edit = QTextEdit()
                    text_edit.setPlainText("[Workbook is empty]")
                    text_edit.setReadOnly(True)
                    layout.addWidget(text_edit)

            elif suffix == '.csv':
                import pandas as pd
                df = pd.read_csv(file_path, nrows=500, encoding='utf-8', on_bad_lines='skip')
                table = self._create_highlighted_table(df, abf_stem, search_term)
                layout.addWidget(table)

            elif suffix == '.txt':
                content = file_path.read_text(encoding='utf-8', errors='replace')
                text_edit = self._create_highlighted_text(content, abf_stem, search_term)
                layout.addWidget(text_edit)

            elif suffix == '.docx':
                try:
                    from docx import Document
                    doc = Document(str(file_path))

                    # Collect paragraphs
                    content_parts = []
                    for p in doc.paragraphs:
                        if p.text.strip():
                            content_parts.append(p.text)

                    # Also collect table contents (where ABF references often appear)
                    for table in doc.tables:
                        table_rows = []
                        for row in table.rows:
                            row_cells = [cell.text.strip() for cell in row.cells]
                            table_rows.append(' | '.join(row_cells))
                        if table_rows:
                            content_parts.append('\n--- Table ---')
                            content_parts.extend(table_rows)
                            content_parts.append('--- End Table ---\n')

                    content = '\n\n'.join(content_parts)
                    text_edit = self._create_highlighted_text(content, abf_stem, search_term)
                    layout.addWidget(text_edit)
                except ImportError:
                    text_edit = QTextEdit()
                    text_edit.setPlainText("python-docx not installed. Cannot preview .docx files.\n\nUse 'Open in External App' to view.")
                    text_edit.setReadOnly(True)
                    text_edit.setStyleSheet("QTextEdit { background-color: #1e1e1e; color: #d4d4d4; border: none; }")
                    layout.addWidget(text_edit)

            else:
                text_edit = QTextEdit()
                text_edit.setPlainText(f"Preview not supported for {suffix} files.\n\nUse 'Open in External App' to view.")
                text_edit.setReadOnly(True)
                text_edit.setStyleSheet("QTextEdit { background-color: #1e1e1e; color: #d4d4d4; border: none; }")
                layout.addWidget(text_edit)

        except Exception as e:
            text_edit = QTextEdit()
            text_edit.setPlainText(f"Error loading preview: {e}")
            text_edit.setReadOnly(True)
            text_edit.setStyleSheet("QTextEdit { background-color: #1e1e1e; color: #d4d4d4; border: none; }")
            layout.addWidget(text_edit)

        return container

    def _create_highlighted_table(self, df, abf_stem: str, search_term: str = None):
        """Create a QTableWidget from DataFrame with dual highlighting.

        Args:
            df: DataFrame to display
            abf_stem: Primary highlight term (ABF filename) - shown in green
            search_term: Optional secondary search term - shown in orange/yellow
        """
        from PyQt6.QtWidgets import QTableWidget, QTableWidgetItem, QHeaderView
        from PyQt6.QtGui import QColor, QBrush

        df = df.fillna('')

        table = QTableWidget()
        table.setRowCount(len(df))
        table.setColumnCount(len(df.columns))
        table.setHorizontalHeaderLabels([str(col) for col in df.columns])

        # Find rows containing the ABF reference (primary) or search term (secondary)
        primary_rows = set()
        secondary_rows = set()
        for row_idx in range(len(df)):
            for col in df.columns:
                val = str(df.iloc[row_idx][col])
                if abf_stem and abf_stem.lower() in val.lower():
                    primary_rows.add(row_idx)
                if search_term and search_term.lower() in val.lower():
                    secondary_rows.add(row_idx)

        # Populate table
        for row_idx in range(len(df)):
            for col_idx, col in enumerate(df.columns):
                value = df.iloc[row_idx, col_idx]
                value_str = str(value) if value != '' else ''
                item = QTableWidgetItem(value_str)

                # Primary highlight (ABF stem) - green background
                if row_idx in primary_rows:
                    item.setBackground(QBrush(QColor(60, 90, 60)))
                    if abf_stem and abf_stem.lower() in value_str.lower():
                        font = item.font()
                        font.setBold(True)
                        item.setFont(font)
                        item.setForeground(QBrush(QColor(150, 255, 150)))

                # Secondary highlight (search term) - orange/yellow
                # If both match, secondary takes precedence for text color
                if search_term and search_term.lower() in value_str.lower():
                    if row_idx not in primary_rows:
                        # Only search term matches - use orange background
                        item.setBackground(QBrush(QColor(90, 70, 40)))
                    font = item.font()
                    font.setBold(True)
                    item.setFont(font)
                    item.setForeground(QBrush(QColor(255, 200, 100)))  # Orange/yellow text

                table.setItem(row_idx, col_idx, item)

        table.setAlternatingRowColors(True)
        table.setStyleSheet("""
            QTableWidget {
                background-color: #1e1e1e;
                alternate-background-color: #252526;
                color: #d4d4d4;
                gridline-color: #3e3e42;
                border: none;
            }
            QHeaderView::section {
                background-color: #2d2d30;
                color: #d4d4d4;
                padding: 6px;
                border: 1px solid #3e3e42;
                font-weight: bold;
            }
        """)

        header = table.horizontalHeader()
        # Resize columns to fit content, with last column stretching to fill space
        header.setSectionResizeMode(QHeaderView.ResizeMode.ResizeToContents)
        header.setStretchLastSection(True)  # Last column fills remaining width
        # Allow wider columns for better readability (up to 500px)
        header.setMaximumSectionSize(500)
        # Minimum size so narrow columns are still readable
        header.setMinimumSectionSize(50)

        # Scroll to first highlighted row (prefer primary, fallback to secondary)
        scroll_rows = primary_rows or secondary_rows
        if scroll_rows:
            first_row = min(scroll_rows)
            table.scrollToItem(table.item(first_row, 0))

        return table

    def _create_highlighted_text(self, content: str, abf_stem: str, search_term: str = None):
        """Create a QTextEdit with dual highlighting and scrolled to first match.

        Args:
            content: Text content to display
            abf_stem: Primary highlight term (ABF filename) - shown in green
            search_term: Optional secondary search term - shown in orange/yellow
        """
        from PyQt6.QtWidgets import QTextEdit
        from PyQt6.QtGui import QTextCursor
        from PyQt6.QtCore import QTimer
        import re

        text_edit = QTextEdit()
        text_edit.setReadOnly(True)

        # Create patterns for both highlight terms
        abf_pattern = re.compile(re.escape(abf_stem), re.IGNORECASE) if abf_stem else None
        search_pattern = re.compile(re.escape(search_term), re.IGNORECASE) if search_term else None

        # Process line by line
        lines = content.split('\n')
        highlighted_lines = []

        for line in lines:
            has_primary = abf_stem and abf_stem.lower() in line.lower()
            has_secondary = search_term and search_term.lower() in line.lower()

            if has_primary or has_secondary:
                line_html = line

                # Apply primary highlighting (green) - ABF stem
                if has_primary and abf_pattern:
                    line_html = abf_pattern.sub(
                        r'<span style="color: #90ff90; font-weight: bold;">\g<0></span>',
                        line_html
                    )

                # Apply secondary highlighting (orange) - search term
                # Need to be careful not to double-highlight if terms overlap
                if has_secondary and search_pattern:
                    # Only highlight search term if it doesn't overlap with ABF stem
                    if not has_primary or (abf_stem.lower() not in search_term.lower() and search_term.lower() not in abf_stem.lower()):
                        line_html = search_pattern.sub(
                            r'<span style="color: #ffcc66; font-weight: bold; background-color: #5a4a2a;">\g<0></span>',
                            line_html
                        )

                # Determine background color based on what matched
                if has_primary and has_secondary:
                    # Both match - use a mixed color
                    bg_color = "#4a5a4a"
                elif has_primary:
                    bg_color = "#3d5a3d"
                else:
                    bg_color = "#5a4a2a"

                highlighted_lines.append(
                    f'<div style="background-color: {bg_color}; padding: 2px 4px; margin: 1px 0;">{line_html}</div>'
                )
            else:
                highlighted_lines.append(f'<div>{line}</div>')

        highlighted = ''.join(highlighted_lines)
        text_edit.setHtml(f"<div style='color: #d4d4d4; font-family: Consolas, monospace; font-size: 10pt;'>{highlighted}</div>")

        text_edit.setStyleSheet("""
            QTextEdit {
                background-color: #1e1e1e;
                border: none;
                font-size: 10pt;
            }
        """)

        # Scroll to first match after widget is shown (prefer primary, fallback to secondary)
        def scroll_to_match():
            # Safety check: widget may have been deleted if user is rapidly typing in search
            try:
                # Check if widget is still valid (accessing any property will throw if deleted)
                if not text_edit.isVisible():
                    return
            except RuntimeError:
                # Widget has been deleted (C/C++ object no longer exists)
                return

            # Find position in original content (case-insensitive)
            match_pos = -1
            if abf_stem:
                match_pos = content.lower().find(abf_stem.lower())
            if match_pos < 0 and search_term:
                match_pos = content.lower().find(search_term.lower())

            if match_pos >= 0:
                # Calculate approximate scroll position as percentage of document
                total_len = len(content)
                if total_len > 0:
                    scroll_percent = match_pos / total_len
                    try:
                        scrollbar = text_edit.verticalScrollBar()
                        max_scroll = scrollbar.maximum()
                        # Scroll to position (with some offset to show context above)
                        target_pos = int(max_scroll * scroll_percent)
                        # Subtract some pixels to show context above the match
                        target_pos = max(0, target_pos - 50)
                        scrollbar.setValue(target_pos)
                    except RuntimeError:
                        # Widget deleted during scroll operation
                        pass

        # Use longer delay to ensure dialog is fully rendered
        QTimer.singleShot(300, scroll_to_match)

        return text_edit

    def _df_contains_stem(self, df, abf_stem: str) -> bool:
        """Check if a DataFrame contains the ABF stem anywhere."""
        for col in df.columns:
            for val in df[col].astype(str):
                if abf_stem.lower() in val.lower():
                    return True
        return False

    def _get_notes_for_abf(self, abf_stem: str) -> list:
        """Find notes files that reference a specific ABF file.

        Args:
            abf_stem: The ABF filename stem (without extension)

        Returns:
            List of note_info dicts that reference this ABF file
        """
        linked = []

        if not hasattr(self, '_notes_files_data') or not self._notes_files_data:
            return linked

        for note_info in self._notes_files_data:
            matches = note_info.get('matches', [])
            if abf_stem in matches:
                linked.append(note_info)

        return linked

    def _get_fuzzy_notes_for_abf(self, abf_stem: str, search_range: int = 5) -> tuple:
        """Find notes files that reference numeric neighbors of an ABF file.

        Used when exact match fails. Searches for files with similar numeric patterns
        within ±search_range of the original filename's numeric portion.

        Args:
            abf_stem: The ABF filename stem (without extension), e.g., "25708007"
            search_range: How far to search (default ±5)

        Returns:
            Tuple of (list of note_info dicts, list of matched_stems that were fuzzy matched)
        """
        import re

        if not hasattr(self, '_notes_files_data') or not self._notes_files_data:
            return [], []

        # Extract numeric portion from filename
        # Handle patterns like: "25708007", "mouse1_25708007", "25708007_trial1"
        # We want the longest numeric sequence (likely the file ID)
        numeric_matches = re.findall(r'\d+', abf_stem)
        if not numeric_matches:
            return [], []

        # Use the longest numeric sequence as the ID to match
        main_numeric = max(numeric_matches, key=len)
        numeric_idx = abf_stem.find(main_numeric)

        # Generate candidate stems within range
        candidates = set()
        try:
            base_num = int(main_numeric)
            for delta in range(-search_range, search_range + 1):
                if delta == 0:
                    continue  # Skip exact match (already checked)
                candidate_num = base_num + delta
                if candidate_num >= 0:
                    # Preserve leading zeros and surrounding text
                    candidate_str = str(candidate_num).zfill(len(main_numeric))
                    candidate_stem = abf_stem[:numeric_idx] + candidate_str + abf_stem[numeric_idx + len(main_numeric):]
                    candidates.add(candidate_stem)
        except ValueError:
            return [], []

        # Search notes files for these candidates
        linked = []
        matched_stems = []

        for note_info in self._notes_files_data:
            matches = note_info.get('matches', [])
            for match in matches:
                if match in candidates:
                    if note_info not in linked:
                        linked.append(note_info)
                    if match not in matched_stems:
                        matched_stems.append(match)

        return linked, matched_stems

    def _update_linked_notes_column(self):
        """Update the linked_notes column in file table based on scanned notes.

        Uses fuzzy matching (±5 numeric range) when exact matches aren't found.
        Fuzzy matches are displayed with a '~' prefix to indicate caution.
        """
        if not hasattr(self, '_file_table_model') or not self._file_table_model:
            return

        if not hasattr(self, '_notes_files_data') or not self._notes_files_data:
            return

        from pathlib import Path

        # Build a reverse index: ABF stem -> count of notes referencing it
        abf_notes_count = {}
        for note_info in self._notes_files_data:
            matches = note_info.get('matches', [])
            for abf_stem in matches:
                abf_notes_count[abf_stem] = abf_notes_count.get(abf_stem, 0) + 1

        fuzzy_match_count = 0

        # Update each row in the file table
        for row in range(self._file_table_model.rowCount()):
            row_data = self._file_table_model.get_row_data(row)
            if row_data:
                file_name = row_data.get('file_name', '')
                if file_name:
                    file_stem = Path(file_name).stem
                    exact_count = abf_notes_count.get(file_stem, 0)

                    # Check for fuzzy match if no exact matches
                    is_fuzzy = False
                    fuzzy_stems = []
                    if exact_count == 0:
                        fuzzy_notes, fuzzy_stems = self._get_fuzzy_notes_for_abf(file_stem)
                        if fuzzy_notes:
                            exact_count = len(fuzzy_notes)
                            is_fuzzy = True
                            fuzzy_match_count += 1

                    # Format display: "~N" for fuzzy, "N" for exact, "0" for no matches
                    if exact_count > 0:
                        display_value = f"~{exact_count}" if is_fuzzy else str(exact_count)
                    else:
                        display_value = '0'

                    # Update model data (including fuzzy info for tooltips)
                    self._file_table_model.set_cell_value(row, 'linked_notes', display_value)
                    self._file_table_model.set_cell_value(row, 'linked_notes_fuzzy', is_fuzzy)
                    self._file_table_model.set_cell_value(row, 'linked_notes_fuzzy_stems', fuzzy_stems)

                    # Also update master file list with detailed info
                    if row < len(self._master_file_list):
                        self._master_file_list[row]['linked_notes'] = exact_count
                        self._master_file_list[row]['linked_notes_fuzzy'] = is_fuzzy
                        self._master_file_list[row]['linked_notes_fuzzy_stems'] = fuzzy_stems

        fuzzy_msg = f" ({fuzzy_match_count} fuzzy)" if fuzzy_match_count > 0 else ""
        print(f"[notes] Updated linked_notes for {self._file_table_model.rowCount()} files{fuzzy_msg}")

    def _on_model_cell_edited(self, row: int, column_key: str, value):
        """Handle cell edits from the model."""
        # Update master file list
        if row < len(self._master_file_list):
            self._master_file_list[row][column_key] = value

            # Update autocomplete history for certain columns
            if column_key == 'experiment' and value:
                self._update_experiment_history(str(value))
            elif column_key == 'strain' and value:
                self._update_autocomplete_history('strain_history', str(value))

            # Autosave project
            self._project_autosave()

    def _on_analyze_row(self, row: int):
        """Handle analyze button click."""
        self._analyze_file_at_row(row)

    def _on_delete_row(self, row: int):
        """Handle delete button click - remove row from table."""
        if row < len(self._master_file_list):
            del self._master_file_list[row]
            self._file_table_model.remove_row(row)
            self._project_autosave()

    def _on_info_row(self, row: int):
        """Handle info button click - show file details."""
        if row < len(self._master_file_list):
            task = self._master_file_list[row]
            file_path = task.get('file_path', 'Unknown')
            protocol = task.get('protocol', 'Unknown')
            channels = task.get('channel_count', 0)
            sweeps = task.get('sweep_count', 0)

            info_text = f"""File: {file_path}
Protocol: {protocol}
Channels: {channels}
Sweeps: {sweeps}"""

            QMessageBox.information(self, "File Information", info_text)

    def _on_header_context_menu(self, pos):
        """Show context menu for column visibility and custom columns."""
        from PyQt6.QtWidgets import QMenu
        from PyQt6.QtCore import QPoint

        header = self.discoveredFilesTable.horizontalHeader()
        menu = QMenu(self)

        # Add show/hide options for each column
        all_cols = self._file_table_model.get_all_column_defs()
        for col_def in all_cols:
            if col_def.key == 'actions':  # Don't allow hiding actions column
                continue
            action = menu.addAction(col_def.header or col_def.key)
            action.setCheckable(True)
            col_idx = self._file_table_model.get_column_index(col_def.key)
            action.setChecked(col_idx >= 0)  # Visible if index >= 0
            action.setData(col_def.key)
            action.triggered.connect(lambda checked, key=col_def.key: self._toggle_column_visibility(key, checked))

        menu.addSeparator()

        # Custom columns submenu
        custom_menu = menu.addMenu("Custom Columns")
        add_custom_action = custom_menu.addAction("+ Add Custom Column...")
        add_custom_action.triggered.connect(self._add_custom_column_dialog)

        # List existing custom columns with remove option
        custom_cols = self._file_table_model.get_custom_columns()
        if custom_cols:
            custom_menu.addSeparator()
            for col_def in custom_cols:
                remove_action = custom_menu.addAction(f"Remove '{col_def.header}'")
                remove_action.triggered.connect(
                    lambda checked, key=col_def.key: self._remove_custom_column(key)
                )

        menu.addSeparator()

        # Reset to default
        reset_action = menu.addAction("Reset Column Order")
        reset_action.triggered.connect(self._reset_column_order)

        # Show menu at cursor position
        menu.exec(header.mapToGlobal(pos))

    def _toggle_column_visibility(self, key: str, visible: bool):
        """Toggle visibility of a column."""
        if visible:
            self._file_table_model.show_column(key)
        else:
            self._file_table_model.hide_column(key)

    def _add_custom_column_dialog(self):
        """Show dialog to add a custom column."""
        from PyQt6.QtWidgets import QInputDialog

        name, ok = QInputDialog.getText(
            self, "Add Custom Column",
            "Enter column name (e.g., 'Notes', 'Coordinates', 'Site'):"
        )

        if ok and name.strip():
            name = name.strip()
            key = self._file_table_model.add_custom_column(name)
            if key:
                self._log_status_message(f"✓ Added custom column: {name}", 2000)
                # Refresh table to show new column
                self._rebuild_table_from_master_list()
                self._auto_fit_table_columns()
                self._project_autosave()
            else:
                self._show_warning("Column Exists", f"A column named '{name}' already exists.")

    def _remove_custom_column(self, key: str):
        """Remove a custom column."""
        from PyQt6.QtWidgets import QMessageBox

        # Confirm removal
        result = QMessageBox.question(
            self, "Remove Custom Column",
            f"Remove this custom column?\n\nThis will delete all data in this column.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )

        if result == QMessageBox.StandardButton.Yes:
            if self._file_table_model.remove_custom_column(key):
                # Also remove from master file list
                for task in self._master_file_list:
                    task.pop(key, None)
                self._log_status_message(f"✓ Removed custom column", 2000)
                self._rebuild_table_from_master_list()
                self._auto_fit_table_columns()
                self._project_autosave()

    def _reset_column_order(self):
        """Reset column order to default."""
        self._file_table_model.reset_column_order()
        # Re-setup delegates after reset
        self._setup_table_delegates(self.discoveredFilesTable)

    def _analyze_file_at_row(self, row: int):
        """Load and analyze file at given row."""
        if row >= len(self._master_file_list):
            return

        task = self._master_file_list[row]
        file_path = task.get('file_path')
        if not file_path:
            return

        from pathlib import Path
        if not Path(file_path).exists():
            self._show_warning("File Not Found", f"File no longer exists:\n{file_path}")
            return

        print(f"[master-list] Opening file for analysis: {file_path}")

        # Track which row is being analyzed
        self._active_master_list_row = row

        # Store pending channel selections from Project Builder
        # These will be used after file loads to override auto-detection
        self._pending_analysis_channel = task.get('channel', '')
        self._pending_stim_channels = task.get('stim_channels', [])

        # Log what we're planning to select
        if self._pending_analysis_channel or self._pending_stim_channels:
            print(f"[master-list] Pre-selecting: analysis={self._pending_analysis_channel}, stim={self._pending_stim_channels}")

        # Load the file
        self.load_file(Path(file_path))

        # Switch to Analysis tab (tab widget is called 'Tabs' in the UI file)
        # Tab 0 = Project Builder, Tab 1 = Analysis
        if hasattr(self, 'Tabs'):
            self.Tabs.setCurrentIndex(1)  # Analysis tab is index 1
            print("[master-list] Switched to Analysis tab")

    def _update_autocomplete_history(self, history_key: str, value: str):
        """Add a value to autocomplete history."""
        from PyQt6.QtCore import QSettings
        settings = QSettings("PhysioMetrics", "BreathAnalysis")
        history = settings.value(f"autocomplete/{history_key}", [])
        if not isinstance(history, list):
            history = []
        if value in history:
            history.remove(value)
        history.insert(0, value)
        history = history[:50]
        settings.setValue(f"autocomplete/{history_key}", history)

    # ========== COMPATIBILITY LAYER ==========
    # These methods provide backward compatibility with old QTableWidget patterns

    def _get_table_cell_value(self, row: int, column_key: str):
        """Get cell value from model by column key (replaces table.item(row, col).text())."""
        return self._file_table_model.get_cell_value(row, column_key)

    def _set_table_cell_value(self, row: int, column_key: str, value):
        """Set cell value in model by column key (replaces table.setItem(row, col, item))."""
        self._file_table_model.set_cell_value(row, column_key, value)
        # Also update master file list
        if row < len(self._master_file_list):
            self._master_file_list[row][column_key] = value

    def _get_table_row_count(self) -> int:
        """Get number of rows in table."""
        return self._file_table_model.rowCount()

    def _clear_table(self):
        """Clear all data from table."""
        self._file_table_model.clear()

    def _format_exports_summary(self, task: dict) -> str:
        """
        Format a compact summary of exports for display in the Exports column.

        Example output: "1 PDF, 3 CSV, 1 NPZ"
        """
        exports = task.get('exports', {})
        if not exports:
            return ''

        parts = []

        # Count CSVs
        csv_count = sum([
            1 if exports.get('timeseries_csv') else 0,
            1 if exports.get('breaths_csv') else 0,
            1 if exports.get('events_csv') else 0,
        ])

        if exports.get('pdf'):
            parts.append('1 PDF')
        if csv_count > 0:
            parts.append(f'{csv_count} CSV')
        if exports.get('npz'):
            parts.append('1 NPZ')
        if exports.get('ml_training'):
            parts.append('ML')
        if exports.get('session_state'):
            parts.append('Session')

        return ', '.join(parts) if parts else ''

    def _format_exports_tooltip(self, task: dict) -> str:
        """
        Format a detailed tooltip for the Exports column.

        Shows what was exported, when, where, and app version.
        """
        exports = task.get('exports', {})
        export_path = task.get('export_path', '')
        export_date = task.get('export_date', '')
        export_version = task.get('export_version', '')

        if not exports and not export_path:
            return 'No exports yet'

        lines = []

        # Header with date
        if export_date:
            # Format ISO date nicely
            try:
                from datetime import datetime
                dt = datetime.fromisoformat(export_date)
                date_str = dt.strftime('%Y-%m-%d %H:%M')
            except:
                date_str = export_date
            lines.append(f"Exported: {date_str}")

        # App version
        if export_version:
            lines.append(f"App version: {export_version}")

        # Export path
        if export_path:
            lines.append(f"Location: {export_path}")

        lines.append('')  # Blank line

        # List of exports
        lines.append("Files saved:")
        if exports.get('pdf'):
            lines.append("  ✓ Summary PDF")
        if exports.get('timeseries_csv'):
            lines.append("  ✓ Timeseries CSV")
        if exports.get('breaths_csv'):
            lines.append("  ✓ Breaths CSV")
        if exports.get('events_csv'):
            lines.append("  ✓ Events CSV")
        if exports.get('npz'):
            lines.append("  ✓ NPZ Bundle")
        if exports.get('session_state'):
            lines.append("  ✓ Session State")
        if exports.get('ml_training'):
            lines.append("  ✓ ML Training Data")

        return '\n'.join(lines)

    def _create_exports_table_item(self, task: dict):
        """Create a QTableWidgetItem for the Exports column with summary and tooltip."""
        from PyQt6.QtWidgets import QTableWidgetItem

        summary = self._format_exports_summary(task)
        tooltip = self._format_exports_tooltip(task)

        item = QTableWidgetItem(summary)
        item.setToolTip(tooltip)

        return item

    def _on_header_sort_clicked(self, column):
        """
        Custom sorting that keeps parent rows with their sub-rows.
        Sorts groups by the parent's value in the clicked column.
        """
        if not self._master_file_list:
            return

        # Track sort direction (toggle on repeated clicks)
        if not hasattr(self, '_sort_column'):
            self._sort_column = -1
            self._sort_ascending = True

        if self._sort_column == column:
            self._sort_ascending = not self._sort_ascending
        else:
            self._sort_column = column
            self._sort_ascending = True

        # Group tasks by parent file path
        groups = {}  # file_path -> {'parent': task, 'sub_rows': [tasks]}
        parent_order = []  # Track order of parents for sorting

        for task in self._master_file_list:
            file_path = str(task.get('file_path', ''))
            is_sub_row = task.get('is_sub_row', False)

            if file_path not in groups:
                groups[file_path] = {'parent': None, 'sub_rows': []}
                parent_order.append(file_path)

            if is_sub_row:
                groups[file_path]['sub_rows'].append(task)
            else:
                groups[file_path]['parent'] = task

        # Get sort key for a group (uses parent's column value)
        def get_sort_key(file_path):
            group = groups[file_path]
            parent = group['parent']
            if not parent:
                # No parent, use first sub-row
                if group['sub_rows']:
                    parent = group['sub_rows'][0]
                else:
                    return ''

            # Get value from the appropriate column
            # Map column index to task field
            column_to_field = {
                0: 'file_name',
                1: 'protocol',
                2: 'channel_count',
                3: 'sweep_count',
                4: 'keywords_display',
                5: 'channel',
                6: 'stim_channel',
                7: 'events_channel',
                8: 'strain',
                9: 'stim_type',
                10: 'power',
                11: 'sex',
                12: 'animal_id',
                13: 'status',
                15: 'export_path',
            }
            field = column_to_field.get(column, 'file_name')
            value = parent.get(field, '')
            # Convert to string for consistent sorting
            return str(value).lower() if value else ''

        # Sort parent_order by the sort key
        parent_order.sort(key=get_sort_key, reverse=not self._sort_ascending)

        # Rebuild _master_file_list with sorted groups
        new_master_list = []
        for file_path in parent_order:
            group = groups[file_path]
            if group['parent']:
                new_master_list.append(group['parent'])
            new_master_list.extend(group['sub_rows'])

        self._master_file_list = new_master_list

        # Rebuild the table
        self._rebuild_table_from_master_list()

        # Show sort indicator
        direction = "↑" if self._sort_ascending else "↓"
        # Get column name from model's headerData
        col_name = self._file_table_model.headerData(column, Qt.Orientation.Horizontal, Qt.ItemDataRole.DisplayRole)
        if not col_name:
            col_name = f"Col {column}"
        self._log_status_message(f"Sorted by {col_name} {direction}", 2000)

    def _rebuild_table_from_master_list(self):
        """Rebuild the table from the current _master_file_list order using Model/View."""
        # With Model/View, we simply update the model's data
        # The view automatically reflects the changes
        # Note: The model's _get_display_value() handles all formatting
        # (sub-row names, exports summary, status icons, etc.)

        # Set data on model - this triggers view update
        self._file_table_model.set_files(self._master_file_list)

        # Track conflict rows for highlighting
        conflict_rows = set()
        for i, task in enumerate(self._master_file_list):
            warnings = task.get('scan_warnings', {})
            if warnings.get('conflicts'):
                conflict_rows.add(i)
        self._file_table_model.set_conflict_rows(conflict_rows)

        print(f"[project-builder] Table rebuilt with {len(self._master_file_list)} rows")

    def _get_experiment_history(self) -> list:
        """Get list of previously used experiment names from QSettings."""
        from PyQt6.QtCore import QSettings
        settings = QSettings("PhysioMetrics", "BreathAnalysis")
        history = settings.value("project_builder/experiment_history", [])
        return history if isinstance(history, list) else []

    def _update_experiment_history(self, experiment_name: str):
        """Add an experiment name to history (max 50 entries, most recent first)."""
        from PyQt6.QtCore import QSettings
        settings = QSettings("PhysioMetrics", "BreathAnalysis")

        history = self._get_experiment_history()

        # Remove if already exists (to move to front)
        if experiment_name in history:
            history.remove(experiment_name)

        # Add to front
        history.insert(0, experiment_name)

        # Keep max 50
        history = history[:50]

        settings.setValue("project_builder/experiment_history", history)

    def mark_active_analysis_complete(self, channel_used: str = None, stim_channel_used: str = None,
                                       events_channel_used: str = None, export_info: dict = None):
        """
        Mark the currently active master list row as completed.
        Called after successful export/save of analysis data.

        Args:
            channel_used: The channel that was analyzed (e.g., "AD0")
            stim_channel_used: The stim channel used (e.g., "AD1")
            events_channel_used: The events channel used (e.g., "AD2")
            export_info: Dict with export metadata and file flags:
                - export_path: Where files were saved
                - export_date: When exported
                - export_version: App version used
                - exports: Dict of boolean flags for each export type
                - strain, stim_type, power, sex, animal_id: Metadata from dialog
        """
        # Get current file path from state
        current_file_path = str(self.state.in_path) if self.state.in_path else None

        # Find matching row(s) by file path if active row not set
        row = self._active_master_list_row

        if row is None and current_file_path:
            # Try to find a row matching the current file
            # Priority: 1) Same file + same channel, 2) Same file + completed (different channel), 3) Same file + pending
            from pathlib import Path
            current_normalized = str(Path(current_file_path).resolve())

            matching_rows = []  # [(row_idx, task), ...]
            for i, task in enumerate(self._master_file_list):
                task_path = task.get('file_path', '')
                if task_path:
                    task_normalized = str(Path(task_path).resolve())
                    if task_normalized == current_normalized:
                        matching_rows.append((i, task))

            if matching_rows:
                # Priority 1: Find row with same channel (to update existing)
                for i, task in matching_rows:
                    if task.get('channel', '') == channel_used:
                        row = i
                        print(f"[master-list] Found matching row {i} with same channel {channel_used}")
                        break

                # Priority 2: Find completed row with different channel (to create sub-row)
                if row is None:
                    for i, task in matching_rows:
                        if task.get('status', 'pending') == 'completed':
                            row = i
                            print(f"[master-list] Found completed row {i} (channel: {task.get('channel', '')})")
                            break

                # Priority 3: Find any pending row (first time analysis)
                if row is None:
                    for i, task in matching_rows:
                        if task.get('status', 'pending') == 'pending':
                            row = i
                            print(f"[master-list] Found pending row {i}")
                            break

                # Fallback: Use first matching row
                if row is None and matching_rows:
                    row = matching_rows[0][0]
                    print(f"[master-list] Using first matching row {row}")

        if row is None:
            print(f"[master-list] No matching row found for file: {current_file_path}")
            return

        if row >= len(self._master_file_list):
            return

        task = self._master_file_list[row]

        # Check if this is a parent row or sub-row
        is_sub_row = task.get('is_sub_row', False)
        existing_channel = task.get('channel', '')
        existing_status = task.get('status', 'pending')

        # Debug logging to understand the flow
        print(f"[master-list] mark_active_analysis_complete called:")
        print(f"  - Row: {row}")
        print(f"  - File: {task.get('file_name', 'unknown')}")
        print(f"  - Is sub-row: {is_sub_row}")
        print(f"  - Existing channel in row: '{existing_channel}'")
        print(f"  - Channel just analyzed: '{channel_used}'")
        print(f"  - Existing status: '{existing_status}'")

        # For parent rows, ALWAYS create sub-rows instead of updating the parent
        # This keeps parent as a "header" row showing file-level info only
        if not is_sub_row:
            print(f"[master-list] ✓ Parent row detected - creating sub-row for {channel_used}")

            # If parent had existing completed analysis, move it to sub-row first
            if existing_channel and existing_status == 'completed':
                # Check if there's already a sub-row for the existing channel
                has_existing_subrow = False
                for t in self._master_file_list:
                    if (t.get('is_sub_row') and
                        str(t.get('file_path')) == str(task.get('file_path')) and
                        t.get('channel') == existing_channel):
                        has_existing_subrow = True
                        break

                if not has_existing_subrow:
                    # Move existing channel data to a sub-row first
                    existing_export_info = {
                        'export_path': task.get('export_path', ''),
                        'export_date': task.get('export_date', ''),
                        'export_version': task.get('export_version', ''),
                        'exports': task.get('exports', {}),
                        'strain': task.get('strain', ''),
                        'stim_type': task.get('stim_type', ''),
                        'power': task.get('power', ''),
                        'sex': task.get('sex', ''),
                        'animal_id': task.get('animal_id', ''),
                    }
                    self._create_sub_row_from_analysis(
                        row, existing_channel,
                        task.get('stim_channel', ''),
                        task.get('events_channel', ''),
                        existing_export_info
                    )
                    print(f"[master-list]   - Moved existing {existing_channel} to sub-row")

            # Check if there's already a sub-row for the channel we're analyzing
            existing_subrow_idx = None
            for idx, t in enumerate(self._master_file_list):
                if (t.get('is_sub_row') and
                    str(t.get('file_path')) == str(task.get('file_path')) and
                    t.get('channel') == channel_used):
                    existing_subrow_idx = idx
                    print(f"[master-list]   - Found existing sub-row at {idx} for {channel_used}")
                    break

            if existing_subrow_idx is not None:
                # Update existing sub-row instead of creating new one
                existing_task = self._master_file_list[existing_subrow_idx]
                existing_task['status'] = 'completed'
                if stim_channel_used:
                    existing_task['stim_channel'] = stim_channel_used
                if events_channel_used:
                    existing_task['events_channel'] = events_channel_used
                if export_info:
                    existing_task['export_path'] = export_info.get('export_path', '')
                    existing_task['export_date'] = export_info.get('export_date', '')
                    existing_task['export_version'] = export_info.get('export_version', '')
                    existing_task['exports'] = export_info.get('exports', existing_task.get('exports', {}))
                    if export_info.get('strain'):
                        existing_task['strain'] = export_info['strain']
                    if export_info.get('stim_type'):
                        existing_task['stim_type'] = export_info['stim_type']
                    if export_info.get('power'):
                        existing_task['power'] = export_info['power']
                    if export_info.get('sex'):
                        existing_task['sex'] = export_info['sex']
                    if export_info.get('animal_id'):
                        existing_task['animal_id'] = export_info['animal_id']
                # Update the model
                if existing_subrow_idx < self._file_table_model.rowCount():
                    self._file_table_model.update_row(existing_subrow_idx, existing_task)
                print(f"[master-list]   - Updated existing sub-row for {channel_used}")
            else:
                # Create new sub-row for the new analysis
                self._create_sub_row_from_analysis(
                    row, channel_used, stim_channel_used, events_channel_used, export_info
                )
                print(f"[master-list]   - Created sub-row for {channel_used}")

            # Apply styling to update colors and clean parent
            self._apply_all_row_styling()

            self._active_master_list_row = None
            return

        # For sub-rows, update the existing sub-row
        task['status'] = 'completed'
        if channel_used:
            task['channel'] = channel_used
        if stim_channel_used:
            task['stim_channel'] = stim_channel_used
        if events_channel_used:
            task['events_channel'] = events_channel_used

        # Update export tracking info
        if export_info:
            task['export_path'] = export_info.get('export_path', '')
            task['export_date'] = export_info.get('export_date', '')
            task['export_version'] = export_info.get('export_version', '')
            task['exports'] = export_info.get('exports', task.get('exports', {}))

            # Update metadata from save dialog
            if export_info.get('strain'):
                task['strain'] = export_info['strain']
            if export_info.get('stim_type'):
                task['stim_type'] = export_info['stim_type']
            if export_info.get('power'):
                task['power'] = export_info['power']
            if export_info.get('sex'):
                task['sex'] = export_info['sex']
            if export_info.get('animal_id'):
                task['animal_id'] = export_info['animal_id']

        # Update the model to refresh the table display
        if row < self._file_table_model.rowCount():
            self._file_table_model.update_row(row, task)

        print(f"[master-list] Marked row {row} as completed:")
        print(f"  - Channel: {channel_used}")
        print(f"  - Stim Ch: {stim_channel_used}")
        print(f"  - Events Ch: {events_channel_used}")
        if export_info:
            print(f"  - Export path: {export_info.get('export_path', 'N/A')}")
            exports = export_info.get('exports', {})
            saved = [k for k, v in exports.items() if v]
            print(f"  - Saved: {', '.join(saved) if saved else 'none'}")

        # Clear active row - analysis is done
        self._active_master_list_row = None

    def _create_sub_row_from_analysis(self, source_row: int, channel_used: str,
                                       stim_channel_used: str, events_channel_used: str,
                                       export_info: dict):
        """
        Create a new sub-row when user analyzes a different channel from an existing analyzed row.

        This prevents overwriting previous analysis when the user analyzes the same file
        but with a different channel.

        Args:
            source_row: Row index of the source task
            channel_used: The channel that was analyzed
            stim_channel_used: The stim channel used
            events_channel_used: The events channel used
            export_info: Export metadata dict from save dialog
        """
        if source_row >= len(self._master_file_list):
            return

        source_task = self._master_file_list[source_row]
        file_path = source_task.get('file_path')

        # Create new task based on source but with new channel and analysis info
        new_task = {
            'file_path': file_path,
            'file_name': source_task.get('file_name', ''),
            'protocol': source_task.get('protocol', ''),
            'channel_count': source_task.get('channel_count', 0),
            'sweep_count': source_task.get('sweep_count', 0),
            'channel_names': source_task.get('channel_names', []),
            'stim_channels': source_task.get('stim_channels', []),
            'path_keywords': source_task.get('path_keywords', {}),
            'keywords_display': source_task.get('keywords_display', ''),
            # Analysis results
            'channel': channel_used,
            'stim_channel': stim_channel_used or '',
            'events_channel': events_channel_used or '',
            # Copy or update metadata
            'strain': export_info.get('strain', '') if export_info else source_task.get('strain', ''),
            'stim_type': export_info.get('stim_type', '') if export_info else source_task.get('stim_type', ''),
            'power': export_info.get('power', '') if export_info else source_task.get('power', ''),
            'sex': export_info.get('sex', '') if export_info else source_task.get('sex', ''),
            'animal_id': export_info.get('animal_id', '') if export_info else '',
            'status': 'completed',
            'is_sub_row': True,
            'parent_file': file_path,
            # Export tracking
            'export_path': export_info.get('export_path', '') if export_info else '',
            'export_date': export_info.get('export_date', '') if export_info else '',
            'export_version': export_info.get('export_version', '') if export_info else '',
            'exports': export_info.get('exports', {}) if export_info else {
                'npz': False,
                'timeseries_csv': False,
                'breaths_csv': False,
                'events_csv': False,
                'pdf': False,
                'session_state': False,
                'ml_training': False,
            }
        }

        # Insert after the source row
        insert_row = source_row + 1
        self._master_file_list.insert(insert_row, new_task)

        # Rebuild table from master list
        self._rebuild_table_from_master_list()

        print(f"[master-list] Created new sub-row at {insert_row} for channel {channel_used}")
        print(f"  - This preserves the previous analysis of {source_task.get('channel', 'unknown')} in row {source_row}")

    def _update_task_with_export_info(self, task: dict, info: dict, row: int):
        """
        Update a task dict with export info from saved data scan.

        Args:
            task: The task dict from _master_file_list
            info: Export info dict from NPZ metadata
            row: Table row index
        """
        # Track conflicts and multiple files
        conflicts = []
        older_files = info.get('older_npz_files', [])

        # Check for conflicts between existing task data and NPZ data
        metadata_fields = [
            ('strain', 'strain', 'Strain'),
            ('stim_type', 'stim_type', 'Stim Type'),
            ('power', 'power', 'Power'),
            ('sex', 'sex', 'Sex'),
            ('animal_id', 'animal_id', 'Animal ID'),
        ]
        for task_key, info_key, display_name in metadata_fields:
            task_val = task.get(task_key, '').strip()
            info_val = (info.get(info_key, '') or '').strip()
            if task_val and info_val and task_val != info_val:
                conflicts.append(f"{display_name}: table='{task_val}' vs NPZ='{info_val}'")

        # Store conflict info in task
        if conflicts or older_files:
            task['scan_warnings'] = {
                'conflicts': conflicts,
                'older_npz_count': len(older_files),
                'older_npz_files': older_files,
            }
            if conflicts:
                print(f"[scan-saved] Conflicts found for row {row}: {conflicts}")
            if older_files:
                print(f"[scan-saved] {len(older_files)} older NPZ file(s) found for row {row}")

        # Update task with export info
        task['export_path'] = info.get('export_path', '')
        task['export_date'] = info.get('export_date', '')
        task['status'] = 'completed'
        task['exports'] = {
            'npz': info.get('npz', False),
            'timeseries_csv': info.get('timeseries_csv', False),
            'breaths_csv': info.get('breaths_csv', False),
            'events_csv': info.get('events_csv', False),
            'pdf': info.get('pdf', False),
            'session_state': info.get('session_state', False),
            'ml_training': info.get('ml_training', False),
        }

        # Update metadata from NPZ if available and not already set (don't overwrite)
        if info.get('strain') and not task.get('strain'):
            task['strain'] = info['strain']
        if info.get('stim_type') and not task.get('stim_type'):
            task['stim_type'] = info['stim_type']
        if info.get('power') and not task.get('power'):
            task['power'] = info['power']
        if info.get('sex') and not task.get('sex'):
            task['sex'] = info['sex']
        if info.get('animal_id') and not task.get('animal_id'):
            task['animal_id'] = info['animal_id']
        if info.get('stim_channel') and not task.get('stim_channel'):
            task['stim_channel'] = info['stim_channel']
        if info.get('events_channel') and not task.get('events_channel'):
            task['events_channel'] = info['events_channel']

        # Update model row - the model will handle display
        if row < self._file_table_model.rowCount():
            self._file_table_model.update_row(row, task)
            # Mark conflict rows for highlighting
            if task.get('scan_warnings', {}).get('conflicts'):
                self._file_table_model.add_conflict_row(row)

    def _create_sub_row_from_saved_data(self, source_task: dict, source_row: int,
                                         channel: str, info: dict):
        """
        Create a new sub-row from scanned saved data for an additional channel.

        Updates _master_file_list only - caller should rebuild table after all updates.

        Args:
            source_task: The source task dict to base the new row on
            source_row: Row index of the source task
            channel: The channel for this saved data
            info: Export info dict from NPZ metadata
        """
        file_path = source_task.get('file_path')

        # Track any older/duplicate NPZ files found
        older_files = info.get('older_npz_files', [])
        scan_warnings = {}
        if older_files:
            scan_warnings = {
                'conflicts': [],
                'older_npz_count': len(older_files),
                'older_npz_files': older_files,
            }

        # Create new task based on source but with new channel and analysis info
        new_task = {
            'file_path': file_path,
            'file_name': source_task.get('file_name', ''),
            'protocol': source_task.get('protocol', ''),
            'channel_count': source_task.get('channel_count', 0),
            'sweep_count': source_task.get('sweep_count', 0),
            'channel_names': source_task.get('channel_names', []),
            'stim_channels': source_task.get('stim_channels', []),
            'path_keywords': source_task.get('path_keywords', {}),
            'keywords_display': source_task.get('keywords_display', ''),
            # Analysis results
            'channel': channel,
            'stim_channel': info.get('stim_channel', ''),
            'events_channel': info.get('events_channel', ''),
            # Metadata from NPZ
            'strain': info.get('strain', '') or source_task.get('strain', ''),
            'stim_type': info.get('stim_type', '') or source_task.get('stim_type', ''),
            'power': info.get('power', '') or source_task.get('power', ''),
            'sex': info.get('sex', '') or source_task.get('sex', ''),
            'animal_id': info.get('animal_id', ''),
            'status': 'completed',
            'is_sub_row': True,
            'parent_file': file_path,
            # Export tracking
            'export_path': info.get('export_path', ''),
            'export_date': info.get('export_date', ''),
            'exports': {
                'npz': info.get('npz', False),
                'timeseries_csv': info.get('timeseries_csv', False),
                'breaths_csv': info.get('breaths_csv', False),
                'events_csv': info.get('events_csv', False),
                'pdf': info.get('pdf', False),
                'session_state': info.get('session_state', False),
                'ml_training': info.get('ml_training', False),
            }
        }
        if scan_warnings:
            new_task['scan_warnings'] = scan_warnings

        # Find where to insert - after all existing rows for this file
        insert_row = source_row + 1
        while insert_row < len(self._master_file_list):
            next_task = self._master_file_list[insert_row]
            if str(next_task.get('file_path')) != str(file_path):
                break
            insert_row += 1

        self._master_file_list.insert(insert_row, new_task)

        print(f"[scan-saved] Created sub-row at {insert_row} for channel {channel}")

    def eventFilter(self, obj, event):
        """Handle events for installed event filters (e.g., table resize)."""
        from PyQt6.QtCore import QEvent, QTimer

        # Check if this is a resize event for the table viewport
        if (hasattr(self, 'discoveredFilesTable') and
            obj == self.discoveredFilesTable.viewport() and
            event.type() == QEvent.Type.Resize):

            # Debounce resize events to avoid excessive recalculation
            if self._table_resize_timer is not None:
                self._table_resize_timer.stop()

            self._table_resize_timer = QTimer()
            self._table_resize_timer.setSingleShot(True)
            self._table_resize_timer.timeout.connect(self._auto_fit_table_columns)
            self._table_resize_timer.start(100)  # 100ms delay

        return super().eventFilter(obj, event)

    def _auto_fit_table_columns(self):
        """
        Auto-fit column widths for QTableView.
        Uses column definitions from the model for sizing hints.
        """
        table = self.discoveredFilesTable
        header = table.horizontalHeader()
        model = self._file_table_model

        if model.columnCount() == 0:
            return

        # Check if full content mode is enabled
        full_content_mode = getattr(self, 'tableFullContentCheckbox', None)
        full_content_mode = full_content_mode.isChecked() if full_content_mode else False

        # Get column definitions from model
        visible_columns = model.get_visible_columns()

        # First, resize to contents to get content-based widths
        table.resizeColumnsToContents()

        if full_content_mode:
            # Full content mode: show all content, enable horizontal scrolling
            table.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
            # Just use content widths, but ensure minimum widths from column defs
            for i, col_def in enumerate(visible_columns):
                current_width = table.columnWidth(i)
                min_width = col_def.min_width
                if current_width < min_width:
                    table.setColumnWidth(i, min_width)
        else:
            # Fit-to-view mode: constrain to visible width
            table.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)

            # Get available width
            available_width = table.viewport().width()
            if available_width < 100:
                available_width = table.width() - 20

            # Calculate column widths based on column definitions
            widths = []
            expandable_indices = []
            total_fixed = 0

            for i, col_def in enumerate(visible_columns):
                content_width = table.columnWidth(i)
                base_width = max(col_def.min_width, min(content_width, col_def.width + 30))

                if col_def.fixed:
                    # Fixed columns use their defined width
                    widths.append(col_def.width)
                    total_fixed += col_def.width
                elif col_def.expandable:
                    widths.append(base_width)
                    expandable_indices.append(i)
                else:
                    widths.append(base_width)

            total_width = sum(widths)

            # Distribute extra space or shrink as needed
            if total_width < available_width and expandable_indices:
                # Distribute extra space to expandable columns
                extra = available_width - total_width
                per_col = extra // len(expandable_indices)
                for idx in expandable_indices:
                    widths[idx] += per_col
            elif total_width > available_width:
                # Shrink proportionally, but respect minimums
                scale = available_width / total_width
                for i, col_def in enumerate(visible_columns):
                    if not col_def.fixed:
                        widths[i] = max(col_def.min_width, int(widths[i] * scale))

            # Apply widths
            for i, width in enumerate(widths):
                table.setColumnWidth(i, width)

    def _apply_row_styling(self, row: int, is_sub_row: bool = False):
        """Apply visual styling - now handled by model via data() method."""
        # Styling is handled by FileTableModel.data() for BackgroundRole, ForegroundRole, FontRole
        pass

    def _apply_all_row_styling(self):
        """Apply styling to all rows - cleans parent rows that have sub-rows."""
        # Styling is now handled by the model, but we still need to clean task dicts
        self._clean_parent_rows_with_subrows()

    def _clean_parent_rows_with_subrows(self):
        """Clear channel-specific fields from parent tasks that have sub-rows."""
        # Find parent rows that have sub-rows
        parent_to_subrows = {}  # parent_file_path -> [sub_row_indices]

        for row, task in enumerate(self._master_file_list):
            if task.get('is_sub_row'):
                parent_path = str(task.get('file_path', ''))
                if parent_path not in parent_to_subrows:
                    parent_to_subrows[parent_path] = []
                parent_to_subrows[parent_path].append(row)

        # Clear fields in parent tasks that have sub-rows
        for row, task in enumerate(self._master_file_list):
            if task.get('is_sub_row'):
                continue

            file_path = str(task.get('file_path', ''))
            if file_path in parent_to_subrows:
                sub_rows = parent_to_subrows[file_path]

                # Count completed sub-rows for status display
                completed = sum(1 for sr in sub_rows
                                if self._master_file_list[sr].get('status') == 'completed')

                # Clear channel-specific fields in the task dict
                task['channel'] = ''
                task['stim_channel'] = ''
                task['events_channel'] = ''
                task['animal_id'] = ''
                task['status'] = f"{completed} ✓" if completed > 0 else ''
                task['exports'] = {}

    def _add_row_action_button(self, row):
        """Add Analyze and '+' buttons to the actions column."""
        from PyQt6.QtWidgets import QPushButton, QWidget, QHBoxLayout

        # Create container widget
        container = QWidget()
        layout = QHBoxLayout(container)
        layout.setContentsMargins(2, 2, 2, 2)
        layout.setSpacing(2)

        # Add 'Analyze' button (opens file in Analysis tab)
        analyze_btn = QPushButton("▶")
        analyze_btn.setMaximumWidth(22)
        analyze_btn.setMaximumHeight(22)
        analyze_btn.setToolTip("Open this file in the Analysis tab")
        analyze_btn.setStyleSheet("""
            QPushButton {
                background-color: #2196F3;
                color: white;
                border: none;
                border-radius: 3px;
                font-weight: bold;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: #1976D2;
            }
        """)
        analyze_btn.clicked.connect(lambda checked, r=row: self._on_analyze_row(r))
        layout.addWidget(analyze_btn)

        # Add '+' button (add another row for different channel/animal)
        add_btn = QPushButton("+")
        add_btn.setMaximumWidth(22)
        add_btn.setMaximumHeight(22)
        add_btn.setToolTip("Add another row for this file (different channel/animal)")
        add_btn.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                border: none;
                border-radius: 3px;
                font-weight: bold;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
        """)
        add_btn.clicked.connect(lambda checked, r=row: self._on_add_row_for_file(r))
        layout.addWidget(add_btn)

        self.discoveredFilesTable.setCellWidget(row, 15, container)

    def _on_analyze_row(self, row):
        """Open the file from the specified row in the Analysis tab."""
        if row >= len(self._master_file_list):
            return

        task = self._master_file_list[row]
        file_path = task.get('file_path')

        if not file_path or not Path(file_path).exists():
            self._show_warning("File Not Found", f"Cannot find file:\n{file_path}")
            return

        print(f"[master-list] Analyzing file from row {row}: {file_path}")

        # Track which row is being analyzed
        self._active_master_list_row = row

        # Store pending channel selections if specified
        self._pending_analysis_channel = task.get('channel', '')
        self._pending_stim_channels = task.get('stim_channels', [])

        # Load the file
        self.load_file(Path(file_path))

        # Switch to Analysis tab (Tab 0 = Project Builder, Tab 1 = Analysis)
        if hasattr(self, 'Tabs'):
            self.Tabs.setCurrentIndex(1)

        # Update status to "in progress"
        task['status'] = 'in_progress'
        # Update the model row to reflect the new status
        if row < self._file_table_model.rowCount():
            self._file_table_model.update_row(row, task)

    def _on_add_row_for_file(self, source_row, force_override=False):
        """Add a new row for the same file (for multi-channel/multi-animal analysis)."""
        if source_row >= len(self._master_file_list):
            return

        source_task = self._master_file_list[source_row]
        file_path = source_task.get('file_path')

        # Find the parent task (may be source_task or a different row)
        parent_task = source_task
        if source_task.get('is_sub_row'):
            # Find the actual parent row
            for task in self._master_file_list:
                if not task.get('is_sub_row') and str(task.get('file_path')) == str(file_path):
                    parent_task = task
                    break

        # Count existing SUB-ROWS for this file (parent row doesn't count as an analysis)
        existing_subrows = []
        for i, task in enumerate(self._master_file_list):
            if task.get('is_sub_row') and str(task.get('file_path')) == str(file_path):
                existing_subrows.append(i)

        # Calculate available analysis channels from PARENT task (has file metadata)
        total_channels = parent_task.get('channel_count', 0)

        # Get stim channels from file metadata
        stim_channels_from_metadata = parent_task.get('stim_channels', [])
        if not isinstance(stim_channels_from_metadata, list):
            stim_channels_from_metadata = [stim_channels_from_metadata] if stim_channels_from_metadata else []

        # Also check existing sub-rows for stim/events channels they use
        stim_events_from_subrows = set()
        for i in existing_subrows:
            task = self._master_file_list[i]
            stim_ch = task.get('stim_channel', '')
            events_ch = task.get('events_channel', '')
            if stim_ch:
                stim_events_from_subrows.add(stim_ch)
            if events_ch:
                stim_events_from_subrows.add(events_ch)

        # Combine all non-analysis channels
        all_stim_channels = set(stim_channels_from_metadata) | stim_events_from_subrows
        num_stim_channels = len(all_stim_channels)
        available_analysis_channels = total_channels - num_stim_channels

        # Check if we can add more rows (compare sub-row count to available channels)
        if len(existing_subrows) >= available_analysis_channels and not force_override:
            from PyQt6.QtWidgets import QMessageBox
            msg = QMessageBox(self)
            msg.setWindowTitle("Channel Limit Reached")
            msg.setText(
                f"This file has {total_channels} channels.\n"
                f"Detected {num_stim_channels} stim/events channel(s): {', '.join(sorted(all_stim_channels)) if all_stim_channels else 'none'}\n"
                f"Maximum {available_analysis_channels} analysis rows allowed.\n"
                f"Already have {len(existing_subrows)} analyzed channel(s)."
            )
            msg.setInformativeText("Do you want to add a row anyway?")
            msg.setStandardButtons(QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            msg.setDefaultButton(QMessageBox.StandardButton.No)
            if msg.exec() != QMessageBox.StandardButton.Yes:
                return
            # User clicked Yes - force override
            force_override = True

        # Find the next available channel (use parent_task for channel metadata)
        channel_names = parent_task.get('channel_names', [])
        if not channel_names:
            # Generate default channel names
            channel_names = [f"AD{i}" for i in range(total_channels)]

        # Filter out stim/events channels to get analysis channels
        analysis_channels = [ch for ch in channel_names if ch not in all_stim_channels]

        # Find which channels are already used (from sub-rows only)
        used_channels = set()
        for i in existing_subrows:
            ch = self._master_file_list[i].get('channel', '')
            if ch:
                used_channels.add(ch)

        # Get next available channel
        next_channel = ''
        for ch in analysis_channels:
            if ch not in used_channels:
                next_channel = ch
                break

        # If no analysis channels available but user forced override, pick any unused channel
        if not next_channel and force_override:
            for ch in channel_names:
                if ch not in used_channels:
                    next_channel = ch
                    break

        # Create a new task based on parent task (for file metadata) and source task (for user metadata)
        new_task = {
            'file_path': file_path,
            'file_name': parent_task.get('file_name', ''),
            'protocol': parent_task.get('protocol', ''),
            'channel_count': total_channels,
            'sweep_count': parent_task.get('sweep_count', 0),
            'channel_names': channel_names,
            'stim_channels': list(all_stim_channels),
            'path_keywords': parent_task.get('path_keywords', {}),
            'keywords_display': parent_task.get('keywords_display', ''),
            # Auto-populated channel
            'channel': next_channel,
            'stim_channel': '',  # User can set this
            'events_channel': '',  # Filled on save
            # Copy metadata from source (may have user edits)
            'strain': source_task.get('strain', ''),
            'stim_type': source_task.get('stim_type', ''),
            'power': source_task.get('power', ''),
            'sex': source_task.get('sex', ''),
            'animal_id': '',  # Different animal - leave blank
            'status': 'pending',
            'is_sub_row': True,
            'parent_file': file_path,
            # Export tracking (filled on save)
            'export_path': '',
            'export_date': '',
            'export_version': '',
            'exports': {
                'npz': False,
                'timeseries_csv': False,
                'breaths_csv': False,
                'events_csv': False,
                'pdf': False,
                'session_state': False,
                'ml_training': False,
            }
        }

        # Insert after the source row
        insert_row = source_row + 1
        self._master_file_list.insert(insert_row, new_task)

        # Rebuild table from master list (model handles all formatting)
        self._rebuild_table_from_master_list()

        remaining = available_analysis_channels - len(existing_subrows) - 1
        print(f"[master-list] Added sub-row at {insert_row} with channel {next_channel} ({remaining} more available)")

        # Autosave after adding row
        self._project_autosave()

    def _add_sub_row_action_buttons(self, row):
        """Add action buttons for sub-rows (Analyze, + and - buttons)."""
        from PyQt6.QtWidgets import QPushButton, QWidget, QHBoxLayout

        container = QWidget()
        layout = QHBoxLayout(container)
        layout.setContentsMargins(2, 2, 2, 2)
        layout.setSpacing(1)

        # Add 'Analyze' button
        analyze_btn = QPushButton("▶")
        analyze_btn.setMaximumWidth(20)
        analyze_btn.setMaximumHeight(20)
        analyze_btn.setToolTip("Open this file in the Analysis tab")
        analyze_btn.setStyleSheet("""
            QPushButton {
                background-color: #2196F3;
                color: white;
                border: none;
                border-radius: 3px;
                font-size: 10px;
            }
            QPushButton:hover { background-color: #1976D2; }
        """)
        analyze_btn.clicked.connect(lambda checked, r=row: self._on_analyze_row(r))
        layout.addWidget(analyze_btn)

        # Add '-' button (remove sub-row)
        remove_btn = QPushButton("-")
        remove_btn.setMaximumWidth(20)
        remove_btn.setMaximumHeight(20)
        remove_btn.setToolTip("Remove this row")
        remove_btn.setStyleSheet("""
            QPushButton {
                background-color: #dc3545;
                color: white;
                border: none;
                border-radius: 3px;
                font-weight: bold;
            }
            QPushButton:hover { background-color: #c82333; }
        """)
        remove_btn.clicked.connect(lambda checked, r=row: self._on_remove_sub_row(r))
        layout.addWidget(remove_btn)

        self.discoveredFilesTable.setCellWidget(row, 15, container)

    def _on_remove_sub_row(self, row):
        """Remove a sub-row from the master list."""
        if row >= len(self._master_file_list):
            return

        task = self._master_file_list[row]

        # Only allow removing sub-rows, not primary rows
        if not task.get('is_sub_row', False):
            self._show_warning("Cannot Remove",
                "Cannot remove the primary row for a file.\n"
                "Use 'Omit' to hide files you don't want to analyze.")
            return

        # Confirm deletion if this row has been analyzed
        if task.get('status') == 'completed':
            from PyQt6.QtWidgets import QMessageBox

            # Get export info for confirmation message
            export_summary = self._format_exports_summary(task)
            export_path = task.get('export_path', '')
            channel = task.get('channel', 'unknown channel')

            msg = QMessageBox(self)
            msg.setIcon(QMessageBox.Icon.Warning)
            msg.setWindowTitle("Delete Analyzed Row?")
            msg.setText(f"This row ({channel}) has been analyzed.")

            details = []
            if export_summary:
                details.append(f"Exports: {export_summary}")
            if export_path:
                details.append(f"Saved to: {export_path}")
            if details:
                msg.setInformativeText("\n".join(details) + "\n\nNote: The exported files will NOT be deleted.")

            msg.setDetailedText(
                "This will remove the row from the Project Builder list only.\n\n"
                "The exported data files on disk remain unchanged. "
                "You can add a new row and re-analyze if needed."
            )
            msg.setStandardButtons(QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.Cancel)
            msg.setDefaultButton(QMessageBox.StandardButton.Cancel)

            if msg.exec() != QMessageBox.StandardButton.Yes:
                return  # User cancelled

        # Remove from data and table
        table = self.discoveredFilesTable
        table.blockSignals(True)

        del self._master_file_list[row]
        table.removeRow(row)

        table.blockSignals(False)

        print(f"[master-list] Removed sub-row at {row}")

    def _on_master_list_context_menu(self, position):
        """Show context menu for bulk editing selected rows."""
        from PyQt6.QtWidgets import QMenu, QInputDialog
        from PyQt6.QtGui import QAction

        table = self.discoveredFilesTable
        selected_rows = set(index.row() for index in table.selectedIndexes())

        if not selected_rows:
            return

        menu = QMenu(self)

        # Bulk edit options
        edit_menu = menu.addMenu(f"Set for {len(selected_rows)} selected rows")

        # Experiment (column 5) - with history suggestions
        experiment_menu = edit_menu.addMenu("Set Experiment")
        exp_history = self._get_experiment_history()
        if exp_history:
            for exp_name in exp_history[:10]:  # Show top 10
                exp_action = QAction(exp_name, self)
                exp_action.triggered.connect(lambda checked, e=exp_name: self._bulk_set_column(selected_rows, 5, 'experiment', "Set Experiment", e))
                experiment_menu.addAction(exp_action)
            experiment_menu.addSeparator()
        custom_exp_action = QAction("Custom...", self)
        custom_exp_action.triggered.connect(lambda: self._bulk_set_column(selected_rows, 5, 'experiment', "Set Experiment"))
        experiment_menu.addAction(custom_exp_action)

        # Strain (column 9)
        strain_action = QAction("Set Strain...", self)
        strain_action.triggered.connect(lambda: self._bulk_set_column(selected_rows, 9, 'strain', "Set Strain"))
        edit_menu.addAction(strain_action)

        # Stim Type (column 10)
        stim_action = QAction("Set Stim Type...", self)
        stim_action.triggered.connect(lambda: self._bulk_set_column(selected_rows, 10, 'stim_type', "Set Stim Type"))
        edit_menu.addAction(stim_action)

        # Power (column 11)
        power_action = QAction("Set Power...", self)
        power_action.triggered.connect(lambda: self._bulk_set_column(selected_rows, 11, 'power', "Set Power"))
        edit_menu.addAction(power_action)

        # Sex (column 12)
        sex_action = QAction("Set Sex...", self)
        sex_action.triggered.connect(lambda: self._bulk_set_column(selected_rows, 12, 'sex', "Set Sex"))
        edit_menu.addAction(sex_action)

        # Animal ID (column 13)
        animal_action = QAction("Set Animal ID...", self)
        animal_action.triggered.connect(lambda: self._bulk_set_column(selected_rows, 13, 'animal_id', "Set Animal ID"))
        edit_menu.addAction(animal_action)

        # Channel (column 6)
        channel_action = QAction("Set Channel...", self)
        channel_action.triggered.connect(lambda: self._bulk_set_column(selected_rows, 6, 'channel', "Set Channel"))
        edit_menu.addAction(channel_action)

        menu.addSeparator()

        # Combine files option (if multiple files selected)
        if len(selected_rows) > 1:
            combine_action = QAction(f"Combine {len(selected_rows)} files for analysis...", self)
            combine_action.triggered.connect(lambda: self._combine_selected_files(selected_rows))
            menu.addAction(combine_action)

        # Check if any selected rows have scan warnings (conflicts or multiple NPZ files)
        rows_with_warnings = []
        rows_with_conflicts = []
        for row in selected_rows:
            if row < len(self._master_file_list):
                task = self._master_file_list[row]
                warnings = task.get('scan_warnings', {})
                if warnings:
                    rows_with_warnings.append(row)
                    if warnings.get('conflicts'):
                        rows_with_conflicts.append(row)

        if rows_with_warnings:
            menu.addSeparator()
            conflict_menu = menu.addMenu(f"⚠ Resolve Warnings ({len(rows_with_warnings)} rows)")

            if rows_with_conflicts:
                # Use NPZ values - overwrite table with NPZ data
                use_npz_action = QAction("Use NPZ values (overwrite table)", self)
                use_npz_action.triggered.connect(lambda: self._resolve_conflicts_use_npz(rows_with_conflicts))
                conflict_menu.addAction(use_npz_action)

                # Keep table values - dismiss the warning
                keep_table_action = QAction("Keep table values (dismiss warning)", self)
                keep_table_action.triggered.connect(lambda: self._resolve_conflicts_keep_table(rows_with_warnings))
                conflict_menu.addAction(keep_table_action)

                conflict_menu.addSeparator()

            # View details
            view_details_action = QAction("View warning details...", self)
            view_details_action.triggered.connect(lambda: self._show_conflict_details(rows_with_warnings))
            conflict_menu.addAction(view_details_action)

            # Clear all warnings
            clear_action = QAction("Clear all warnings", self)
            clear_action.triggered.connect(lambda: self._clear_scan_warnings(rows_with_warnings))
            conflict_menu.addAction(clear_action)

        # Export options
        menu.addSeparator()
        export_menu = menu.addMenu("Export Table")

        export_selected_action = QAction(f"Export {len(selected_rows)} selected rows to CSV...", self)
        export_selected_action.triggered.connect(lambda: self._export_table_to_csv(selected_rows))
        export_menu.addAction(export_selected_action)

        export_all_action = QAction("Export all rows to CSV...", self)
        export_all_action.triggered.connect(lambda: self._export_table_to_csv(None))
        export_menu.addAction(export_all_action)

        menu.exec(table.viewport().mapToGlobal(position))

    def _export_table_to_csv(self, selected_rows=None):
        """Export table data to CSV file.

        Args:
            selected_rows: Set of row indices to export, or None for all rows
        """
        import csv
        from PyQt6.QtWidgets import QFileDialog

        # Determine which rows to export
        if selected_rows is None:
            rows_to_export = list(range(len(self._master_file_list)))
        else:
            rows_to_export = sorted(selected_rows)

        if not rows_to_export:
            self._show_warning("No Data", "No rows to export.")
            return

        # Get visible columns
        visible_cols = self._file_table_model.get_visible_columns()

        # Build header row (skip actions column)
        headers = []
        col_keys = []
        for col_def in visible_cols:
            if col_def.key != 'actions':
                headers.append(col_def.header or col_def.key)
                col_keys.append(col_def.key)

        # Get save file path
        default_name = f"{getattr(self, '_current_project_name', None) or 'project'}_table.csv"
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export Table to CSV",
            str(Path(self._project_directory) / default_name) if self._project_directory else default_name,
            "CSV files (*.csv);;All files (*.*)"
        )

        if not file_path:
            return

        try:
            with open(file_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)

                # Write header
                writer.writerow(headers)

                # Write data rows
                for row_idx in rows_to_export:
                    if row_idx >= len(self._master_file_list):
                        continue
                    task = self._master_file_list[row_idx]

                    row_data = []
                    for key in col_keys:
                        value = task.get(key, '')

                        # Format special values
                        if key == 'file_name' and task.get('is_sub_row'):
                            channel = task.get('channel', '')
                            value = f"↳ {channel}"
                        elif key == 'exports' and isinstance(value, dict):
                            # Format exports dict
                            parts = []
                            csv_count = sum([
                                1 if value.get('timeseries_csv') else 0,
                                1 if value.get('breaths_csv') else 0,
                                1 if value.get('events_csv') else 0,
                            ])
                            if value.get('pdf'):
                                parts.append('PDF')
                            if csv_count > 0:
                                parts.append(f'{csv_count} CSV')
                            if value.get('npz'):
                                parts.append('NPZ')
                            if value.get('session_state'):
                                parts.append('Session')
                            value = ', '.join(parts) if parts else ''
                        elif key == 'file_path':
                            value = str(value) if value else ''

                        row_data.append(str(value) if value is not None else '')

                    writer.writerow(row_data)

            self._log_status_message(f"✓ Exported {len(rows_to_export)} rows to {Path(file_path).name}", 3000)
            self._show_info("Export Complete", f"Exported {len(rows_to_export)} rows to:\n{file_path}")

        except Exception as e:
            self._show_error("Export Failed", f"Failed to export table:\n{e}")
            print(f"[export] Error: {e}")

    def _bulk_set_column(self, rows, column, field_name, dialog_title, preset_value=None):
        """Set a column value for multiple rows at once.

        Args:
            rows: Set of row indices to update
            column: Column index in the table
            field_name: Field name in the task dict
            dialog_title: Title for the input dialog
            preset_value: If provided, use this value without showing dialog
        """
        from PyQt6.QtWidgets import QInputDialog

        if preset_value is not None:
            value = preset_value
            ok = True
        else:
            # Get existing values to suggest
            existing_values = set()
            for row in rows:
                if row < len(self._master_file_list):
                    val = self._master_file_list[row].get(field_name, '')
                    if val:
                        existing_values.add(val)

            # Default to first existing value or empty
            default_value = list(existing_values)[0] if existing_values else ''

            value, ok = QInputDialog.getText(
                self, dialog_title,
                f"Enter value for {len(rows)} rows:",
                text=default_value
            )

        if ok:
            for row in rows:
                if row < len(self._master_file_list):
                    task = self._master_file_list[row]
                    task[field_name] = value
                    # Update the model row to reflect the change
                    if row < self._file_table_model.rowCount():
                        self._file_table_model.update_row(row, task)

            # Update experiment history if this is the experiment field
            if field_name == 'experiment' and value:
                self._update_experiment_history(value)

            print(f"[master-list] Bulk set {field_name} = '{value}' for {len(rows)} rows")

    def _resolve_conflicts_use_npz(self, rows):
        """Resolve conflicts by overwriting table values with NPZ values."""
        resolved_count = 0
        for row in rows:
            if row >= len(self._master_file_list):
                continue

            task = self._master_file_list[row]
            warnings = task.get('scan_warnings', {})
            conflicts = warnings.get('conflicts', [])

            if not conflicts:
                continue

            # Parse conflicts and apply NPZ values
            # Conflicts are in format: "Field: table='X' vs NPZ='Y'"
            for conflict in conflicts:
                try:
                    # Extract field name and NPZ value
                    field_part = conflict.split(':')[0].strip()
                    npz_part = conflict.split("NPZ='")[1].split("'")[0]

                    # Map display name to field
                    field_map = {
                        'Strain': 'strain',
                        'Stim Type': 'stim_type',
                        'Power': 'power',
                        'Sex': 'sex',
                        'Animal ID': 'animal_id',
                    }

                    if field_part in field_map:
                        field_name = field_map[field_part]
                        task[field_name] = npz_part
                        print(f"[conflict-resolve] Row {row}: Set {field_name} = '{npz_part}' from NPZ")
                except Exception as e:
                    print(f"[conflict-resolve] Error parsing conflict '{conflict}': {e}")

            # Clear the warnings after resolving
            task.pop('scan_warnings', None)
            resolved_count += 1

            # Update the model row to reflect changes
            self._update_row_status_icon(row, task)

        self._log_status_message(f"✓ Resolved conflicts for {resolved_count} rows (used NPZ values)", 3000)
        self._update_resolve_conflicts_button()

    def _resolve_conflicts_keep_table(self, rows):
        """Resolve conflicts by keeping table values and dismissing warnings."""
        resolved_count = 0
        for row in rows:
            if row >= len(self._master_file_list):
                continue

            task = self._master_file_list[row]
            if task.get('scan_warnings'):
                task.pop('scan_warnings', None)
                resolved_count += 1
                self._update_row_status_icon(row, task)

        self._log_status_message(f"✓ Dismissed warnings for {resolved_count} rows (kept table values)", 3000)
        self._update_resolve_conflicts_button()

    def _clear_scan_warnings(self, rows):
        """Clear all scan warnings for selected rows."""
        cleared_count = 0
        for row in rows:
            if row >= len(self._master_file_list):
                continue

            task = self._master_file_list[row]
            if task.get('scan_warnings'):
                task.pop('scan_warnings', None)
                cleared_count += 1
                self._update_row_status_icon(row, task)

        self._log_status_message(f"✓ Cleared warnings for {cleared_count} rows", 3000)
        self._update_resolve_conflicts_button()

    def _show_conflict_details(self, rows):
        """Show a dialog with detailed conflict information."""
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QTextEdit, QPushButton, QHBoxLayout

        dialog = QDialog(self)
        dialog.setWindowTitle("Scan Warning Details")
        dialog.setMinimumSize(600, 400)

        layout = QVBoxLayout(dialog)

        text_edit = QTextEdit()
        text_edit.setReadOnly(True)

        # Build details text
        details = []
        for row in rows:
            if row >= len(self._master_file_list):
                continue

            task = self._master_file_list[row]
            warnings = task.get('scan_warnings', {})
            if not warnings:
                continue

            file_name = task.get('file_name', 'Unknown')
            channel = task.get('channel', '')
            details.append(f"═══ Row {row}: {file_name} ({channel}) ═══")

            conflicts = warnings.get('conflicts', [])
            if conflicts:
                details.append("⚠ Data Conflicts:")
                for c in conflicts:
                    details.append(f"   • {c}")

            older_files = warnings.get('older_npz_files', [])
            if older_files:
                details.append(f"📁 {len(older_files)} Older NPZ File(s):")
                for older in older_files:
                    details.append(f"   • {Path(older['file']).name}")
                    details.append(f"     Date: {older['date']}")
                    details.append(f"     Path: {older['file']}")

            details.append("")

        text_edit.setText('\n'.join(details) if details else "No warning details found.")
        layout.addWidget(text_edit)

        # Buttons
        btn_layout = QHBoxLayout()

        btn_use_npz = QPushButton("Use NPZ Values")
        btn_use_npz.clicked.connect(lambda: (self._resolve_conflicts_use_npz(rows), dialog.accept()))
        btn_layout.addWidget(btn_use_npz)

        btn_keep_table = QPushButton("Keep Table Values")
        btn_keep_table.clicked.connect(lambda: (self._resolve_conflicts_keep_table(rows), dialog.accept()))
        btn_layout.addWidget(btn_keep_table)

        btn_close = QPushButton("Close")
        btn_close.clicked.connect(dialog.reject)
        btn_layout.addWidget(btn_close)

        layout.addLayout(btn_layout)
        dialog.exec()

    def _update_row_status_icon(self, row, task):
        """Update the status icon for a row based on its current state.

        The model's data() method handles the actual status icon display.
        This method just triggers a refresh of the row.
        """
        if row >= self._file_table_model.rowCount():
            return

        # Update the model row to refresh the display
        self._file_table_model.update_row(row, task)

    def _combine_selected_files(self, rows):
        """Combine selected files for multi-file analysis."""
        from PyQt6.QtWidgets import QMessageBox

        # Get file paths for selected rows
        file_paths = []
        for row in sorted(rows):
            if row < len(self._master_file_list):
                task = self._master_file_list[row]
                file_path = task.get('file_path')
                if file_path and Path(file_path).exists():
                    file_paths.append(Path(file_path))

        if len(file_paths) < 2:
            self._show_warning("Cannot Combine", "Need at least 2 valid files to combine.")
            return

        # Confirm with user
        file_names = [p.name for p in file_paths]
        msg = QMessageBox(self)
        msg.setIcon(QMessageBox.Icon.Question)
        msg.setWindowTitle("Combine Files")
        msg.setText(f"Combine {len(file_paths)} files for analysis?")
        msg.setInformativeText("Files will be concatenated in order:\n• " + "\n• ".join(file_names[:5]))
        if len(file_names) > 5:
            msg.setInformativeText(msg.informativeText() + f"\n... and {len(file_names) - 5} more")
        msg.setStandardButtons(QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.Cancel)

        if msg.exec() == QMessageBox.StandardButton.Yes:
            # Load files using multi-file loader
            print(f"[master-list] Combining {len(file_paths)} files for analysis")
            self.load_multiple_files(file_paths)
            # Switch to Analysis tab (Tab 0 = Project Builder, Tab 1 = Analysis)
            if hasattr(self, 'Tabs'):
                self.Tabs.setCurrentIndex(1)

    def on_project_browse_directory(self):
        """Browse for directory containing recordings."""
        from PyQt6.QtWidgets import QFileDialog

        directory = QFileDialog.getExistingDirectory(
            self,
            "Select Directory with Recordings",
            "",
            QFileDialog.Option.ShowDirsOnly
        )

        if directory:
            self._project_directory = directory
            self.directoryPathEdit.setText(directory)
            self._log_status_message(f"Selected directory: {directory}", 2000)
            print(f"[project-builder] Selected directory: {directory}")

    def on_project_scan_files(self):
        """Scan directory for new files - additive mode, preserves existing data."""
        if not self._project_directory:
            self._show_warning("No Directory", "Please select a directory first using 'Browse Directory'.")
            return

        from core import project_builder
        from PyQt6.QtWidgets import QProgressDialog
        from PyQt6.QtCore import QThread, pyqtSignal

        # Prevent multiple concurrent scans
        if hasattr(self, '_metadata_thread') and self._metadata_thread and self._metadata_thread.isRunning():
            self._show_warning("Scan In Progress", "A scan is already running. Please wait for it to complete.")
            return

        progress = None
        try:
            # Disable scan button during operation
            self.scanFilesButton.setEnabled(False)

            recursive = self.recursiveCheckBox.isChecked()

            # Get selected file types from checkboxes
            file_types = self._get_selected_file_types()

            # PHASE 1: Quick file discovery
            progress = QProgressDialog("Finding files...", "Cancel", 0, 0, self)
            progress.setWindowTitle("Scanning Directory")
            progress.setWindowModality(Qt.WindowModality.WindowModal)
            progress.setMinimumDuration(0)
            progress.setValue(0)
            progress.show()
            QApplication.processEvents()

            files = project_builder.discover_files(
                self._project_directory, recursive=recursive, file_types=file_types
            )
            abf_files = files['abf_files']
            smrx_files = files.get('smrx_files', [])
            edf_files = files.get('edf_files', [])
            notes_files = files.get('notes_files', [])

            # Combine all data files
            all_data_files = abf_files + smrx_files + edf_files

            if progress.wasCanceled():
                progress.close()
                self.scanFilesButton.setEnabled(True)
                self._log_status_message("Scan cancelled", 2000)
                return

            progress.close()

            # Store notes files for AI integration (future use)
            self._discovered_notes_files = notes_files
            if notes_files:
                print(f"[project-builder] Found {len(notes_files)} notes files")

            # PHASE 2: Build set of existing file paths to avoid duplicates
            # Use string paths for comparison (normalized)
            existing_paths = set()
            for task in self._master_file_list:
                fp = task.get('file_path')
                if fp:
                    # Normalize path for comparison
                    existing_paths.add(str(Path(fp).resolve()))

            # Find new files only (across all types)
            new_data_files = []
            for data_path in all_data_files:
                normalized = str(data_path.resolve())
                if normalized not in existing_paths:
                    new_data_files.append(data_path)

            if not new_data_files:
                total_existing = len(all_data_files)

                # Check if any existing files have incomplete metadata (need reloading)
                files_needing_metadata = []
                for i, task in enumerate(self._master_file_list):
                    protocol = task.get('protocol', '')
                    if protocol in ('Loading...', '', 'Unknown') or not protocol:
                        file_path = task.get('file_path')
                        if file_path:
                            files_needing_metadata.append((i, Path(file_path)))

                if files_needing_metadata:
                    # Reload metadata for incomplete files
                    self._log_status_message(f"Reloading metadata for {len(files_needing_metadata)} files...", 0)
                    self._reload_incomplete_metadata(files_needing_metadata)
                else:
                    self._log_status_message(f"No new files found ({total_existing} files already in list)", 3000)
                    self.scanFilesButton.setEnabled(True)
                return

            # Count by type for summary
            new_abf = len([f for f in new_data_files if f.suffix.lower() == '.abf'])
            new_smrx = len([f for f in new_data_files if f.suffix.lower() == '.smrx'])
            new_edf = len([f for f in new_data_files if f.suffix.lower() == '.edf'])

            # Extract path keywords for each new file
            from core.fast_abf_reader import extract_path_keywords

            # Track starting position for metadata loading offset
            start_row = len(self._master_file_list)
            new_files_data = []

            for i, data_path in enumerate(new_data_files):
                # Note: file_size is obtained during metadata loading to avoid network I/O here
                file_type = data_path.suffix.lower()[1:]  # 'abf', 'smrx', or 'edf'

                # Extract keywords from path
                path_info = extract_path_keywords(data_path, Path(self._project_directory))

                # Format keywords for display - show relative path as the primary keyword
                # This makes it easy to see where the file is located relative to the project
                keywords_display = []

                # Add relative path (the main keyword showing file location)
                if path_info.get('relative_path'):
                    keywords_display.append(path_info['relative_path'])
                elif path_info['subdirs']:
                    # Fallback to subdirs if relative_path not available
                    keywords_display.append('/'.join(path_info['subdirs']))

                # Also add power levels and animal IDs if detected
                if path_info['power_levels']:
                    keywords_display.extend(path_info['power_levels'])
                if path_info['animal_ids']:
                    keywords_display.extend([f"ID:{id}" for id in path_info['animal_ids']])

                # Auto-fill power and animal ID from path keywords
                power_auto = path_info['power_levels'][0] if path_info['power_levels'] else ''
                animal_id_auto = path_info['animal_ids'][0] if path_info['animal_ids'] else ''

                file_info = {
                    'file_path': data_path,
                    'file_name': data_path.name,
                    'file_type': file_type,  # Track file type
                    'protocol': 'Loading...',
                    'channel_count': 0,
                    'sweep_count': 0,
                    'stim_channels': [],  # Detected stimulus channels (e.g., ['AD1', 'AD2'])
                    'stim_frequency': '',  # Detected stim frequency (e.g., '30Hz')
                    'path_keywords': path_info,
                    'keywords_display': ', '.join(keywords_display) if keywords_display else '',
                    # Metadata fields (filled on save)
                    'channel': '',  # Analyzed channel - filled on save
                    'stim_channel': '',  # Stim channel used - filled on save
                    'events_channel': '',  # Events channel used - filled on save
                    'strain': '',
                    'stim_type': '',  # Auto-filled from stim_frequency
                    'power': power_auto,
                    'sex': '',
                    'animal_id': animal_id_auto,
                    'status': 'pending',
                    # Export tracking (filled on save)
                    'export_path': '',
                    'export_date': '',
                    'export_version': '',
                    'exports': {
                        'npz': False,
                        'timeseries_csv': False,
                        'breaths_csv': False,
                        'events_csv': False,
                        'pdf': False,
                        'session_state': False,
                        'ml_training': False,
                    }
                }
                new_files_data.append(file_info)
                self._master_file_list.append(file_info)

            # Rebuild table from master list (model handles all rendering)
            self._rebuild_table_from_master_list()
            self._auto_fit_table_columns()

            # Update summary - show total files by type
            total_files = len(self._master_file_list)
            type_counts = []
            if new_abf:
                type_counts.append(f"{new_abf} ABF")
            if new_smrx:
                type_counts.append(f"{new_smrx} SMRX")
            if new_edf:
                type_counts.append(f"{new_edf} EDF")
            type_summary = ", ".join(type_counts) if type_counts else "0 files"

            summary_text = f"Summary: {total_files} files total | Added {type_summary} | Loading metadata..."
            self.summaryLabel.setText(summary_text)
            self._log_status_message(f"Added {len(new_data_files)} new files, loading metadata...", 3000)

            # Store reference to new files data for background metadata update
            self._discovered_files_data = new_files_data

            # PHASE 3: Load metadata in background thread (only for new files)
            self._start_background_metadata_loading(new_data_files, start_row_offset=start_row)

        except Exception as e:
            if progress:
                progress.close()
            self.scanFilesButton.setEnabled(True)  # Re-enable on error
            self._show_error("Scan Error", f"Failed to scan directory:\n{e}")
            print(f"[project-builder] Error: {e}")
            import traceback
            traceback.print_exc()

    def _start_background_metadata_loading(self, data_files, start_row_offset=0):
        """Start background thread to load metadata using parallel processing.

        Args:
            data_files: List of data file paths to load (.abf, .smrx, .edf)
            start_row_offset: Row offset for updating table (for additive scans)
        """
        from PyQt6.QtCore import QThread, pyqtSignal

        # Store the row offset for use in the update callback
        self._metadata_row_offset = start_row_offset

        class MetadataThread(QThread):
            # Batch updates: send list of results instead of individual items
            batch_progress = pyqtSignal(list, int, int)  # [(index, metadata), ...], total, completed_so_far
            finished = pyqtSignal(set)  # protocols

            def __init__(self, files):
                super().__init__()
                self.files = files
                self.should_stop = False

            def run(self):
                # Use generic metadata reader that handles all file types
                from core.fast_abf_reader import read_metadata_parallel
                protocols = set()
                batch = []
                batch_size = 25  # Update UI every 25 files to reduce signal traffic
                completed_count = [0]  # Use list to allow modification in nested function

                def callback(index, total, metadata):
                    if self.should_stop:
                        return

                    if metadata:
                        protocols.add(metadata.get('protocol', 'Unknown'))

                    # Collect results in batches
                    batch.append((index, metadata))
                    completed_count[0] += 1

                    # Emit batch when it reaches batch_size or ALL files are done
                    # Note: index is original order, completed_count tracks actual progress
                    if len(batch) >= batch_size or completed_count[0] == total:
                        self.batch_progress.emit(batch[:], total, completed_count[0])
                        batch.clear()

                try:
                    # Use parallel processing with 4 workers (handles ABF, SMRX, EDF)
                    read_metadata_parallel(self.files, progress_callback=callback, max_workers=4)

                    # Emit any remaining items in batch (safety check)
                    if batch:
                        self.batch_progress.emit(batch[:], len(self.files), completed_count[0])
                        batch.clear()

                    self.finished.emit(protocols)
                except Exception as e:
                    print(f"[project-builder] Error during parallel loading: {e}")
                    import traceback
                    traceback.print_exc()
                    self.finished.emit(protocols)  # Still emit finish signal

        # Show progress bar
        self.projectProgressBar.setVisible(True)
        self.projectProgressBar.setValue(0)
        self.projectProgressBar.setFormat(f"Loading metadata: 0/{len(data_files)} (0%)")

        self._metadata_thread = MetadataThread(data_files)
        self._metadata_thread.batch_progress.connect(self._update_file_metadata_batch)
        self._metadata_thread.finished.connect(self._metadata_finished)
        self._metadata_thread.start()
        print(f"[project-builder] Started background loading for {len(data_files)} files (offset={start_row_offset})")

    def _update_file_metadata_batch(self, batch, total, completed):
        """Update table cells with a batch of loaded metadata (called from main thread via signal)."""
        # Get the row offset (for additive scans)
        row_offset = getattr(self, '_metadata_row_offset', 0)

        # Process all items in the batch
        for index, metadata in batch:
            if metadata:
                # Update internal data structures
                if index < len(self._discovered_files_data):
                    self._discovered_files_data[index]['protocol'] = metadata['protocol']
                    self._discovered_files_data[index]['channel_count'] = metadata.get('channel_count', 0)
                    self._discovered_files_data[index]['sweep_count'] = metadata.get('sweep_count', 0)
                    self._discovered_files_data[index]['stim_channels'] = metadata.get('stim_channels', [])
                    self._discovered_files_data[index]['stim_frequency'] = metadata.get('stim_frequency', '')
                    # Auto-fill stim_type from detected frequency
                    if metadata.get('stim_frequency') and not self._discovered_files_data[index].get('stim_type'):
                        self._discovered_files_data[index]['stim_type'] = metadata.get('stim_frequency')

                # Master list index includes the offset
                master_idx = row_offset + index
                if master_idx < len(self._master_file_list):
                    task = self._master_file_list[master_idx]
                    task['protocol'] = metadata['protocol']
                    task['channel_count'] = metadata.get('channel_count', 0)
                    task['sweep_count'] = metadata.get('sweep_count', 0)
                    task['stim_channels'] = metadata.get('stim_channels', [])
                    task['stim_frequency'] = metadata.get('stim_frequency', '')
                    # Auto-fill stim_type from detected frequency
                    if metadata.get('stim_frequency') and not task.get('stim_type'):
                        task['stim_type'] = metadata.get('stim_frequency')
                    # Auto-fill events_channel from stim_channels
                    if metadata.get('stim_channels') and not task.get('events_channel'):
                        task['events_channel'] = ', '.join(metadata.get('stim_channels', []))

                    # Update the model row to reflect the changes
                    if master_idx < self._file_table_model.rowCount():
                        self._file_table_model.update_row(master_idx, task)

        # Update progress bar and status using actual completed count
        if batch:
            progress_pct = int(completed / total * 100)
            self.projectProgressBar.setValue(progress_pct)
            self.projectProgressBar.setFormat(f"Loading metadata: {completed}/{total} ({progress_pct}%)")
            self._log_status_message(f"Loading metadata... {completed}/{total}", 0)

    def _metadata_finished(self, protocols):
        """Called when background loading completes."""
        # Auto-fit columns with padding
        self._auto_fit_table_columns()

        # Apply row styling (parent vs sub-row colors)
        self._apply_all_row_styling()

        # Hide progress bar
        self.projectProgressBar.setVisible(False)
        self.projectProgressBar.setValue(0)

        # Count total files in master list
        total_files = len(self._master_file_list)
        new_files = len(self._discovered_files_data)

        summary_text = f"Summary: {total_files} ABF files | {len(protocols)} protocols"
        self.summaryLabel.setText(summary_text)
        self._log_status_message(f"✓ Loaded metadata for {new_files} files ({total_files} total)", 3000)
        print(f"[project-builder] Complete! {len(protocols)} protocols: {sorted(protocols)}")

        # Re-enable scan button
        self.scanFilesButton.setEnabled(True)

        # Clear any active filter to show all rows
        if hasattr(self, 'tableFilterEdit') and self.tableFilterEdit.text():
            self._on_table_filter_changed()

        # Reset row offset
        self._metadata_row_offset = 0

        # Autosave after scan completes
        self._project_autosave()

        # Populate the Consolidation tab's source list with project files
        self._populate_consolidation_source_list()

    def _reload_incomplete_metadata(self, files_with_indices: list):
        """
        Reload metadata for files that have incomplete data (e.g., protocol='Loading...').

        Args:
            files_with_indices: List of (row_index, file_path) tuples
        """
        from PyQt6.QtCore import QThread, pyqtSignal

        if not files_with_indices:
            return

        # Extract just the file paths for the parallel reader
        file_paths = [fp for _, fp in files_with_indices]
        row_indices = [idx for idx, _ in files_with_indices]

        class ReloadMetadataThread(QThread):
            batch_progress = pyqtSignal(list, int, int)  # [(orig_index, row_index, metadata), ...], total, completed
            finished = pyqtSignal(set)

            def __init__(self, file_paths, row_indices):
                super().__init__()
                self.file_paths = file_paths
                self.row_indices = row_indices

            def run(self):
                from core.fast_abf_reader import read_file_metadata_fast
                protocols = set()
                results = []

                for i, (file_path, row_idx) in enumerate(zip(self.file_paths, self.row_indices)):
                    try:
                        metadata = read_file_metadata_fast(file_path)
                        if metadata:
                            protocols.add(metadata.get('protocol', 'Unknown'))
                        results.append((i, row_idx, metadata))
                    except Exception as e:
                        print(f"[metadata-reload] Error reading {file_path}: {e}")
                        results.append((i, row_idx, None))

                    # Emit progress every 10 files
                    if len(results) % 10 == 0 or i == len(self.file_paths) - 1:
                        self.batch_progress.emit(results[:], len(self.file_paths), len(results))
                        results.clear()

                # Emit any remaining
                if results:
                    self.batch_progress.emit(results, len(self.file_paths), len(self.file_paths))

                self.finished.emit(protocols)

        def on_reload_batch(batch, total, completed):
            for orig_idx, row_idx, metadata in batch:
                if metadata and row_idx < len(self._master_file_list):
                    task = self._master_file_list[row_idx]
                    task['protocol'] = metadata.get('protocol', 'Unknown')
                    task['channel_count'] = metadata.get('channel_count', 0)
                    task['sweep_count'] = metadata.get('sweep_count', 0)
                    task['stim_channels'] = metadata.get('stim_channels', [])
                    task['stim_frequency'] = metadata.get('stim_frequency', '')

                    # Update table display
                    if row_idx < self._file_table_model.rowCount():
                        self._file_table_model.update_row(row_idx, task)

            # Update progress
            progress_pct = int(completed / total * 100)
            self.projectProgressBar.setValue(progress_pct)
            self.projectProgressBar.setFormat(f"Reloading metadata: {completed}/{total} ({progress_pct}%)")

        def on_reload_finished(protocols):
            self.projectProgressBar.setVisible(False)
            self.scanFilesButton.setEnabled(True)
            self._log_status_message(f"✓ Reloaded metadata ({len(protocols)} protocols found)", 3000)
            self._project_autosave()

        # Show progress bar
        self.projectProgressBar.setVisible(True)
        self.projectProgressBar.setValue(0)
        self.projectProgressBar.setFormat(f"Reloading metadata: 0/{len(file_paths)} (0%)")

        self._reload_thread = ReloadMetadataThread(file_paths, row_indices)
        self._reload_thread.batch_progress.connect(on_reload_batch)
        self._reload_thread.finished.connect(on_reload_finished)
        self._reload_thread.start()

    # NOTE: Old experiment-based methods removed (on_project_add_files, _add_files_to_experiment)
    # We now use the master file list approach where all files are in a flat table with metadata columns

    def on_project_clear_files(self):
        """Clear the discovered files table and master file list with confirmation."""
        # Count files to show in warning
        file_count = len(self._master_file_list) if self._master_file_list else 0

        if file_count == 0:
            self._log_status_message("No files to clear", 1500)
            return

        # Show confirmation dialog
        from PyQt6.QtWidgets import QMessageBox
        msg = QMessageBox(self)
        msg.setIcon(QMessageBox.Icon.Warning)
        msg.setWindowTitle("Clear All Files?")
        msg.setText(f"Are you sure you want to clear {file_count} files from the list?")
        msg.setInformativeText(
            "This will remove all files from the table.\n"
            "Your original data files will NOT be deleted.\n\n"
            "Any unsaved metadata changes will be lost."
        )
        msg.setStandardButtons(QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        msg.setDefaultButton(QMessageBox.StandardButton.No)

        if msg.exec() != QMessageBox.StandardButton.Yes:
            return

        # User confirmed - clear everything
        self._file_table_model.clear()
        self._discovered_files_data = []
        self._master_file_list = []
        self.summaryLabel.setText("Summary: No files scanned")
        # Hide progress bar
        self.projectProgressBar.setVisible(False)
        self.projectProgressBar.setValue(0)
        # Clear filter
        if hasattr(self, 'tableFilterEdit'):
            self.tableFilterEdit.clear()
        if hasattr(self, 'filterCountLabel'):
            self.filterCountLabel.setText("")
        self._log_status_message("Cleared all files", 1500)

    # =========================================================================
    # Notes Files Tab Methods
    # =========================================================================

    def _init_notes_files_model(self):
        """Initialize the notes files table model."""
        from PyQt6.QtGui import QStandardItemModel, QStandardItem
        from PyQt6.QtCore import Qt

        if not hasattr(self, 'notesFilesTable'):
            return

        self._notes_model = QStandardItemModel()
        # Columns: Use, File Name, Actions, Type, Matches, Location, Size, Modified
        self._notes_model.setHorizontalHeaderLabels(['Use', 'File Name', '', 'Type', 'Matches', 'Location', 'Size', 'Modified'])
        self.notesFilesTable.setModel(self._notes_model)

        # Set up Actions column delegate (Folder/Open/Preview buttons)
        from core.file_table_delegates import NotesActionsDelegate
        self._notes_actions_delegate = NotesActionsDelegate(self.notesFilesTable)
        self._notes_actions_delegate.folder_clicked.connect(self._on_notes_action_folder)
        self._notes_actions_delegate.open_clicked.connect(self._on_notes_action_open)
        self._notes_actions_delegate.preview_clicked.connect(self._on_notes_action_preview)
        self.notesFilesTable.setItemDelegateForColumn(2, self._notes_actions_delegate)

        # Enable mouse tracking for hover effects
        self.notesFilesTable.setMouseTracking(True)

        # Set column widths - auto-fit to content with reasonable minimums
        from PyQt6.QtWidgets import QHeaderView
        header = self.notesFilesTable.horizontalHeader()
        header.setStretchLastSection(True)
        # Use ResizeToContents for most columns, with minimum widths
        header.setSectionResizeMode(0, QHeaderView.ResizeMode.Fixed)  # Use checkbox - fixed
        header.setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)  # File Name - stretch
        header.setSectionResizeMode(2, QHeaderView.ResizeMode.Fixed)  # Actions - fixed
        header.setSectionResizeMode(3, QHeaderView.ResizeMode.ResizeToContents)  # Type
        header.setSectionResizeMode(4, QHeaderView.ResizeMode.ResizeToContents)  # Matches
        header.setSectionResizeMode(5, QHeaderView.ResizeMode.ResizeToContents)  # Location
        header.setSectionResizeMode(6, QHeaderView.ResizeMode.ResizeToContents)  # Size
        header.setSectionResizeMode(7, QHeaderView.ResizeMode.ResizeToContents)  # Modified
        self.notesFilesTable.setColumnWidth(0, 50)   # Use (checkbox) - fixed width for "Use" header
        self.notesFilesTable.setColumnWidth(2, 82)   # Actions (Folder/Open/Preview) - 3 buttons

        # Double-click to open
        self.notesFilesTable.doubleClicked.connect(self.on_notes_open)

        # Connect model changes to track checkbox state
        self._notes_model.itemChanged.connect(self._on_notes_item_changed)

        # Hide the bottom Open/Preview buttons since they're now in the table
        if hasattr(self, 'openNoteButton'):
            self.openNoteButton.setVisible(False)
        if hasattr(self, 'previewNoteButton'):
            self.previewNoteButton.setVisible(False)

    def _add_notes_action_buttons(self):
        """Add additional action buttons to the Notes tab programmatically."""
        from PyQt6.QtWidgets import QPushButton, QHBoxLayout, QWidget

        # Find the notes actions layout or create one
        if not hasattr(self, 'notesFilesTable'):
            return

        # Try to find existing action buttons layout by name
        # The notesActionsLayout is a QHBoxLayout containing the action buttons
        if hasattr(self, 'linkNoteButton'):
            # Find the notesActionsLayout - it's the layout containing linkNoteButton
            parent_layout = None
            if hasattr(self, 'notesActionsLayout'):
                # Direct access if available from UI
                parent_layout = self.notesActionsLayout
            else:
                # Fallback: find QHBoxLayout containing the button
                parent_widget = self.linkNoteButton.parentWidget()
                if parent_widget:
                    layout = parent_widget.layout()
                    if layout:
                        # Look for the QHBoxLayout that contains linkNoteButton
                        for i in range(layout.count()):
                            item = layout.itemAt(i)
                            if item and item.layout():
                                inner_layout = item.layout()
                                # Check if this layout contains linkNoteButton
                                for j in range(inner_layout.count()):
                                    widget = inner_layout.itemAt(j).widget()
                                    if widget == self.linkNoteButton:
                                        parent_layout = inner_layout
                                        break
                                if parent_layout:
                                    break
            if parent_layout:
                # Common button style matching UI file buttons
                button_style = """
                    QPushButton {
                        background-color: #6c757d;
                        color: white;
                        border-radius: 4px;
                        padding: 4px 8px;
                    }
                    QPushButton:hover {
                        background-color: #5a6268;
                    }
                    QPushButton:pressed {
                        background-color: #545b62;
                    }
                """

                # Add Scan References button (after Link to File, before spacer)
                # Layout order: Search(0), Open(1), Preview(2), Link(3), [insert here], Spacer
                self.scanReferencesButton = QPushButton("Scan for Refs")
                self.scanReferencesButton.setToolTip("Scan notes files for references to data file names")
                self.scanReferencesButton.setMinimumSize(70, 26)
                self.scanReferencesButton.setMaximumWidth(100)
                self.scanReferencesButton.setStyleSheet(button_style)
                self.scanReferencesButton.clicked.connect(self.on_notes_scan_references)
                parent_layout.insertWidget(4, self.scanReferencesButton)

                # Add Toggle Unmatched button
                self.toggleUnmatchedButton = QPushButton("Hide Unmatched")
                self.toggleUnmatchedButton.setToolTip("Hide/show files with no data file matches")
                self.toggleUnmatchedButton.setMinimumSize(100, 26)
                self.toggleUnmatchedButton.setMaximumWidth(120)
                self.toggleUnmatchedButton.setStyleSheet(button_style)
                self.toggleUnmatchedButton.clicked.connect(self.on_notes_toggle_unmatched)
                parent_layout.insertWidget(5, self.toggleUnmatchedButton)

                # Add Build AI Index button (blue for distinction)
                ai_button_style = """
                    QPushButton {
                        background-color: #0d6efd;
                        color: white;
                        border: 1px solid #0b5ed7;
                        border-radius: 4px;
                        padding: 4px 8px;
                    }
                    QPushButton:hover {
                        background-color: #0b5ed7;
                    }
                    QPushButton:pressed {
                        background-color: #0a58ca;
                    }
                """
                self.buildAIIndexButton = QPushButton("🤖 AI Index")
                self.buildAIIndexButton.setToolTip("Build AI-powered metadata lookup table from selected notes")
                self.buildAIIndexButton.setMinimumSize(80, 26)
                self.buildAIIndexButton.setMaximumWidth(100)
                self.buildAIIndexButton.setStyleSheet(ai_button_style)
                self.buildAIIndexButton.clicked.connect(self.on_notes_build_ai_index)
                self.buildAIIndexButton.setVisible(False)  # Hidden - not currently used
                parent_layout.insertWidget(6, self.buildAIIndexButton)

                # Move filter box and count label from search row to action row
                # Find them in the search layout and re-parent to actions layout
                if hasattr(self, 'notesFilterEdit') and hasattr(self, 'notesFilterCountLabel'):
                    filter_edit = self.notesFilterEdit
                    count_label = self.notesFilterCountLabel

                    # Remove from current parent layout
                    old_layout = filter_edit.parentWidget().layout() if filter_edit.parentWidget() else None
                    if old_layout:
                        # Find and remove from old layout
                        for i in range(old_layout.count()):
                            item = old_layout.itemAt(i)
                            if item and item.widget() == filter_edit:
                                old_layout.takeAt(i)
                                break
                        for i in range(old_layout.count()):
                            item = old_layout.itemAt(i)
                            if item and item.widget() == count_label:
                                old_layout.takeAt(i)
                                break

                    # Insert into action row (after the buttons, before spacer)
                    # Buttons are at 0-6, spacer is at end
                    insert_pos = parent_layout.count() - 1  # Before spacer
                    if insert_pos < 0:
                        insert_pos = 0
                    parent_layout.insertWidget(insert_pos, filter_edit)
                    parent_layout.insertWidget(insert_pos + 1, count_label)
                    print("[notes] Moved filter box to action row")

                print("[notes] Added Scan Refs and Unmatched buttons to action row")

    def on_notes_browse(self):
        """Browse for a folder to search for notes files."""
        from PyQt6.QtWidgets import QFileDialog

        # Default to parent of project directory if set
        start_dir = ""
        if self._notes_directory:
            start_dir = str(self._notes_directory)
        elif self._project_directory:
            start_dir = str(Path(self._project_directory).parent)

        folder = QFileDialog.getExistingDirectory(
            self,
            "Select Notes Folder",
            start_dir
        )

        if folder:
            self._notes_directory = Path(folder)
            if hasattr(self, 'notesFolderEdit'):
                self.notesFolderEdit.setText(str(folder))
            self._log_status_message(f"Notes folder set to: {folder}", 2000)

    def _set_default_notes_directory(self):
        """Set default notes directory to parent of project directory."""
        if self._project_directory:
            # Default to parent folder
            parent = Path(self._project_directory).parent
            self._notes_directory = parent
            if hasattr(self, 'notesFolderEdit'):
                self.notesFolderEdit.setText(str(parent))

    def on_notes_search(self):
        """Search for notes files in project folder and parent directories.

        Search pattern:
        - Project folder + all subfolders (recursive)
        - Parent folder (immediate files only, no subdirs)
        - Grandparent folder (immediate files only, no subdirs)
        """
        from PyQt6.QtGui import QStandardItem
        from PyQt6.QtCore import Qt
        import os
        from datetime import datetime

        if not self._project_directory:
            self._log_status_message("Please select a project directory first", 2000)
            return

        project_path = Path(self._project_directory)

        # Get selected file types
        extensions = []
        if hasattr(self, 'notesWordCheckbox') and self.notesWordCheckbox.isChecked():
            extensions.extend(['.docx', '.doc'])
        if hasattr(self, 'notesExcelCheckbox') and self.notesExcelCheckbox.isChecked():
            extensions.extend(['.xlsx', '.xls'])
        if hasattr(self, 'notesCsvCheckbox') and self.notesCsvCheckbox.isChecked():
            extensions.append('.csv')
        if hasattr(self, 'notesTxtCheckbox') and self.notesTxtCheckbox.isChecked():
            extensions.append('.txt')
        if hasattr(self, 'notesPdfCheckbox') and self.notesPdfCheckbox.isChecked():
            extensions.append('.pdf')

        if not extensions:
            self._log_status_message("Please select at least one file type", 2000)
            return

        # Clear existing data
        self._notes_model.removeRows(0, self._notes_model.rowCount())
        self._notes_files_data = []

        found_files = set()  # Use set to avoid duplicates

        # Folders to exclude (analysis output folders)
        excluded_folders = {'pleth_app_analysis', 'pleth_app_analysis'}  # case-insensitive check

        # File suffixes to exclude (analysis output files)
        excluded_suffixes = (
            '_breaths.csv', '_means_by_time.csv', '_timeseries.csv',
            '_events.csv', '_summary.pdf', '_session.npz', '.pleth.npz',
            '_ml_training.npz', '_stim_aligned.csv'
        )

        def is_excluded_file(file_path: Path) -> bool:
            """Check if file should be excluded from notes search."""
            # Check if in excluded folder
            for part in file_path.parts:
                if part.lower() in excluded_folders:
                    return True
            # Check if has excluded suffix
            name_lower = file_path.name.lower()
            for suffix in excluded_suffixes:
                if name_lower.endswith(suffix):
                    return True
            return False

        # 1. Search project folder recursively (includes all subfolders)
        for ext in extensions:
            for f in project_path.rglob(f'*{ext}'):
                if not is_excluded_file(f):
                    found_files.add(f)

        # 2. Search parent folder (immediate files only, no subdirs)
        parent_path = project_path.parent
        if parent_path.exists() and parent_path != project_path:
            for ext in extensions:
                for f in parent_path.glob(f'*{ext}'):
                    if not is_excluded_file(f):
                        found_files.add(f)

        # 3. Search grandparent folder (immediate files only, no subdirs)
        grandparent_path = parent_path.parent
        if grandparent_path.exists() and grandparent_path != parent_path:
            for ext in extensions:
                for f in grandparent_path.glob(f'*{ext}'):
                    if not is_excluded_file(f):
                        found_files.add(f)

        # Update the notes folder display to show search scope
        if hasattr(self, 'notesFolderEdit'):
            self.notesFolderEdit.setText(f"{project_path.name} + parent folders")
        self._notes_directory = str(project_path)

        # Populate table
        for file_path in sorted(found_files):
            try:
                stat = file_path.stat()
                size_kb = stat.st_size / 1024
                if size_kb < 1024:
                    size_str = f"{size_kb:.1f} KB"
                else:
                    size_str = f"{size_kb/1024:.1f} MB"
                mod_time = datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M')

                # Determine location relative to project
                if file_path.parent == project_path or project_path in file_path.parents:
                    try:
                        rel_path = file_path.relative_to(project_path)
                        location = str(rel_path.parent) if rel_path.parent != Path('.') else '.'
                    except ValueError:
                        location = str(file_path.parent)
                elif file_path.parent == parent_path:
                    location = '../'
                elif file_path.parent == grandparent_path:
                    location = '../../'
                else:
                    location = str(file_path.parent)

                # Create checkbox item for "Use" column
                use_item = QStandardItem()
                use_item.setCheckable(True)
                use_item.setCheckState(Qt.CheckState.Checked)  # Default to checked
                use_item.setEditable(False)

                # File name item with path stored
                name_item = QStandardItem(file_path.name)
                name_item.setData(str(file_path), Qt.ItemDataRole.UserRole)

                # Actions column - empty (rendered by delegate)
                actions_item = QStandardItem("")
                actions_item.setEditable(False)

                # Type column - file extension
                type_item = QStandardItem(file_path.suffix.lower())

                # Matches column - will be populated by scan
                matches_item = QStandardItem("—")
                matches_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)

                row = [
                    use_item,                    # 0: Use
                    name_item,                   # 1: File Name
                    actions_item,                # 2: Actions (buttons)
                    type_item,                   # 3: Type
                    matches_item,                # 4: Matches
                    QStandardItem(location),     # 5: Location
                    QStandardItem(size_str),     # 6: Size
                    QStandardItem(mod_time),     # 7: Modified
                ]

                self._notes_model.appendRow(row)

                # Store file metadata for project save
                self._notes_files_data.append({
                    'name': file_path.name,
                    'path': str(file_path),
                    'type': file_path.suffix.lower(),
                    'location': location,
                    'size': size_str,
                    'modified': mod_time,
                    'use_as_notes': True,  # Track checkbox state
                    'matches': [],  # Will store matched ABF filenames
                    'match_count': 0
                })
            except Exception as e:
                print(f"Error processing {file_path}: {e}")

        # Update summary and filter count
        count = len(self._notes_files_data)
        self._log_status_message(f"Found {count} notes files", 2000)
        if hasattr(self, 'notesFilterCountLabel'):
            self.notesFilterCountLabel.setText(f"{count} files")

    def _on_notes_filter_changed(self):
        """Filter notes table rows based on search text."""
        if not hasattr(self, 'notesFilterEdit') or not hasattr(self, 'notesFilesTable'):
            return

        filter_text = self.notesFilterEdit.text().strip().lower()
        table = self.notesFilesTable
        total_count = self._notes_model.rowCount()
        visible_count = 0

        # If no filter, show all rows
        if not filter_text:
            for row in range(total_count):
                table.setRowHidden(row, False)
            visible_count = total_count
        else:
            # Filter rows - search across name, type, and location columns
            for row in range(total_count):
                match = False
                for col in range(self._notes_model.columnCount()):
                    item = self._notes_model.item(row, col)
                    if item and filter_text in item.text().lower():
                        match = True
                        break
                table.setRowHidden(row, not match)
                if match:
                    visible_count += 1

        # Update filter count label
        if hasattr(self, 'notesFilterCountLabel'):
            if filter_text:
                self.notesFilterCountLabel.setText(f"{visible_count} of {total_count}")
            else:
                self.notesFilterCountLabel.setText(f"{total_count} files")

    def _on_notes_item_changed(self, item):
        """Handle changes to notes table items (e.g., checkbox state)."""
        from PyQt6.QtCore import Qt

        # Only handle checkbox column (column 0)
        if item.column() != 0:
            return

        row = item.row()
        if row < len(self._notes_files_data):
            is_checked = item.checkState() == Qt.CheckState.Checked
            self._notes_files_data[row]['use_as_notes'] = is_checked
            print(f"[notes] File '{self._notes_files_data[row]['name']}' use_as_notes = {is_checked}")

    def on_notes_scan_references(self):
        """Scan notes files marked with 'Use' checkbox to find references to ABF data files.

        Searches file contents for ABF filenames (with or without .abf extension).
        Updates the 'Matches' column with count of found references.
        Uses background thread to keep UI responsive for network drives.
        Only scans files where use_as_notes is True.

        Optimization: Copies files to local temp directory first for faster network access.
        """
        from PyQt6.QtCore import Qt, QThread, pyqtSignal
        from PyQt6.QtWidgets import QProgressDialog
        import tempfile
        import shutil

        if not self._notes_files_data:
            self._log_status_message("No notes files to scan. Search for files first.", 2000)
            return

        # Filter to only files with 'Use' checkbox checked
        files_to_scan = [(i, f) for i, f in enumerate(self._notes_files_data) if f.get('use_as_notes', False)]

        if not files_to_scan:
            self._log_status_message("No files selected for scanning. Check the 'Use' boxes for files to scan.", 2000)
            return

        # Get list of ABF filenames from the file table
        abf_filenames = self._get_project_abf_filenames()
        if not abf_filenames:
            self._log_status_message("No data files in project to match against.", 2000)
            return

        print(f"[notes-scan] Scanning {len(files_to_scan)} selected notes files for {len(abf_filenames)} ABF references...")

        # Create worker thread for scanning
        class ScanWorker(QThread):
            # stage (0=copy, 1=scan), progress_idx, original_idx, filename, matches
            progress_update = pyqtSignal(int, int, int, str, list)
            status_update = pyqtSignal(str)  # Status message
            finished = pyqtSignal(int, int)  # files_with_matches, total_matches

            def __init__(self, files_to_scan, abf_names, scan_func):
                super().__init__()
                self.files_to_scan = files_to_scan  # List of (original_index, file_info)
                self.abf_names = abf_names
                self.scan_func = scan_func
                self._cancelled = False
                self.cache_dir = None

            def cancel(self):
                self._cancelled = True

            def run(self):
                import tempfile
                import shutil
                import time
                import hashlib
                from pathlib import Path

                total_matches = 0
                files_with_matches = 0

                # Use persistent cache directory (not deleted after scan)
                cache_base = Path(tempfile.gettempdir()) / "physiometrics_notes_cache"
                cache_base.mkdir(exist_ok=True)
                self.cache_dir = cache_base

                print(f"[notes-scan] ========== BENCHMARK START ==========")
                print(f"[notes-scan] Cache directory: {cache_base}")
                print(f"[notes-scan] Files to scan: {len(self.files_to_scan)}")
                print(f"[notes-scan] ABF names to match: {len(self.abf_names)}")

                overall_start = time.perf_counter()

                # Stage 1: Copy files locally (only if not already cached)
                self.status_update.emit("Checking cache / copying files...")
                local_files = []  # (original_idx, local_path, file_info)

                copy_start = time.perf_counter()
                copied_count = 0
                cached_count = 0

                for progress_idx, (original_idx, file_info) in enumerate(self.files_to_scan):
                    if self._cancelled:
                        break

                    src_path = Path(file_info['path'])

                    # Create cache filename based on path hash + original name
                    path_hash = hashlib.md5(str(src_path).encode()).hexdigest()[:8]
                    cache_name = f"{path_hash}_{src_path.name}"
                    cache_path = cache_base / cache_name

                    try:
                        # Check if cached copy exists and is up-to-date
                        if cache_path.exists():
                            src_mtime = src_path.stat().st_mtime
                            cache_mtime = cache_path.stat().st_mtime
                            if cache_mtime >= src_mtime:
                                # Use cached copy
                                local_files.append((original_idx, str(cache_path), file_info))
                                cached_count += 1
                                self.progress_update.emit(0, progress_idx, original_idx, f"(cached) {file_info['name']}", [])
                                continue

                        # Copy file to cache
                        file_start = time.perf_counter()
                        shutil.copy2(src_path, cache_path)
                        file_time = (time.perf_counter() - file_start) * 1000
                        print(f"[notes-scan]   Copy: {file_info['name']} - {file_time:.1f}ms")

                        local_files.append((original_idx, str(cache_path), file_info))
                        copied_count += 1
                        self.progress_update.emit(0, progress_idx, original_idx, file_info['name'], [])

                    except Exception as e:
                        print(f"[notes-scan] Failed to copy {src_path.name}: {e}")

                copy_time = time.perf_counter() - copy_start
                print(f"[notes-scan] COPY STAGE: {copy_time:.2f}s total ({copied_count} copied, {cached_count} from cache)")

                if self._cancelled:
                    return

                # Stage 2: Scan local copies (parallel with fallback to sequential)
                self.status_update.emit("Scanning files (parallel)...")

                scan_start = time.perf_counter()
                scan_results = {}  # original_idx -> (file_info, matches, scan_time)

                def scan_single_file(args):
                    """Worker function for parallel scanning."""
                    original_idx, local_path, file_info = args
                    file_start = time.perf_counter()
                    try:
                        matches = self.scan_func(local_path, self.abf_names)
                    except Exception as e:
                        print(f"[notes-scan] Error scanning {file_info['name']}: {e}")
                        matches = []
                    file_time = (time.perf_counter() - file_start) * 1000
                    return (original_idx, file_info, matches, file_time, local_path)

                # Try parallel processing first
                use_parallel = True
                try:
                    from concurrent.futures import ThreadPoolExecutor, as_completed
                    import os

                    # Use number of CPUs, but cap at 8 to avoid overwhelming the system
                    max_workers = min(os.cpu_count() or 4, 8)
                    print(f"[notes-scan] Using parallel scanning with {max_workers} workers")

                    with ThreadPoolExecutor(max_workers=max_workers) as executor:
                        # Submit all scan tasks
                        futures = {executor.submit(scan_single_file, (orig_idx, path, info)): orig_idx
                                   for orig_idx, path, info in local_files}

                        completed = 0
                        for future in as_completed(futures):
                            if self._cancelled:
                                executor.shutdown(wait=False, cancel_futures=True)
                                break

                            original_idx, file_info, matches, file_time, local_path = future.result()
                            suffix = Path(local_path).suffix.lower()
                            print(f"[notes-scan]   Scan: {file_info['name']} ({suffix}) - {file_time:.1f}ms - {len(matches)} matches")

                            if matches:
                                files_with_matches += 1
                                total_matches += len(matches)

                            # Emit progress (completed count as progress_idx)
                            self.progress_update.emit(1, completed, original_idx, file_info['name'], matches)
                            completed += 1

                except Exception as parallel_error:
                    # Fallback to sequential scanning
                    print(f"[notes-scan] Parallel scanning failed ({parallel_error}), falling back to sequential...")
                    self.status_update.emit("Scanning files (sequential fallback)...")
                    use_parallel = False

                    for progress_idx, (original_idx, local_path, file_info) in enumerate(local_files):
                        if self._cancelled:
                            break

                        # Benchmark individual file scan
                        file_start = time.perf_counter()
                        try:
                            matches = self.scan_func(local_path, self.abf_names)
                        except Exception as e:
                            print(f"[notes-scan] Error scanning {file_info['name']}: {e}")
                            matches = []
                        file_time = (time.perf_counter() - file_start) * 1000

                        suffix = Path(local_path).suffix.lower()
                        print(f"[notes-scan]   Scan: {file_info['name']} ({suffix}) - {file_time:.1f}ms - {len(matches)} matches")

                        if matches:
                            files_with_matches += 1
                            total_matches += len(matches)

                        self.progress_update.emit(1, progress_idx, original_idx, file_info['name'], matches)

                scan_time = time.perf_counter() - scan_start
                mode_str = "parallel" if use_parallel else "sequential (fallback)"
                overall_time = time.perf_counter() - overall_start

                print(f"[notes-scan] SCAN STAGE ({mode_str}): {scan_time:.2f}s total")
                print(f"[notes-scan] ========== BENCHMARK END ==========")
                print(f"[notes-scan] TOTAL TIME: {overall_time:.2f}s")
                print(f"[notes-scan] Results: {files_with_matches} files with {total_matches} matches")

                # Note: Cache is NOT deleted - files persist for faster re-scans

                self.finished.emit(files_with_matches, total_matches)

        # Create progress dialog
        progress = QProgressDialog("Copying files locally...", "Cancel", 0, len(files_to_scan) * 2, self)
        progress.setWindowModality(Qt.WindowModality.WindowModal)
        progress.setMinimumDuration(0)
        progress.setValue(0)

        # Create and store worker
        self._scan_worker = ScanWorker(files_to_scan, abf_filenames, self._scan_file_for_abf_references)

        def on_status(msg):
            progress.setLabelText(msg)

        def on_progress(stage, progress_idx, original_idx, filename, matches):
            if progress.wasCanceled():
                self._scan_worker.cancel()
                return

            # Stage 0 = copying, Stage 1 = scanning
            total_progress = progress_idx + (len(files_to_scan) if stage == 1 else 0)
            progress.setValue(total_progress + 1)

            if stage == 0:
                progress.setLabelText(f"Copying ({progress_idx+1}/{len(files_to_scan)}): {filename}")
            else:
                progress.setLabelText(f"Scanning ({progress_idx+1}/{len(files_to_scan)}): {filename}")

                # Update file info using original index
                if original_idx < len(self._notes_files_data):
                    self._notes_files_data[original_idx]['matches'] = matches
                    self._notes_files_data[original_idx]['match_count'] = len(matches)

                # Update table display using original index
                if original_idx < self._notes_model.rowCount():
                    matches_item = self._notes_model.item(original_idx, 4)  # Column 4 = Matches
                    if matches_item:
                        count = len(matches)
                        matches_item.setText(str(count) if count > 0 else "0")

        def on_finished(files_with_matches, total_matches):
            progress.close()
            msg = f"Scan complete: {files_with_matches} files with {total_matches} total references"
            print(f"[notes-scan] {msg}")
            self._log_status_message(msg, 3000)

            # Update linked_notes column in file table
            self._update_linked_notes_column()

            # Auto-hide unmatched files after scan
            if not hasattr(self, '_notes_hide_unmatched'):
                self._notes_hide_unmatched = False
            self._notes_hide_unmatched = True
            self._apply_hide_unmatched(show_status=False)  # Don't show redundant status

        # Connect signals
        self._scan_worker.status_update.connect(on_status)
        self._scan_worker.progress_update.connect(on_progress)
        self._scan_worker.finished.connect(on_finished)

        # Handle cancel
        progress.canceled.connect(self._scan_worker.cancel)

        # Start scanning
        self._scan_worker.start()

    def _get_project_abf_filenames(self) -> list:
        """Get list of ABF filenames from the project file table.

        Returns list of filename stems (without .abf extension) for matching.
        """
        abf_names = []

        # Try to get from file table model
        if hasattr(self, 'project_builder') and self.project_builder:
            pb = self.project_builder
            if hasattr(pb, 'file_table_model') and pb.file_table_model:
                model = pb.file_table_model
                for row in range(model.rowCount()):
                    if row < len(model._files):
                        file_name = model._files[row].get('file_name', '')
                        if file_name:
                            # Store stem (without extension) for flexible matching
                            stem = Path(file_name).stem
                            abf_names.append(stem)

        # Also get from master file list if available
        if hasattr(self, '_master_file_list') and self._master_file_list:
            for file_info in self._master_file_list:
                # file_info is a dict with 'file_path' and 'file_name' keys
                if isinstance(file_info, dict):
                    file_name = file_info.get('file_name', '')
                    if file_name:
                        stem = Path(file_name).stem
                        if stem not in abf_names:
                            abf_names.append(stem)
                elif isinstance(file_info, (str, Path)):
                    # Fallback for string paths
                    stem = Path(file_info).stem
                    if stem not in abf_names:
                        abf_names.append(stem)

        return abf_names

    def _scan_file_for_abf_references(self, file_path: str, abf_filenames: list) -> list:
        """Scan a single file for references to ABF filenames.

        Args:
            file_path: Path to the notes file
            abf_filenames: List of ABF filename stems to search for

        Returns:
            List of matched ABF filename stems found in the file
        """
        import re

        file_path = Path(file_path)
        if not file_path.exists():
            return []

        matches = []
        suffix = file_path.suffix.lower()

        try:
            content = ""

            # Read content based on file type
            if suffix in ['.txt', '.csv']:
                with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                    content = f.read()

            elif suffix == '.xlsx':
                # Use openpyxl in read-only mode with batch regex matching
                # Much faster than per-cell substring checks (O(1) regex vs O(n) checks per cell)
                try:
                    from openpyxl import load_workbook

                    # Build a single regex pattern for all ABF names (case-insensitive)
                    # This is much faster than checking each name individually
                    escaped_names = [re.escape(name) for name in abf_filenames]
                    pattern = re.compile('|'.join(escaped_names), re.IGNORECASE)

                    # Create lookup from lowercase to original case
                    name_lookup = {name.lower(): name for name in abf_filenames}
                    found_set = set()

                    wb = load_workbook(file_path, read_only=True, data_only=True)

                    for sheet in wb.worksheets:
                        # Batch all cell values into one string per sheet for faster regex
                        sheet_text_parts = []
                        for row in sheet.iter_rows(values_only=True):
                            for cell in row:
                                if cell is not None:
                                    sheet_text_parts.append(str(cell))

                        # Single regex search on concatenated sheet text
                        if sheet_text_parts:
                            sheet_text = ' '.join(sheet_text_parts)
                            for match in pattern.findall(sheet_text):
                                found_set.add(name_lookup.get(match.lower(), match))

                        # Early exit if all names found
                        if len(found_set) >= len(abf_filenames):
                            break

                    wb.close()
                    return list(found_set)

                except Exception as e:
                    print(f"[notes-scan] Could not read Excel file {file_path.name}: {e}")
                    return []

            elif suffix == '.xls':
                # Old .xls format - use pandas (slower but works)
                try:
                    import pandas as pd
                    xl = pd.ExcelFile(file_path)
                    parts = []
                    for sheet_name in xl.sheet_names:
                        df = pd.read_excel(xl, sheet_name=sheet_name, dtype=str)
                        parts.append(df.fillna('').astype(str).values.flatten())
                    content = ' '.join(' '.join(p) for p in parts)
                except Exception as e:
                    print(f"[notes-scan] Could not read .xls file {file_path.name}: {e}")
                    return []

            elif suffix in ['.docx']:
                try:
                    from docx import Document
                    doc = Document(file_path)
                    paragraphs = [p.text for p in doc.paragraphs]
                    # Also get table contents
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                paragraphs.append(cell.text)
                    content = ' '.join(paragraphs)
                except ImportError:
                    print(f"[notes-scan] python-docx not installed, skipping {file_path.name}")
                    return []
                except Exception as e:
                    print(f"[notes-scan] Could not read Word file {file_path.name}: {e}")
                    return []

            elif suffix == '.doc':
                # Old .doc format not easily supported
                print(f"[notes-scan] Skipping old .doc format: {file_path.name}")
                return []

            elif suffix == '.pdf':
                # Would need PyPDF2 or similar
                print(f"[notes-scan] PDF scanning not implemented: {file_path.name}")
                return []

            # Search for ABF filename references
            content_lower = content.lower()
            for abf_name in abf_filenames:
                # Search for the filename (case-insensitive)
                # Match with or without .abf extension
                if abf_name.lower() in content_lower:
                    if abf_name not in matches:
                        matches.append(abf_name)

        except Exception as e:
            print(f"[notes-scan] Error scanning {file_path.name}: {e}")

        return matches

    def on_notes_toggle_unmatched(self):
        """Toggle visibility of notes files with no ABF matches."""
        if not hasattr(self, '_notes_hide_unmatched'):
            self._notes_hide_unmatched = False

        self._notes_hide_unmatched = not self._notes_hide_unmatched
        self._apply_hide_unmatched()

    def _apply_hide_unmatched(self, show_status=True):
        """Apply the current hide/show unmatched setting.

        Args:
            show_status: Whether to show status message (False when auto-hiding after scan)
        """
        table = self.notesFilesTable
        hidden_count = 0

        for row in range(self._notes_model.rowCount()):
            if row < len(self._notes_files_data):
                match_count = self._notes_files_data[row].get('match_count', 0)
                if match_count == 0 and self._notes_hide_unmatched:
                    table.setRowHidden(row, True)
                    hidden_count += 1
                else:
                    table.setRowHidden(row, False)

        # Update button text to reflect current state
        if hasattr(self, 'toggleUnmatchedButton'):
            if self._notes_hide_unmatched:
                self.toggleUnmatchedButton.setText("Show Unmatched")
            else:
                self.toggleUnmatchedButton.setText("Hide Unmatched")

        if show_status:
            if self._notes_hide_unmatched:
                self._log_status_message(f"Hidden {hidden_count} files with no matches", 2000)
            else:
                self._log_status_message("Showing all files", 2000)

    def on_notes_build_ai_index(self):
        """Build AI-powered metadata index from selected notes files.

        Sends notes file contents to AI to extract structured metadata
        for each data file, creating a lookup table for auto-fill suggestions.
        """
        from PyQt6.QtWidgets import QMessageBox

        # Get files marked for use
        use_files = [f for f in self._notes_files_data if f.get('use_as_notes', False)]

        if not use_files:
            QMessageBox.warning(
                self,
                "No Notes Selected",
                "Please check the 'Use' column for notes files you want to include in the AI index."
            )
            return

        # Get ABF filenames and their metadata for context
        abf_filenames = self._get_project_abf_filenames()

        # TODO: Implement AI indexing
        # 1. Read content from selected notes files
        # 2. Get keywords/metadata from project table
        # 3. Send to AI (Sonnet) to build structured lookup table
        # 4. Save as project_metadata.json

        QMessageBox.information(
            self,
            "AI Index - Coming Soon",
            f"Ready to index {len(use_files)} notes files.\n\n"
            f"This feature will:\n"
            f"• Read selected notes files\n"
            f"• Send to AI for metadata extraction\n"
            f"• Build lookup table for {len(abf_filenames)} data files\n\n"
            f"Implementation in progress..."
        )

    def on_notes_open(self):
        """Open selected notes file in default application."""
        import os
        import subprocess
        import sys
        from PyQt6.QtCore import Qt

        if not hasattr(self, 'notesFilesTable'):
            return

        indexes = self.notesFilesTable.selectionModel().selectedRows()
        if not indexes:
            self._log_status_message("Please select a file to open", 1500)
            return

        # Get file path from first selected row (column 1 = File Name with path in UserRole)
        row = indexes[0].row()
        path_item = self._notes_model.item(row, 1)  # Column 1 is File Name
        if path_item:
            file_path = path_item.data(Qt.ItemDataRole.UserRole)
            if file_path and Path(file_path).exists():
                try:
                    if sys.platform == 'win32':
                        os.startfile(file_path)
                    elif sys.platform == 'darwin':
                        subprocess.run(['open', file_path])
                    else:
                        subprocess.run(['xdg-open', file_path])
                    self._log_status_message(f"Opened: {Path(file_path).name}", 1500)
                except Exception as e:
                    self._log_status_message(f"Error opening file: {e}", 3000)

    def on_notes_preview(self):
        """Preview contents of selected notes file using the shared preview dialog."""
        from PyQt6.QtCore import Qt

        if not hasattr(self, 'notesFilesTable'):
            return

        indexes = self.notesFilesTable.selectionModel().selectedRows()
        if not indexes:
            self._log_status_message("Please select a file to preview", 1500)
            return

        row = indexes[0].row()
        path_item = self._notes_model.item(row, 1)  # Column 1 = File Name with path in UserRole
        if not path_item:
            return

        file_path = Path(path_item.data(Qt.ItemDataRole.UserRole))
        if not file_path.exists():
            self._log_status_message("File not found", 1500)
            return

        # Use shared preview dialog (no highlighting since not searching for ABF)
        self._show_notes_preview_dialog(
            files=[{'name': file_path.name, 'path': str(file_path)}],
            title=f"Preview: {file_path.name}"
        )

    def _on_notes_action_folder(self, row: int):
        """Open containing folder for notes file at given row."""
        import os
        import subprocess
        import sys
        from PyQt6.QtCore import Qt

        if not hasattr(self, '_notes_model'):
            return

        path_item = self._notes_model.item(row, 1)  # Column 1 is File Name with path in UserRole
        if path_item:
            file_path = path_item.data(Qt.ItemDataRole.UserRole)
            if file_path and Path(file_path).exists():
                folder_path = Path(file_path).parent
                try:
                    if sys.platform == 'win32':
                        # Open folder and select the file
                        subprocess.run(['explorer', '/select,', str(file_path)])
                    elif sys.platform == 'darwin':
                        subprocess.run(['open', '-R', str(file_path)])
                    else:
                        subprocess.run(['xdg-open', str(folder_path)])
                    self._log_status_message(f"Opened folder: {folder_path.name}", 1500)
                except Exception as e:
                    self._log_status_message(f"Error opening folder: {e}", 3000)

    def _on_notes_action_open(self, row: int):
        """Open notes file at given row (called from delegate button click)."""
        import os
        import subprocess
        import sys
        from PyQt6.QtCore import Qt

        if not hasattr(self, '_notes_model'):
            return

        path_item = self._notes_model.item(row, 1)  # Column 1 is File Name with path in UserRole
        if path_item:
            file_path = path_item.data(Qt.ItemDataRole.UserRole)
            if file_path and Path(file_path).exists():
                try:
                    if sys.platform == 'win32':
                        os.startfile(file_path)
                    elif sys.platform == 'darwin':
                        subprocess.run(['open', file_path])
                    else:
                        subprocess.run(['xdg-open', file_path])
                    self._log_status_message(f"Opened: {Path(file_path).name}", 1500)
                except Exception as e:
                    self._log_status_message(f"Error opening file: {e}", 3000)

    def _on_notes_action_preview(self, row: int):
        """Preview notes file at given row (called from delegate button click)."""
        from PyQt6.QtCore import Qt

        if not hasattr(self, '_notes_model'):
            return

        path_item = self._notes_model.item(row, 1)  # Column 1 = File Name with path in UserRole
        if not path_item:
            return

        file_path = Path(path_item.data(Qt.ItemDataRole.UserRole))
        if not file_path.exists():
            self._log_status_message("File not found", 1500)
            return

        # Use shared preview dialog
        self._show_notes_preview_dialog(
            files=[{'name': file_path.name, 'path': str(file_path)}],
            title=f"Preview: {file_path.name}"
        )

    def on_notes_link(self):
        """Link selected notes file to a data file in the project."""
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QLabel, QListWidget, QDialogButtonBox
        from PyQt6.QtCore import Qt

        if not hasattr(self, 'notesFilesTable'):
            return

        indexes = self.notesFilesTable.selectionModel().selectedRows()
        if not indexes:
            self._log_status_message("Please select a notes file to link", 1500)
            return

        row = indexes[0].row()
        path_item = self._notes_model.item(row, 1)  # Column 1 = File Name with path in UserRole
        if not path_item:
            return

        notes_path = Path(path_item.data(Qt.ItemDataRole.UserRole))

        # Get list of data files
        if not self._master_file_list:
            self._log_status_message("No data files in project. Scan for files first.", 2000)
            return

        # Create link dialog
        dialog = QDialog(self)
        dialog.setWindowTitle(f"Link: {notes_path.name}")
        dialog.resize(500, 400)
        layout = QVBoxLayout(dialog)

        layout.addWidget(QLabel(f"Link '{notes_path.name}' to which data file?"))

        file_list = QListWidget()
        for task in self._master_file_list:
            if not task.get('is_sub_row'):
                file_list.addItem(task.get('file_name', str(task.get('file_path', 'Unknown'))))
        layout.addWidget(file_list)

        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        buttons.accepted.connect(dialog.accept)
        buttons.rejected.connect(dialog.reject)
        layout.addWidget(buttons)

        if dialog.exec() == QDialog.DialogCode.Accepted:
            selected = file_list.currentItem()
            if selected:
                # TODO: Store the link in metadata
                self._log_status_message(f"Linked '{notes_path.name}' to '{selected.text()}'", 2000)
            else:
                self._log_status_message("No file selected", 1500)

    def _setup_project_name_combo(self):
        """Set up the unified project name combo (shows name, loads recent, editable for rename)."""
        if not hasattr(self, 'projectNameCombo'):
            print("[project-builder] WARNING: projectNameCombo not found in UI")
            return

        combo = self.projectNameCombo
        self._project_combo_updating = False  # Prevent recursive updates

        # Add a small dropdown indicator button after the combo
        # Find the combo's parent layout and add a subtle arrow button
        from PyQt6.QtWidgets import QPushButton
        parent_widget = combo.parentWidget()
        if parent_widget and parent_widget.layout():
            layout = parent_widget.layout()
            # Find combo's index in layout
            combo_idx = -1
            for i in range(layout.count()):
                item = layout.itemAt(i)
                if item and item.widget() == combo:
                    combo_idx = i
                    break

            if combo_idx >= 0:
                # Add dropdown arrow button right after combo
                dropdown_btn = QPushButton("▼")
                dropdown_btn.setFixedSize(24, 28)
                dropdown_btn.setToolTip("Select project")
                dropdown_btn.setStyleSheet("""
                    QPushButton {
                        background-color: transparent;
                        color: #cccccc;
                        border: none;
                        font-size: 10px;
                        padding: 0;
                        margin-left: -4px;
                    }
                    QPushButton:hover {
                        color: #ffffff;
                    }
                """)
                dropdown_btn.setFlat(True)
                dropdown_btn.clicked.connect(lambda: combo.showPopup())
                layout.insertWidget(combo_idx + 1, dropdown_btn)

        # Populate with "None" + recent projects
        self._populate_project_name_combo()

        # Connect selection change (for loading projects)
        combo.currentIndexChanged.connect(self._on_project_combo_selected)

        # Connect text editing (for renaming)
        combo.lineEdit().editingFinished.connect(self._on_project_name_edited)

        print("[project-builder] Set up projectNameCombo with None + recent projects")

    def _populate_project_name_combo(self, current_name: str = None):
        """Populate the project name combo with None + recent projects."""
        if not hasattr(self, 'projectNameCombo'):
            return

        combo = self.projectNameCombo
        self._project_combo_updating = True
        combo.blockSignals(True)
        combo.clear()

        # Add "None" option (index 0)
        combo.addItem("No Project", None)

        # Add recent projects
        recent_projects = self.project_manager.get_recent_projects()
        for project_info in recent_projects:
            combo.addItem(project_info['name'], project_info['path'])

        # Set current selection
        if current_name:
            # Find and select the current project
            idx = combo.findText(current_name)
            if idx >= 0:
                combo.setCurrentIndex(idx)
            else:
                # Project name not in list - set as editable text
                combo.setCurrentText(current_name)
        else:
            combo.setCurrentIndex(0)  # "No Project"

        combo.blockSignals(False)
        self._project_combo_updating = False

    def _on_project_combo_selected(self, index: int):
        """Handle selection from project name combo dropdown."""
        if self._project_combo_updating:
            return

        combo = self.projectNameCombo

        if index == 0:
            # "None" selected - clear current project
            self._clear_current_project()
            return

        # Get project path from item data
        project_path = combo.itemData(index)
        if not project_path:
            return

        project_path = Path(project_path)
        if not project_path.exists():
            self._log_status_message(f"Project file not found: {project_path.name}", 3000)
            self._populate_project_name_combo()  # Refresh list
            return

        # IMPORTANT: Cancel any pending autosave before loading new project
        self._cancel_pending_autosave()

        try:
            project_data = self.project_manager.load_project(project_path)
            self._populate_ui_from_project(project_data)
            self._log_status_message(f"Loaded project: {project_data['project_name']}", 2000)
        except Exception as e:
            self._show_error("Load Failed", f"Failed to load project:\n{e}")
            self._populate_project_name_combo()  # Reset

    def _on_project_name_edited(self):
        """Handle project name editing (rename)."""
        if self._project_combo_updating:
            return

        combo = self.projectNameCombo
        new_name = combo.currentText().strip()

        if not new_name or new_name == "No Project":
            # Empty or invalid name - revert
            self._populate_project_name_combo(getattr(self, '_current_project_name', None))
            return

        # Get current project name
        old_name = getattr(self, '_current_project_name', None)

        if new_name == old_name:
            return  # No change

        if old_name:
            # Rename existing project
            self._current_project_name = new_name
            self._log_status_message(f"Project renamed to: {new_name}", 2000)
            # Schedule autosave to persist the rename
            self._schedule_project_autosave()
        else:
            # Setting name for new unsaved project
            self._current_project_name = new_name
            self._log_status_message(f"Project name set: {new_name}", 2000)

    def _clear_current_project(self):
        """Clear the current project (when None is selected)."""
        self._current_project_name = None
        self._project_directory = None
        self._master_file_list = []
        self._discovered_files_data = []
        self._notes_files_data = []

        # Clear UI
        if hasattr(self, 'directoryPathEdit'):
            self.directoryPathEdit.clear()
        if hasattr(self, '_file_table_model'):
            self._file_table_model.clear_all()
        if hasattr(self, '_notes_model'):
            self._notes_model.removeRows(0, self._notes_model.rowCount())
        if hasattr(self, 'summaryLabel'):
            self.summaryLabel.setText("Summary: No project loaded")

        self._log_status_message("Project closed", 2000)

    def _connect_project_builder_buttons(self):
        """Connect signals for Project Builder buttons defined in .ui file."""
        # Connect Scan Saved Data button
        if hasattr(self, 'scanSavedDataButton'):
            self.scanSavedDataButton.clicked.connect(self.on_project_scan_saved_data)

        # Connect Resolve Conflicts button
        if hasattr(self, 'resolveConflictsButton'):
            self.resolveConflictsButton.clicked.connect(self.on_resolve_all_conflicts)

        # Connect Show Full Content checkbox
        if hasattr(self, 'tableFullContentCheckbox'):
            self.tableFullContentCheckbox.stateChanged.connect(self._on_table_column_mode_changed)

        # Connect chatbot panel widgets
        if hasattr(self, 'chatSendButton'):
            self.chatSendButton.clicked.connect(self._on_chat_send)
        if hasattr(self, 'chatInputEdit'):
            self.chatInputEdit.returnPressed.connect(self._on_chat_send)
        if hasattr(self, 'chatSettingsButton'):
            self.chatSettingsButton.clicked.connect(self._open_ai_settings)
        if hasattr(self, 'chatClearButton'):
            self.chatClearButton.clicked.connect(self._clear_chat_history)
        if hasattr(self, 'chatStopButton'):
            self.chatStopButton.clicked.connect(self._on_chat_stop)

        # Initialize model selection combo
        if hasattr(self, 'modelSelectCombo'):
            self._init_model_selector()

        # Connect notebook widgets
        if hasattr(self, 'runCodeButton'):
            self.runCodeButton.clicked.connect(self._on_run_code)
        if hasattr(self, 'clearOutputButton'):
            self.clearOutputButton.clicked.connect(self._on_clear_code_output)

        # Add Pop Out and Save Figure buttons to notebook header
        self._add_notebook_extra_buttons()

        # Connect table filter widgets
        if hasattr(self, 'tableFilterEdit'):
            self.tableFilterEdit.textChanged.connect(self._on_table_filter_changed)
        if hasattr(self, 'filterColumnCombo'):
            self.filterColumnCombo.currentIndexChanged.connect(self._on_table_filter_changed)

        # Set up main horizontal splitter (left column vs AI assistant, 75%/25%)
        if hasattr(self, 'mainContentSplitter'):
            self.mainContentSplitter.setSizes([750, 250])

    def _on_table_column_mode_changed(self, state):
        """Handle toggle of the 'Show Full Content' checkbox."""
        self._auto_fit_table_columns()

    def _init_model_selector(self):
        """Initialize the AI model selection dropdown with only AVAILABLE models."""
        from PyQt6.QtCore import QSettings
        settings = QSettings("PhysioMetrics", "BreathAnalysis")

        # Only show models that are actually configured/available
        # Format: (display_name, provider, model_id)
        self._ai_models = []

        # 1. Add Ollama models if running and installed
        ollama_available = False
        try:
            from core.ai_client import get_ollama_models, check_ollama_running
            if check_ollama_running():
                installed_models = get_ollama_models()
                if installed_models:
                    ollama_available = True
                    for model in installed_models[:8]:  # Limit to 8 models
                        self._ai_models.append((f"🆓 Ollama: {model}", "ollama", model))
        except Exception as e:
            print(f"[ai] Could not fetch Ollama models: {e}")

        # 2. Add Claude models only if API key is configured
        claude_key = settings.value("ai/claude_api_key", "")
        if claude_key:
            self._ai_models.append(("Claude: Haiku (fast)", "claude", "claude-3-5-haiku-latest"))
            self._ai_models.append(("Claude: Sonnet (smart)", "claude", "claude-sonnet-4-20250514"))

        # 3. Add OpenAI models only if API key is configured
        openai_key = settings.value("ai/openai_api_key", "")
        if openai_key:
            self._ai_models.append(("GPT-4o", "openai", "gpt-4o"))
            self._ai_models.append(("GPT-4o mini (fast)", "openai", "gpt-4o-mini"))

        # 4. If nothing available, show setup hint
        if not self._ai_models:
            self._ai_models.append(("⚙️ Click gear to set up AI", "setup_hint", ""))

        # Populate combo
        self.modelSelectCombo.clear()
        for display_name, _, _ in self._ai_models:
            self.modelSelectCombo.addItem(display_name)

        # Load saved selection
        settings = QSettings("PhysioMetrics", "BreathAnalysis")
        saved_provider = settings.value("ai/provider", "ollama")
        saved_model = settings.value("ai/selected_model", "llama3.2")

        # Find and select the saved model
        for i, (_, provider, model_id) in enumerate(self._ai_models):
            if provider == saved_provider and model_id == saved_model:
                self.modelSelectCombo.setCurrentIndex(i)
                break

        # Connect change signal to save selection
        self.modelSelectCombo.currentIndexChanged.connect(self._on_model_changed)

    def _on_model_changed(self, index):
        """Save selected model to settings."""
        if hasattr(self, '_ai_models') and 0 <= index < len(self._ai_models):
            _, provider, model_id = self._ai_models[index]
            settings = QSettings("PhysioMetrics", "BreathAnalysis")
            settings.setValue("ai/provider", provider)
            settings.setValue("ai/selected_model", model_id)

    def _get_selected_model(self) -> tuple:
        """Get the currently selected (provider, model_id) tuple."""
        if hasattr(self, 'modelSelectCombo') and hasattr(self, '_ai_models'):
            index = self.modelSelectCombo.currentIndex()
            if 0 <= index < len(self._ai_models):
                _, provider, model_id = self._ai_models[index]
                return (provider, model_id)
        # Default to Ollama
        return ("ollama", "llama3.2")

    def _get_model_display_name(self) -> str:
        """Get a short display name for the current model."""
        if hasattr(self, 'modelSelectCombo') and hasattr(self, '_ai_models'):
            index = self.modelSelectCombo.currentIndex()
            if 0 <= index < len(self._ai_models):
                display_name, provider, model_id = self._ai_models[index]
                # Extract short name
                if provider == "ollama":
                    return model_id.split(":")[0].capitalize()
                elif 'haiku' in model_id.lower():
                    return "Haiku"
                elif 'sonnet' in model_id.lower():
                    return "Sonnet"
                elif 'gpt-4o-mini' in model_id.lower():
                    return "GPT-4o mini"
                elif 'gpt-4o' in model_id.lower():
                    return "GPT-4o"
                else:
                    return model_id.split("-")[0].capitalize()
        return "AI"

    def _on_table_filter_changed(self, *args):
        """Filter table rows based on search text and selected column."""
        if not hasattr(self, 'tableFilterEdit') or not hasattr(self, 'filterColumnCombo'):
            return

        filter_text = self.tableFilterEdit.text().strip().lower()
        column_filter = self.filterColumnCombo.currentText()

        # Map column filter to data keys
        column_map = {
            'All Columns': None,  # Search all
            'File Name': 'file_name',
            'Protocol': 'protocol',
            'Animal ID': 'animal_id',
            'Strain': 'strain',
            'Keywords': 'keywords_display',
        }
        target_column = column_map.get(column_filter)

        # Get table view (it's called discoveredFilesTable in the UI)
        if not hasattr(self, 'discoveredFilesTable'):
            return

        table = self.discoveredFilesTable
        visible_count = 0
        total_count = len(self._master_file_list) if self._master_file_list else 0

        # If no filter, show all rows
        if not filter_text:
            for row in range(self._file_table_model.rowCount()):
                table.setRowHidden(row, False)
            if hasattr(self, 'filterCountLabel'):
                self.filterCountLabel.setText("")
            return

        # Filter rows
        for row in range(self._file_table_model.rowCount()):
            if row >= len(self._master_file_list):
                continue

            task = self._master_file_list[row]
            match = False

            if target_column:
                # Search specific column
                value = str(task.get(target_column, '')).lower()
                match = filter_text in value
            else:
                # Search all relevant columns - include file path for better searching
                searchable_columns = ['file_name', 'protocol', 'animal_id', 'strain',
                                     'keywords_display', 'power', 'sex', 'stim_type',
                                     'experiment', 'file_path']
                for col in searchable_columns:
                    value = str(task.get(col, '')).lower()
                    if filter_text in value:
                        match = True
                        break

            table.setRowHidden(row, not match)
            if match:
                visible_count += 1

        # Update count label
        if hasattr(self, 'filterCountLabel'):
            if visible_count == total_count:
                self.filterCountLabel.setText("")
            else:
                self.filterCountLabel.setText(f"({visible_count} of {total_count})")

    def _on_chat_send(self):
        """Handle sending a message in the chatbot panel."""
        if not hasattr(self, 'chatInputEdit') or not hasattr(self, 'chatHistoryText'):
            return

        message = self.chatInputEdit.text().strip()
        if not message:
            return

        # Add user message to chat history
        self.chatHistoryText.append(f"<b style='color: #569cd6;'>You:</b> {message}")
        self.chatInputEdit.clear()

        # Get selected provider and model from dropdown
        provider, model = self._get_selected_model()

        # Check if provider is available
        from PyQt6.QtCore import QSettings
        settings = QSettings("PhysioMetrics", "BreathAnalysis")

        if provider == "setup_hint":
            # User selected the "Click gear to set up AI" hint
            self.chatHistoryText.append(
                f"<b style='color: #f48771;'>AI Setup Required:</b><br><br>"
                f"<b>Option 1 - Ollama (FREE, runs locally):</b><br>"
                f"• Download from: <a href='https://ollama.com/download'>https://ollama.com/download</a><br>"
                f"• Run: <code>ollama pull llama3.2</code><br><br>"
                f"<b>Option 2 - Cloud API (requires account):</b><br>"
                f"• Click ⚙️ gear icon to add API key<br>"
                f"• Supports Claude and GPT models"
            )
            self._scroll_chat_to_bottom()
            return

        if provider == "ollama":
            # Ollama doesn't need API key - check if server is running
            try:
                from core.ai_client import check_ollama_running
                if not check_ollama_running():
                    self.chatHistoryText.append(
                        f"<b style='color: #f44747;'>Error:</b> Ollama is not running.<br>"
                        f"<i>Download FREE from: <a href='https://ollama.com/download'>https://ollama.com/download</a><br>"
                        f"Then run: ollama pull {model}</i>"
                    )
                    self._scroll_chat_to_bottom()
                    return
                # Use Ollama (no API key needed)
                self._send_to_ai_api(message, provider, api_key=None, model=model)
            except ImportError as e:
                self.chatHistoryText.append(
                    f"<b style='color: #f44747;'>Error:</b> {e}"
                )
                self._scroll_chat_to_bottom()
        else:
            # Cloud provider - needs API key
            api_key = settings.value(f"ai/{provider}_api_key", "")

            if api_key:
                # Use real AI API
                self._send_to_ai_api(message, provider, api_key, model=model)
            else:
                # No API key configured
                self.chatHistoryText.append(
                    f"<b style='color: #f44747;'>Error:</b> No API key configured for {provider}.<br>"
                    f"<i>Click the ⚙️ gear icon to configure, or select an Ollama model (FREE!).</i>"
                )
                self._scroll_chat_to_bottom()

    def _send_to_ai_api(self, message: str, provider: str, api_key: str, model: str = None):
        """Send message to actual AI API in background thread."""
        from PyQt6.QtCore import QThread, pyqtSignal

        # Soft token warning threshold (warn but allow continuing)
        TOKEN_WARNING_THRESHOLD = 500000  # 500k tokens - just a warning

        # Initialize conversation history if not exists
        if not hasattr(self, '_chat_conversation_history'):
            self._chat_conversation_history = []

        # Check token usage and warn (but allow continuing)
        current_tokens = getattr(self, '_total_tokens_used', 0)
        if current_tokens >= TOKEN_WARNING_THRESHOLD and not getattr(self, '_token_warning_shown', False):
            from PyQt6.QtWidgets import QMessageBox
            msg = QMessageBox(self)
            msg.setIcon(QMessageBox.Icon.Information)
            msg.setWindowTitle("High Token Usage")
            msg.setText(f"You have used {current_tokens:,} tokens in this session.")
            msg.setInformativeText(
                "This is just a heads up about your usage.\n\n"
                "Click 'Continue' to keep chatting, or 'Clear Chat' to start fresh."
            )
            msg.setStandardButtons(QMessageBox.StandardButton.Ok | QMessageBox.StandardButton.Cancel)
            msg.button(QMessageBox.StandardButton.Ok).setText("Continue")
            msg.button(QMessageBox.StandardButton.Cancel).setText("Clear Chat")

            if msg.exec() == QMessageBox.StandardButton.Cancel:
                self._clear_chat_history()
                return

            self._token_warning_shown = True  # Don't show again until cleared

        # Show thinking indicator with model name
        model_name = self._get_model_display_name()
        self.chatHistoryText.append(f"<b style='color: #4ec9b0;'>{model_name}:</b> <i>Thinking...</i>")
        self._scroll_chat_to_bottom()

        # Disable send button, enable stop button while processing
        if hasattr(self, 'chatSendButton'):
            self.chatSendButton.setEnabled(False)
        if hasattr(self, 'chatStopButton'):
            self.chatStopButton.setEnabled(True)

        # Build context about current files
        context = self._build_ai_context()

        # Add user message to conversation history
        self._chat_conversation_history.append({
            "role": "user",
            "content": message
        })

        # Prepare messages for API (copy to avoid mutation)
        messages_for_api = list(self._chat_conversation_history)

        # Use the model passed as parameter (already extracted from dropdown)
        selected_model = model

        # Define native tools for Claude (Ollama may not support tools)
        tools = self._get_ai_tool_definitions() if provider != "ollama" else []

        # Create worker thread with native tool support
        class AIWorker(QThread):
            # response, is_error, input_tokens, output_tokens, tool_calls (JSON string)
            finished = pyqtSignal(str, bool, int, int, str)

            def __init__(self, provider, api_key, messages, context, model, tools):
                super().__init__()
                self.provider = provider
                self.api_key = api_key
                self.messages = messages
                self.context = context
                self.model = model
                self.tools = tools

            def run(self):
                import time
                import json
                max_retries = 3
                base_delay = 2  # seconds

                for attempt in range(max_retries):
                    try:
                        from core.ai_client import AIClient

                        client = AIClient(provider=self.provider, api_key=self.api_key, model=self.model)

                        # Comprehensive system prompt with data documentation
                        system_prompt = f"""You are an AI assistant for PhysioMetrics, a respiratory signal analysis application.
You help users analyze plethysmography data and generate Python code for plotting and analysis.

=== DEMO REQUEST ===

If the user says "demo", "example", "show me", "sample plot", or similar:

First, use list_available_files() to find a file with exports, 10 sweeps, and "30hz" in name/protocol.

Then: "Create a demo plot showing breathing frequency response to optogenetic stimulation. Use a single panel with dark theme (plt.style.use('dark_background')), plot individual sweeps in light gray, mean frequency in cyan with SEM shading, and highlight stimulation periods with steelblue rectangles. Make it publication-ready with white text and clean styling."

Helper functions to use:
- load_means_csv(filename) - DataFrame with frequency_sweep0, frequency_sweep1, ..., frequency_mean, frequency_sem
- load_bundle_npz(filename) - bundle with stim timing
- get_stim_spans(bundle) - returns [(start_time, end_time), ...]
- add_stim_shading(ax, stim_spans, color='steelblue', alpha=0.4) - adds stim rectangles

=== SEARCH STRATEGY ===
ALWAYS call get_searchable_values() FIRST before searching! This shows you:
- Exact protocol names, status values, animal IDs, etc. that exist in the data
- Which files have exports

IMPORTANT: If the user asks for something that doesn't exist in the searchable values:
- DON'T search for it (you'll get 0 results)
- TELL the user what's actually available and suggest alternatives

=== QUICK CONTEXT ===
{self.context}

=== EXPORTED DATA FILE STRUCTURE ===

PhysioMetrics exports several file types. Use the export_path from file context above.

**1. _means_by_time.csv** - Time-series metrics (best for plotting over time)
Columns:
- `time`: Time in seconds (relative to stim start if stim present)
- `sweep`: Sweep index (0, 1, 2, ...)
- For each metric: `<metric>_sweep0`, `<metric>_sweep1`, etc., `<metric>_mean`, `<metric>_sem`
- Normalized versions: `<metric>_norm_mean`, `<metric>_norm_sem` (baseline-normalized)

Key metrics available:
- `frequency` or `if` - Instantaneous frequency (Hz, breaths/min)
- `amp_insp` - Inspiratory amplitude
- `amp_exp` - Expiratory amplitude
- `ti` - Inspiratory time (seconds)
- `te` - Expiratory time (seconds)
- `ttot` - Total breath cycle time (seconds)
- `area_insp` - Inspiratory area under curve
- `area_exp` - Expiratory area under curve
- `duty_cycle` - Ti/Ttot ratio (0-1)
- `ve` - Minute ventilation (frequency × amplitude)

**2. _breaths.csv** - Per-breath data (one row per breath)
Columns:
- `sweep_idx`, `breath_idx` - Identifiers
- `time_onset`, `time_peak` - Timing (seconds)
- All metrics above, plus:
- `is_in_stim` - Boolean, True if breath occurred during stimulation
- `is_eupnea` - Boolean, True if classified as eupnea
- `is_sniffing` - Boolean, True if classified as sniffing

**3. _bundle.npz** - Complete data bundle (numpy format)
Load with: `data = np.load('file_bundle.npz', allow_pickle=True)`
Contains:
- `t_downsampled` - Time array (downsampled)
- `trace_downsampled` - Signal trace per sweep
- `stim_spans_json` - JSON string with stim timing
- Metric arrays for each sweep

**4. Stimulation Timing**
From _bundle.npz, stim_spans are stored as JSON:
```python
import json
stim_spans = json.loads(str(data['stim_spans_json']))
# stim_spans is dict: {{sweep_idx: [(start_time, end_time), ...]}}
```
From _breaths.csv, use `is_in_stim` column.

=== HELPER FUNCTIONS (available in Code Notebook) ===

The Code Notebook has these helper functions pre-loaded:

```python
# List all files with exports
print(list_available_files())

# Load data by filename (partial match)
df = load_means_csv('xxx.abf')      # Returns DataFrame from _means_by_time.csv
df = load_breaths_csv('xxx.abf')    # Returns DataFrame from _breaths.csv
data = load_bundle_npz('xxx.abf')   # Returns numpy NpzFile

# Get stim timing from bundle
stim_spans = get_stim_spans(data)   # Returns {{sweep_idx: [(start, end), ...]}}

# Add stim shading to plot
add_stim_shading(ax, stim_spans, sweep_idx=0, color='blue', alpha=0.2)

# Get export paths
paths = get_export_paths('xxx')     # Returns {{filename: export_path}}
```

=== CODE TEMPLATES ===

**Plot frequency vs time with mean±SEM and stim shading (RECOMMENDED):**
```python
# Load data using helper function
df = load_means_csv('xxx.abf')  # Replace xxx.abf with actual filename
bundle = load_bundle_npz('xxx.abf')
stim_spans = get_stim_spans(bundle)

fig, ax = plt.subplots(figsize=(12, 6))

# Overlay all sweeps in gray
sweep_cols = [c for c in df.columns if c.startswith('frequency_sweep')]
for col in sweep_cols:
    ax.plot(df['time'], df[col], alpha=0.3, color='gray', linewidth=0.5)

# Plot mean±SEM
ax.fill_between(df['time'],
                df['frequency_mean'] - df['frequency_sem'],
                df['frequency_mean'] + df['frequency_sem'],
                alpha=0.3, color='blue', label='Mean ± SEM')
ax.plot(df['time'], df['frequency_mean'], 'b-', linewidth=2, label='Mean')

# Add blue stim shading
add_stim_shading(ax, stim_spans, color='blue', alpha=0.2, label='Laser ON')

ax.set_xlabel('Time (s)')
ax.set_ylabel('Frequency (Hz)')
ax.legend()
ax.set_title('Breathing Frequency Over Time')
```

**Compare eupnea vs sniffing:**
```python
df = load_breaths_csv('xxx.abf')
eupnea = df[df['is_eupnea'] == True]
sniffing = df[df['is_sniffing'] == True]

fig, ax = plt.subplots(1, 2, figsize=(10, 4))
ax[0].hist(eupnea['frequency'], bins=30, alpha=0.7, label='Eupnea')
ax[0].hist(sniffing['frequency'], bins=30, alpha=0.7, label='Sniffing')
ax[0].legend()
ax[0].set_xlabel('Frequency (Hz)')
ax[1].boxplot([eupnea['ti'].dropna(), sniffing['ti'].dropna()], labels=['Eupnea', 'Sniffing'])
ax[1].set_ylabel('Ti (s)')
```

**Multi-metric comparison:**
```python
df = load_means_csv('xxx.abf')
metrics = ['frequency', 'amp_insp', 'ti', 'duty_cycle']

fig, axes = plt.subplots(2, 2, figsize=(12, 8))
for ax, metric in zip(axes.flat, metrics):
    mean_col = f'{{metric}}_mean'
    sem_col = f'{{metric}}_sem'
    if mean_col in df.columns:
        ax.fill_between(df['time'], df[mean_col]-df[sem_col], df[mean_col]+df[sem_col], alpha=0.3)
        ax.plot(df['time'], df[mean_col])
        ax.set_ylabel(metric)
        ax.set_xlabel('Time (s)')
plt.tight_layout()
```

=== OTHER COMMANDS (include in your response text) ===
- [LOAD_PROJECT: project_name] - Load a saved project
- [UPDATE_META: filename | field=value] - Update file metadata

WHEN GENERATING CODE:
- ALWAYS use actual filenames from the context or search results (NOT placeholder 'xxx.abf')
- Helper functions accept: full paths, filenames, or partial name matches
- For raw Windows paths, use raw strings: r'C:\\path\\to\\file.csv'
- Import: pandas as pd, matplotlib.pyplot as plt, numpy as np
- Code runs in a Code Notebook - user clicks "Run" to execute
- Plots will render inline in the output area

Be concise. Use markdown. Put code in ```python ``` blocks."""

                        # Use native tool use
                        response = client.chat_with_tools(
                            messages=self.messages,
                            tools=self.tools,
                            system_prompt=system_prompt,
                            max_tokens=4096,  # Increased for longer code responses
                            temperature=0.7
                        )

                        # Extract token usage
                        input_tokens = response.usage.get('input_tokens', 0) if response.usage else 0
                        output_tokens = response.usage.get('output_tokens', 0) if response.usage else 0

                        # Convert tool_calls to JSON string for signal
                        tool_calls_json = json.dumps(response.tool_calls) if response.tool_calls else ""

                        self.finished.emit(response.content, False, input_tokens, output_tokens, tool_calls_json)
                        return  # Success, exit retry loop

                    except Exception as e:
                        error_str = str(e).lower()
                        is_rate_limit = '429' in str(e) or 'rate_limit' in error_str or 'rate limit' in error_str

                        if is_rate_limit and attempt < max_retries - 1:
                            # Rate limit - wait and retry with exponential backoff
                            delay = base_delay * (2 ** attempt)  # 2, 4, 8 seconds
                            time.sleep(delay)
                            continue  # Retry
                        else:
                            # Final attempt failed or non-rate-limit error
                            self.finished.emit(str(e), True, 0, 0, "")
                            return

        self._chat_worker = AIWorker(provider, api_key, messages_for_api, context, selected_model, tools)
        self._chat_worker.finished.connect(self._on_ai_response)
        self._chat_worker.start()

    def _on_ai_response(self, response: str, is_error: bool, input_tokens: int = 0, output_tokens: int = 0, tool_calls_json: str = ""):
        """Handle response from AI API with native tool support."""
        import json

        # Remove "Thinking..." message
        cursor = self.chatHistoryText.textCursor()
        cursor.movePosition(cursor.MoveOperation.End)
        cursor.movePosition(cursor.MoveOperation.StartOfBlock, cursor.MoveMode.KeepAnchor)
        cursor.removeSelectedText()
        cursor.deletePreviousChar()  # Remove newline

        if is_error:
            # Detect specific error types for better messages
            error_lower = response.lower()
            if '429' in response or 'rate_limit' in error_lower or 'rate limit' in error_lower:
                # Rate limit error - give helpful guidance
                self.chatHistoryText.append(
                    f"<b style='color: #f48771;'>Rate Limit:</b> Too many requests.<br>"
                    "<i>Wait 30-60 seconds before sending another message, or try using Haiku (faster/lower limits).</i>"
                )
            elif '401' in response or 'unauthorized' in error_lower or 'invalid' in error_lower and 'key' in error_lower:
                # Authentication error
                self.chatHistoryText.append(
                    f"<b style='color: #f48771;'>Auth Error:</b> {response}<br>"
                    "<i>Check your API key in settings (⚙)</i>"
                )
            elif 'timeout' in error_lower or 'timed out' in error_lower:
                # Timeout error
                self.chatHistoryText.append(
                    f"<b style='color: #f48771;'>Timeout:</b> Request took too long.<br>"
                    "<i>Try again with a shorter question, or check your network connection.</i>"
                )
            else:
                # Generic error
                self.chatHistoryText.append(
                    f"<b style='color: #f48771;'>Error:</b> {response}<br>"
                    "<i>Check your API key in settings (⚙) or try again.</i>"
                )
            # Remove the failed user message from history
            if hasattr(self, '_chat_conversation_history') and self._chat_conversation_history:
                self._chat_conversation_history.pop()
        else:
            # Update token count
            if not hasattr(self, '_total_tokens_used'):
                self._total_tokens_used = 0
            self._total_tokens_used += input_tokens + output_tokens
            self._update_token_display()

            # Check for NATIVE tool calls
            tool_calls = json.loads(tool_calls_json) if tool_calls_json else []

            if tool_calls:
                # Show that tools are being executed
                self.chatHistoryText.append(
                    f"<i style='color: #dcdcaa;'>🔍 Searching locally...</i>"
                )

                # Execute each tool call and collect results
                tool_results = []
                for tc in tool_calls:
                    tool_name = tc.get('name', '')
                    tool_input = tc.get('input', {})
                    tool_id = tc.get('id', '')

                    result = self._execute_native_tool(tool_name, tool_input)
                    tool_results.append({
                        'id': tool_id,
                        'name': tool_name,
                        'input': tool_input,
                        'result': result
                    })

                    # Show summary in chat (include search params for debugging)
                    if 'error' in result:
                        self.chatHistoryText.append(
                            f"<i style='color: #f48771;'>✗ {tool_name}: {result['error']}</i>"
                        )
                    else:
                        if 'count' in result:
                            # Show what was searched for
                            search_params = ', '.join(f"{k}={v}" for k, v in tool_input.items() if v) if tool_input else "all"
                            self.chatHistoryText.append(
                                f"<i style='color: #89d185;'>✓ {tool_name}({search_params}): Found {result.get('count', 0)} results</i>"
                            )
                        elif 'protocols' in result:
                            self.chatHistoryText.append(
                                f"<i style='color: #89d185;'>✓ {tool_name}: Found {len(result['protocols'])} protocols</i>"
                            )
                        elif 'animals' in result:
                            self.chatHistoryText.append(
                                f"<i style='color: #89d185;'>✓ {tool_name}: Found {len(result['animals'])} animals</i>"
                            )
                        elif 'file_count' in result:
                            self.chatHistoryText.append(
                                f"<i style='color: #89d185;'>✓ {tool_name}: {result['file_count']} files in project</i>"
                            )
                        else:
                            self.chatHistoryText.append(
                                f"<i style='color: #89d185;'>✓ {tool_name}: Done</i>"
                            )

                # Make follow-up request with native tool results
                self._send_native_tool_followup(response, tool_calls, tool_results)
                return  # Don't process this response further - wait for follow-up

            # Add AI response to conversation history
            if hasattr(self, '_chat_conversation_history'):
                self._chat_conversation_history.append({
                    "role": "assistant",
                    "content": response
                })

            # Check if response contains code blocks and extract them
            code_extracted = self._extract_code_from_response(response)
            model_name = self._get_model_display_name()

            # Debug logging
            has_backticks = '```' in response
            print(f"[AI Response Debug] Response length: {len(response)}, has ```: {has_backticks}, code extracted: {len(code_extracted) if code_extracted else 0} chars")

            if code_extracted:
                # Put code in the notebook
                has_widget = hasattr(self, 'codeInputEdit')
                print(f"[AI Response Debug] codeInputEdit exists: {has_widget}")
                if has_widget:
                    self.codeInputEdit.setPlainText(code_extracted)
                    print(f"[AI Response Debug] Code placed in notebook ({len(code_extracted)} chars)")
                else:
                    print(f"[AI Response Debug] WARNING: codeInputEdit widget not found!")
                # Format response with code stripped (since it's in notebook)
                formatted = self._format_ai_response(response, strip_code=True)
                formatted += "<br><br><i style='color: #4ec9b0;'>📓 Code placed in notebook below. Click ▶ Run to execute.</i>"
                self.chatHistoryText.append(f"<b style='color: #4ec9b0;'>{model_name}:</b><br>{formatted}")
            else:
                # Format response with code blocks styled (if any)
                if has_backticks:
                    print(f"[AI Response Debug] WARNING: Response has ``` but no code extracted!")
                    print(f"[AI Response Debug] First 500 chars: {response[:500]}")
                formatted = self._format_ai_response(response, strip_code=False)
                self.chatHistoryText.append(f"<b style='color: #4ec9b0;'>{model_name}:</b><br>{formatted}")

            # Execute any AI commands in the response
            command_results = self._execute_ai_commands(response)
            if command_results:
                for result in command_results:
                    if result['success']:
                        self.chatHistoryText.append(
                            f"<i style='color: #89d185;'>✓ {result['command']}: {result['message']}</i>"
                        )
                    else:
                        self.chatHistoryText.append(
                            f"<i style='color: #f48771;'>✗ {result['command']}: {result['message']}</i>"
                        )

        self._scroll_chat_to_bottom()

        # Re-enable send button, disable stop button
        if hasattr(self, 'chatSendButton'):
            self.chatSendButton.setEnabled(True)
        if hasattr(self, 'chatStopButton'):
            self.chatStopButton.setEnabled(False)

    def _send_native_tool_followup(self, original_response: str, tool_calls: list, tool_results: list):
        """
        Send a follow-up request with native tool results so AI can generate final response.
        Uses proper Claude tool_result format.
        """
        import json
        from PyQt6.QtCore import QThread, pyqtSignal

        # Build the assistant message with tool_use blocks (reconstruct what Claude sent)
        assistant_content = []
        if original_response:
            assistant_content.append({"type": "text", "text": original_response})
        for tc in tool_calls:
            assistant_content.append({
                "type": "tool_use",
                "id": tc.get('id', ''),
                "name": tc.get('name', ''),
                "input": tc.get('input', {})
            })

        # Build the user message with tool_result blocks
        user_content = []
        for tr in tool_results:
            result_data = tr.get('result', {})
            is_error = 'error' in result_data
            user_content.append({
                "type": "tool_result",
                "tool_use_id": tr.get('id', ''),
                "content": json.dumps(result_data, indent=2, default=str),
                "is_error": is_error
            })

        # Add to conversation history with proper structure
        if hasattr(self, '_chat_conversation_history'):
            self._chat_conversation_history.append({
                "role": "assistant",
                "content": assistant_content
            })
            self._chat_conversation_history.append({
                "role": "user",
                "content": user_content
            })

        # Get API settings from the model selector
        provider, selected_model = self._get_selected_model()
        settings = QSettings("PhysioMetrics", "BreathAnalysis")
        api_key = settings.value(f"ai/{provider}_api_key", "") if provider != "ollama" else None
        tools = self._get_ai_tool_definitions() if provider != "ollama" else []

        if provider != "ollama" and not api_key:
            self.chatHistoryText.append("<b style='color: #f48771;'>Error:</b> No API key configured")
            return

        # Create follow-up worker using native tools
        class FollowupWorker(QThread):
            finished = pyqtSignal(str, bool, int, int, str)  # Match main worker signature

            def __init__(self, provider, api_key, messages, model, tools):
                super().__init__()
                self.provider = provider
                self.api_key = api_key
                self.messages = messages
                self.model = model
                self.tools = tools

            def run(self):
                try:
                    from core.ai_client import AIClient
                    client = AIClient(provider=self.provider, api_key=self.api_key, model=self.model)

                    system_prompt = """You received tool results. Now provide a helpful response to the user.
If they asked for code, generate Python code in ```python ``` blocks.
Use the file paths and data from the tool results.
Be concise and helpful."""

                    # Use chat_with_tools for proper format handling
                    response = client.chat_with_tools(
                        messages=self.messages,
                        tools=self.tools,
                        system_prompt=system_prompt,
                        max_tokens=4096,  # Increased for longer code responses
                        temperature=0.7
                    )
                    input_tokens = response.usage.get('input_tokens', 0) if response.usage else 0
                    output_tokens = response.usage.get('output_tokens', 0) if response.usage else 0
                    tool_calls_json = json.dumps(response.tool_calls) if response.tool_calls else ""
                    self.finished.emit(response.content, False, input_tokens, output_tokens, tool_calls_json)
                except Exception as e:
                    self.finished.emit(str(e), True, 0, 0, "")

        # Prepare messages (include conversation history)
        messages_for_api = list(self._chat_conversation_history) if hasattr(self, '_chat_conversation_history') else []

        self._chat_worker = FollowupWorker(provider, api_key, messages_for_api, selected_model, tools)
        self._chat_worker.finished.connect(self._on_ai_response)
        self._chat_worker.start()

    def _on_chat_stop(self):
        """Stop the current AI request."""
        if hasattr(self, '_chat_worker') and self._chat_worker and self._chat_worker.isRunning():
            # Terminate the worker thread
            self._chat_worker.terminate()
            self._chat_worker.wait(1000)

            # Remove "Thinking..." message
            cursor = self.chatHistoryText.textCursor()
            cursor.movePosition(cursor.MoveOperation.End)
            cursor.movePosition(cursor.MoveOperation.StartOfBlock, cursor.MoveMode.KeepAnchor)
            cursor.removeSelectedText()
            cursor.deletePreviousChar()

            self.chatHistoryText.append("<span style='color: #f48771;'>⏹ Request stopped by user.</span>")

            # Remove the pending user message from history
            if hasattr(self, '_chat_conversation_history') and self._chat_conversation_history:
                self._chat_conversation_history.pop()

            # Re-enable send button, disable stop button
            if hasattr(self, 'chatSendButton'):
                self.chatSendButton.setEnabled(True)
            if hasattr(self, 'chatStopButton'):
                self.chatStopButton.setEnabled(False)

            self._scroll_chat_to_bottom()

    def _clear_chat_history(self):
        """Clear the chat conversation history and display."""
        self._chat_conversation_history = []
        self._total_tokens_used = 0
        self._token_warning_shown = False  # Reset warning flag
        self._update_token_display()
        if hasattr(self, 'chatHistoryText'):
            self.chatHistoryText.clear()

    def _update_token_display(self):
        """Update the token usage display in the chat header."""
        if hasattr(self, 'tokenUsageLabel'):
            tokens = getattr(self, '_total_tokens_used', 0)
            if tokens >= 1000:
                display = f"{tokens / 1000:.1f}k tokens"
            else:
                display = f"{tokens} tokens"
            self.tokenUsageLabel.setText(display)

            # Estimate cost (approximate rates as of 2025)
            # Claude Sonnet: ~$3/1M input, ~$15/1M output (avg ~$9/1M combined)
            # GPT-4o: ~$2.50/1M input, ~$10/1M output (avg ~$6/1M combined)
            # Using conservative estimate of ~$10 per 1M tokens
            est_cost = tokens * 0.00001  # $10 per 1M tokens

            # Update tooltip with more detail
            self.tokenUsageLabel.setToolTip(
                f"Total tokens used: {tokens:,}\n"
                f"Estimated cost: ${est_cost:.4f}\n"
                f"(actual cost varies by model)\n\n"
                f"Note: API providers don't offer\n"
                f"balance check APIs. Check your\n"
                f"account dashboard for usage limits."
            )

    def _extract_code_from_response(self, response: str) -> str:
        """Extract Python code blocks from AI response."""
        import re

        all_code_blocks = []

        # Most reliable method: Split by ``` and take alternating sections
        # This handles all edge cases with the markdown format
        if '```' in response:
            parts = response.split('```')
            print(f"[AI Code Extraction] Split response into {len(parts)} parts")

            for i, part in enumerate(parts):
                if i % 2 == 1:  # Odd indices are inside code blocks
                    # The first line might be a language identifier
                    lines = part.split('\n', 1)  # Split only on first newline
                    first_line = lines[0].strip().lower() if lines else ''

                    # Check if first line is a language identifier
                    if first_line in ['python', 'py', 'python3', '']:
                        # Take everything after the first line
                        code = lines[1].strip() if len(lines) > 1 else ''
                    elif first_line.startswith('python'):
                        # Handle cases like "python\n" or "python "
                        code = lines[1].strip() if len(lines) > 1 else ''
                    else:
                        # First line is part of the code
                        code = part.strip()

                    # Only add if it looks like meaningful code
                    if code and len(code) > 20:
                        # Quick check it looks like Python
                        python_indicators = ['import ', 'def ', 'class ', 'for ', 'while ',
                                            'plt.', 'pd.', 'np.', 'print(', 'from ', '= ']
                        if any(kw in code for kw in python_indicators):
                            all_code_blocks.append(code)
                            print(f"[AI Code Extraction] Found code block {i}: {len(code)} chars")

        # Combine all code blocks with newlines
        if all_code_blocks:
            combined = '\n\n'.join(all_code_blocks)

            # Check for signs of truncation
            truncation_signs = []
            # Unbalanced quotes
            single_quotes = combined.count("'") - combined.count("\\'")
            double_quotes = combined.count('"') - combined.count('\\"')
            if single_quotes % 2 != 0:
                truncation_signs.append("odd number of single quotes")
            if double_quotes % 2 != 0:
                truncation_signs.append("odd number of double quotes")
            # Unbalanced brackets/parens
            if combined.count('(') != combined.count(')'):
                truncation_signs.append("unbalanced parentheses")
            if combined.count('[') != combined.count(']'):
                truncation_signs.append("unbalanced brackets")
            if combined.count('{') != combined.count('}'):
                truncation_signs.append("unbalanced braces")
            # Code ends mid-statement
            last_line = combined.strip().split('\n')[-1].strip()
            if last_line and not last_line.endswith((':',')',')',']','}',',','"',"'")):
                if '=' in last_line or last_line.startswith(('def ', 'class ', 'if ', 'for ', 'while ')):
                    truncation_signs.append(f"incomplete last line: '{last_line[:50]}...'")

            if truncation_signs:
                print(f"[AI Code Extraction] WARNING: Code may be truncated! Signs: {', '.join(truncation_signs)}")
                print(f"[AI Code Extraction] Last 100 chars: {repr(combined[-100:])}")

            print(f"[AI Code Extraction] SUCCESS: {len(all_code_blocks)} block(s), {len(combined)} total chars")
            return combined

        # Debug: log if no code found but response looks like it has code
        if '```' in response:
            print(f"[AI Code Extraction] Warning: Response contains ``` but no code extracted. Response preview: {response[:200]}")

        return ""

    def _format_ai_response(self, response: str, strip_code: bool = False) -> str:
        """
        Format AI response for display in chat.

        - Converts newlines to <br> for proper line breaks
        - Formats or strips code blocks
        - Handles basic markdown (bold, italic, lists)
        """
        import re

        formatted = response

        # If stripping code (because it went to notebook), remove code blocks entirely
        if strip_code:
            # Remove ```python...``` blocks
            formatted = re.sub(r'```python\s*.*?\s*```', '<i>[Code sent to notebook]</i>', formatted, flags=re.DOTALL | re.IGNORECASE)
            # Remove generic ```...``` blocks that look like code
            formatted = re.sub(r'```\s*.*?\s*```', '<i>[Code sent to notebook]</i>', formatted, flags=re.DOTALL)
        else:
            # Format code blocks with styled pre tags
            def format_code_block(match):
                code = match.group(1).strip()
                # Escape HTML in code
                code = code.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                return f'<pre style="background-color: #1e1e1e; padding: 8px; border-radius: 4px; overflow-x: auto; font-size: 9pt;">{code}</pre>'

            formatted = re.sub(r'```python\s*(.*?)\s*```', format_code_block, formatted, flags=re.DOTALL | re.IGNORECASE)
            formatted = re.sub(r'```\s*(.*?)\s*```', format_code_block, formatted, flags=re.DOTALL)

        # Convert markdown bold **text** to <b>text</b>
        formatted = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', formatted)

        # Convert markdown italic *text* to <i>text</i> (but not inside code)
        formatted = re.sub(r'(?<!\*)\*([^*]+?)\*(?!\*)', r'<i>\1</i>', formatted)

        # Convert markdown bullet lists
        formatted = re.sub(r'^[\-\*]\s+(.+)$', r'• \1', formatted, flags=re.MULTILINE)

        # Convert numbered lists (1. 2. 3.)
        formatted = re.sub(r'^(\d+)\.\s+(.+)$', r'\1. \2', formatted, flags=re.MULTILINE)

        # Convert newlines to <br> for display (but preserve pre blocks)
        # Split by pre tags, only convert newlines outside pre
        parts = re.split(r'(<pre.*?</pre>)', formatted, flags=re.DOTALL)
        for i, part in enumerate(parts):
            if not part.startswith('<pre'):
                parts[i] = part.replace('\n', '<br>')
        formatted = ''.join(parts)

        return formatted

    def _build_ai_context(self) -> str:
        """Build comprehensive context string about current files for AI."""
        import json
        lines = []

        # Data files - detailed view
        num_files = len(self._master_file_list) if self._master_file_list else 0
        if num_files > 0:
            # Summary counts
            abf_count = sum(1 for f in self._master_file_list if str(f.get('file_path', '')).lower().endswith('.abf'))
            smrx_count = sum(1 for f in self._master_file_list if str(f.get('file_path', '')).lower().endswith('.smrx'))
            edf_count = sum(1 for f in self._master_file_list if str(f.get('file_path', '')).lower().endswith('.edf'))

            # Count by status
            completed = sum(1 for f in self._master_file_list if f.get('status') == 'completed')
            pending = sum(1 for f in self._master_file_list if f.get('status', 'pending') == 'pending')

            # Count files with exports
            with_exports = sum(1 for f in self._master_file_list if f.get('exports'))

            lines.append(f"=== PROJECT DATA SUMMARY ===")
            lines.append(f"Total rows: {num_files} ({abf_count} ABF, {smrx_count} SMRX, {edf_count} EDF)")
            lines.append(f"Status: {completed} completed, {pending} pending")
            lines.append(f"Files with exports: {with_exports}")

            # Protocols
            protocols = set(f.get('protocol', '') for f in self._master_file_list if f.get('protocol'))
            if protocols:
                lines.append(f"Protocols: {', '.join(sorted(protocols))}")

            # Detailed file list (limit to avoid token bloat)
            MAX_FILES_DETAILED = 30  # Show full details for first N files
            lines.append(f"\n=== DETAILED FILE LIST (showing {min(num_files, MAX_FILES_DETAILED)} of {num_files} rows) ===")

            # Group files by parent (for sub-row organization)
            current_parent = None
            for i, task in enumerate(self._master_file_list[:MAX_FILES_DETAILED]):
                file_name = task.get('file_name', 'Unknown')
                is_sub_row = task.get('is_sub_row', False)

                # Build file info line
                info_parts = []

                # Basic info
                status = task.get('status', 'pending')
                status_icon = "✓" if status == 'completed' else "○"

                if is_sub_row:
                    prefix = f"  └─ {status_icon} [Sub-row]"
                else:
                    prefix = f"{status_icon}"
                    current_parent = file_name

                info_parts.append(f"{prefix} {file_name}")

                # File path (source data file)
                file_path = task.get('file_path', '')
                if file_path:
                    info_parts.append(f"path={file_path}")

                # Channel info
                channel = task.get('channel', '')
                if channel:
                    info_parts.append(f"channel={channel}")

                # Protocol
                protocol = task.get('protocol', '')
                if protocol:
                    info_parts.append(f"protocol={protocol}")

                # Export info - this is what the user asked about!
                exports = task.get('exports', {})
                if exports:
                    export_str = ", ".join([f"{count} {fmt.upper()}" for fmt, count in exports.items() if count > 0])
                    if export_str:
                        info_parts.append(f"exports=[{export_str}]")

                # Export path (where CSV/NPZ files are saved)
                export_path = task.get('export_path', '')
                if export_path:
                    info_parts.append(f"export_dir={export_path}")

                # Metadata columns
                animal_id = task.get('animal_id', '')
                if animal_id:
                    info_parts.append(f"animal={animal_id}")

                strain = task.get('strain', '')
                if strain:
                    info_parts.append(f"strain={strain}")

                sex = task.get('sex', '')
                if sex:
                    info_parts.append(f"sex={sex}")

                power = task.get('power', '')
                if power:
                    info_parts.append(f"power={power}")

                stim_type = task.get('stim_type', '')
                if stim_type:
                    info_parts.append(f"stim={stim_type}")

                # Scan warnings
                warnings = task.get('scan_warnings', {})
                if warnings:
                    info_parts.append(f"warnings={len(warnings)}")

                lines.append(" | ".join(info_parts))

            if num_files > MAX_FILES_DETAILED:
                lines.append(f"... and {num_files - MAX_FILES_DETAILED} more files (ask about specific files by name or filter)")
            lines.append(f"=== END FILE LIST ===")
        else:
            lines.append("No data files currently loaded.")
            lines.append("User should scan a directory or load a project first.")

        # Scan for CSV files and read headers (limit to first few unique export dirs)
        if num_files > 0:
            export_dirs = set()
            for task in self._master_file_list:
                ep = task.get('export_path', '')
                if ep:
                    export_dirs.add(ep)

            if export_dirs:
                lines.append(f"\n=== EXPORTED CSV FILES & HEADERS ===")
                csv_files_found = []

                for export_dir in list(export_dirs)[:3]:  # Limit to 3 directories
                    try:
                        from pathlib import Path
                        export_path = Path(export_dir)
                        if export_path.exists():
                            # Find CSV files
                            csv_files = list(export_path.glob("*.csv"))[:5]  # Limit per dir
                            for csv_file in csv_files:
                                try:
                                    # Read just the header line
                                    with open(csv_file, 'r') as f:
                                        header_line = f.readline().strip()
                                    headers = header_line.split(',')
                                    csv_files_found.append({
                                        'path': str(csv_file),
                                        'name': csv_file.name,
                                        'headers': headers[:20]  # Limit columns shown
                                    })
                                except Exception:
                                    pass
                    except Exception:
                        pass

                if csv_files_found:
                    # Show unique header patterns
                    header_patterns = {}
                    for cf in csv_files_found:
                        pattern = tuple(cf['headers'])
                        if pattern not in header_patterns:
                            header_patterns[pattern] = []
                        header_patterns[pattern].append(cf['name'])

                    lines.append(f"Found {len(csv_files_found)} CSV files in export directories")
                    lines.append("")

                    for headers, files in list(header_patterns.items())[:3]:  # Show up to 3 patterns
                        lines.append(f"CSV Pattern (used by {len(files)} files, e.g., {files[0]}):")
                        lines.append(f"  Columns: {', '.join(headers)}")
                        lines.append("")

                    # List CSV files with paths (limited)
                    lines.append("CSV file paths:")
                    for cf in csv_files_found[:8]:  # Limit to 8 files
                        lines.append(f"  - {cf['path']}")
                    if len(csv_files_found) > 8:
                        lines.append(f"  ... and {len(csv_files_found) - 8} more")
                else:
                    lines.append("No CSV files found in export directories yet.")

        # Notes files
        notes_files = getattr(self, '_discovered_notes_files', [])
        if notes_files:
            lines.append(f"\n=== NOTES FILES ({len(notes_files)}) ===")
            for nf in notes_files:
                lines.append(f"  - {nf.name} ({nf})")
        else:
            lines.append("\nNo notes files found.")

        # Project info
        project_dir = getattr(self, '_project_directory', None)
        if project_dir:
            lines.append(f"\nProject directory: {project_dir}")

        # Available project files
        lines.append(f"\n=== AVAILABLE PROJECT FILES ===")
        try:
            recent_projects = self.project_manager.get_recent_projects()
            if recent_projects:
                lines.append(f"Recent projects ({len(recent_projects)}):")
                for proj in recent_projects:
                    lines.append(f"  - {proj['name']} (path: {proj['path']})")
            else:
                lines.append("No recent projects found.")
        except Exception:
            lines.append("Could not retrieve project list.")

        # AI Commands documentation
        lines.append(f"\n=== AI COMMANDS ===")
        lines.append("You can execute these commands by including them in your response:")
        lines.append("")
        lines.append("PROJECT MANAGEMENT:")
        lines.append("  [NEW_PROJECT: project_name] - Create a new project with the given name")
        lines.append("  [SAVE_PROJECT] - Save the current project")
        lines.append("  [RENAME_PROJECT: new_name] - Rename the current project")
        lines.append("  [LOAD_PROJECT: project_name] - Load a project by name")
        lines.append("")
        lines.append("FILE SCANNING:")
        lines.append("  [SCAN_DIRECTORY: path] - Scan a directory for data files")
        lines.append("  [SCAN_SAVED_DATA] - Scan for existing exported CSV/NPZ files")
        lines.append("  [SET_FILE_TYPES: abf,smrx,edf] - Set which file types to scan for")
        lines.append("")
        lines.append("TABLE FILTERING:")
        lines.append("  [FILTER_ROWS: search_text] - Filter table rows by search text")
        lines.append("  [CLEAR_FILTER] - Clear the table filter to show all rows")
        lines.append("  [SET_FILTER_COLUMN: column_name] - Set which column to filter (All, File Name, Protocol, Animal ID, Strain, Keywords)")
        lines.append("")
        lines.append("METADATA UPDATES:")
        lines.append("  [UPDATE_META: filename | field=value, field2=value2] - Update metadata")
        lines.append("    Valid fields: animal_id, strain, sex, power, stim_type, experiment")
        lines.append("  Example: [UPDATE_META: 0801_testday7_10.edf | animal_id=mouse_42, sex=M]")

        return "\n".join(lines)

    def _scroll_chat_to_bottom(self):
        """Scroll chat history to bottom."""
        scrollbar = self.chatHistoryText.verticalScrollBar()
        scrollbar.setValue(scrollbar.maximum())

    # ==================== LOCAL AI TOOLS ====================
    # These tools are called by the AI to search/query data locally
    # without sending the full dataset over the API

    def _get_ai_tool_definitions(self) -> list:
        """
        Get tool definitions for native Claude tool use.
        These define the schema for each tool the AI can call.
        """
        return [
            {
                "name": "search_files",
                "description": "Search project files by combining multiple criteria in ONE call. You can specify ANY combination of filters - they are applied together (AND logic). For example: protocol='30Hz' + has_exports=true + min_sweeps=10 finds 30Hz files with exports and at least 10 sweeps.",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "protocol": {
                            "type": "string",
                            "description": "Protocol name to filter by (partial match, e.g., '30Hz', 'baseline')"
                        },
                        "animal_id": {
                            "type": "string",
                            "description": "Animal ID to filter by"
                        },
                        "strain": {
                            "type": "string",
                            "description": "Strain to filter by (e.g., 'VgatCre', 'C57')"
                        },
                        "sex": {
                            "type": "string",
                            "description": "Sex to filter by ('M' or 'F')"
                        },
                        "status": {
                            "type": "string",
                            "description": "Status to filter by ('completed' or 'pending')"
                        },
                        "has_exports": {
                            "type": "boolean",
                            "description": "If true, only return files that have exported CSV/NPZ data"
                        },
                        "min_sweeps": {
                            "type": "integer",
                            "description": "Minimum number of sweeps required (e.g., 10 for files with at least 10 sweeps)"
                        },
                        "stim_type": {
                            "type": "string",
                            "description": "Stimulation type to filter by"
                        },
                        "power": {
                            "type": "string",
                            "description": "Laser/stim power to filter by (e.g., '10mW')"
                        },
                        "keyword": {
                            "type": "string",
                            "description": "General keyword to search in file name, protocol, and notes"
                        }
                    }
                }
            },
            {
                "name": "list_all_files",
                "description": "FALLBACK: List ALL files in the project with compact info. Use this when search isn't finding what you need, or to see everything available. Returns file names, protocols, sweep counts, status, and export info for every file.",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "include_exports": {
                            "type": "boolean",
                            "description": "Include export details (default: true)"
                        },
                        "limit": {
                            "type": "integer",
                            "description": "Max files to return (default: 50, max: 100)"
                        }
                    }
                }
            },
            {
                "name": "get_csv_columns",
                "description": "Get the column headers from CSV export files for a specific data file or export directory.",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "file_path": {
                            "type": "string",
                            "description": "Path or name of the data file to get CSV columns for"
                        },
                        "export_path": {
                            "type": "string",
                            "description": "Direct path to export directory containing CSV files"
                        }
                    }
                }
            },
            {
                "name": "list_protocols",
                "description": "List all unique protocols in the project with file counts for each.",
                "input_schema": {
                    "type": "object",
                    "properties": {}
                }
            },
            {
                "name": "list_animals",
                "description": "List all unique animal IDs in the project with file counts for each.",
                "input_schema": {
                    "type": "object",
                    "properties": {}
                }
            },
            {
                "name": "project_summary",
                "description": "Get a summary overview of the current project including file counts, status breakdown, and available protocols.",
                "input_schema": {
                    "type": "object",
                    "properties": {}
                }
            },
            {
                "name": "get_searchable_values",
                "description": "IMPORTANT: Call this FIRST before searching! Returns all unique values for each searchable field so you know exactly what terms exist in the data (exact protocol names, status values, animal IDs, etc). Use these exact values in your search queries.",
                "input_schema": {
                    "type": "object",
                    "properties": {}
                }
            }
        ]

    def _execute_native_tool(self, tool_name: str, tool_input: dict) -> dict:
        """Execute a native tool call and return the result."""
        tool_method = getattr(self, f'_ai_tool_{tool_name}', None)
        if tool_method:
            try:
                return tool_method(**tool_input)
            except Exception as e:
                return {"error": str(e)}
        else:
            return {"error": f"Unknown tool: {tool_name}"}

    def _ai_tool_search_files(self, **kwargs) -> dict:
        """
        Search files by criteria. Returns matching files with details.
        Kwargs can include: protocol, animal_id, strain, sex, status,
                           has_exports, min_sweeps, stim_type, power, keyword
        """
        if not self._master_file_list:
            return {"matches": [], "count": 0, "message": "No files loaded"}

        # First, build a map of file_path -> aggregated info from all rows (parent + sub-rows)
        # This allows us to aggregate exports from sub-rows to their parent file
        file_info_map = {}  # file_path -> aggregated info

        for task in self._master_file_list:
            file_path = str(task.get('file_path', ''))
            if not file_path:
                continue

            if file_path not in file_info_map:
                # Initialize with parent info
                file_info_map[file_path] = {
                    'file_name': task.get('file_name', ''),
                    'file_path': file_path,
                    'protocol': task.get('protocol', ''),
                    'status': task.get('status', 'pending'),
                    'animal_id': task.get('animal_id', ''),
                    'strain': task.get('strain', ''),
                    'sex': task.get('sex', ''),
                    'power': task.get('power', ''),
                    'stim_type': task.get('stim_type', ''),
                    'sweeps': task.get('sweeps', task.get('sweep_count', '')),
                    'keywords': task.get('keywords', ''),
                    'notes': task.get('notes', ''),
                    'exports': {},
                    'export_path': '',
                    'channels': [],
                }

            info = file_info_map[file_path]

            # Track channels (from sub-rows)
            channel = task.get('channel', '')
            if channel and channel not in info['channels']:
                info['channels'].append(channel)

            # Aggregate exports from all rows (sub-rows have the actual export data)
            task_exports = task.get('exports', {})
            if isinstance(task_exports, dict):
                for key, val in task_exports.items():
                    if val is True or val == 'true' or val == 'True' or val == 1:
                        info['exports'][key] = True

            # Update export_path if this row has one
            if task.get('export_path') and not info['export_path']:
                info['export_path'] = task.get('export_path')

            # Update status if any row is completed
            if task.get('status') == 'completed':
                info['status'] = 'completed'

            # Aggregate metadata (prefer non-empty values)
            for field in ['animal_id', 'strain', 'sex', 'power', 'stim_type']:
                if task.get(field) and not info.get(field):
                    info[field] = task.get(field)

        # Now search the aggregated file info
        matches = []
        for file_path, info in file_info_map.items():
            match = True

            # Protocol filter (partial match, case-insensitive)
            if 'protocol' in kwargs and kwargs['protocol']:
                proto = str(info.get('protocol', '')).lower()
                if kwargs['protocol'].lower() not in proto:
                    match = False

            # Animal ID filter
            if 'animal_id' in kwargs and kwargs['animal_id']:
                animal = str(info.get('animal_id', '')).lower()
                if kwargs['animal_id'].lower() not in animal:
                    match = False

            # Strain filter
            if 'strain' in kwargs and kwargs['strain']:
                strain = str(info.get('strain', '')).lower()
                if kwargs['strain'].lower() not in strain:
                    match = False

            # Sex filter
            if 'sex' in kwargs and kwargs['sex']:
                sex = str(info.get('sex', '')).lower()
                if kwargs['sex'].lower() != sex:
                    match = False

            # Status filter - flexible matching for "completed"
            if 'status' in kwargs and kwargs['status']:
                status_val = str(info.get('status', '')).lower().strip()
                search_status = kwargs['status'].lower().strip()
                # Treat various values as "completed"
                completed_values = {'completed', 'complete', 'done', '✓', 'true', 'yes', '1', 'checked'}
                pending_values = {'pending', 'incomplete', '', 'false', 'no', '0', 'unchecked'}

                if search_status in completed_values or search_status == 'completed':
                    # User wants completed files
                    if status_val not in completed_values:
                        match = False
                elif search_status in pending_values or search_status == 'pending':
                    # User wants pending files
                    if status_val not in pending_values:
                        match = False
                else:
                    # Exact/partial match for other values
                    if search_status not in status_val:
                        match = False

            # Has exports filter - check aggregated exports
            if 'has_exports' in kwargs and kwargs['has_exports']:
                exports = info.get('exports', {})
                export_path = info.get('export_path', '')

                # Check if exports dict has any truthy values
                has_export_data = any(
                    val is True or val == 'true' or val == 'True' or val == 1
                    for val in exports.values()
                ) if isinstance(exports, dict) else False

                # Check if export_path exists
                has_export_path = bool(export_path)

                if not (has_export_data or has_export_path):
                    match = False

            # Min sweeps filter
            if 'min_sweeps' in kwargs and kwargs['min_sweeps']:
                sweeps = info.get('sweeps', 0)
                # Handle string or int sweep counts
                try:
                    sweep_count = int(sweeps) if sweeps else 0
                except (ValueError, TypeError):
                    sweep_count = 0
                if sweep_count < int(kwargs['min_sweeps']):
                    match = False

            # Stim type filter
            if 'stim_type' in kwargs and kwargs['stim_type']:
                stim = str(info.get('stim_type', '')).lower()
                if kwargs['stim_type'].lower() not in stim:
                    match = False

            # Power filter
            if 'power' in kwargs and kwargs['power']:
                power = str(info.get('power', '')).lower()
                if kwargs['power'].lower() not in power:
                    match = False

            # Keyword filter (searches file name, protocol, notes)
            if 'keyword' in kwargs and kwargs['keyword']:
                kw = kwargs['keyword'].lower()
                searchable = ' '.join([
                    str(info.get('file_name', '')),
                    str(info.get('protocol', '')),
                    str(info.get('notes', '')),
                    str(info.get('keywords', ''))
                ]).lower()
                if kw not in searchable:
                    match = False

            if match:
                # Build compact file info for output
                file_info = {
                    'file_name': info.get('file_name', ''),
                    'file_path': info.get('file_path', ''),
                    'protocol': info.get('protocol', ''),
                    'status': info.get('status', 'pending'),
                }
                # Add channels if present
                if info.get('channels'):
                    file_info['channels'] = info['channels']
                # Add export info if present
                if info.get('exports'):
                    file_info['exports'] = info['exports']
                    file_info['export_path'] = info.get('export_path', '')
                # Add other metadata if present
                for field in ['animal_id', 'strain', 'sex', 'power', 'stim_type', 'sweeps']:
                    if info.get(field):
                        file_info[field] = info.get(field)
                matches.append(file_info)

        return {
            "matches": matches[:20],  # Limit results
            "count": len(matches),
            "truncated": len(matches) > 20
        }

    def _ai_tool_get_csv_columns(self, file_path: str = None, export_path: str = None) -> dict:
        """Get CSV column headers for a specific file or export directory."""
        from pathlib import Path

        if export_path:
            search_path = Path(export_path)
        elif file_path:
            # Find export path for this file
            for task in self._master_file_list:
                if task.get('file_path') == file_path or task.get('file_name') == file_path:
                    search_path = Path(task.get('export_path', ''))
                    break
            else:
                return {"error": f"File not found: {file_path}"}
        else:
            return {"error": "Provide file_path or export_path"}

        if not search_path.exists():
            return {"error": f"Path not found: {search_path}"}

        csv_info = []
        for csv_file in list(search_path.glob("*.csv"))[:5]:
            try:
                with open(csv_file, 'r') as f:
                    header = f.readline().strip()
                columns = [c.strip() for c in header.split(',')]
                csv_info.append({
                    "file": csv_file.name,
                    "path": str(csv_file),
                    "columns": columns
                })
            except Exception as e:
                csv_info.append({"file": csv_file.name, "error": str(e)})

        return {"csv_files": csv_info, "count": len(csv_info)}

    def _ai_tool_list_protocols(self) -> dict:
        """List all unique protocols in the project."""
        if not self._master_file_list:
            return {"protocols": [], "message": "No files loaded"}

        protocols = {}
        for task in self._master_file_list:
            if task.get('is_sub_row'):
                continue
            proto = task.get('protocol', '')
            if proto:
                protocols[proto] = protocols.get(proto, 0) + 1

        return {
            "protocols": [{"name": k, "count": v} for k, v in sorted(protocols.items())],
            "total": len(protocols)
        }

    def _ai_tool_list_animals(self) -> dict:
        """List all unique animal IDs in the project."""
        if not self._master_file_list:
            return {"animals": [], "message": "No files loaded"}

        animals = {}
        for task in self._master_file_list:
            if task.get('is_sub_row'):
                continue
            animal = task.get('animal_id', '')
            if animal:
                animals[animal] = animals.get(animal, 0) + 1

        return {
            "animals": [{"id": k, "count": v} for k, v in sorted(animals.items())],
            "total": len(animals)
        }

    def _ai_tool_project_summary(self) -> dict:
        """Get a summary of the current project without full file list."""
        if not self._master_file_list:
            return {"message": "No files loaded", "file_count": 0}

        # Aggregate data from all rows (parent + sub-rows) by file_path
        file_info_map = {}
        for task in self._master_file_list:
            file_path = str(task.get('file_path', ''))
            if not file_path:
                continue

            if file_path not in file_info_map:
                file_info_map[file_path] = {
                    'status': task.get('status', 'pending'),
                    'protocol': task.get('protocol', ''),
                    'animal_id': task.get('animal_id', ''),
                    'has_exports': False,
                    'has_export_path': False,
                }

            info = file_info_map[file_path]

            # Aggregate exports from sub-rows
            task_exports = task.get('exports', {})
            if isinstance(task_exports, dict):
                for val in task_exports.values():
                    if val is True or val == 'true' or val == 'True' or val == 1:
                        info['has_exports'] = True
                        break

            # Update export_path if this row has one
            if task.get('export_path'):
                info['has_export_path'] = True

            # Update status if any row is completed
            if task.get('status') == 'completed':
                info['status'] = 'completed'

        file_count = len(file_info_map)

        # Count status values
        status_counts = {}
        for info in file_info_map.values():
            status = str(info.get('status', 'none')).strip()
            status_counts[status] = status_counts.get(status, 0) + 1

        # Count files with exports
        with_exports = sum(1 for info in file_info_map.values() if info['has_exports'])
        with_export_path = sum(1 for info in file_info_map.values() if info['has_export_path'])

        protocols = set(info.get('protocol', '') for info in file_info_map.values() if info.get('protocol'))
        animals = set(info.get('animal_id', '') for info in file_info_map.values() if info.get('animal_id'))

        return {
            "file_count": file_count,
            "status_breakdown": status_counts,  # Shows actual status values
            "with_exports": with_exports,
            "with_export_path": with_export_path,
            "protocols": list(protocols)[:10],
            "animals": list(animals)[:10],
            "project_dir": str(getattr(self, '_project_directory', 'Not set'))
        }

    def _ai_tool_list_all_files(self, include_exports: bool = True, limit: int = 50) -> dict:
        """
        FALLBACK: List all files in the project with compact info.
        Use when search isn't working or to see everything.
        """
        if not self._master_file_list:
            return {"files": [], "count": 0, "message": "No files loaded"}

        # Clamp limit
        limit = min(max(1, limit), 100)

        # Aggregate data from all rows by file_path (same as search)
        file_info_map = {}
        for task in self._master_file_list:
            file_path = str(task.get('file_path', ''))
            if not file_path:
                continue

            if file_path not in file_info_map:
                file_info_map[file_path] = {
                    'file_name': task.get('file_name', ''),
                    'protocol': task.get('protocol', ''),
                    'sweeps': task.get('sweeps', task.get('sweep_count', '')),
                    'status': task.get('status', 'pending'),
                    'channels': [],
                    'exports': {},
                    'export_path': '',
                }

            info = file_info_map[file_path]

            # Track channels
            channel = task.get('channel', '')
            if channel and channel not in info['channels']:
                info['channels'].append(channel)

            # Aggregate exports
            task_exports = task.get('exports', {})
            if isinstance(task_exports, dict):
                for key, val in task_exports.items():
                    if val is True or val == 'true' or val == 'True' or val == 1:
                        info['exports'][key] = True

            if task.get('export_path') and not info['export_path']:
                info['export_path'] = task.get('export_path')

            if task.get('status') == 'completed':
                info['status'] = 'completed'

        # Build compact file list
        files = []
        for file_path, info in list(file_info_map.items())[:limit]:
            file_entry = {
                'name': info['file_name'],
                'protocol': info['protocol'],
                'sweeps': info['sweeps'],
                'status': info['status'],
            }
            if info['channels']:
                file_entry['channels'] = info['channels']
            if include_exports and info['exports']:
                # Format exports compactly
                export_types = [k for k, v in info['exports'].items() if v]
                if export_types:
                    file_entry['exports'] = export_types
                    file_entry['export_path'] = info['export_path']
            files.append(file_entry)

        return {
            "files": files,
            "count": len(file_info_map),
            "showing": len(files),
            "truncated": len(file_info_map) > limit
        }

    def _ai_tool_get_searchable_values(self) -> dict:
        """
        Get all unique values for each searchable field.
        This helps the AI know exactly what terms to search for.
        """
        if not self._master_file_list:
            return {"message": "No files loaded", "values": {}}

        # Collect unique values for each searchable field
        values = {
            "protocols": set(),
            "statuses": set(),
            "animal_ids": set(),
            "strains": set(),
            "sexes": set(),
            "stim_types": set(),
            "powers": set(),
            "channels": set(),
        }

        # Aggregate data from all rows (parent + sub-rows) by file_path
        file_info_map = {}
        for task in self._master_file_list:
            file_path = str(task.get('file_path', ''))
            if not file_path:
                continue

            if file_path not in file_info_map:
                file_info_map[file_path] = {
                    'file_name': task.get('file_name', ''),
                    'has_exports': False,
                }

            info = file_info_map[file_path]

            # Collect each field from all rows (parent and sub-rows)
            if task.get('protocol'):
                values['protocols'].add(str(task['protocol']))
            if task.get('status'):
                values['statuses'].add(str(task['status']))
            if task.get('animal_id'):
                values['animal_ids'].add(str(task['animal_id']))
            if task.get('strain'):
                values['strains'].add(str(task['strain']))
            if task.get('sex'):
                values['sexes'].add(str(task['sex']))
            if task.get('stim_type'):
                values['stim_types'].add(str(task['stim_type']))
            if task.get('power'):
                values['powers'].add(str(task['power']))
            if task.get('channel'):
                values['channels'].add(str(task['channel']))

            # Track files with exports (aggregate from all rows including sub-rows)
            task_exports = task.get('exports', {})
            if isinstance(task_exports, dict):
                for val in task_exports.values():
                    if val is True or val == 'true' or val == 'True' or val == 1:
                        info['has_exports'] = True
                        break
            if task.get('export_path'):
                info['has_exports'] = True

        # Build list of files with exports
        files_with_exports = [
            info['file_name'] for info in file_info_map.values()
            if info['has_exports']
        ]

        # Convert sets to sorted lists
        result = {}
        for key, val_set in values.items():
            sorted_vals = sorted([v for v in val_set if v])  # Remove empty strings
            if sorted_vals:
                result[key] = sorted_vals

        # Add export info
        result['files_with_exports'] = files_with_exports[:20]  # Limit to 20
        result['total_files'] = len(file_info_map)
        result['total_with_exports'] = len(files_with_exports)

        return result

    # ==================== END LOCAL AI TOOLS ====================

    def _execute_ai_commands(self, response: str) -> list:
        """
        Parse AI response for commands and execute them.
        Returns list of executed command results.
        """
        import re
        results = []

        # PROJECT MANAGEMENT COMMANDS

        # [NEW_PROJECT: project_name]
        new_proj_pattern = r'\[NEW_PROJECT:\s*([^\]]+)\]'
        for match in re.findall(new_proj_pattern, response, re.IGNORECASE):
            result = self._ai_new_project(match.strip())
            results.append(result)

        # [SAVE_PROJECT]
        if re.search(r'\[SAVE_PROJECT\]', response, re.IGNORECASE):
            result = self._ai_save_project()
            results.append(result)

        # [RENAME_PROJECT: new_name]
        rename_pattern = r'\[RENAME_PROJECT:\s*([^\]]+)\]'
        for match in re.findall(rename_pattern, response, re.IGNORECASE):
            result = self._ai_rename_project(match.strip())
            results.append(result)

        # [LOAD_PROJECT: project_name]
        load_pattern = r'\[LOAD_PROJECT:\s*([^\]]+)\]'
        for match in re.findall(load_pattern, response, re.IGNORECASE):
            result = self._ai_load_project(match.strip())
            results.append(result)

        # FILE SCANNING COMMANDS

        # [SCAN_DIRECTORY: path]
        scan_dir_pattern = r'\[SCAN_DIRECTORY:\s*([^\]]+)\]'
        for match in re.findall(scan_dir_pattern, response, re.IGNORECASE):
            result = self._ai_scan_directory(match.strip())
            results.append(result)

        # [SCAN_SAVED_DATA]
        if re.search(r'\[SCAN_SAVED_DATA\]', response, re.IGNORECASE):
            result = self._ai_scan_saved_data()
            results.append(result)

        # [SET_FILE_TYPES: abf,smrx,edf]
        file_types_pattern = r'\[SET_FILE_TYPES:\s*([^\]]+)\]'
        for match in re.findall(file_types_pattern, response, re.IGNORECASE):
            result = self._ai_set_file_types(match.strip())
            results.append(result)

        # TABLE FILTERING COMMANDS

        # [FILTER_ROWS: search_text]
        filter_pattern = r'\[FILTER_ROWS:\s*([^\]]+)\]'
        for match in re.findall(filter_pattern, response, re.IGNORECASE):
            result = self._ai_filter_rows(match.strip())
            results.append(result)

        # [CLEAR_FILTER]
        if re.search(r'\[CLEAR_FILTER\]', response, re.IGNORECASE):
            result = self._ai_clear_filter()
            results.append(result)

        # [SET_FILTER_COLUMN: column_name]
        filter_col_pattern = r'\[SET_FILTER_COLUMN:\s*([^\]]+)\]'
        for match in re.findall(filter_col_pattern, response, re.IGNORECASE):
            result = self._ai_set_filter_column(match.strip())
            results.append(result)

        # METADATA COMMANDS

        # [UPDATE_META: filename | field=value, field2=value2]
        meta_pattern = r'\[UPDATE_META:\s*([^|]+)\s*\|\s*([^\]]+)\]'
        for filename, fields_str in re.findall(meta_pattern, response, re.IGNORECASE):
            result = self._ai_update_metadata(filename.strip(), fields_str.strip())
            results.append(result)

        return results

    def _ai_load_project(self, project_name: str) -> dict:
        """Load a project by name (called by AI command)."""
        try:
            recent_projects = self.project_manager.get_recent_projects()
            matching_project = None

            # Find project by name (case-insensitive partial match)
            for proj in recent_projects:
                if project_name.lower() in proj['name'].lower():
                    matching_project = proj
                    break

            if not matching_project:
                return {
                    'command': 'LOAD_PROJECT',
                    'success': False,
                    'message': f"Project '{project_name}' not found in recent projects."
                }

            # Load the project
            from pathlib import Path
            project_path = Path(matching_project['path'])
            project_data = self.project_manager.load_project(project_path)

            # Populate UI with loaded data (this handles master_file_list, table refresh, etc.)
            self._populate_ui_from_project(project_data)

            # Update status
            self._log_status_message(f"✓ Loaded project: {matching_project['name']}", 3000)

            return {
                'command': 'LOAD_PROJECT',
                'success': True,
                'message': f"Successfully loaded project '{matching_project['name']}' with {len(self._master_file_list)} files."
            }

        except Exception as e:
            return {
                'command': 'LOAD_PROJECT',
                'success': False,
                'message': f"Error loading project: {str(e)}"
            }

    def _ai_update_metadata(self, filename: str, fields_str: str) -> dict:
        """Update metadata for a file (called by AI command)."""
        try:
            # Valid metadata fields
            valid_fields = {'animal_id', 'strain', 'sex', 'power', 'stim_type', 'experiment'}

            # Parse field=value pairs
            updates = {}
            for pair in fields_str.split(','):
                pair = pair.strip()
                if '=' in pair:
                    field, value = pair.split('=', 1)
                    field = field.strip().lower()
                    value = value.strip()
                    if field in valid_fields:
                        updates[field] = value

            if not updates:
                return {
                    'command': 'UPDATE_META',
                    'success': False,
                    'message': f"No valid fields to update. Valid fields: {', '.join(valid_fields)}"
                }

            # Find the file(s) matching the filename
            matched_rows = []
            for i, task in enumerate(self._master_file_list):
                task_filename = task.get('file_name', '')
                if filename.lower() in task_filename.lower():
                    matched_rows.append(i)

            if not matched_rows:
                return {
                    'command': 'UPDATE_META',
                    'success': False,
                    'message': f"No file matching '{filename}' found in current project."
                }

            # Apply updates to all matched rows
            for row in matched_rows:
                for field, value in updates.items():
                    self._set_table_cell_value(row, field, value)

            # Report success
            fields_updated = ', '.join([f"{k}={v}" for k, v in updates.items()])
            return {
                'command': 'UPDATE_META',
                'success': True,
                'message': f"Updated {len(matched_rows)} file(s) matching '{filename}': {fields_updated}"
            }

        except Exception as e:
            return {
                'command': 'UPDATE_META',
                'success': False,
                'message': f"Error updating metadata: {str(e)}"
            }

    def _ai_new_project(self, project_name: str) -> dict:
        """Create a new project with the given name (called by AI command)."""
        try:
            from pathlib import Path

            # Clear existing data first
            self._master_file_list = []
            self._discovered_files_data = []
            self._file_table_model.clear()

            # Set the project name in the UI
            self._current_project_name = project_name
            if hasattr(self, 'projectNameCombo'):
                self.projectNameCombo.setCurrentText(project_name)

            # Create a default project directory (user can change it later)
            default_dir = Path.home() / "PhysioMetrics" / project_name
            default_dir.mkdir(parents=True, exist_ok=True)
            self._project_directory = str(default_dir)

            if hasattr(self, 'directoryPathEdit'):
                self.directoryPathEdit.setText(str(default_dir))

            self._log_status_message(f"✓ Created new project: {project_name}", 3000)

            return {
                'command': 'NEW_PROJECT',
                'success': True,
                'message': f"Created new project '{project_name}' with directory: {default_dir}"
            }

        except Exception as e:
            return {
                'command': 'NEW_PROJECT',
                'success': False,
                'message': f"Error creating project: {str(e)}"
            }

    def _ai_save_project(self) -> dict:
        """Save the current project (called by AI command)."""
        try:
            # Call the existing save method
            self.on_project_save()

            return {
                'command': 'SAVE_PROJECT',
                'success': True,
                'message': "Project saved successfully."
            }

        except Exception as e:
            return {
                'command': 'SAVE_PROJECT',
                'success': False,
                'message': f"Error saving project: {str(e)}"
            }

    def _ai_rename_project(self, new_name: str) -> dict:
        """Rename the current project (called by AI command)."""
        try:
            old_name = getattr(self, '_current_project_name', '') or ''
            self._current_project_name = new_name
            if hasattr(self, 'projectNameCombo'):
                self.projectNameCombo.setCurrentText(new_name)

            self._log_status_message(f"✓ Renamed project to: {new_name}", 3000)

            return {
                'command': 'RENAME_PROJECT',
                'success': True,
                'message': f"Renamed project from '{old_name}' to '{new_name}'"
            }

        except Exception as e:
            return {
                'command': 'RENAME_PROJECT',
                'success': False,
                'message': f"Error renaming project: {str(e)}"
            }

    def _ai_scan_directory(self, directory_path: str) -> dict:
        """Scan a directory for data files (called by AI command)."""
        try:
            from pathlib import Path

            dir_path = Path(directory_path)
            if not dir_path.exists():
                return {
                    'command': 'SCAN_DIRECTORY',
                    'success': False,
                    'message': f"Directory does not exist: {directory_path}"
                }

            # Update the directory path in UI
            if hasattr(self, 'directoryPathEdit'):
                self.directoryPathEdit.setText(directory_path)

            self._project_directory = directory_path

            # Trigger the scan (this runs async)
            self.on_project_scan_files()

            return {
                'command': 'SCAN_DIRECTORY',
                'success': True,
                'message': f"Started scanning directory: {directory_path}"
            }

        except Exception as e:
            return {
                'command': 'SCAN_DIRECTORY',
                'success': False,
                'message': f"Error scanning directory: {str(e)}"
            }

    def _ai_scan_saved_data(self) -> dict:
        """Scan for existing exported CSV/NPZ files (called by AI command)."""
        try:
            # Trigger the saved data scan
            self.on_project_scan_saved_data()

            return {
                'command': 'SCAN_SAVED_DATA',
                'success': True,
                'message': "Started scanning for saved data files."
            }

        except Exception as e:
            return {
                'command': 'SCAN_SAVED_DATA',
                'success': False,
                'message': f"Error scanning saved data: {str(e)}"
            }

    def _ai_set_file_types(self, types_str: str) -> dict:
        """Set which file types to scan for (called by AI command)."""
        try:
            # Parse the comma-separated file types
            file_types = [t.strip().lower() for t in types_str.split(',')]

            # Update checkboxes based on file types
            changed = []
            if hasattr(self, 'scanAbfCheckbox'):
                state = 'abf' in file_types
                self.scanAbfCheckbox.setChecked(state)
                if state:
                    changed.append('ABF')

            if hasattr(self, 'scanSmrxCheckbox'):
                state = 'smrx' in file_types
                self.scanSmrxCheckbox.setChecked(state)
                if state:
                    changed.append('SMRX')

            if hasattr(self, 'scanEdfCheckbox'):
                state = 'edf' in file_types
                self.scanEdfCheckbox.setChecked(state)
                if state:
                    changed.append('EDF')

            if hasattr(self, 'scanNotesCheckbox'):
                state = 'notes' in file_types or 'xlsx' in file_types or 'txt' in file_types
                self.scanNotesCheckbox.setChecked(state)
                if state:
                    changed.append('Notes')

            return {
                'command': 'SET_FILE_TYPES',
                'success': True,
                'message': f"Set file types to scan: {', '.join(changed) if changed else 'None'}"
            }

        except Exception as e:
            return {
                'command': 'SET_FILE_TYPES',
                'success': False,
                'message': f"Error setting file types: {str(e)}"
            }

    def _ai_filter_rows(self, search_text: str) -> dict:
        """Filter table rows by search text (called by AI command)."""
        try:
            if hasattr(self, 'tableFilterEdit'):
                self.tableFilterEdit.setText(search_text)
                # Trigger the filter
                self._on_table_filter_changed()

            # Count visible rows
            visible_count = 0
            total_count = self._file_table_model.rowCount()
            table = self.discoveredFilesTable
            for row in range(total_count):
                if not table.isRowHidden(row):
                    visible_count += 1

            return {
                'command': 'FILTER_ROWS',
                'success': True,
                'message': f"Filtered table with '{search_text}': showing {visible_count} of {total_count} rows"
            }

        except Exception as e:
            return {
                'command': 'FILTER_ROWS',
                'success': False,
                'message': f"Error filtering rows: {str(e)}"
            }

    def _ai_clear_filter(self) -> dict:
        """Clear the table filter to show all rows (called by AI command)."""
        try:
            if hasattr(self, 'tableFilterEdit'):
                self.tableFilterEdit.clear()
                # Trigger the filter update
                self._on_table_filter_changed()

            total_count = self._file_table_model.rowCount()

            return {
                'command': 'CLEAR_FILTER',
                'success': True,
                'message': f"Cleared filter. Showing all {total_count} rows."
            }

        except Exception as e:
            return {
                'command': 'CLEAR_FILTER',
                'success': False,
                'message': f"Error clearing filter: {str(e)}"
            }

    def _ai_set_filter_column(self, column_name: str) -> dict:
        """Set which column to filter (called by AI command)."""
        try:
            if not hasattr(self, 'filterColumnCombo'):
                return {
                    'command': 'SET_FILTER_COLUMN',
                    'success': False,
                    'message': "Filter column selector not available."
                }

            # Map column names to combo box indices
            column_map = {
                'all': 0,
                'file name': 1,
                'filename': 1,
                'protocol': 2,
                'animal id': 3,
                'animal_id': 3,
                'strain': 4,
                'keywords': 5
            }

            col_lower = column_name.lower().strip()
            if col_lower in column_map:
                self.filterColumnCombo.setCurrentIndex(column_map[col_lower])

                return {
                    'command': 'SET_FILTER_COLUMN',
                    'success': True,
                    'message': f"Set filter column to: {column_name}"
                }
            else:
                available = "All, File Name, Protocol, Animal ID, Strain, Keywords"
                return {
                    'command': 'SET_FILTER_COLUMN',
                    'success': False,
                    'message': f"Unknown column '{column_name}'. Available: {available}"
                }

        except Exception as e:
            return {
                'command': 'SET_FILTER_COLUMN',
                'success': False,
                'message': f"Error setting filter column: {str(e)}"
            }

    def _process_chat_message_local(self, message: str) -> str:
        """Process chat message locally (no API) based on keyword matching."""
        msg_lower = message.lower()

        # Get current data state
        num_data_files = len(self._master_file_list) if self._master_file_list else 0
        notes_files = getattr(self, '_discovered_notes_files', [])
        num_notes = len(notes_files) if notes_files else 0

        # Check for file-related queries
        if any(word in msg_lower for word in ['file', 'files', 'data', 'see', 'list', 'show', 'what']):
            if 'note' in msg_lower:
                # Query about notes files
                if num_notes == 0:
                    return ("I don't see any notes files yet. Please scan a directory first using "
                            "'Scan for New Files' with the 'Notes' checkbox enabled.")
                else:
                    notes_list = "<br>".join([f"• {f.name}" for f in notes_files[:10]])
                    extra = f"<br>... and {num_notes - 10} more" if num_notes > 10 else ""
                    return (f"I can see <b>{num_notes} notes file(s)</b>:<br>{notes_list}{extra}<br><br>"
                            "Once AI integration is configured (click ⚙), I'll be able to read and "
                            "extract metadata from these files!")

            elif any(word in msg_lower for word in ['data', 'abf', 'smrx', 'edf', 'file']):
                # Query about data files
                if num_data_files == 0:
                    return ("I don't see any data files yet. Please scan a directory first using "
                            "'Scan for New Files'.")
                else:
                    # Summarize by file type
                    abf_count = sum(1 for f in self._master_file_list if str(f.get('file_path', '')).lower().endswith('.abf'))
                    smrx_count = sum(1 for f in self._master_file_list if str(f.get('file_path', '')).lower().endswith('.smrx'))
                    edf_count = sum(1 for f in self._master_file_list if str(f.get('file_path', '')).lower().endswith('.edf'))

                    # Get unique protocols
                    protocols = set(f.get('protocol', 'Unknown') for f in self._master_file_list)
                    protocols_str = ", ".join(sorted(protocols)[:5])
                    if len(protocols) > 5:
                        protocols_str += f" (+{len(protocols)-5} more)"

                    # Sample file names
                    sample_files = [f.get('file_name', 'Unknown') for f in self._master_file_list[:5]]
                    samples_str = "<br>".join([f"• {name}" for name in sample_files])
                    extra = f"<br>... and {num_data_files - 5} more" if num_data_files > 5 else ""

                    return (f"I can see <b>{num_data_files} data file(s)</b>:<br>"
                            f"• ABF: {abf_count}<br>• SMRX: {smrx_count}<br>• EDF: {edf_count}<br><br>"
                            f"<b>Protocols found:</b> {protocols_str}<br><br>"
                            f"<b>Sample files:</b><br>{samples_str}{extra}")

        # Check for plot requests
        if any(word in msg_lower for word in ['plot', 'graph', 'chart', 'visualize', 'draw']):
            if num_data_files == 0:
                return "I'd love to help you plot data, but no files are loaded yet. Please scan a directory first."

            # Generate sample code in the notebook
            sample_code = '''# Example: Load and plot data from first file
from pathlib import Path
import pyabf
import matplotlib.pyplot as plt

# Get the first file path
file_path = master_file_list[0]['file_path']
print(f"Loading: {file_path}")

# Load and plot (for ABF files)
if str(file_path).lower().endswith('.abf'):
    abf = pyabf.ABF(str(file_path))
    abf.setSweep(0)
    plt.figure(figsize=(10, 4))
    plt.plot(abf.sweepX, abf.sweepY)
    plt.xlabel('Time (s)')
    plt.ylabel(abf.sweepLabelY)
    plt.title(f'{Path(file_path).name}')
    plt.tight_layout()
    plt.show()
'''
            if hasattr(self, 'codeInputEdit'):
                self.codeInputEdit.setPlainText(sample_code)

            return ("I've generated sample plotting code in the Code Notebook below! "
                    "Click <b>▶ Run</b> to execute it.<br><br>"
                    "The code will load and plot the first file in your list. "
                    "You can modify it to plot different files or customize the visualization.")

        # Check for help queries
        if any(word in msg_lower for word in ['help', 'can you', 'what can', 'how']):
            return (f"<b>What I can see right now:</b><br>"
                    f"• Data files: {num_data_files}<br>"
                    f"• Notes files: {num_notes}<br><br>"
                    "<b>What I can do:</b><br>"
                    "• Tell you about loaded files (ask: 'what files do you see?')<br>"
                    "• List notes files (ask: 'show me the notes files')<br>"
                    "• Generate plotting code (ask: 'plot the data')<br><br>"
                    "<b>Coming soon</b> (configure API key via ⚙):<br>"
                    "• Read and summarize notes files<br>"
                    "• Extract metadata automatically<br>"
                    "• Generate custom analysis code<br>"
                    "• Answer questions about your data")

        # Default response
        return (f"I'm here to help! I can currently see {num_data_files} data file(s) and {num_notes} notes file(s).<br><br>"
                "Try asking me:<br>"
                "• 'What files do you see?'<br>"
                "• 'Show me the notes files'<br>"
                "• 'Plot the data'<br>"
                "• 'Help'<br><br>"
                "For full AI capabilities, click the ⚙ button to configure your API key.")

    def _check_code_safety(self, code: str) -> tuple:
        """
        Check code for dangerous operations using AST parsing.

        RELAXED POLICY: Allows file reading and plotting, blocks only destructive operations.

        Returns:
            (is_safe, blocked_reasons, warnings)
            - is_safe: True if code passes all checks
            - blocked_reasons: List of reasons code was blocked (hard reject)
            - warnings: List of warnings (user can override)
        """
        import ast

        blocked_reasons = []
        warnings = []

        # Blocked imports - only truly dangerous ones (system manipulation, network, code execution)
        blocked_imports = {
            'subprocess',  # Running shell commands
            'shutil',      # File deletion/moving
            'ctypes',      # Low-level memory access
            'multiprocessing',  # Process spawning
            'socket',      # Network connections
            'ftplib', 'smtplib',  # Network protocols
            'pty', 'pipes',  # Terminal/process control
            'signal',      # Signal handling
        }

        # Allowed imports (explicitly safe for data analysis)
        # os - needed for os.path operations (reading paths)
        # sys - may be needed for some libraries
        # pickle - needed to load saved data
        # tempfile - harmless for temp files

        # Blocked function calls - only destructive ones
        blocked_calls = {
            'eval', 'exec', '__import__', 'compile',  # Code execution
            'breakpoint',  # Debugger
            'exit', 'quit',  # Exit handlers
        }

        # Warning patterns (user can proceed with confirmation)
        warning_patterns = ['.to_csv', '.to_excel', '.to_parquet', '.to_json']

        try:
            tree = ast.parse(code)
        except SyntaxError as e:
            return False, [f"Syntax error: {e}"], []

        class SafetyVisitor(ast.NodeVisitor):
            def __init__(self):
                self.blocked = []
                self.warned = []

            def visit_Import(self, node):
                for alias in node.names:
                    module_name = alias.name.split('.')[0]
                    if module_name in blocked_imports:
                        self.blocked.append(f"Blocked import: '{alias.name}' (system/network access)")
                self.generic_visit(node)

            def visit_ImportFrom(self, node):
                if node.module:
                    module_name = node.module.split('.')[0]
                    if module_name in blocked_imports:
                        self.blocked.append(f"Blocked import: 'from {node.module}' (system/network access)")
                self.generic_visit(node)

            def visit_Call(self, node):
                # Check for blocked function calls
                if isinstance(node.func, ast.Name):
                    if node.func.id in blocked_calls:
                        self.blocked.append(f"Blocked call: '{node.func.id}()' (code execution)")
                elif isinstance(node.func, ast.Attribute):
                    # Check for dangerous/destructive method calls
                    attr = node.func.attr
                    if attr in ['unlink', 'rmdir', 'remove', 'rmtree']:
                        self.blocked.append(f"Blocked method: '.{attr}()' (file deletion)")
                    elif attr == 'system':
                        self.blocked.append(f"Blocked method: '.system()' (shell command)")
                    # Check for warning patterns (writing files)
                    for pattern in warning_patterns:
                        if pattern.lstrip('.') == attr:
                            self.warned.append(f"File write operation: '.{attr}()'")
                self.generic_visit(node)

        visitor = SafetyVisitor()
        visitor.visit(tree)

        blocked_reasons = visitor.blocked
        warnings = visitor.warned

        # Warn (but don't block) on file writes
        code_str = code.lower()
        if 'open(' in code_str and ("'w'" in code_str or '"w"' in code_str or "'a'" in code_str or '"a"' in code_str):
            warnings.append("File open with write/append mode detected")

        is_safe = len(blocked_reasons) == 0
        return is_safe, blocked_reasons, warnings

    def _get_sandbox_directory(self) -> str:
        """Get or create the sandbox directory for code execution file operations."""
        from pathlib import Path
        import os

        # Use a sandbox folder in user's app data
        if os.name == 'nt':  # Windows
            base = Path(os.environ.get('LOCALAPPDATA', Path.home()))
        else:  # Linux/Mac
            base = Path.home() / '.local' / 'share'

        sandbox = base / 'PhysioMetrics' / 'code_sandbox'
        sandbox.mkdir(parents=True, exist_ok=True)
        return str(sandbox)

    def _on_run_code(self):
        """Execute code from the notebook code input with safety checks and timeout."""
        if not hasattr(self, 'codeInputEdit') or not hasattr(self, 'codeOutputText'):
            return

        code = self.codeInputEdit.toPlainText().strip()
        if not code:
            self.codeOutputText.append("<span style='color: #f48771;'>No code to execute.</span>")
            return

        # AST-based security check
        is_safe, blocked_reasons, warnings = self._check_code_safety(code)

        # Hard reject if blocked operations found
        if not is_safe:
            from PyQt6.QtWidgets import QMessageBox
            msg = QMessageBox(self)

            # Check if it's a syntax error vs security block
            is_syntax_error = any("Syntax error" in r for r in blocked_reasons)

            if is_syntax_error:
                msg.setIcon(QMessageBox.Icon.Warning)
                msg.setWindowTitle("Code Syntax Error")
                msg.setText("The code has a syntax error and cannot be executed:")
                msg.setInformativeText("\n".join(f"  ✗ {r}" for r in blocked_reasons[:5]))
                msg.setDetailedText(
                    "Check the code for:\n"
                    "• Missing colons after def/if/for/while\n"
                    "• Unmatched parentheses or brackets\n"
                    "• Incorrect indentation\n"
                    "• Missing quotes in strings\n\n"
                    "Try fixing the syntax error and running again."
                )
            else:
                msg.setIcon(QMessageBox.Icon.Critical)
                msg.setWindowTitle("Code Blocked - Security Violation")
                msg.setText("This code contains blocked operations and cannot be executed:")
                msg.setInformativeText("\n".join(f"  ✗ {r}" for r in blocked_reasons[:5]))
                msg.setDetailedText(
                    "For security, the following are blocked:\n"
                    "• Imports: subprocess, shutil, ctypes, socket, multiprocessing\n"
                    "• Calls: eval(), exec(), __import__(), compile()\n"
                    "• Methods: .unlink(), .rmdir(), .remove(), .rmtree(), .system()\n\n"
                    "ALLOWED: os, sys, pandas, numpy, scipy, matplotlib, open() for reading\n\n"
                    "The code notebook is for data analysis.\n"
                    "Use built-in Export functions to save results."
                )
            msg.exec()
            error_type = "syntax error" if is_syntax_error else "prohibited operations"
            self.codeOutputText.append(f"<span style='color: #f48771;'>⛔ Code blocked - {error_type}.</span>")
            return

        # Warn if potentially risky operations found (user can proceed)
        if warnings:
            from PyQt6.QtWidgets import QMessageBox
            msg = QMessageBox(self)
            msg.setIcon(QMessageBox.Icon.Warning)
            msg.setWindowTitle("Potentially Risky Code")
            msg.setText("This code contains operations that may write files:")
            msg.setInformativeText("\n".join(f"  ⚠ {w}" for w in warnings[:5]))
            msg.setDetailedText(
                "These operations can create files but are sandboxed.\n"
                "Any files created will be in the sandbox directory.\n\n"
                "Are you sure you want to run this code?"
            )
            msg.setStandardButtons(QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            msg.setDefaultButton(QMessageBox.StandardButton.No)

            if msg.exec() != QMessageBox.StandardButton.Yes:
                self.codeOutputText.append("<span style='color: #f48771;'>Execution cancelled by user.</span>")
                return

        self.codeOutputText.append("<span style='color: #569cd6;'>>>> Running code...</span>")

        # Set up execution namespace
        import numpy as np
        import matplotlib
        import matplotlib.pyplot as plt
        import io
        import base64
        from contextlib import redirect_stdout, redirect_stderr

        # Use non-interactive backend to prevent external windows
        original_backend = matplotlib.get_backend()
        matplotlib.use('Agg')
        plt.close('all')  # Close any existing figures

        # Try to import pandas
        pd = None
        try:
            import pandas as pd
        except ImportError:
            pass

        # Helper functions for data access
        def _normalize_search_term(filename_contains):
            """Extract clean search term from path, filename, or fragment."""
            if not filename_contains:
                return None
            # Handle full paths - extract just the base name
            search_term = Path(filename_contains).stem.lower()
            # Remove common suffixes like _bundle, _means_by_time, etc.
            for suffix in ['_bundle', '_means_by_time', '_breaths', '_pleth']:
                if search_term.endswith(suffix):
                    search_term = search_term[:-len(suffix)]
            return search_term

        def get_export_paths(filename_contains=None):
            """Get export paths for files. Returns dict of {filename: export_path}.

            Args:
                filename_contains: Can be full path, filename, or partial name to match
            """
            result = {}
            search_term = _normalize_search_term(filename_contains)

            for f in self._master_file_list or []:
                fn = f.get('file_name', '')
                export_path = f.get('export_path', '')
                if export_path:
                    fn_stem = Path(fn).stem.lower()
                    if search_term is None or search_term in fn.lower() or search_term in fn_stem:
                        result[fn] = export_path
            return result

        def load_means_csv(filename_contains):
            """Load _means_by_time.csv for a file. Returns DataFrame.

            Args:
                filename_contains: Filename, partial name, or full path to match
            """
            # If it's a direct path to a CSV, load it
            p = Path(filename_contains)
            if p.suffix.lower() == '.csv' and p.exists():
                return pd.read_csv(p)

            paths = get_export_paths(filename_contains)
            if not paths:
                raise FileNotFoundError(f"No exports found matching '{_normalize_search_term(filename_contains)}'")
            export_dir = list(paths.values())[0]
            csv_files = list(Path(export_dir).glob('*_means_by_time.csv'))
            if not csv_files:
                raise FileNotFoundError(f"No _means_by_time.csv in {export_dir}")
            return pd.read_csv(csv_files[0])

        def load_breaths_csv(filename_contains):
            """Load _breaths.csv for a file. Returns DataFrame.

            Args:
                filename_contains: Filename, partial name, or full path to match
            """
            p = Path(filename_contains)
            if p.suffix.lower() == '.csv' and p.exists():
                return pd.read_csv(p)

            paths = get_export_paths(filename_contains)
            if not paths:
                raise FileNotFoundError(f"No exports found matching '{_normalize_search_term(filename_contains)}'")
            export_dir = list(paths.values())[0]
            csv_files = list(Path(export_dir).glob('*_breaths.csv'))
            if not csv_files:
                raise FileNotFoundError(f"No _breaths.csv in {export_dir}")
            return pd.read_csv(csv_files[0])

        def load_bundle_npz(filename_contains):
            """Load _bundle.npz for a file. Returns numpy NpzFile object.

            Args:
                filename_contains: Filename, partial name, or full path to match
            """
            p = Path(filename_contains)
            if p.suffix.lower() == '.npz' and p.exists():
                return np.load(p, allow_pickle=True)

            paths = get_export_paths(filename_contains)
            if not paths:
                raise FileNotFoundError(f"No exports found matching '{_normalize_search_term(filename_contains)}'")
            export_dir = list(paths.values())[0]
            npz_files = list(Path(export_dir).glob('*_bundle.npz'))
            if not npz_files:
                raise FileNotFoundError(f"No _bundle.npz in {export_dir}")
            return np.load(npz_files[0], allow_pickle=True)

        def get_stim_spans(bundle_data):
            """Extract stim spans from bundle.npz data. Returns dict {sweep_idx: [(start, end), ...]}."""
            import json as _json
            if 'stim_spans_json' in bundle_data:
                spans_str = str(bundle_data['stim_spans_json'])
                return {int(k): v for k, v in _json.loads(spans_str).items()}
            return {}

        def add_stim_shading(ax, stim_spans, sweep_idx=0, color='blue', alpha=0.2, label='Stim'):
            """Add blue shading for stimulation periods."""
            spans = stim_spans.get(sweep_idx, [])
            for i, (start, end) in enumerate(spans):
                ax.axvspan(start, end, alpha=alpha, color=color,
                          label=label if i == 0 else None)

        def list_available_files():
            """List files with exports available."""
            result = []
            for f in self._master_file_list or []:
                fn = f.get('file_name', '')
                export_path = f.get('export_path', '')
                if export_path:
                    result.append(f"{fn} -> {export_path}")
            return result

        # Create namespace with common data science imports and helpers
        exec_namespace = {
            'np': np,
            'plt': plt,
            'pd': pd,
            'Path': Path,
            'json': __import__('json'),
            'master_file_list': self._master_file_list,
            # Helper functions
            'get_export_paths': get_export_paths,
            'load_means_csv': load_means_csv,
            'load_breaths_csv': load_breaths_csv,
            'load_bundle_npz': load_bundle_npz,
            'get_stim_spans': get_stim_spans,
            'add_stim_shading': add_stim_shading,
            'list_available_files': list_available_files,
        }

        # Capture stdout
        stdout_capture = io.StringIO()
        stderr_capture = io.StringIO()

        # Disable run button during execution
        if hasattr(self, 'runCodeButton'):
            self.runCodeButton.setEnabled(False)

        # Process events to update UI before potentially long execution
        QApplication.processEvents()

        # Execute directly on main thread
        success = True
        try:
            with redirect_stdout(stdout_capture), redirect_stderr(stderr_capture):
                exec(code, exec_namespace)
        except Exception as e:
            stderr_capture.write(f"{type(e).__name__}: {e}")
            success = False

        # Display text results
        stdout = stdout_capture.getvalue()
        stderr = stderr_capture.getvalue()

        if stdout:
            self.codeOutputText.append(f"<pre style='color: #d4d4d4;'>{stdout}</pre>")
        if stderr:
            color = '#f48771' if not success else '#dcdcaa'
            self.codeOutputText.append(f"<pre style='color: {color};'>{stderr}</pre>")

        # Capture and embed any matplotlib figures inline
        figs = [plt.figure(i) for i in plt.get_fignums()]
        if figs:
            # Initialize figure storage if needed
            if not hasattr(self, '_notebook_figures'):
                self._notebook_figures = []

            self.codeOutputText.append(f"<span style='color: #569cd6;'>📊 {len(figs)} figure(s) generated:</span>")
            for i, fig in enumerate(figs):
                # Save figure to bytes buffer
                buf = io.BytesIO()
                fig.savefig(buf, format='png', dpi=100, facecolor='#1e1e1e',
                           edgecolor='none', bbox_inches='tight')
                buf.seek(0)
                fig_bytes = buf.read()
                buf.close()

                # Store figure data for Pop Out / Save functionality
                self._notebook_figures.append({
                    'bytes': fig_bytes,
                    'figsize': fig.get_size_inches().tolist(),
                    'index': len(self._notebook_figures) + 1
                })

                # Convert to base64 for embedding in HTML
                img_data = base64.b64encode(fig_bytes).decode('utf-8')

                # Embed as inline image
                img_html = f'''
                <div style="margin: 8px 0; padding: 4px; border: 1px solid #3e3e42; border-radius: 4px; background: #252526;">
                    <img src="data:image/png;base64,{img_data}" style="max-width: 100%; height: auto;" />
                    <div style="color: #808080; font-size: 9pt; margin-top: 4px;">Figure {i+1} - Use "Pop Out" to view larger or "Save" to export</div>
                </div>
                '''
                self.codeOutputText.append(img_html)

            plt.close('all')
        elif success and not stdout and not stderr:
            self.codeOutputText.append("<span style='color: #4ec9b0;'>✓ Code executed successfully (no output).</span>")

        # Restore original matplotlib backend
        try:
            matplotlib.use(original_backend)
        except Exception:
            pass  # May fail if backend doesn't support switching

        # Scroll to bottom
        scrollbar = self.codeOutputText.verticalScrollBar()
        scrollbar.setValue(scrollbar.maximum())

        # Re-enable run button
        if hasattr(self, 'runCodeButton'):
            self.runCodeButton.setEnabled(True)

    def _on_clear_code_output(self):
        """Clear the notebook output area."""
        if hasattr(self, 'codeOutputText'):
            self.codeOutputText.clear()
        # Also clear stored figures
        self._notebook_figures = []

    def _add_notebook_extra_buttons(self):
        """Add Pop Out and Save buttons to the notebook header."""
        from PyQt6.QtWidgets import QPushButton

        # Store figures for pop-out functionality
        self._notebook_figures = []

        # Find the header layout (where Run and Clear buttons are)
        if not hasattr(self, 'clearOutputButton') or not self.clearOutputButton.parent():
            return

        parent_layout = self.clearOutputButton.parent().layout()
        if not parent_layout:
            return

        # Button style matching existing buttons
        button_style = """
            QPushButton {
                background-color: #6c757d;
                color: white;
                border-radius: 4px;
                padding: 4px 8px;
            }
            QPushButton:hover {
                background-color: #5a6268;
            }
            QPushButton:pressed {
                background-color: #545b62;
            }
        """

        # Pop Out button - opens figures in external windows
        self.popOutFigureButton = QPushButton("🪟 Pop Out")
        self.popOutFigureButton.setToolTip("Open the last figure in an external window")
        self.popOutFigureButton.setMinimumSize(80, 24)
        self.popOutFigureButton.setStyleSheet(button_style)
        self.popOutFigureButton.clicked.connect(self._on_pop_out_figure)
        parent_layout.addWidget(self.popOutFigureButton)

        # Save Figure button
        self.saveFigureButton = QPushButton("💾 Save")
        self.saveFigureButton.setToolTip("Save the last figure to a file")
        self.saveFigureButton.setMinimumSize(70, 24)
        self.saveFigureButton.setStyleSheet(button_style)
        self.saveFigureButton.clicked.connect(self._on_save_figure)
        parent_layout.addWidget(self.saveFigureButton)

        print("[notebook] Added Pop Out and Save Figure buttons")

    def _on_pop_out_figure(self):
        """Open the last generated figure in an external matplotlib window."""
        import matplotlib
        import matplotlib.pyplot as plt

        if not hasattr(self, '_notebook_figures') or not self._notebook_figures:
            self._log_status_message("No figures to pop out. Run code that generates a plot first.", 2000)
            return

        # Temporarily switch to interactive backend
        try:
            matplotlib.use('TkAgg')  # Or 'Qt5Agg' depending on system
        except Exception:
            pass

        # Recreate the last figure
        fig_data = self._notebook_figures[-1]
        fig, ax = plt.subplots(figsize=fig_data.get('figsize', (8, 6)))

        # If we stored the figure bytes, display it
        if 'bytes' in fig_data:
            import io
            from PIL import Image
            img = Image.open(io.BytesIO(fig_data['bytes']))
            ax.imshow(img)
            ax.axis('off')
            ax.set_title(f"Figure {len(self._notebook_figures)}")

        plt.show()
        self._log_status_message("Figure opened in external window", 2000)

    def _on_save_figure(self):
        """Save the last generated figure to a file."""
        from PyQt6.QtWidgets import QFileDialog

        if not hasattr(self, '_notebook_figures') or not self._notebook_figures:
            self._log_status_message("No figures to save. Run code that generates a plot first.", 2000)
            return

        # Get save path from user
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Save Figure",
            "figure.png",
            "PNG Image (*.png);;PDF Document (*.pdf);;SVG Image (*.svg);;All Files (*)"
        )

        if not file_path:
            return

        # Write the figure bytes
        fig_data = self._notebook_figures[-1]
        if 'bytes' in fig_data:
            try:
                with open(file_path, 'wb') as f:
                    f.write(fig_data['bytes'])
                self._log_status_message(f"Figure saved to {file_path}", 2000)
            except Exception as e:
                self._log_status_message(f"Error saving figure: {e}", 3000)
        else:
            self._log_status_message("No figure data to save", 2000)

    def _get_selected_file_types(self) -> list:
        """Get list of file types to scan based on checkbox states."""
        file_types = []
        if hasattr(self, 'scanAbfCheckbox') and self.scanAbfCheckbox.isChecked():
            file_types.append('abf')
        if hasattr(self, 'scanSmrxCheckbox') and self.scanSmrxCheckbox.isChecked():
            file_types.append('smrx')
        if hasattr(self, 'scanEdfCheckbox') and self.scanEdfCheckbox.isChecked():
            file_types.append('edf')
        if hasattr(self, 'scanNotesCheckbox') and self.scanNotesCheckbox.isChecked():
            file_types.append('notes')
        return file_types if file_types else None  # None means scan all types

    def _open_ai_settings(self):
        """Open the AI Settings dialog for configuring AI integration."""
        try:
            from dialogs.ai_settings_dialog import AISettingsDialog

            # With Model/View, get data directly from model or master list
            files_metadata = []
            for task in self._master_file_list:
                metadata = {
                    'file_name': task.get('file_name', ''),
                    'protocol': task.get('protocol', ''),
                    'keywords_display': task.get('keywords_display', ''),
                    'experiment': task.get('experiment', ''),
                    'file_path': str(task.get('file_path', '')),
                }
                files_metadata.append(metadata)

            dialog = AISettingsDialog(self, files_metadata=files_metadata)
            dialog.exec()

        except ImportError as e:
            self._show_warning("AI Module Not Found",
                             f"Could not load AI settings dialog:\n{e}\n\n"
                             "Make sure the dialogs/ai_settings_dialog.py file exists.")
        except Exception as e:
            self._show_error("Error", f"Failed to open AI settings:\n{e}")

    def on_project_scan_saved_data(self, scan_folder: Path = None, silent: bool = False):
        """
        Scan for existing saved data files and auto-populate the table.

        Args:
            scan_folder: If provided, only scan this specific folder (faster).
                        If None, scan all Pleth_App_analysis folders in project.
            silent: If True, don't show progress dialog (for quick rescans after save).
        """
        if not self._master_file_list:
            if not silent:
                self._show_warning("No Files", "Please scan for ABF files first.")
            return

        from PyQt6.QtWidgets import QProgressDialog
        import re

        # Show progress dialog (unless silent mode)
        progress = None
        if not silent:
            progress = QProgressDialog("Scanning for saved data...", "Cancel", 0, 100, self)
            progress.setWindowTitle("Scanning Saved Data")
            progress.setWindowModality(Qt.WindowModality.WindowModal)
            progress.setMinimumDuration(0)
            progress.setValue(0)
            progress.show()
            QApplication.processEvents()

        # Look for Pleth_App_analysis folders
        if scan_folder:
            # Focused scan: just use the provided folder
            analysis_folders = [scan_folder] if scan_folder.exists() else []
            print(f"[scan-saved] Focused scan of: {scan_folder}")
        else:
            # Full scan: search all folders
            base_dir = Path(self._project_directory) if self._project_directory else Path.cwd()
            analysis_folders = list(base_dir.glob("**/Pleth_App_analysis"))

        if not analysis_folders:
            if progress:
                progress.close()
            if not silent:
                self._show_info("No Saved Data", "No 'Pleth_App_analysis' folders found.\n\nAnalyzed data is saved to this folder.")
            return

        # Build a mapping of ABF names to their saved data
        saved_data_map = {}  # abf_stem -> {channel -> export_info}

        if progress:
            progress.setLabelText(f"Scanning {len(analysis_folders)} analysis folders...")
            progress.setValue(10)
            QApplication.processEvents()

        # Collect ALL NPZ files first, then sort by modification time (newest first)
        # This ensures newer files take priority over older ones in subfolders
        all_npz_files = []
        for folder in analysis_folders:
            for npz_file in folder.glob("*_bundle.npz"):
                try:
                    mtime = npz_file.stat().st_mtime
                except:
                    mtime = 0
                all_npz_files.append((npz_file, folder, mtime))

        # Sort by modification time, newest first
        all_npz_files.sort(key=lambda x: x[2], reverse=True)
        print(f"[scan-saved] Found {len(all_npz_files)} NPZ bundle files, processing newest first")

        for npz_file, folder, mtime in all_npz_files:
            # Parse the filename to extract ABF name and channel
            stem = npz_file.stem.replace('_bundle', '')

            # Try to read metadata from NPZ file for reliable source file info
            npz_metadata = None
            source_file_from_npz = None
            channel_from_npz = None

            try:
                import numpy as np
                import json
                with np.load(npz_file, allow_pickle=True) as data:
                    # Bundle NPZ stores meta as JSON string in 'meta_json'
                    if 'meta_json' in data:
                        meta_str = str(data['meta_json'])
                        npz_metadata = json.loads(meta_str)
                        # Get source file and channel from metadata
                        if isinstance(npz_metadata, dict):
                            source_file_from_npz = npz_metadata.get('abf_path', '')
                            channel_from_npz = npz_metadata.get('analyze_channel', '')
                    # ML training NPZ stores source_file directly
                    elif 'source_file' in data:
                        source_file_from_npz = str(data['source_file'])
                        npz_metadata = {'source_file': source_file_from_npz}
            except Exception as e:
                print(f"[scan-saved] Could not read metadata from {npz_file.name}: {e}")

            # Try to match to an ABF file using multiple methods
            # First, try exact metadata match (most reliable) across ALL tasks
            matched_task = None
            match_method = None

            if source_file_from_npz:
                source_stem = Path(source_file_from_npz).stem
                for task in self._master_file_list:
                    if task.get('is_sub_row'):
                        continue  # Only match to parent rows
                    abf_path = Path(task.get('file_path', ''))
                    if abf_path.stem == source_stem:
                        matched_task = task
                        match_method = "NPZ metadata"
                        break

            # Fallback: Pattern match - ABF name as distinct segment (not just substring)
            # Only use if no metadata match found
            if not matched_task:
                for task in self._master_file_list:
                    if task.get('is_sub_row'):
                        continue  # Only match to parent rows
                    abf_path = Path(task.get('file_path', ''))
                    abf_name = abf_path.stem
                    if not abf_name:
                        continue

                    # Check for ABF name as a distinct segment (bounded by _ or start/end)
                    # This prevents "580" matching "580_10mw" when we want "2024_03_21_0018"
                    pattern = r'(^|_)' + re.escape(abf_name) + r'(_|$)'
                    if re.search(pattern, stem):
                        matched_task = task
                        match_method = "pattern"
                        break

            if matched_task:
                abf_name = Path(matched_task.get('file_path', '')).stem
                print(f"[scan-saved] Matched {npz_file.name} to {abf_name} via {match_method}")

                # Found a match - extract more info
                key = str(matched_task.get('file_path'))

                # Determine channel - prefer NPZ metadata, fall back to filename regex
                channel = channel_from_npz or ''
                if not channel:
                    channel_match = re.search(r'_(AD\d+)_', stem) or re.search(r'_(AD\d+)$', stem)
                    channel = channel_match.group(1) if channel_match else ''

                if key not in saved_data_map:
                    saved_data_map[key] = {}

                # Skip if we already have data for this file+channel (since we're processing newest first)
                if channel in saved_data_map[key]:
                    # Track this as an older/duplicate NPZ file
                    if 'older_npz_files' not in saved_data_map[key][channel]:
                        saved_data_map[key][channel]['older_npz_files'] = []
                    saved_data_map[key][channel]['older_npz_files'].append({
                        'file': str(npz_file),
                        'date': datetime.fromtimestamp(mtime).strftime('%Y-%m-%d %H:%M'),
                    })
                    print(f"[scan-saved] Skipping older {npz_file.name} - already have newer data for {abf_name} {channel}")
                    continue  # Skip to next NPZ file

                export_info = {
                    'export_path': str(folder),
                    'export_date': '',
                    'npz': True,
                    'timeseries_csv': (folder / f"{stem}_means_by_time.csv").exists(),
                    'breaths_csv': (folder / f"{stem}_breaths.csv").exists(),
                    'events_csv': (folder / f"{stem}_events.csv").exists(),
                    'pdf': (folder / f"{stem}_summary.pdf").exists(),
                    'session_state': (folder / f"{stem}_session.npz").exists(),
                    'ml_training': False,  # Would need to check ML folder
                    'npz_metadata': npz_metadata,  # Store for later use
                }

                # Use the mtime we already collected during sorting
                from datetime import datetime
                export_info['export_date'] = datetime.fromtimestamp(mtime).strftime('%Y-%m-%d %H:%M')

                # Extract additional metadata from ui_meta if available
                if npz_metadata:
                    ui_meta = npz_metadata.get('ui_meta', {})
                    if isinstance(ui_meta, dict):
                        export_info['strain'] = ui_meta.get('strain', '')
                        export_info['stim_type'] = ui_meta.get('stim', '')  # Dialog saves as 'stim'
                        export_info['power'] = ui_meta.get('power', '')
                        export_info['sex'] = ui_meta.get('sex', '')
                        export_info['animal_id'] = ui_meta.get('animal', '')  # Dialog saves as 'animal', not 'animal_id'
                        # Also get stim and events channels from NPZ metadata (not ui_meta)
                        export_info['stim_channel'] = npz_metadata.get('stim_chan', '')
                        export_info['events_channel'] = npz_metadata.get('event_channel', '')

                saved_data_map[key][channel] = export_info
                print(f"[scan-saved] Using {npz_file.name} (newest) for {abf_name} {channel}")

        if progress:
            progress.setValue(50)
            QApplication.processEvents()

        # Debug: Print what was matched
        print(f"[scan-saved] saved_data_map summary:")
        for file_key, channels in saved_data_map.items():
            file_name = Path(file_key).name
            channel_list = list(channels.keys())
            print(f"  {file_name}: {channel_list}")

        if not saved_data_map:
            if progress:
                progress.close()
            if not silent:
                self._show_info("No Matches", "No saved data files matched the current ABF files.")
            return

        # Update the master list with found saved data
        # We need to handle multiple channels per file by creating sub-rows
        updated_count = 0
        new_rows_created = 0

        # Track which file+channel combinations already exist in the master list
        existing_file_channels = set()
        for task in self._master_file_list:
            fp = str(task.get('file_path', ''))
            ch = task.get('channel', '')
            if fp and ch:
                existing_file_channels.add((fp, ch))

        # First pass: update existing sub-rows that have matching channels
        # Parent rows will have sub-rows created in second pass
        processed_files = set()
        for row, task in enumerate(self._master_file_list):
            file_key = str(task.get('file_path'))
            if file_key not in saved_data_map:
                continue

            existing_channel = task.get('channel', '')
            is_sub_row = task.get('is_sub_row', False)
            export_data = saved_data_map[file_key]

            # Only update if this is a sub-row with a matching channel
            if is_sub_row and existing_channel and existing_channel in export_data:
                info = export_data[existing_channel]
                self._update_task_with_export_info(task, info, row)
                updated_count += 1
                processed_files.add(file_key)
                print(f"[scan-saved] Updated sub-row {row}: {task.get('file_name')} channel {existing_channel}")

            # Parent rows: Don't update them directly - all channels will get sub-rows in second pass
            elif not is_sub_row:
                processed_files.add(file_key)  # Mark as processed so we know to create sub-rows

            # Update progress
            if progress and row % 10 == 0:
                progress.setValue(50 + int(25 * row / len(self._master_file_list)))
                QApplication.processEvents()

        # Second pass: create new sub-rows for additional channels not yet in the list
        # Collect info about what to add (file_key, channel, info) - we'll find row indices at insertion time
        rows_to_add = []  # [(file_key, channel, info), ...]

        for file_key, export_data in saved_data_map.items():
            # Check which channels need new rows for this file
            for channel, info in export_data.items():
                if (file_key, channel) not in existing_file_channels:
                    rows_to_add.append((file_key, channel, info))
                    existing_file_channels.add((file_key, channel))  # Mark as will be added

        # Add new sub-rows for additional channels
        # Find the correct row index at insertion time (not pre-computed) to handle index shifts
        print(f"[scan-saved] Second pass: {len(rows_to_add)} sub-rows to add")
        for file_key, channel, info in rows_to_add:
            # Find the PARENT task (not sub-row) for this file to get full metadata
            parent_task = None
            parent_row = None
            for row, task in enumerate(self._master_file_list):
                if str(task.get('file_path')) == file_key and not task.get('is_sub_row', False):
                    parent_task = task
                    parent_row = row
                    break

            # If no parent found, use any matching task (fallback)
            if not parent_task:
                for row, task in enumerate(self._master_file_list):
                    if str(task.get('file_path')) == file_key:
                        parent_task = task
                        parent_row = row
                        break

            if parent_task:
                print(f"[scan-saved] Adding sub-row for {Path(file_key).name} channel {channel} at row {parent_row}")
                self._create_sub_row_from_saved_data(parent_task, parent_row, channel, info)
                new_rows_created += 1
                updated_count += 1
            else:
                print(f"[scan-saved] WARNING: Could not find source task for {file_key}")

        if progress:
            progress.setValue(90)
            QApplication.processEvents()

        # Rebuild table from updated master list
        self._rebuild_table_from_master_list()

        if progress:
            progress.close()

        # Count warnings
        warnings_count = 0
        conflicts_count = 0
        older_files_count = 0
        for task in self._master_file_list:
            warnings = task.get('scan_warnings', {})
            if warnings:
                warnings_count += 1
                if warnings.get('conflicts'):
                    conflicts_count += 1
                if warnings.get('older_npz_count', 0) > 0:
                    older_files_count += 1

        msg = f"✓ Found saved data for {updated_count} analyses"
        if new_rows_created > 0:
            msg += f" ({new_rows_created} new rows created)"
        if warnings_count > 0:
            warning_details = []
            if conflicts_count > 0:
                warning_details.append(f"{conflicts_count} conflicts")
            if older_files_count > 0:
                warning_details.append(f"{older_files_count} with multiple NPZ files")
            msg += f" ⚠ {', '.join(warning_details)}"
        self._log_status_message(msg, 5000 if warnings_count > 0 else 3000)
        print(f"[project-builder] Found saved data for {updated_count} analyses across {len(analysis_folders)} folders")
        if warnings_count > 0:
            print(f"[project-builder] Warnings: {conflicts_count} conflicts, {older_files_count} with older NPZ files")

        # Enable/disable the Resolve Conflicts button based on warnings
        if hasattr(self, 'resolveConflictsButton'):
            self.resolveConflictsButton.setEnabled(warnings_count > 0)
            if warnings_count > 0:
                self.resolveConflictsButton.setText(f"⚠ Resolve Conflicts ({warnings_count})")
            else:
                self.resolveConflictsButton.setText("⚠ Resolve Conflicts")

    def on_resolve_all_conflicts(self):
        """Show dialog to resolve all conflicts at once."""
        # Find all rows with warnings
        rows_with_warnings = []
        rows_with_conflicts = []
        for i, task in enumerate(self._master_file_list):
            warnings = task.get('scan_warnings', {})
            if warnings:
                rows_with_warnings.append(i)
                if warnings.get('conflicts'):
                    rows_with_conflicts.append(i)

        if not rows_with_warnings:
            self._show_info("No Warnings", "No scan warnings or conflicts to resolve.")
            return

        # Show the conflict details dialog with all warning rows
        self._show_conflict_details(rows_with_warnings)

        # Update button state after resolution
        self._update_resolve_conflicts_button()

    def _update_resolve_conflicts_button(self):
        """Update the Resolve Conflicts button text and enabled state."""
        if not hasattr(self, 'resolveConflictsButton'):
            return

        warnings_count = sum(1 for task in self._master_file_list if task.get('scan_warnings'))

        self.resolveConflictsButton.setEnabled(warnings_count > 0)
        if warnings_count > 0:
            self.resolveConflictsButton.setText(f"⚠ Resolve Conflicts ({warnings_count})")
        else:
            self.resolveConflictsButton.setText("⚠ Resolve Conflicts")

    # NOTE: Old experiment methods removed (on_project_add_experiment, on_project_remove_experiment,
    # on_project_export_experiment). These buttons are hidden and we now use the master file list.

    def on_edit_project_name(self):
        """Edit the current project name (legacy - combo is now directly editable)."""
        from PyQt6.QtWidgets import QInputDialog

        current_name = getattr(self, '_current_project_name', '') or "Untitled Project"

        new_name, ok = QInputDialog.getText(
            self, "Edit Project Name", "Enter project name:",
            text=current_name
        )

        if ok and new_name.strip():
            self._current_project_name = new_name.strip()
            if hasattr(self, 'projectNameCombo'):
                self.projectNameCombo.setCurrentText(new_name.strip())
            self._log_status_message(f"Project renamed to: {new_name.strip()}", 2000)

    def on_project_new(self):
        """Create a new project - prompt for name and directory."""
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QHBoxLayout, QLabel, QLineEdit, QPushButton, QFileDialog

        # Create dialog
        dialog = QDialog(self)
        dialog.setWindowTitle("New Project")
        dialog.setMinimumWidth(500)

        layout = QVBoxLayout(dialog)

        # Project name
        layout.addWidget(QLabel("Project Name:"))
        name_edit = QLineEdit()
        name_edit.setPlaceholderText("Enter project name...")
        layout.addWidget(name_edit)

        layout.addSpacing(10)

        # Directory selection
        layout.addWidget(QLabel("Data Directory:"))
        dir_layout = QHBoxLayout()
        dir_edit = QLineEdit()
        dir_edit.setPlaceholderText("Select directory containing data files...")
        dir_edit.setReadOnly(True)
        dir_layout.addWidget(dir_edit)

        browse_btn = QPushButton("Browse...")
        browse_btn.setMaximumWidth(100)
        def browse_dir():
            directory = QFileDialog.getExistingDirectory(dialog, "Select Data Directory")
            if directory:
                dir_edit.setText(directory)
                # Auto-fill project name if empty
                if not name_edit.text().strip():
                    name_edit.setText(Path(directory).name)
        browse_btn.clicked.connect(browse_dir)
        dir_layout.addWidget(browse_btn)
        layout.addLayout(dir_layout)

        layout.addSpacing(20)

        # Buttons
        button_layout = QHBoxLayout()
        button_layout.addStretch()

        cancel_btn = QPushButton("Cancel")
        cancel_btn.clicked.connect(dialog.reject)
        button_layout.addWidget(cancel_btn)

        create_btn = QPushButton("Create Project")
        create_btn.setDefault(True)
        create_btn.clicked.connect(dialog.accept)
        button_layout.addWidget(create_btn)

        layout.addLayout(button_layout)

        # Show dialog
        if dialog.exec() == QDialog.DialogCode.Accepted:
            project_name = name_edit.text().strip()
            directory = dir_edit.text().strip()

            if not project_name:
                self._show_warning("Missing Information", "Please enter a project name.")
                return

            if not directory:
                self._show_warning("Missing Information", "Please select a data directory.")
                return

            # Clear everything
            self._file_table_model.clear()
            self._discovered_files_data = []
            self._master_file_list = []

            # Set new project info
            self._project_directory = directory
            self.directoryPathEdit.setText(directory)
            self._current_project_name = project_name
            if hasattr(self, 'projectNameCombo'):
                self.projectNameCombo.setCurrentText(project_name)
            self.summaryLabel.setText("Summary: No files scanned")

            self._log_status_message(f"✓ New project created: {project_name}", 2000)

    def _cancel_pending_autosave(self):
        """Cancel any pending autosave operation."""
        if hasattr(self, '_autosave_timer') and self._autosave_timer:
            self._autosave_timer.stop()
            print("[autosave] Cancelled pending autosave")

    def _project_autosave(self):
        """Silently autosave the project if it has a name and directory."""
        if not self._project_directory or not self._master_file_list:
            return

        project_name = getattr(self, '_current_project_name', '') or ''
        if not project_name.strip():
            return  # Don't autosave unnamed projects

        try:
            # Debounce autosave - use a timer to avoid saving on every keystroke
            if not hasattr(self, '_autosave_timer'):
                self._autosave_timer = QTimer()
                self._autosave_timer.setSingleShot(True)
                self._autosave_timer.timeout.connect(self._do_autosave)

            # Reset timer - will save 1 second after last change
            self._autosave_timer.start(1000)
        except Exception as e:
            print(f"[autosave] Error setting up autosave: {e}")

    def _do_autosave(self):
        """Actually perform the autosave."""
        if not self._project_directory or not self._master_file_list:
            return

        project_name = getattr(self, '_current_project_name', '') or ''
        if not project_name.strip():
            return

        try:
            self.project_manager.save_project(
                project_name,
                Path(self._project_directory),
                self._master_file_list,
                [],  # No experiments in new workflow
                notes_directory=self._notes_directory,
                notes_files=self._notes_files_data
            )
            print(f"[autosave] Project autosaved: {project_name}")
        except Exception as e:
            print(f"[autosave] Error: {e}")

    def on_project_save(self):
        """Save current project to data directory."""
        if not self._project_directory:
            self._show_warning("No Directory", "Please select a directory first.")
            return

        if not self._discovered_files_data:
            self._show_warning("No Files", "Please scan for files first.")
            return

        # Get project name
        project_name = getattr(self, '_current_project_name', '') or ''
        if not project_name.strip():
            # Prompt for project name
            from PyQt6.QtWidgets import QInputDialog
            project_name, ok = QInputDialog.getText(
                self, "Save Project", "Enter project name:",
                text=Path(self._project_directory).name
            )
            if not ok or not project_name.strip():
                return
            project_name = project_name.strip()
            self._current_project_name = project_name
            if hasattr(self, 'projectNameCombo'):
                self.projectNameCombo.setCurrentText(project_name)

        try:
            # Use master file list as the source of truth
            # (it contains all metadata edits the user made)
            files_to_save = self._master_file_list if self._master_file_list else self._discovered_files_data
            print(f"[project-save] Saving {len(files_to_save)} files from master list")

            # Save project (no experiments - using flat file list now)
            project_path = self.project_manager.save_project(
                project_name,
                Path(self._project_directory),
                files_to_save,
                [],  # No experiments in new workflow
                notes_directory=self._notes_directory,
                notes_files=self._notes_files_data
            )

            self._show_info("Project Saved", f"Project saved to:\n{project_path}")
            self._log_status_message(f"✓ Project saved: {project_name}", 3000)

            # Update recent projects dropdown
            self._populate_load_project_combo()

        except Exception as e:
            self._show_error("Save Failed", f"Failed to save project:\n{e}")
            print(f"[project] Error saving: {e}")
            import traceback
            traceback.print_exc()

    def on_project_load(self, index):
        """Load a project from recent projects list (legacy - now handled by _on_project_combo_selected)."""
        # This method is kept for backward compatibility but is no longer connected
        # Project loading is now handled by _on_project_combo_selected via projectNameCombo
        pass

    def _apply_recovered_metadata(self):
        """Apply recovered metadata from corrupted file to rescanned files. Returns count of applied fields."""
        if not hasattr(self, '_pending_recovery_metadata') or not self._pending_recovery_metadata:
            return 0

        recovered = self._pending_recovery_metadata
        applied_count = 0
        fields_applied = 0

        for task in self._master_file_list:
            file_path = task.get('file_path')
            if not file_path:
                continue

            # Try to match by filename or relative path
            file_name = Path(file_path).name

            # Look for a match in recovered metadata
            matched_meta = None
            for recovered_path, meta in recovered.items():
                if recovered_path.endswith(file_name) or file_name in recovered_path:
                    matched_meta = meta
                    break

            if matched_meta:
                file_updated = False
                # Apply recovered metadata (only if not already set)
                for key in ['strain', 'animal_id', 'power', 'sex', 'channel', 'channels']:
                    if matched_meta.get(key) and not task.get(key):
                        task[key] = matched_meta[key]
                        fields_applied += 1
                        file_updated = True
                if file_updated:
                    applied_count += 1

        # Clear pending metadata
        self._pending_recovery_metadata = None

        if applied_count > 0:
            # Refresh the table to show recovered data
            self._rebuild_table_from_master_list()
            self._log_status_message(f"✓ Recovered metadata for {applied_count} files", 3000)
            print(f"[recovery] Applied {fields_applied} metadata fields to {applied_count} files")

        return fields_applied

    def _apply_recovered_metadata_and_show_summary(self):
        """Apply recovered metadata and show summary dialog."""
        fields_applied = self._apply_recovered_metadata()

        # Get recovery stats
        stats = getattr(self, '_pending_recovery_stats', {
            'method': 'rescan',
            'files_recovered': len(self._master_file_list),
            'metadata_fields_recovered': fields_applied,
            'ai_tokens_used': 0,
            'errors': [],
        })
        stats['metadata_fields_recovered'] = fields_applied
        stats['files_recovered'] = len(self._master_file_list)

        # Get project name
        project_name = getattr(self, '_current_project_name', '') or "Unknown"

        # Clear pending stats
        self._pending_recovery_stats = None

        # Show summary
        self._show_recovery_summary(stats, project_name)

    def _show_recovery_summary(self, stats: dict, project_name: str):
        """Show a summary dialog after project recovery."""
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QLabel, QTextEdit, QPushButton, QHBoxLayout

        dialog = QDialog(self)
        dialog.setWindowTitle("Recovery Complete")
        dialog.setMinimumSize(450, 300)
        layout = QVBoxLayout(dialog)

        # Header
        header = QLabel(f"<b>Project '{project_name}' recovered successfully!</b>")
        layout.addWidget(header)

        # Build summary text
        summary_lines = []

        # Method used
        method_names = {
            'backup_merge': 'Backup restore with change merge',
            'rescan': 'Directory rescan with metadata recovery',
            'unknown': 'Recovery',
        }
        summary_lines.append(f"<b>Recovery method:</b> {method_names.get(stats.get('method', 'unknown'), 'Unknown')}")
        summary_lines.append("")

        # Stats
        summary_lines.append("<b>Results:</b>")
        summary_lines.append(f"  • Files in project: {stats.get('files_recovered', 0)}")
        summary_lines.append(f"  • Metadata fields recovered: {stats.get('metadata_fields_recovered', 0)}")

        if stats.get('ai_tokens_used', 0) > 0:
            summary_lines.append(f"  • AI tokens used: {stats['ai_tokens_used']}")

        # Errors
        if stats.get('errors'):
            summary_lines.append("")
            summary_lines.append("<b>Warnings/Errors:</b>")
            for error in stats['errors']:
                summary_lines.append(f"  ⚠ {error}")

        # What was lost (if rescan method)
        if stats.get('method') == 'rescan':
            summary_lines.append("")
            summary_lines.append("<b>Note:</b> Recovery via rescan may have lost some metadata")
            summary_lines.append("(channels, custom fields) that wasn't in the corrupted portion.")
            summary_lines.append("Check your project and re-enter any missing information.")

        summary_text = QTextEdit()
        summary_text.setReadOnly(True)
        summary_text.setHtml("<br>".join(summary_lines))
        layout.addWidget(summary_text)

        # OK button
        button_layout = QHBoxLayout()
        button_layout.addStretch()
        ok_btn = QPushButton("OK")
        ok_btn.clicked.connect(dialog.accept)
        button_layout.addWidget(ok_btn)
        layout.addLayout(button_layout)

        dialog.exec()

    def _get_configured_ai_client(self):
        """Get AI client if configured, or None if not set up."""
        try:
            settings = QSettings("PhysioMetrics", "BreathAnalysis")
            provider = settings.value("ai/provider", "claude")
            api_key = settings.value(f"ai/{provider}_api_key", "")

            if not api_key:
                return None

            from core.ai_client import AIClient
            return AIClient(provider=provider, api_key=api_key)
        except Exception as e:
            print(f"[ai] Could not initialize AI client: {e}")
            return None

    def _extract_metadata_from_corrupted(self, project_path: Path) -> dict:
        """Extract whatever metadata we can from a corrupted JSON file."""
        try:
            with open(project_path, 'r', encoding='utf-8', errors='ignore') as f:
                corrupted_content = f.read()

            import re
            recovered_metadata = {}

            # First, try to find file_path entries
            file_path_pattern = r'"file_path"\s*:\s*"([^"]+)"'
            file_paths = re.findall(file_path_pattern, corrupted_content)

            for fp in set(file_paths):  # dedupe
                if fp:
                    recovered_metadata[fp] = {}

            # Now extract individual fields for each file
            # We search for each field near its file_path
            field_patterns = {
                'strain': r'"strain"\s*:\s*"([^"]*)"',
                'animal_id': r'"animal_id"\s*:\s*"([^"]*)"',
                'power': r'"power"\s*:\s*"([^"]*)"',
                'sex': r'"sex"\s*:\s*"([^"]*)"',
                'channel': r'"channel"\s*:\s*"([^"]*)"',
                'channels': r'"channels"\s*:\s*"([^"]*)"',
                'notes': r'"notes"\s*:\s*"([^"]*)"',
            }

            # Find each file entry block and extract fields
            # Look for blocks that start with file_path
            entry_pattern = r'\{\s*"file_path"\s*:\s*"([^"]+)"([^}]*)\}'
            entry_matches = re.findall(entry_pattern, corrupted_content, re.DOTALL)

            for file_path, entry_content in entry_matches:
                if file_path not in recovered_metadata:
                    recovered_metadata[file_path] = {}

                # Extract each field from this entry's content
                for field_name, field_pattern in field_patterns.items():
                    match = re.search(field_pattern, entry_content)
                    if match and match.group(1):
                        recovered_metadata[file_path][field_name] = match.group(1)

            # Count how many entries have actual metadata
            entries_with_data = sum(1 for meta in recovered_metadata.values() if any(meta.values()))

            print(f"[recovery] Extracted {len(recovered_metadata)} file entries ({entries_with_data} with metadata)")
            return recovered_metadata

        except Exception as e:
            print(f"[recovery] Failed to extract data: {e}")
            return {}

    def _perform_project_recovery(self, project_path: Path, project_name: str,
                                   data_directory: Path, recovered_metadata: dict,
                                   use_ai: bool = False):
        """Perform the actual project recovery with optional AI enhancement."""

        backup_path = project_path.with_suffix('.physiometrics.bak')
        backup_data = None

        # First, try to load the backup file as our baseline
        if backup_path.exists():
            try:
                with open(backup_path, 'r') as f:
                    backup_data = json.load(f)
                print(f"[recovery] Loaded backup file with {len(backup_data.get('files', []))} files")
                self._log_status_message("Found valid backup - comparing with corrupted file...", 2000)
            except Exception as e:
                print(f"[recovery] Backup file also corrupted: {e}")
                backup_data = None

        # If we have both backup and AI, do smart diff-based recovery
        if use_ai and backup_data and recovered_metadata:
            self._log_status_message("Using AI to recover changes since last save...", 3000)
            ai_client = self._get_configured_ai_client()
            if ai_client:
                try:
                    # Build a map of backup file metadata for comparison
                    backup_files = {f.get('file_path', ''): f for f in backup_data.get('files', [])}

                    # Find what's different in corrupted vs backup
                    changes = []
                    for file_path, meta in recovered_metadata.items():
                        backup_entry = backup_files.get(file_path, {})
                        # Check if this file has different metadata than backup
                        if meta != {k: backup_entry.get(k, '') for k in ['strain', 'animal_id', 'power', 'sex']}:
                            changes.append({
                                'file': file_path,
                                'backup_had': {k: backup_entry.get(k, '') for k in ['strain', 'animal_id', 'power', 'sex'] if backup_entry.get(k)},
                                'corrupted_has': {k: v for k, v in meta.items() if v}
                            })

                    if changes:
                        # Read the corrupted section for context
                        try:
                            with open(project_path, 'r', errors='ignore') as f:
                                corrupted_content = f.read()
                            # Find where corruption likely occurred (truncated JSON)
                            last_complete = corrupted_content.rfind('},')
                            corrupted_tail = corrupted_content[last_complete:] if last_complete > 0 else corrupted_content[-500:]
                        except:
                            corrupted_tail = "(could not read)"

                        prompt = f"""A JSON project file was corrupted during save. I have:
1. A valid backup from before the corruption
2. The corrupted file with partial data

The user was editing {len(changes)} file entries when corruption occurred.

Files that changed (comparing backup vs corrupted):
{json.dumps(changes[:10], indent=2)}  # Limit to 10 for token efficiency

The corrupted section ends with:
{corrupted_tail[:500]}

Based on the partial data in 'corrupted_has', what were the likely intended values?
Return ONLY a JSON object mapping file_path to the corrected metadata fields.
Only include fields you can confidently infer from the partial data."""

                        response = ai_client.complete(prompt)

                        # Parse AI suggestions
                        try:
                            content = response.content
                            json_start = content.find('{')
                            json_end = content.rfind('}') + 1
                            if json_start >= 0 and json_end > json_start:
                                ai_fixes = json.loads(content[json_start:json_end])
                                ai_fixed = 0
                                for fp, meta in ai_fixes.items():
                                    if fp in recovered_metadata:
                                        for key, value in meta.items():
                                            if value:
                                                recovered_metadata[fp][key] = value
                                                ai_fixed += 1
                                print(f"[recovery] AI repaired {ai_fixed} fields from diff analysis")
                        except json.JSONDecodeError:
                            print("[recovery] Could not parse AI diff-repair response")
                    else:
                        print("[recovery] No changes detected between backup and corrupted file")

                except Exception as e:
                    print(f"[recovery] AI diff-based recovery failed: {e}")

        # If we have backup data but no AI, merge backup metadata into recovered
        elif backup_data and not use_ai:
            backup_files = {f.get('file_path', ''): f for f in backup_data.get('files', [])}
            merged_count = 0
            for file_path in recovered_metadata:
                if file_path in backup_files:
                    backup_entry = backup_files[file_path]
                    # Use backup values for any missing fields
                    for key in ['strain', 'animal_id', 'power', 'sex']:
                        if backup_entry.get(key) and not recovered_metadata[file_path].get(key):
                            recovered_metadata[file_path][key] = backup_entry[key]
                            merged_count += 1
            if merged_count > 0:
                print(f"[recovery] Merged {merged_count} fields from backup file")

        # Fallback: AI without backup - just analyze file paths for patterns
        elif use_ai and recovered_metadata and not backup_data:
            self._log_status_message("No backup available - AI analyzing file patterns...", 3000)
            ai_client = self._get_configured_ai_client()
            if ai_client:
                try:
                    # Only send file paths (not full metadata) for efficiency
                    file_paths = list(recovered_metadata.keys())[:20]  # Limit for tokens
                    prompt = f"""Analyze these file paths from a neuroscience experiment and suggest metadata.

File paths:
{json.dumps(file_paths, indent=2)}

Look for patterns like:
- Animal IDs (numbers like 25729)
- Power levels (e.g., "10mW")
- Strain names (e.g., "VgatCre", "C57")

Return ONLY a JSON object mapping file_path to suggested {{strain, animal_id, power}} fields.
Only include fields where you're confident."""

                    response = ai_client.complete(prompt)
                    content = response.content
                    json_start = content.find('{')
                    json_end = content.rfind('}') + 1
                    if json_start >= 0 and json_end > json_start:
                        ai_suggestions = json.loads(content[json_start:json_end])
                        for fp, meta in ai_suggestions.items():
                            if fp in recovered_metadata:
                                for key, value in meta.items():
                                    if value and not recovered_metadata[fp].get(key):
                                        recovered_metadata[fp][key] = value
                except Exception as e:
                    print(f"[recovery] AI pattern analysis failed: {e}")

        # Track recovery stats for summary
        recovery_stats = {
            'method': 'unknown',
            'files_recovered': 0,
            'metadata_fields_recovered': 0,
            'ai_tokens_used': getattr(self, '_last_ai_tokens', 0),
            'errors': [],
        }

        # STRATEGY 1: If we have valid backup data, use it directly (preserves ALL metadata)
        if backup_data:
            self._log_status_message("Restoring from backup with recovered changes...", 2000)
            recovery_stats['method'] = 'backup_merge'

            try:
                # Start with backup data as our base
                # Deep copy to avoid modifying original
                import copy
                project_data = copy.deepcopy(backup_data)

                # Build filename-based map for matching (backup has relative paths, recovered has absolute)
                # Map filename -> index in project_data['files']
                backup_by_filename = {}
                for i, f in enumerate(project_data.get('files', [])):
                    fp = f.get('file_path', '')
                    if fp:
                        filename = Path(fp).name
                        backup_by_filename[filename] = i

                # Merge any recovered metadata changes from corrupted file
                # recovered_metadata has absolute paths as keys
                for recovered_path, meta in recovered_metadata.items():
                    filename = Path(recovered_path).name
                    if filename in backup_by_filename:
                        idx = backup_by_filename[filename]
                        # Only update fields that have values in recovered data
                        # AND are different from backup (these are the user's recent changes)
                        for key, value in meta.items():
                            if value:  # Only apply non-empty values
                                old_value = project_data['files'][idx].get(key, '')
                                if value != old_value:
                                    project_data['files'][idx][key] = value
                                    recovery_stats['metadata_fields_recovered'] += 1
                                    print(f"[recovery] Merged change: {filename}.{key} = '{value}'")

                recovery_stats['files_recovered'] = len(project_data.get('files', []))

                # Convert relative paths back to absolute for save_project
                # (save_project expects absolute paths and converts them to relative)
                for file_entry in project_data.get('files', []):
                    fp = file_entry.get('file_path', '')
                    if fp and not Path(fp).is_absolute():
                        abs_path = (data_directory / fp).resolve()
                        file_entry['file_path'] = abs_path

                # Save the merged data
                self.project_manager.save_project(
                    project_name=project_name,
                    data_directory=data_directory,
                    files_data=project_data.get('files', []),
                    experiments=project_data.get('experiments', []),
                    notes_directory=project_data.get('notes_directory'),
                    notes_files=project_data.get('notes_files', [])
                )

                # Load the recovered project
                saved_data = self.project_manager.load_project(project_path)
                self._populate_ui_from_project(saved_data)

                print(f"[recovery] Restored from backup with {recovery_stats['metadata_fields_recovered']} field updates")
                print(f"[recovery] Total files in project: {recovery_stats['files_recovered']}")

            except Exception as e:
                recovery_stats['errors'].append(f"Backup merge failed: {e}")
                print(f"[recovery] Backup merge failed, falling back to rescan: {e}")
                import traceback
                traceback.print_exc()
                backup_data = None  # Fall through to rescan

        # STRATEGY 2: No valid backup - must rescan and apply recovered metadata
        if not backup_data:
            recovery_stats['method'] = 'rescan'

            # Delete the corrupted file
            try:
                project_path.unlink()
                print(f"[recovery] Deleted corrupted file: {project_path}")
            except Exception as e:
                print(f"[recovery] Could not delete corrupted file: {e}")
                recovery_stats['errors'].append(f"Could not delete corrupted file: {e}")

            # Set up the project
            self._current_project_name = project_name
            if hasattr(self, 'projectNameCombo'):
                self.projectNameCombo.setCurrentText(project_name)
            self._project_directory = str(data_directory)
            self.directoryPathEdit.setText(self._project_directory)
            self._master_file_list = []
            self._discovered_files_data = []

            # Store recovered metadata and stats to apply after rescan
            self._pending_recovery_metadata = recovered_metadata
            self._pending_recovery_stats = recovery_stats

            # Trigger rescan
            self._log_status_message(f"Rescanning '{project_name}'...", 3000)
            self.on_project_scan_files()

            # Apply recovered metadata after scan completes (with callback for summary)
            if recovered_metadata:
                QTimer.singleShot(3000, self._apply_recovered_metadata_and_show_summary)
            return  # Summary will be shown after rescan completes

        # Show recovery summary for backup-based recovery
        self._show_recovery_summary(recovery_stats, project_name)

    def _get_text_diff(self, backup_path: Path, corrupted_path: Path) -> dict:
        """Get a proper text diff between backup and corrupted file."""
        import difflib

        result = {
            'has_diff': False,
            'diff_lines': [],
            'diff_summary': '',
            'additions': [],      # Lines added in corrupted
            'deletions': [],      # Lines removed from backup
            'backup_content': '',
            'corrupted_content': '',
        }

        try:
            with open(backup_path, 'r', encoding='utf-8', errors='ignore') as f:
                backup_lines = f.readlines()
                result['backup_content'] = ''.join(backup_lines)

            with open(corrupted_path, 'r', encoding='utf-8', errors='ignore') as f:
                corrupted_lines = f.readlines()
                result['corrupted_content'] = ''.join(corrupted_lines)

            # Generate unified diff
            diff = list(difflib.unified_diff(
                backup_lines, corrupted_lines,
                fromfile='backup', tofile='corrupted',
                lineterm=''
            ))

            result['diff_lines'] = diff
            result['has_diff'] = len(diff) > 0

            # Categorize changes
            for line in diff:
                if line.startswith('+') and not line.startswith('+++'):
                    result['additions'].append(line[1:].strip())
                elif line.startswith('-') and not line.startswith('---'):
                    result['deletions'].append(line[1:].strip())

            # Create human-readable summary
            result['diff_summary'] = '\n'.join(diff[:50])  # First 50 lines of diff

            print(f"[diff] Found {len(result['additions'])} additions, {len(result['deletions'])} deletions")

        except Exception as e:
            print(f"[diff] Error comparing files: {e}")

        return result

    def _ai_smart_repair(self, backup_path: Path, corrupted_path: Path,
                         diff_info: dict, error_msg: str) -> tuple:
        """
        Use AI to intelligently repair a corrupted JSON file.

        Strategy: Start with backup (valid JSON), have AI identify valid new data
        from the diff, then merge those changes into the backup.

        Returns (repaired_content, explanation) or (None, error_message).
        """
        ai_client = self._get_configured_ai_client()
        if not ai_client:
            return None, "AI not configured"

        # Load the backup as our base (it's valid JSON)
        try:
            backup_data = json.loads(diff_info['backup_content'])
        except:
            return None, "Backup file is not valid JSON"

        # Extract just the additions from the diff (new data in corrupted file)
        additions = diff_info.get('additions', [])
        if not additions:
            return None, "No additions found in diff"

        # Format additions for AI analysis
        additions_text = '\n'.join(additions[:100])  # Limit for tokens

        prompt = f"""Analyze changes made to a JSON project file that caused corruption.

## JSON Parse Error:
{error_msg}

## Lines ADDED to the file (these are the changes since last backup):
```
{additions_text}
```

## Your Task:
These additions contain BOTH:
1. **VALID NEW DATA** - intentional edits like strain="VgatCre", animal_id="25729", etc.
2. **CORRUPTION** - accidental changes that broke the JSON (extra newlines, broken paths, typos)

Analyze each addition and categorize it:
- VALID: Looks like intentional metadata (field values, new entries)
- CORRUPT: Looks like an error (broken syntax, truncated text, random characters)

## Output Format:
Return a JSON object with this structure:
```json
{{
  "analysis": "Brief explanation of what you found",
  "valid_changes": [
    {{"file_pattern": "filename or pattern", "field": "strain", "value": "VgatCre"}},
    {{"file_pattern": "filename", "field": "animal_id", "value": "25729"}}
  ],
  "corrupted_changes": [
    {{"description": "what was corrupted and why"}}
  ]
}}
```

Only include changes you're confident about. For valid_changes, use file_pattern to match which file entry should be updated (can be partial filename)."""

        try:
            self._log_status_message("AI analyzing changes...", 5000)
            response = ai_client.complete(prompt)
            content = response.content.strip()

            # Extract JSON from response
            json_content = None
            if '```json' in content:
                start = content.find('```json') + 7
                end = content.rfind('```')
                json_content = content[start:end].strip()
            elif '```' in content:
                start = content.find('```') + 3
                end = content.rfind('```')
                json_content = content[start:end].strip()
            else:
                json_start = content.find('{')
                if json_start >= 0:
                    json_content = content[json_start:]

            if not json_content:
                return None, "AI did not return analysis"

            # Parse AI's analysis
            try:
                ai_analysis = json.loads(json_content)
            except json.JSONDecodeError as e:
                print(f"[ai-repair] Could not parse AI analysis: {e}")
                return None, f"Could not parse AI analysis: {e}"

            explanation = ai_analysis.get('analysis', 'AI analyzed the changes')
            valid_changes = ai_analysis.get('valid_changes', [])
            corrupted_changes = ai_analysis.get('corrupted_changes', [])

            if not valid_changes:
                explanation += "\n\nNo valid changes identified - restoring from backup only."
                repaired_json = json.dumps(backup_data, indent=2)
                return repaired_json, explanation

            # Record original file count to ensure we don't accidentally change structure
            original_file_count = len(backup_data.get('files', []))

            # Group changes by file pattern to apply consistently to multi-channel files
            changes_by_pattern = {}
            for change in valid_changes:
                file_pattern = change.get('file_pattern', '')
                field = change.get('field', '')
                value = change.get('value', '')
                channel = change.get('channel', None)  # Optional channel specifier

                if not file_pattern or not field:
                    continue

                if file_pattern not in changes_by_pattern:
                    changes_by_pattern[file_pattern] = []
                changes_by_pattern[file_pattern].append({
                    'field': field,
                    'value': value,
                    'channel': channel
                })

            # Apply changes to backup data
            applied_count = 0
            files_updated = set()

            for file_pattern, changes in changes_by_pattern.items():
                # Find ALL matching file entries (handles multi-channel files)
                matching_indices = []
                for i, file_entry in enumerate(backup_data.get('files', [])):
                    file_path = file_entry.get('file_path', '')
                    filename = Path(file_path).name.lower()
                    pattern_lower = file_pattern.lower()

                    # Match by filename (more precise than substring)
                    if filename == pattern_lower or pattern_lower in filename:
                        matching_indices.append(i)

                if not matching_indices:
                    print(f"[ai-repair] No match found for pattern: {file_pattern}")
                    continue

                # Apply changes to matching entries
                for change in changes:
                    field = change['field']
                    value = change['value']
                    channel = change.get('channel')

                    for idx in matching_indices:
                        file_entry = backup_data['files'][idx]
                        entry_channel = file_entry.get('channel', file_entry.get('channels', ''))

                        # If channel specified, only update matching channel
                        if channel and str(channel) not in str(entry_channel):
                            continue

                        # Apply the change
                        file_entry[field] = value
                        applied_count += 1
                        files_updated.add(Path(file_entry.get('file_path', '')).name)
                        print(f"[ai-repair] Applied: {Path(file_entry.get('file_path', '')).name}[{entry_channel}].{field} = '{value}'")

            # Verify we haven't changed the file structure
            final_file_count = len(backup_data.get('files', []))
            if final_file_count != original_file_count:
                print(f"[ai-repair] WARNING: File count changed from {original_file_count} to {final_file_count}")
                return None, f"Repair would change file count from {original_file_count} to {final_file_count} - aborting"

            explanation += f"\n\nApplied {applied_count} changes to {len(files_updated)} files."
            if corrupted_changes:
                explanation += f"\nIgnored {len(corrupted_changes)} corrupted changes."
            explanation += f"\nFile structure preserved ({original_file_count} entries)."

            # Generate repaired JSON
            repaired_json = json.dumps(backup_data, indent=2)

            # Validate it
            try:
                json.loads(repaired_json)
                file_count = len(backup_data.get('files', []))
                print(f"[ai-repair] Produced valid JSON with {file_count} files, {applied_count} changes applied")
                return repaired_json, explanation
            except json.JSONDecodeError as e:
                return None, f"Generated invalid JSON: {e}"

        except Exception as e:
            print(f"[ai-repair] AI repair failed: {e}")
            import traceback
            traceback.print_exc()
            return None, f"AI repair failed: {e}"

    def _analyze_corruption_cause(self, project_path: Path, error_msg: str) -> dict:
        """Analyze the corrupted file to determine the likely cause of corruption."""
        result = {
            'type': 'unknown',
            'description': 'Unknown corruption',
            'likely_cause': 'Unable to determine cause',
            'corrupted_section': '',
        }

        try:
            with open(project_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()

            file_size = len(content)

            # Check for truncation (file ends abruptly)
            content_stripped = content.rstrip()
            if not content_stripped.endswith('}'):
                result['type'] = 'truncation'
                result['description'] = 'File appears truncated (incomplete JSON)'
                result['likely_cause'] = 'Save operation was interrupted (app crash, power loss, or switching projects during autosave)'
                # Get the last 200 chars to show where it cut off
                result['corrupted_section'] = content_stripped[-200:] if len(content_stripped) > 200 else content_stripped

            # Check for invalid control characters
            elif 'invalid control character' in error_msg.lower():
                result['type'] = 'invalid_chars'
                result['description'] = 'File contains invalid control characters'
                result['likely_cause'] = 'Binary data or encoding issue during save (possibly disk error or memory corruption)'
                # Find the problematic area
                import re
                bad_chars = re.findall(r'[\x00-\x1f]', content)
                if bad_chars:
                    result['corrupted_section'] = f"Found {len(bad_chars)} invalid characters"

            # Check for missing brackets/braces
            elif content.count('{') != content.count('}'):
                result['type'] = 'unbalanced_braces'
                open_count = content.count('{')
                close_count = content.count('}')
                result['description'] = f'Unbalanced braces ({{ {open_count} vs }} {close_count})'
                result['likely_cause'] = 'Partial write - save was interrupted before completion'
                result['corrupted_section'] = content[-200:] if len(content) > 200 else content

            # Check for duplicate keys or malformed structure
            elif 'expecting' in error_msg.lower():
                result['type'] = 'malformed_json'
                result['description'] = 'Malformed JSON structure'
                result['likely_cause'] = 'Data serialization error or concurrent write conflict'
                # Extract line number from error if available
                import re
                line_match = re.search(r'line (\d+)', error_msg)
                if line_match:
                    line_num = int(line_match.group(1))
                    lines = content.split('\n')
                    if line_num <= len(lines):
                        start = max(0, line_num - 3)
                        end = min(len(lines), line_num + 2)
                        result['corrupted_section'] = '\n'.join(f"{i+1}: {lines[i]}" for i in range(start, end))

            else:
                # Generic corruption
                result['type'] = 'generic'
                result['description'] = 'JSON parsing error'
                result['likely_cause'] = 'File was modified or corrupted'
                result['corrupted_section'] = content[-200:] if len(content) > 200 else content

        except Exception as e:
            result['likely_cause'] = f'Could not analyze file: {e}'

        return result

    def _handle_corrupted_project(self, project_path: Path, project_name: str, error_msg: str):
        """Handle a corrupted project file with a recovery preview dialog."""
        from PyQt6.QtWidgets import QMessageBox, QDialog, QVBoxLayout, QLabel, QTextEdit, QPushButton, QHBoxLayout, QGroupBox

        data_directory = project_path.parent
        backup_path = project_path.with_suffix('.physiometrics.bak')
        has_backup = backup_path.exists()

        # First, analyze what can be recovered
        self._log_status_message("Analyzing corrupted file...", 2000)

        # Analyze the corruption cause
        corruption_info = self._analyze_corruption_cause(project_path, error_msg)

        # Get proper text diff if backup exists
        diff_info = None
        if has_backup:
            diff_info = self._get_text_diff(backup_path, project_path)

        # Extract metadata from corrupted file
        recovered_metadata = self._extract_metadata_from_corrupted(project_path)

        # Check if AI is configured
        has_ai = self._get_configured_ai_client() is not None

        # Build recovery preview dialog
        dialog = QDialog(self)
        dialog.setWindowTitle("Project Recovery")
        dialog.setMinimumSize(600, 500)
        layout = QVBoxLayout(dialog)

        # Header
        header = QLabel(f"<b>The project '{project_name}' is corrupted.</b>")
        layout.addWidget(header)

        # Check if backup is valid (can be loaded as JSON)
        backup_valid = False
        backup_file_count = 0
        backup_data = None
        if has_backup:
            try:
                with open(backup_path, 'r') as f:
                    backup_data = json.load(f)
                backup_valid = True
                backup_file_count = len(backup_data.get('files', []))
            except:
                backup_valid = False

        # Show corruption diagnosis
        summary_parts = []
        summary_parts.append(f"<b>Corruption type:</b> {corruption_info['description']}")
        summary_parts.append(f"<b>Likely cause:</b> {corruption_info['likely_cause']}")
        summary_parts.append("")

        # Recovery options summary
        if has_backup and backup_valid:
            summary_parts.append(f"✓ Valid backup found ({backup_file_count} files)")
        elif has_backup:
            summary_parts.append("⚠ Backup file exists but is also corrupted")
        else:
            summary_parts.append("✗ No backup file found")

        summary_parts.append(f"✓ Found {len(recovered_metadata)} file entries in corrupted data")

        if has_ai:
            if backup_valid:
                summary_parts.append("✓ AI configured (will compare backup vs corrupted to find changes)")
            else:
                summary_parts.append("✓ AI configured (will analyze file patterns)")
        else:
            summary_parts.append("○ AI not configured (optional)")

        # Explain recovery strategy
        summary_parts.append("")
        summary_parts.append("<b>Recovery strategy:</b>")
        if backup_valid and has_ai:
            summary_parts.append("1. Load backup as baseline")
            summary_parts.append("2. Find changes in corrupted file")
            summary_parts.append("3. AI repairs corrupted changes (token-efficient)")
            summary_parts.append("4. Merge with backup data")
        elif backup_valid:
            summary_parts.append("1. Load backup as baseline")
            summary_parts.append("2. Merge any recoverable changes from corrupted file")
        elif has_ai:
            summary_parts.append("1. Rescan directory for files")
            summary_parts.append("2. AI analyzes file paths for metadata patterns")
        else:
            summary_parts.append("1. Rescan directory for files")
            summary_parts.append("2. Apply any metadata found in corrupted file")

        summary_label = QLabel("\n".join(summary_parts))
        layout.addWidget(summary_label)

        # Show REAL text diff between backup and corrupted
        if diff_info and diff_info['has_diff']:
            preview_group = QGroupBox(f"Text Diff: {len(diff_info['additions'])} additions, {len(diff_info['deletions'])} deletions")
        elif has_backup:
            preview_group = QGroupBox("Detected Changes (Backup → Corrupted)")
        else:
            preview_group = QGroupBox("Recovered Data Preview")
        preview_layout = QVBoxLayout(preview_group)

        preview_text = QTextEdit()
        preview_text.setReadOnly(True)
        preview_text.setMaximumHeight(200)
        preview_text.setStyleSheet("font-family: Consolas, monospace; font-size: 9pt;")

        preview_lines = []

        # Use actual text diff if available
        if diff_info and diff_info['has_diff']:
            preview_lines.append("Changes detected (- = backup, + = corrupted):")
            preview_lines.append("")
            # Show first 30 diff lines
            for line in diff_info['diff_lines'][:30]:
                preview_lines.append(line.rstrip())
            if len(diff_info['diff_lines']) > 30:
                preview_lines.append(f"... and {len(diff_info['diff_lines']) - 30} more lines")

        elif diff_info and not diff_info['has_diff']:
            preview_lines.append("✓ No text differences between backup and corrupted file.")
            preview_lines.append("  This shouldn't happen - the files appear identical.")

        elif has_backup and not diff_info:
            preview_lines.append("Could not compute diff between files.")

        elif recovered_metadata:
            # No backup - just show what we extracted
            preview_lines.append("No backup available for comparison.")
            preview_lines.append(f"Extracted {len(recovered_metadata)} file entries from corrupted data:")
            preview_lines.append("")
            for file_path, meta in list(recovered_metadata.items())[:8]:
                meta_str = ", ".join(f"{k}={v}" for k, v in meta.items() if v)
                preview_lines.append(f"• {Path(file_path).name}: {meta_str or '(no metadata)'}")
            if len(recovered_metadata) > 8:
                preview_lines.append(f"... and {len(recovered_metadata) - 8} more files")
        else:
            preview_lines.append("No metadata could be extracted from the corrupted file.")
            preview_lines.append("A fresh scan will discover all files but user-entered data will be lost.")

        preview_text.setPlainText("\n".join(preview_lines))
        preview_layout.addWidget(preview_text)
        layout.addWidget(preview_group)

        # AI enhancement option (if not configured)
        if not has_ai and recovered_metadata:
            ai_note = QLabel(
                "<i>Tip: Configure an AI API key in Settings → AI Integration to help\n"
                "recover additional metadata from file path patterns.</i>"
            )
            ai_note.setStyleSheet("color: #888;")
            layout.addWidget(ai_note)

        # Buttons
        button_layout = QHBoxLayout()

        # AI Smart Repair button - most powerful option when AI + backup are available
        ai_repair_btn = None
        if has_ai and has_backup and diff_info and diff_info['has_diff']:
            ai_repair_btn = QPushButton("AI Smart Repair")
            ai_repair_btn.setStyleSheet("background-color: #2196F3; color: white; font-weight: bold; padding: 8px 16px;")
            ai_repair_btn.setToolTip("AI analyzes the diff to fix corruption while preserving your new data")
            button_layout.addWidget(ai_repair_btn)

        recover_btn = QPushButton("Recover Project")
        recover_btn.setStyleSheet("background-color: #4CAF50; color: white; font-weight: bold; padding: 8px 16px;")
        if has_ai and has_backup:
            recover_btn.setToolTip("Restore from backup and merge any recoverable changes")

        # Only show backup restore option if backup exists AND is valid
        if backup_valid:
            backup_btn = QPushButton("Restore Backup Only")
            backup_btn.setToolTip(f"Restore from backup ({backup_file_count} files, loses ALL recent changes)")
            button_layout.addWidget(backup_btn)

        cancel_btn = QPushButton("Cancel")

        button_layout.addStretch()
        button_layout.addWidget(recover_btn)
        button_layout.addWidget(cancel_btn)
        layout.addLayout(button_layout)

        # Button handlers
        result = {'action': None}

        def on_ai_repair():
            result['action'] = 'ai_repair'
            dialog.accept()

        def on_recover():
            result['action'] = 'recover'
            dialog.accept()

        def on_backup():
            result['action'] = 'backup'
            dialog.accept()

        def on_cancel():
            result['action'] = 'cancel'
            dialog.reject()

        if ai_repair_btn:
            ai_repair_btn.clicked.connect(on_ai_repair)
        recover_btn.clicked.connect(on_recover)
        cancel_btn.clicked.connect(on_cancel)
        if backup_valid:
            backup_btn.clicked.connect(on_backup)

        dialog.exec()

        # Handle result
        if result['action'] == 'ai_repair' and diff_info:
            # AI Smart Repair - let AI analyze diff and fix the file
            repaired_json, explanation = self._ai_smart_repair(
                backup_path, project_path, diff_info, error_msg
            )

            if repaired_json:
                # Show what AI found and ask for confirmation
                confirm = QMessageBox.question(
                    self, "AI Repair Complete",
                    f"AI Analysis:\n{explanation}\n\nApply this repair?",
                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
                )

                if confirm == QMessageBox.StandardButton.Yes:
                    try:
                        # Write the repaired JSON
                        with open(project_path, 'w') as f:
                            f.write(repaired_json)
                        print(f"[ai-repair] Wrote repaired file: {project_path}")

                        # Load it
                        project_data = self.project_manager.load_project(project_path)
                        self._populate_ui_from_project(project_data)
                        self._log_status_message(f"✓ AI repaired project '{project_name}'", 3000)
                    except Exception as e:
                        self._show_error("Repair Failed", f"Could not apply AI repair:\n{e}")
            else:
                self._show_error("AI Repair Failed", f"AI could not repair the file:\n{explanation}")

        elif result['action'] == 'recover':
            self._perform_project_recovery(
                project_path, project_name, data_directory,
                recovered_metadata, use_ai=has_ai
            )

        elif result['action'] == 'backup' and backup_valid:
            # Restore from backup file (we already validated it can be loaded)
            try:
                import shutil
                shutil.copy2(backup_path, project_path)
                print(f"[recovery] Restored from backup: {backup_path}")

                # Load the restored file (should work since we validated it)
                project_data = self.project_manager.load_project(project_path)
                self._populate_ui_from_project(project_data)
                self._log_status_message(f"✓ Restored project '{project_name}' from backup ({backup_file_count} files)", 3000)
            except Exception as e:
                self._show_error("Restore Failed", f"Could not restore from backup:\n{e}")

        # else: Cancel - do nothing

    def _populate_load_project_combo(self):
        """Populate the project name combo (legacy wrapper)."""
        # Use the new unified method
        current_name = getattr(self, '_current_project_name', None)
        self._populate_project_name_combo(current_name)

    def _populate_ui_from_project(self, project_data):
        """Populate UI with loaded project data."""
        # Set project name and directory
        project_name = project_data['project_name']
        self._current_project_name = project_name
        if hasattr(self, 'projectNameCombo'):
            self.projectNameCombo.setCurrentText(project_name)
        self._project_directory = str(project_data['data_directory'])
        self.directoryPathEdit.setText(self._project_directory)

        # Populate discovered files data
        self._discovered_files_data = project_data['files']

        # Count protocols and check if keywords need extraction
        protocols = set()
        need_keywords = any('keywords_display' not in f for f in self._discovered_files_data)

        if need_keywords:
            total_files = len(self._discovered_files_data)
            # Show progress bar
            self.projectProgressBar.setVisible(True)
            self.projectProgressBar.setValue(0)
            self.projectProgressBar.setFormat(f"Extracting keywords: 0/{total_files} (0%)")

            self._log_status_message(f"Extracting path keywords for {total_files} files...", 0)
            QApplication.processEvents()  # Update status message

            # Extract keywords for files that don't have them
            from core.fast_abf_reader import extract_path_keywords
            base_dir = Path(self._project_directory)

            for idx, file_data in enumerate(self._discovered_files_data):
                if 'keywords_display' not in file_data:
                    # Extract keywords
                    file_path = file_data.get('file_path')
                    if file_path and not Path(file_path).is_absolute():
                        file_path = base_dir / file_path

                    if file_path:
                        path_info = extract_path_keywords(Path(file_path), base_dir)

                        # Format keywords for display - show relative path as the primary keyword
                        keywords_display = []

                        # Add relative path (the main keyword showing file location)
                        if path_info.get('relative_path'):
                            keywords_display.append(path_info['relative_path'])
                        elif path_info['subdirs']:
                            # Fallback to subdirs if relative_path not available
                            keywords_display.append('/'.join(path_info['subdirs']))

                        # Also add power levels and animal IDs if detected
                        if path_info['power_levels']:
                            keywords_display.extend(path_info['power_levels'])
                        if path_info['animal_ids']:
                            keywords_display.extend([f"ID:{id}" for id in path_info['animal_ids']])

                        file_data['path_keywords'] = path_info
                        file_data['keywords_display'] = ', '.join(keywords_display) if keywords_display else ''

                # Update progress every 50 files
                if idx % 50 == 0 or idx == total_files - 1:
                    progress_pct = int((idx + 1) / total_files * 100)
                    self.projectProgressBar.setValue(progress_pct)
                    self.projectProgressBar.setFormat(f"Extracting keywords: {idx + 1}/{total_files} ({progress_pct}%)")
                    QApplication.processEvents()

        # Show progress for data preparation
        total_files = len(self._discovered_files_data)
        self.projectProgressBar.setVisible(True)
        self.projectProgressBar.setValue(0)
        self.projectProgressBar.setFormat(f"Loading files: 0/{total_files} (0%)")
        QApplication.processEvents()

        # Ensure all files have required fields
        for i, file_data in enumerate(self._discovered_files_data):
            # Ensure file has all needed fields
            if 'channel' not in file_data:
                file_data['channel'] = ''
            if 'stim_channel' not in file_data:
                file_data['stim_channel'] = ''
            if 'events_channel' not in file_data:
                file_data['events_channel'] = ''
            if 'stim_channels' not in file_data:
                file_data['stim_channels'] = []
            if 'stim_frequency' not in file_data:
                file_data['stim_frequency'] = ''
            if 'strain' not in file_data:
                file_data['strain'] = ''
            if 'stim_type' not in file_data:
                file_data['stim_type'] = ''
            if 'power' not in file_data:
                # Try to auto-fill from keywords
                path_kw = file_data.get('path_keywords', {})
                file_data['power'] = path_kw.get('power_levels', [''])[0] if path_kw.get('power_levels') else ''
            if 'sex' not in file_data:
                file_data['sex'] = ''
            if 'animal_id' not in file_data:
                # Try to auto-fill from keywords
                path_kw = file_data.get('path_keywords', {})
                file_data['animal_id'] = path_kw.get('animal_ids', [''])[0] if path_kw.get('animal_ids') else ''
            if 'status' not in file_data:
                file_data['status'] = 'pending'
            # Export tracking fields
            if 'export_path' not in file_data:
                file_data['export_path'] = ''
            if 'export_date' not in file_data:
                file_data['export_date'] = ''
            if 'export_version' not in file_data:
                file_data['export_version'] = ''
            if 'exports' not in file_data:
                file_data['exports'] = {
                    'npz': False,
                    'timeseries_csv': False,
                    'breaths_csv': False,
                    'events_csv': False,
                    'pdf': False,
                    'session_state': False,
                    'ml_training': False,
                }

            if file_data.get('protocol'):
                protocols.add(file_data['protocol'])

            # Update progress every 50 files
            if i % 50 == 0 or i == total_files - 1:
                progress_pct = int((i + 1) / total_files * 100)
                self.projectProgressBar.setValue(progress_pct)
                self.projectProgressBar.setFormat(f"Loading files: {i + 1}/{total_files} ({progress_pct}%)")
                QApplication.processEvents()

        # Sync master file list with discovered files
        self._master_file_list = [dict(f) for f in self._discovered_files_data]

        # Rebuild the table using the model
        self._rebuild_table_from_master_list()

        # Auto-fit columns
        self._auto_fit_table_columns()

        # Hide progress bar
        self.projectProgressBar.setVisible(False)
        self.projectProgressBar.setValue(0)

        # Update summary
        summary_text = f"Summary: {len(self._discovered_files_data)} ABF files | {len(protocols)} protocols"
        self.summaryLabel.setText(summary_text)

        # Restore notes data if available
        if project_data.get('notes_directory'):
            self._notes_directory = project_data['notes_directory']
            if hasattr(self, 'notesFolderEdit'):
                self.notesFolderEdit.setText(self._notes_directory)
        if project_data.get('notes_files'):
            self._notes_files_data = project_data['notes_files']
            # Populate notes table
            self._populate_notes_table_from_data()

        # Populate consolidation source list with analyzed files from project
        self._populate_consolidation_source_list()

        # Note: Old experiment structure is ignored - we now use the flat master file list
        # If old project had experiments, their task data could be migrated here if needed

    def _populate_notes_table_from_data(self):
        """Populate notes table from stored notes files data."""
        if not hasattr(self, '_notes_model') or not self._notes_files_data:
            return

        self._notes_model.removeRows(0, self._notes_model.rowCount())

        from PyQt6.QtGui import QStandardItem
        from PyQt6.QtCore import Qt

        for file_info in self._notes_files_data:
            # Create checkbox item for "Use" column
            use_item = QStandardItem()
            use_item.setCheckable(True)
            use_as_notes = file_info.get('use_as_notes', True)
            use_item.setCheckState(Qt.CheckState.Checked if use_as_notes else Qt.CheckState.Unchecked)
            use_item.setEditable(False)

            # File name with path stored
            name_item = QStandardItem(file_info.get('name', ''))
            name_item.setData(file_info.get('path', ''), Qt.ItemDataRole.UserRole)

            # Matches column
            match_count = file_info.get('match_count', 0)
            matches_item = QStandardItem(str(match_count) if match_count > 0 else "—")
            matches_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)

            type_item = QStandardItem(file_info.get('type', ''))
            location_item = QStandardItem(file_info.get('location', ''))
            size_item = QStandardItem(file_info.get('size', ''))
            modified_item = QStandardItem(file_info.get('modified', ''))

            self._notes_model.appendRow([use_item, name_item, matches_item, type_item, location_item, size_item, modified_item])

    # NOTE: _load_experiments_from_project method removed - no longer using experiment tree

    def _ask_locate_project(self, project_name, expected_path):
        """
        Ask user if they want to locate a missing project file.

        Returns:
            "locate" if user wants to locate, "cancel" otherwise
        """
        from PyQt6.QtWidgets import QMessageBox
        msg = QMessageBox(self)
        msg.setIcon(QMessageBox.Icon.Warning)
        msg.setWindowTitle("Project Not Found")
        msg.setText(f"Cannot find project '{project_name}'")
        msg.setInformativeText(f"Expected location:\n{expected_path}\n\nWould you like to locate the project file?")
        msg.setStandardButtons(QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.Cancel)
        msg.setDefaultButton(QMessageBox.StandardButton.Yes)

        result = msg.exec()
        if result == QMessageBox.StandardButton.Yes:
            return "locate"
        return "cancel"

if __name__ == "__main__":
    from PyQt6.QtWidgets import QSplashScreen, QProgressBar, QVBoxLayout, QLabel, QWidget
    from PyQt6.QtGui import QPixmap
    from PyQt6.QtCore import Qt, QTimer

    app = QApplication(sys.argv)

    # Create splash screen
    # Try to load icon (with fallback path handling)
    splash_paths = [
        Path(__file__).parent / "images" / "plethapp_splash_dark-01.png",
        Path(__file__).parent / "images" / "plethapp_splash.png",
        Path(__file__).parent / "images" / "plethapp_thumbnail_dark_round.ico",
        Path(__file__).parent / "assets" / "plethapp_thumbnail_dark_round.ico",
    ]

    splash_pix = None
    for splash_path in splash_paths:
        if splash_path.exists():
            splash_pix = QPixmap(str(splash_path))
            break

    if splash_pix is None or splash_pix.isNull():
        # Fallback: create simple splash with text
        splash_pix = QPixmap(200, 150)
        splash_pix.fill(Qt.GlobalColor.darkGray)

    # Scale to smaller size for faster display
    splash_pix = splash_pix.scaled(150, 150, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)

    splash = QSplashScreen(splash_pix, Qt.WindowType.WindowStaysOnTopHint)
    splash.setWindowFlags(Qt.WindowType.WindowStaysOnTopHint | Qt.WindowType.FramelessWindowHint)

    # Add loading message
    splash.showMessage(
        "Loading PhysioMetrics...",
        Qt.AlignmentFlag.AlignBottom | Qt.AlignmentFlag.AlignCenter,
        Qt.GlobalColor.white
    )
    splash.show()
    app.processEvents()

    # Create main window (this is where the loading time happens)
    w = MainWindow()

    # Close splash and show main window
    splash.finish(w)
    w.show()

    sys.exit(app.exec())